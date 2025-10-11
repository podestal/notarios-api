from django.db import connection
from django.http import HttpResponse, JsonResponse
from docx import Document
import io
from decimal import Decimal
from datetime import datetime
from rest_framework.response import Response
from notaria import models
from .utils import (
    NumberToLetterConverter,
    DocumentFormatter,
    PlaceholderProcessor,
    DocxTemplateProcessor,
    DataValidator,
    TemplateManager,
)


class TestamentosDocumentService:
    """
    Service to generate Testamentos documents
    Mirrors: PHP testamento() function
    Reuses all utilities from utils.py
    """

    def __init__(
        self,
        letras=None,
        formatter=None,
        placeholder_processor=None,
        docx_template_processor=None,
        data_validator=None,
        template_manager=None,
    ):
        """
        Initialize with dependency injection - SAME as other services
        All utilities are reusable!
        """
        self.letras = letras or NumberToLetterConverter()
        self.formatter = DocumentFormatter(self.letras)
        self.placeholder_processor = placeholder_processor or PlaceholderProcessor()
        self.docx_template_processor = docx_template_processor or DocxTemplateProcessor()
        self.data_validator = data_validator or DataValidator()
        self.template_manager = template_manager or TemplateManager()

    def generate_testamentos_document(self, template_id, kardex, action, mode):
        """
        Main entry point for generating Testamentos documents
        Mirrors: PHP testamento() function
        
        FLOW: (Same as Escrituras, No Contenciosos, and Garantias)
        1. Get template information
        2. Download template from R2 (or existing document for actualizar)
        3. Fetch data from database
        4. Process data into sections
        5. Replace placeholders
        6. Upload to R2
        7. Return HTTP response
        """
        
        # Handle "actualizar" action
        if action == "actualizar":
            return self._update_existing_document(kardex, mode)

        # STEP 1: Get template info
        template_info = self._get_template_info(template_id)

        # STEP 2: Download template from R2
        template_bytes = self.template_manager.get_template_from_r2(
            template_id, template_info["filename"]
        )

        # STEP 3: Fetch data from database
        raw_data = self._consulta_testamentos(kardex, action, template_id)

        # STEP 4: Format data using REUSABLE utilities
        data_documento = self.formatter.format_document_data(raw_data)
        data_vehiculos = self.formatter.format_vehicle_data(raw_data)
        data_pagos = self.formatter.format_payment_data(raw_data)
        data_escrituracion = self.formatter.format_escrituracion_data(raw_data)
        data_contratantes = self.formatter.format_contractor_data(raw_data)
        data_company = self.formatter.format_company_data(raw_data)

        # STEP 5: Combine all data
        final_data = self.formatter.combine_all_data(
            data_documento,
            data_vehiculos,
            data_pagos,
            data_escrituracion,
            data_contratantes,
            data_company,
        )

        # STEP 6: Replace placeholders (try docxtpl, fallback to PlaceholderProcessor)
        try:
            processed_bytes = self.docx_template_processor.replace_placeholders(
                template_bytes, final_data
            )
            if processed_bytes:
                buffer = io.BytesIO(processed_bytes)
            else:
                raise Exception("DocxTemplateProcessor returned None")
        except Exception as e:
            print(f"DEBUG: Falling back to PlaceholderProcessor for Testamentos")
            buffer = io.BytesIO(template_bytes)
            doc = Document(buffer)

            self.placeholder_processor.replace_placeholders(doc, final_data)
            self.placeholder_processor.clean_unfilled_placeholders(doc)

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

        # STEP 7: Upload to R2
        self.template_manager.upload_document_to_r2(buffer, kardex)
        
        # STEP 8: Return HTTP response
        filename = f"__PROY__{kardex}.docx"
        return self._create_response_from_buffer(buffer, filename, kardex, mode)

    def _update_existing_document(self, kardex, mode):
        """
        Update existing Testamentos document with escrituracion data
        Mirrors: PHP actualizar action
        REUSES: TemplateManager.update_document_escrituracion()
        """
        # STEP 1: Validate numescritura exists
        kardex_obj = models.Kardex.objects.get(kardex=kardex)
        if not kardex_obj.numescritura:
            raise ValueError("ERROR: FALTA GRABAR NUMERO DE ACTA")
        
        # STEP 2: Get escrituracion data
        raw_data = self._consulta_testamentos(kardex, "actualizar", None)
        data_escrituracion = self.formatter.format_escrituracion_data(raw_data)
        
        # STEP 3: Update document using REUSABLE generic method
        output_buffer = self.template_manager.update_document_escrituracion(
            kardex, 
            data_escrituracion, 
            self.placeholder_processor
        )
        
        # STEP 4: Return HTTP response
        filename = f"__PROY__{kardex}.docx"
        return self._create_response_from_buffer(output_buffer, filename, kardex, mode)

    def _get_template_info(self, template_id):
        """Get template information from database"""
        template = models.TplTemplate.objects.get(pktemplate=template_id)
        return {"filename": template.filename}

    def _create_response_from_buffer(self, buffer, filename, kardex, mode):
        """Create HTTP response from buffer"""
        if mode == "open":
            response = JsonResponse(
                {
                    "status": "success",
                    "mode": "open",
                    "filename": filename,
                    "kardex": kardex,
                    "message": "Document generated and ready to open in Word",
                }
            )
            response["Access-Control-Allow-Origin"] = "*"
            return response
        else:
            response = HttpResponse(
                buffer.read(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = f'inline; filename="{filename}"'
            response["Content-Length"] = str(buffer.getbuffer().nbytes)
            response["Access-Control-Allow-Origin"] = "*"
            return response

    def _consulta_testamentos(self, num_kardex, action, template_id):
        """
        SQL Query for Testamentos documents
        Mirrors: PHP consulta_transferencia() in testamento.php
        
        NOTE: The query structure is IDENTICAL to other services
        """
        
        # Get idtipoacto from kardex
        idtipoacto = models.Kardex.objects.get(kardex=num_kardex).idtipkar
        
        # Same query as other services
        query = """
            SELECT k.idkardex as id_kardex,
                k.kardex,
                k.numescritura as numero_escritura,
                k.fechaescritura as fecha_escritura,
                k.txa_minuta as registro_escritura,
                CURRENT_DATE() as fecha_generado,
                k.fechaconclusion as fecha_conclusion,
                k.numminuta as numero_minuta,
                k.kardexconexo as kardex_conexo,
                k.folioini as folio_inicial, 
                k.foliofin as folio_final, 
                k.papelini as papel_inicial, 
                k.papelfin as papel_final,
                (SELECT desacto FROM tiposdeacto WHERE idtipoacto=%s) as acto,
                (SELECT fileName FROM tpl_template WHERE pkTemplate=%s) as plantilla,
                (SELECT urlTemplate FROM tpl_template WHERE pkTemplate=%s) as url_plantilla,
                k.fechaingreso as fecha_ingreso,
                k.responsable_new as usuario,
                abo.razonsocial as abogado,
                abo.matricula as matricula,
                usu.dni as dni_usuario,
                GROUP_CONCAT(c2.idcliente) as id_cliente,
                GROUP_CONCAT(IF(c2.conyuge='','NO',c2.conyuge)) as id_conyuge,
                GROUP_CONCAT(cxa.idcontratante) as id_contratante,
                GROUP_CONCAT(TRIM(CONCAT(IFNULL(c2.prinom, ''), ' ', IFNULL(c2.segnom, ''), IF(c2.segnom='','',' ') ,IFNULL(c2.apepat, ''), ' ',IFNULL(c2.apemat, ''),
            IFNULL(c2.razonsocial, '')))) AS nombres,
                GROUP_CONCAT(cxa.uif) as uif,
                GROUP_CONCAT(ac.condicion) as condicion,
                GROUP_CONCAT(IF(n.descripcion IS NULL OR n.descripcion='','EMPRESA',n.descripcion)) as nacionalidad,
                GROUP_CONCAT(td.destipdoc) as tipo_documento,
                GROUP_CONCAT(c2.numdoc) AS numero_documento,
                GROUP_CONCAT(UPPER(c2.profesion_plantilla)) AS ocupacion,
                GROUP_CONCAT(IF(tec.desestcivil IS NULL OR tec.desestcivil='','EMPRESA',tec.desestcivil)) as estado_civil,
                GROUP_CONCAT(IF(c2.tipper='N',c2.direccion,c2.domfiscal) SEPARATOR ',,') as direccion,
                GROUP_CONCAT(IFNULL(u.codpto, '')) as codigo_departamento,
                GROUP_CONCAT(IFNULL(u.coddis, '')) as codigo_distrito,
                GROUP_CONCAT(IFNULL(u.codprov, '')) as codigo_provincia,
                GROUP_CONCAT(IFNULL(IF(SUBSTRING_INDEX(c2.ubigeo_plantilla, '/', -1)='',u.nomdis,SUBSTRING_INDEX(c2.ubigeo_plantilla, '/', -1)),(IFNULL(u.nomdis, '')))) AS distrito,
                GROUP_CONCAT(IFNULL(u.nomprov, '')) as provincia,
                GROUP_CONCAT(IFNULL(u.nomdpto, '')) as departamento,
                GROUP_CONCAT(c2.sexo) AS sexo,
                GROUP_CONCAT(c2.tipper) as tipo_persona,
                GROUP_CONCAT(IF(cn.firma = '0', 'NO', 'SI')) AS firma,
                GROUP_CONCAT(cn.firma) as n_firma,
                GROUP_CONCAT(cn.tiporepresentacion) AS tipo_representacion,
                dv.numplaca AS placa, 
                dv.marca AS marca, 
                dv.clase AS clase,
                dv.anofab AS anio,
                dv.numserie AS serie, 
                dv.color AS color,
                dv.motor AS motor, 
                dv.modelo AS modelo, 
                dv.carroceria AS carroceria,
                dv.pregistral as partida,
                dv.fecinsc AS fecha_inscripcion,
                dv.combustible AS combustible,
                UPPER(sr.dessede) AS sede,
                UPPER(sr.num_zona) AS numero_zona,
                pat.importetrans AS precio , 
                pat.idmon AS moneda,
                pat.exhibiomp, 
                pat.idoppago, 
                uif.descripcion AS medio_pago, 
                mon.simbolo as simbolo_moneda,
                mon.desmon as descripcion_moneda,
                mp.desmpagos as descripcion_medio_pago,
                mp.sunat as sunat_medio_pago,
                GROUP_CONCAT(cnr.idcontratanterp) as id_empresa,
                GROUP_CONCAT(TRIM(CONCAT(IFNULL(cr2.prinom, ''), ' ', IFNULL(cr2.segnom, ''), ' ',IFNULL(cr2.apepat, ''), ' ',IFNULL(cr2.apemat, ''),
            IFNULL(cr2.razonsocial, '')))) AS nombre_empresa,
                GROUP_CONCAT(cr2.tipper) as tipo_persona_empresa,
                GROUP_CONCAT(acr.condicion) as condicion_empresa,
                GROUP_CONCAT(tdr.destipdoc) as tipo_documento_empresa,
                GROUP_CONCAT(cr2.numdoc) AS numero_documento_empresa,
                GROUP_CONCAT(cr2.domfiscal) as domicilio_empresa,
                GROUP_CONCAT(ur.nomdis) as distrito_empresa,
                GROUP_CONCAT(ur.nomprov) as provincia_empresa,
                GROUP_CONCAT(ur.nomdpto) as departamento_empresa,
                GROUP_CONCAT(srr2.zona_depar SEPARATOR ',,') as oficina_registral,
                GROUP_CONCAT(cr2.numpartida) as numero_partida,
                dmp.foperacion as fecha_operacion,
                dmp.documentos as documentos,
                ban.desbanco as banco
            FROM kardex as k
            LEFT JOIN tb_abogado as abo on abo.idabogado=k.idabogado
            LEFT JOIN usuarios as usu on usu.idusuario=k.idusuario
            LEFT JOIN contratantesxacto as cxa on cxa.kardex=k.kardex
            LEFT JOIN actocondicion as ac ON cxa.idcondicion=ac.idcondicion
            LEFT JOIN contratantes cn ON cxa.idcontratante = cn.idcontratante
            LEFT JOIN cliente2 as c2 on c2.idcontratante=cxa.idcontratante
            LEFT JOIN nacionalidades as n on n.idnacionalidad=c2.nacionalidad
            LEFT JOIN tipodocumento as td ON td.idtipdoc = c2.idtipdoc
            LEFT OUTER JOIN tipoestacivil as tec ON tec.idestcivil = c2.idestcivil
            LEFT OUTER JOIN ubigeo as u ON u.coddis = c2.idubigeo
            LEFT JOIN detallevehicular as dv ON dv.kardex=k.kardex
            LEFT JOIN sedesregistrales as sr ON sr.idsedereg=dv.idsedereg
            LEFT JOIN patrimonial as pat ON pat.kardex=k.kardex
            LEFT JOIN fpago_uif as uif ON uif.id_fpago = pat.fpago
            LEFT JOIN monedas as mon ON mon.idmon = pat.idmon
            LEFT JOIN detallemediopago as dmp ON pat.kardex = dmp.kardex
            LEFT JOIN mediospago as mp ON dmp.codmepag = mp.codmepag
            LEFT JOIN contratantes as cnr ON cxa.idcontratante = cnr.idcontratante
            LEFT JOIN contratantesxacto as cxar on cxar.idcontratante=cnr.idcontratanterp
            LEFT JOIN actocondicion as acr ON acr.idcondicion=cxar.idcondicion
            LEFT JOIN cliente2 as cr2 on cr2.idcontratante=cnr.idcontratanterp
            left JOIN nacionalidades as nr on nr.idnacionalidad=cr2.nacionalidad
            LEFT JOIN sedesregistrales as srr2 on srr2.idsedereg=cr2.idsedereg
            LEFT JOIN tipodocumento as tdr ON tdr.idtipdoc = cr2.idtipdoc
            LEFT OUTER JOIN tipoestacivil as tecr ON tecr.idestcivil = cr2.idestcivil
            LEFT OUTER JOIN ubigeo as ur ON ur.coddis = cr2.idubigeo
            LEFT JOIN  bancos as ban ON ban.idbancos=dmp.idbancos
            WHERE k.kardex=%s and (c2.tipper='N')
            GROUP BY k.idkardex, dmp.detmp LIMIT 1
        """
        with connection.cursor() as cursor:
            cursor.execute(query, [idtipoacto, template_id, template_id, num_kardex])
            desc = cursor.description
            row = cursor.fetchone()
            if not row:
                return None
            result = dict(zip([col[0] for col in desc], row))
            
            # Also query company data if needed
            company_query = """
                SELECT 
                    c2.razonsocial as nombre_empresa_constitucion,
                    c2.domfiscal as domicilio_empresa_constitucion,
                    c2.tipper as tipo_persona_empresa_constitucion,
                    c2.numdoc as numero_documento_empresa_constitucion,
                    c2.numpartida as numero_partida_constitucion
                FROM contratantesxacto as cxa
                LEFT JOIN cliente2 as c2 on c2.idcontratante=cxa.idcontratante
                WHERE cxa.kardex=%s and c2.tipper='J'
                LIMIT 1
            """
            cursor.execute(company_query, [num_kardex])
            company_row = cursor.fetchone()
            
            if company_row and company_row[0]:
                result['nombre_empresa_constitucion'] = company_row[0]
                result['domicilio_empresa_constitucion'] = company_row[1]
                result['tipo_persona_empresa_constitucion'] = company_row[2]
                result['numero_documento_empresa_constitucion'] = company_row[3]
                result['numero_partida_constitucion'] = company_row[4]
            
            return result


class TestamentosReportService:
    """Service for generating testamentos reports matching PHP script format"""

    def _sanitize_cell_value(self, value):
        """Sanitize cell values to prevent Excel corruption"""
        if value is None:
            return ""

        # Convert to string and strip whitespace
        val_str = str(value).strip()

        # Remove control characters that can cause XML corruption
        import re

        val_str = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]", "", val_str)

        # Limit to Excel's cell character limit
        val_str = val_str[:32767]

        return val_str

    def _format_date_in_spanish(self, date_input):
        """Convert date to Spanish format like 'LUNES, 15 DE ENERO DEL 2025'"""
        try:
            # Handle both datetime objects and date strings
            if hasattr(date_input, "strftime"):
                date_obj = date_input
            else:
                from datetime import datetime

                # Try different formats
                if "-" in str(date_input) and len(str(date_input).split("-")[0]) == 4:
                    date_obj = datetime.strptime(str(date_input), "%Y-%m-%d")
                elif "/" in str(date_input):
                    date_obj = datetime.strptime(str(date_input), "%d/%m/%Y")
                else:
                    return str(date_input)

            # Spanish day names
            dias = ["LUNES", "MARTES", "MIÉRCOLES", "JUEVES", "VIERNES", "SÁBADO", "DOMINGO"]
            # Spanish month names
            meses = [
                "ENERO",
                "FEBRERO",
                "MARZO",
                "ABRIL",
                "MAYO",
                "JUNIO",
                "JULIO",
                "AGOSTO",
                "SEPTIEMBRE",
                "OCTUBRE",
                "NOVIEMBRE",
                "DICIEMBRE",
            ]

            dia_semana = dias[date_obj.weekday()]
            dia = date_obj.day
            mes = meses[date_obj.month - 1]
            anio = date_obj.year

            return f"{dia_semana}, {dia} DE {mes} DEL {anio}"
        except:
            return str(date_input)

    def _extract_year_from_date(self, date_input):
        """Extract year from date string DD/MM/YYYY or YYYY-MM-DD or datetime object"""
        try:
            # Handle datetime objects
            if hasattr(date_input, "year"):
                return str(date_input.year)

            # Handle strings
            date_str = str(date_input)
            # Try to parse as YYYY-MM-DD first
            if "-" in date_str and len(date_str.split("-")[0]) == 4:
                return date_str.split("-")[0]
            # Try to parse as DD/MM/YYYY
            elif "/" in date_str:
                return date_str.split("/")[-1]
            else:
                from datetime import datetime

                return str(datetime.now().year)
        except:
            from datetime import datetime

            return str(datetime.now().year)

    def _format_date_for_display(self, date_obj):
        """Format date like PHP script logic"""
        try:
            if hasattr(date_obj, "strftime"):
                return date_obj.strftime("%d/%m/%Y")
            return str(date_obj)
        except:
            return str(date_obj)

    def _get_notary_info(self):
        """Get notary configuration info"""
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT nombre, apellido, telefono, correo, ruc, direccion, distrito FROM confinotario"
            )
            result = cursor.fetchone()
            if result:
                return {
                    "nombre": f"{result[0]} {result[1]}",
                    "telefono": result[2],
                    "correo": result[3],
                    "ruc": result[4],
                    "direccion": result[5],
                    "distrito": result[6],
                }
            return {
                "nombre": "NOTARIO",
                "telefono": "(051) 326609",
                "correo": "",
                "ruc": "10024231572",
                "direccion": "JR.BOLIVAR NRO. 340",
                "distrito": "JULIACA",
            }

    def _get_report_data(self, desde, hasta):
        """Fetch data for the report matching PHP query"""
        try:
            from django.db import connection
            from datetime import datetime

            import time

            start_time = time.time()

            # Convert dates to proper format for Django ORM
            desde_dt = None
            hasta_dt = None

            try:
                if isinstance(desde, str):
                    if "-" in desde and len(desde.split("-")[0]) == 4:
                        desde_dt = datetime.strptime(desde, "%Y-%m-%d")
                    else:
                        desde_dt = datetime.strptime(desde, "%d/%m/%Y")
                else:
                    desde_dt = desde
            except (ValueError, TypeError) as e:
                return []

            try:
                if isinstance(hasta, str):
                    if "-" in hasta and len(hasta.split("-")[0]) == 4:
                        hasta_dt = datetime.strptime(hasta, "%Y-%m-%d")
                    else:
                        hasta_dt = datetime.strptime(hasta, "%d/%m/%Y")
                else:
                    hasta_dt = hasta
            except (ValueError, TypeError) as e:
                return []

            # Validate dates are not None
            if desde_dt is None or hasta_dt is None:
                return []

            # Set session variables for optimization
            with connection.cursor() as cursor:
                cursor.execute("SET SESSION group_concat_max_len = 1000000")
                cursor.execute("SET SESSION sql_mode = 'NO_AUTO_VALUE_ON_ZERO'")
                cursor.execute("SET SESSION sort_buffer_size = 2097152")

                # Try to create indexes if they don't exist (for performance)
                try:
                    cursor.execute(
                        "CREATE INDEX IF NOT EXISTS idx_kardex_tipkar_fecha ON kardex(idtipkar, fechaescritura, nc)"
                    )
                    cursor.execute(
                        "CREATE INDEX IF NOT EXISTS idx_kardex_fecha ON kardex(fechaescritura)"
                    )
                except Exception as e:
                    pass

            # Ultra-simple query to get basic data first
            query = """
                SELECT 
                    k.fechaescritura,
                    k.kardex,
                    k.contrato,
                    k.numescritura,
                    k.numminuta,
                    k.folioini,
                    k.numescritura as numescritura2
                FROM kardex as k 
                WHERE k.idtipkar='5' 
                    AND k.fechaescritura <> '' 
                    AND k.fechaescritura >= %s
                    AND k.fechaescritura <= %s
                ORDER BY k.fechaescritura ASC, k.numescritura ASC, k.numminuta ASC
                LIMIT 2000
            """

            with connection.cursor() as cursor:
                cursor.execute(
                    query, [desde_dt.strftime("%Y-%m-%d"), hasta_dt.strftime("%Y-%m-%d")]
                )
                testamentos = []
                rows = cursor.fetchall()

                if len(rows) > 0:
                    print(f"DEBUG: Sample row: {rows[0]}")
                else:

                    # Let's check if there are any testamentos at all
                    cursor.execute("SELECT COUNT(*) FROM kardex WHERE idtipkar='5'")
                    count_result = cursor.fetchone()

                    # Let's also check what idtipkar values exist
                    cursor.execute("SELECT DISTINCT idtipkar FROM kardex ORDER BY idtipkar")
                    tipkar_results = cursor.fetchall()

                    # Let's check testamentos without date filter
                    cursor.execute("SELECT COUNT(*) FROM kardex WHERE idtipkar='5' AND nc=0")
                    count_nc_result = cursor.fetchone()

                    # Let's see what dates we actually have
                    cursor.execute(
                        """
                        SELECT 
                            fechaescritura,
                            kardex,
                            numescritura,
                            DATE(fechaescritura) as fecha_parsed,
                            YEAR(fechaescritura) as anio
                        FROM kardex 
                        WHERE idtipkar='5' AND nc=0 
                        ORDER BY fechaescritura
                    """
                    )
                    all_dates = cursor.fetchall()

                    # Let's check the query with actual values
                    test_query = f"""
                        SELECT COUNT(*) 
                        FROM kardex 
                        WHERE idtipkar='5' 
                            AND nc=0 
                            AND fechaescritura <> ''
                            AND DATE(fechaescritura) >= DATE('{desde_dt.strftime("%Y-%m-%d")}')
                            AND DATE(fechaescritura) <= DATE('{hasta_dt.strftime("%Y-%m-%d")}')
                    """
                    cursor.execute(test_query)
                    test_count = cursor.fetchone()

                    # Let's see what dates the testamentos actually have
                    cursor.execute(
                        "SELECT fechaescritura, numescritura FROM kardex WHERE idtipkar='5' AND nc=0 ORDER BY fechaescritura LIMIT 10"
                    )
                    date_results = cursor.fetchall()

                    # Check specifically for 2022
                    cursor.execute(
                        "SELECT COUNT(*) FROM kardex WHERE idtipkar='5' AND nc=0 AND YEAR(fechaescritura) = 2022"
                    )
                    count_2022 = cursor.fetchone()

                    # Check date format issue - maybe fechaescritura is a string?
                    cursor.execute(
                        "SELECT DATE_FORMAT(fechaescritura, '%Y-%m-%d'), numescritura FROM kardex WHERE idtipkar='5' AND nc=0 AND YEAR(fechaescritura) = 2022"
                    )
                    date_2022_results = cursor.fetchall()

                    # Let's see what years the testamentos actually have
                    cursor.execute(
                        "SELECT DISTINCT YEAR(fechaescritura) as year, COUNT(*) FROM kardex WHERE idtipkar='5' AND nc=0 GROUP BY YEAR(fechaescritura) ORDER BY year"
                    )
                    all_years = cursor.fetchall()

                # Process each kardex with optimized processing
                processing_start = time.time()
                testamentos = []
                for row in rows:
                    kardex = row[1]

                    # Get contractors for this specific kardex using the PHP logic
                    # Testador/Otorgante query
                    testador_query = """
                        SELECT cliente2.nombre 
                        FROM contratantesxacto 
                        INNER JOIN cliente2 ON cliente2.idcontratante = contratantesxacto.idcontratante
                        INNER JOIN actocondicion ON contratantesxacto.idcondicion = actocondicion.idcondicion
                        WHERE contratantesxacto.kardex = %s 
                            AND (actocondicion.condicion LIKE '%%TESTADOR%%' 
                                 OR actocondicion.condicion LIKE '%%OTORGANTE%%')
                    """

                    # Beneficiario/Otorgado query
                    beneficiario_query = """
                        SELECT cliente2.nombre 
                        FROM contratantesxacto 
                        INNER JOIN cliente2 ON cliente2.idcontratante = contratantesxacto.idcontratante
                        INNER JOIN actocondicion ON contratantesxacto.idcondicion = actocondicion.idcondicion
                        WHERE contratantesxacto.kardex = %s 
                            AND (actocondicion.condicion LIKE '%%BENEFICIARIO%%' 
                                 OR actocondicion.condicion LIKE '%%OTORGADO%%')
                    """

                    # Execute queries
                    cursor.execute(testador_query, [kardex])
                    testadores = cursor.fetchall()

                    cursor.execute(beneficiario_query, [kardex])
                    beneficiarios = cursor.fetchall()

                    # Process contractors based on PHP logic
                    conteo1 = len(testadores)
                    conteo2 = len(beneficiarios)

                    testador_names = []
                    beneficiario_names = []

                    # PHP logic: if conteo1>0 && conteo2==0: show testadores
                    # if conteo2>0 && conteo1==0: show beneficiarios
                    # if conteo1>0 && conteo2>0: show testadores
                    if conteo1 > 0 and conteo2 == 0:
                        testador_names = [t[0] for t in testadores]
                    elif conteo2 > 0 and conteo1 == 0:
                        beneficiario_names = [b[0] for b in beneficiarios]
                    elif conteo1 > 0 and conteo2 > 0:
                        testador_names = [t[0] for t in testadores]
                    else:
                        pass

                    # Clean contract name like in PHP - optimized string operations
                    contrato_raw = row[2] or ""
                    contrato_clean = contrato_raw.replace("/", "").upper()

                    # Optimized date formatting
                    fecha_str = row[0]
                    fecha_formatted = (
                        f"{fecha_str[8:10]}/{fecha_str[5:7]}/{fecha_str[0:4]}"
                        if len(fecha_str) >= 10
                        else fecha_str
                    )

                    # Handle minuta like in PHP
                    minuta = row[4] if row[4] else "S/M"

                    testament_data = {
                        "numero_escritura": row[3],
                        "fecha": fecha_formatted,
                        "testador": ", ".join(testador_names).upper(),
                        "beneficiario": ", ".join(beneficiario_names).upper(),
                        "contrato": contrato_clean,
                        "folio": row[5],
                        "minuta": minuta,
                    }

                    testamentos.append(testament_data)

                return testamentos

        except Exception as e:
            import traceback

            traceback.print_exc()
            return []

    def generate_excel_report(self, desde, hasta):
        """Generate Excel report matching PHP script format"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            import io

            import time

            report_start = time.time()

            # Get data
            report_data = self._get_report_data(desde, hasta)

            notary_info = self._get_notary_info()

            anio = self._extract_year_from_date(hasta)

            # Create workbook and worksheet
            wb = Workbook()
            ws = wb.active
            ws.title = "TESTAMENTOS"

            # Styles
            title_font = Font(name="Arial", size=18.5, bold=True)
            header_font = Font(name="Arial", size=12, bold=True)
            data_font = Font(name="Arial", size=12)
            center_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            left_alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            right_alignment = Alignment(horizontal="right", vertical="top", wrap_text=True)

            # Borders
            thin_border = Border(
                left=Side(border_style="thin"),
                right=Side(border_style="thin"),
                top=Side(border_style="thin"),
                bottom=Side(border_style="thin"),
            )
            no_border = Border(
                left=Side(style=None),
                right=Side(style=None),
                top=Side(style=None),
                bottom=Side(style=None),
            )

            # Title section
            ws.merge_cells("A1:F1")
            ws["A1"] = "INDICE CRONOLOGICO - TESTAMENTOS"
            ws["A1"].font = title_font
            ws["A1"].alignment = center_alignment
            ws["A1"].border = no_border

            ws.merge_cells("A2:F2")
            ws["A2"] = f"AÑO {anio}"
            ws["A2"].font = title_font
            ws["A2"].alignment = center_alignment
            ws["A2"].border = no_border

            # Notary info section
            row = 4
            ws[f"A{row}"] = "NOTARIA"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"C{row}"] = f': {self._sanitize_cell_value(notary_info["nombre"])}'
            ws[f"C{row}"].font = data_font
            ws[f"C{row}"].border = no_border

            row += 1
            ws[f"A{row}"] = "DIRECCION"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"C{row}"] = f': {notary_info["direccion"]}'
            ws[f"C{row}"].font = data_font
            ws[f"C{row}"].border = no_border
            ws[f"D{row}"] = "TELEFONO"
            ws[f"D{row}"].font = header_font
            ws[f"D{row}"].border = no_border
            ws[f"E{row}"] = f': {notary_info["telefono"]}'
            ws[f"E{row}"].font = data_font
            ws[f"E{row}"].border = no_border

            row += 1
            ws[f"A{row}"] = "DEPARTAMENTO"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"C{row}"] = ": PUNO"
            ws[f"C{row}"].font = data_font
            ws[f"C{row}"].border = no_border
            ws[f"D{row}"] = "RUC"
            ws[f"D{row}"].font = header_font
            ws[f"D{row}"].border = no_border
            ws[f"E{row}"] = f': {notary_info["ruc"]}'
            ws[f"E{row}"].font = data_font
            ws[f"E{row}"].border = no_border

            row += 1
            ws[f"A{row}"] = "PROVINCIA"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"C{row}"] = ": SAN ROMAN"
            ws[f"C{row}"].font = data_font
            ws[f"C{row}"].border = no_border
            ws[f"D{row}"] = "DESDE"
            ws[f"D{row}"].font = header_font
            ws[f"D{row}"].border = no_border
            ws[f"E{row}"] = f": {self._format_date_in_spanish(desde).upper()}"
            ws[f"E{row}"].font = data_font
            ws[f"E{row}"].border = no_border

            row += 1
            ws[f"A{row}"] = "DISTRITO"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"C{row}"] = f': {notary_info["distrito"]}'
            ws[f"C{row}"].font = data_font
            ws[f"C{row}"].border = no_border
            ws[f"D{row}"] = "HASTA"
            ws[f"D{row}"].font = header_font
            ws[f"D{row}"].border = no_border
            ws[f"E{row}"] = f": {self._format_date_in_spanish(hasta).upper()}"
            ws[f"E{row}"].font = data_font
            ws[f"E{row}"].border = no_border

            # Data table headers
            row += 2
            headers = [
                "N° ESC.",
                "FECHA",
                "TESTADOR",
                "A FAVOR",
                "ACTO",
                "FOJA",
            ]
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = header_font
                cell.alignment = center_alignment
                cell.border = thin_border

            # Data rows
            for i, data_row in enumerate(report_data, 1):
                row += 1

                # Row data
                row_data = [
                    data_row["numero_escritura"],
                    data_row["fecha"],
                    self._sanitize_cell_value(data_row["testador"]),
                    self._sanitize_cell_value(data_row["beneficiario"]),
                    self._sanitize_cell_value(data_row["contrato"]),
                    data_row["folio"],
                ]

                for col, value in enumerate(row_data, 1):
                    cell = ws.cell(row=row, column=col, value=value)
                    cell.font = data_font
                    cell.border = thin_border

                    # Alignment based on column
                    if col in [1, 2, 6]:  # N° ESC., FECHA, FOJA
                        cell.alignment = center_alignment
                    else:
                        cell.alignment = left_alignment

            # Auto-adjust column widths
            for column in ws.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width

            # Create response
            from django.http import StreamingHttpResponse

            def file_iterator():
                buffer = io.BytesIO()
                wb.save(buffer)
                buffer.seek(0)
                yield buffer.getvalue()
                buffer.close()

            response = StreamingHttpResponse(
                file_iterator(),
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            response["Content-Disposition"] = (
                f"attachment; filename=INDICE_CRONOLOGICO_TESTAMENTOS_{anio}.xlsx"
            )
            response["Access-Control-Allow-Origin"] = "*"

            return response

        except Exception as e:
            import traceback

            traceback.print_exc()
            from django.http import HttpResponse

            return HttpResponse(f"Error generating Excel report: {e}", status=500)

    def generate_word_report(self, desde, hasta):
        """Generate Word report matching PHP script format"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            import io

            import time

            report_start = time.time()

            # Get data
            report_data = self._get_report_data(desde, hasta)

            notary_info = self._get_notary_info()

            anio = self._extract_year_from_date(hasta)

            # Create document
            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            # Title
            title = doc.add_heading("INDICE CRONOLOGICO - TESTAMENTOS", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            subtitle = doc.add_heading(f"AÑO {anio}", 0)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add spacing
            doc.add_paragraph()

            # Notary info table
            info_table = doc.add_table(rows=5, cols=6)
            info_table.style = "Table Grid"

            # Row 1: NOTARIA
            row1 = info_table.rows[0]
            row1.cells[0].text = "NOTARIA"
            row1.cells[0].paragraphs[0].runs[0].bold = True
            row1.cells[2].text = f': {notary_info["nombre"]}'

            # Row 2: DIRECCION
            row2 = info_table.rows[1]
            row2.cells[0].text = "DIRECCION"
            row2.cells[0].paragraphs[0].runs[0].bold = True
            row2.cells[2].text = f': {notary_info["direccion"]}'
            row2.cells[3].text = "TELEFONO"
            row2.cells[3].paragraphs[0].runs[0].bold = True
            row2.cells[4].text = f': {notary_info["telefono"]}'

            # Row 3: DEPARTAMENTO
            row3 = info_table.rows[2]
            row3.cells[0].text = "DEPARTAMENTO"
            row3.cells[0].paragraphs[0].runs[0].bold = True
            row3.cells[2].text = ": PUNO"
            row3.cells[3].text = "RUC"
            row3.cells[3].paragraphs[0].runs[0].bold = True
            row3.cells[4].text = f': {notary_info["ruc"]}'

            # Row 4: PROVINCIA
            row4 = info_table.rows[3]
            row4.cells[0].text = "PROVINCIA"
            row4.cells[0].paragraphs[0].runs[0].bold = True
            row4.cells[2].text = ": SAN ROMAN"
            row4.cells[3].text = "DESDE"
            row4.cells[3].paragraphs[0].runs[0].bold = True
            row4.cells[4].text = f": {self._format_date_in_spanish(desde).upper()}"

            # Row 5: DISTRITO
            row5 = info_table.rows[4]
            row5.cells[0].text = "DISTRITO"
            row5.cells[0].paragraphs[0].runs[0].bold = True
            row5.cells[2].text = f': {notary_info["distrito"]}'
            row5.cells[3].text = "HASTA"
            row5.cells[3].paragraphs[0].runs[0].bold = True
            row5.cells[4].text = f": {self._format_date_in_spanish(hasta).upper()}"

            # Add spacing
            doc.add_paragraph()

            # Data table
            if report_data:
                data_table = doc.add_table(rows=1, cols=6)
                data_table.style = "Table Grid"

                # Headers
                header_row = data_table.rows[0]
                headers = [
                    "N° ESC.",
                    "FECHA",
                    "TESTADOR",
                    "A FAVOR",
                    "ACTO",
                    "FOJA",
                ]
                for i, header in enumerate(headers):
                    header_row.cells[i].text = header
                    header_row.cells[i].paragraphs[0].runs[0].bold = True

            # Data rows
            for i, data_row in enumerate(report_data, 1):
                row = data_table.add_row()
                row.cells[0].text = str(data_row["numero_escritura"])
                row.cells[1].text = data_row["fecha"]
                row.cells[2].text = self._sanitize_cell_value(data_row["testador"])
                row.cells[3].text = self._sanitize_cell_value(data_row["beneficiario"])
                row.cells[4].text = self._sanitize_cell_value(data_row["contrato"])
                row.cells[5].text = str(data_row["folio"])

            # Create response
            from django.http import StreamingHttpResponse

            def file_iterator():
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                yield buffer.getvalue()
                buffer.close()

            response = StreamingHttpResponse(
                file_iterator(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = (
                f"attachment; filename=INDICE_CRONOLOGICO_TESTAMENTOS_{anio}.docx"
            )
            response["Access-Control-Allow-Origin"] = "*"

            return response

        except Exception as e:
            import traceback

            traceback.print_exc()
            from django.http import HttpResponse

            return HttpResponse(f"Error generating Word report: {e}", status=500)
