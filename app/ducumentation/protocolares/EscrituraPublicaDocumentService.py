from django.db import connection
from django.http import HttpResponse, JsonResponse
from docx import Document
import io
import time
from decimal import Decimal
from datetime import datetime
import locale
import re
from docx.shared import RGBColor
import gc
from typing import Dict
from functools import lru_cache
import os
from rest_framework.response import Response
from notaria import models
from .utils import (
    NumberToLetterConverter,
    DocumentFormatter,
    PlaceholderProcessor,
    DocxTemplateProcessor,
    DataValidator,
    TemplateManager,
)


class EscrituraDocumentService:
    """
    Service to generate escritura publica documents based on PHP legacy script
    """

    def __init__(
        self,
        letras=None,
        formatter=None,
        placeholder_processor=None,
        docx_template_processor=None,
        data_validator=None,
        template_manager=None,
    ):
        """
        Initialize with dependency injection

        PARAMETERS:
        - letras: NumberToLetterConverter instance (optional)
        - formatter: DocumentFormatter instance (optional)
        - processor: PlaceholderProcessor instance (optional)
        - docx_template_processor: DocxTemplateProcessor instance (optional)
        - validator: DataValidator instance (optional)
        - template_manager: TemplateManager instance (optional)

        If not provided, will create new instances
        """

        self.letras = letras or NumberToLetterConverter()
        self.formatter = DocumentFormatter(self.letras)
        self.placeholder_processor = placeholder_processor or PlaceholderProcessor()
        self.docx_template_processor = docx_template_processor or DocxTemplateProcessor()
        self.data_validator = data_validator or DataValidator()
        self.template_manager = template_manager or TemplateManager()


    def generate_escritura_publica_document(self, template_id, kardex, action, mode):
        """
        Main entry point for generating escritura publica documents

        FLOW:
        1. Get template information from database (filename, etc.)
        2. Download the template from R2 storage (or existing document for actualizar)
        3. Fetch data from database (consulta_escritura)
        4. Process data into sections (documento, vehiculos, pagos, contratantes)
        5. Replace placeholders in template
        6. Upload to R2
        7. Return HTTP response

        PARAMETERS:
        - template_id: ID of the template in the database
        - kardex: Document identifier (e.g., "KAR6508-2025")
        - action: Action type ("generate", "actualizar", "parte")
        - mode: Response mode ("download" or "open")
        """
        
        # Handle "actualizar" action - update existing document
        if action == "actualizar":
            return self._update_existing_document(kardex, mode)

        # STEP 1: Get template info
        template_info = self._get_template_info(template_id)

        # STEP 2: Download template from R2
        template_bytes = self.template_manager.get_template_from_r2(
            template_id, template_info["filename"]
        )

        # STEP 3: Fetch ALL data from database (mirrors PHP consulta_escritura)
        # TODO: Implement this - fetch all data in ONE query
        raw_data = self._consulta_escritura(kardex, action, template_id)
        
        # Debug: Check what contractors were found
        print(f"DEBUG: Contractors data for {kardex}:")
        print(f"DEBUG: condicion = {raw_data.get('condicion', 'NOT_FOUND')}")
        print(f"DEBUG: nombres = {raw_data.get('nombres', 'NOT_FOUND')}")
        print(f"DEBUG: sexo = {raw_data.get('sexo', 'NOT_FOUND')}")

        data_documento = self.formatter.format_document_data(raw_data)
        data_vehiculos = self.formatter.format_vehicle_data(raw_data)
        data_pagos = self.formatter.format_payment_data(raw_data)
        data_escrituracion = self.formatter.format_escrituracion_data(raw_data)
        data_contratantes = self.formatter.format_contractor_data(raw_data)
        data_company = self.formatter.format_company_data(raw_data)

        final_data = self.formatter.combine_all_data(
            data_documento,
            data_vehiculos,
            data_pagos,
            data_escrituracion,
            data_contratantes,
            data_company,
        )

        # Try using python-docx-template first (docxtpl)
        # Falls back to PlaceholderProcessor if docxtpl fails
        try:
            processed_bytes = self.docx_template_processor.replace_placeholders(template_bytes, final_data)
            if processed_bytes:
                buffer = io.BytesIO(processed_bytes)
            else:
                raise Exception("DocxTemplateProcessor returned None")
        except Exception as e:
            # Fallback to PlaceholderProcessor
            # This happens when template contains Jinja2 keywords in placeholder values
            buffer = io.BytesIO(template_bytes)
            doc = Document(buffer)

            self.placeholder_processor.replace_placeholders(doc, final_data)
            self.placeholder_processor.clean_unfilled_placeholders(doc)

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

        # STEP 9: Upload to R2
        self.template_manager.upload_document_to_r2(buffer, kardex)
        
        # STEP 10: Return HTTP response
        filename = f"__PROY__{kardex}.docx"

        return self._create_response_from_buffer(buffer, filename, kardex, mode)

    def _update_existing_document(self, kardex, mode):
        """
        Update existing document with escrituracion data
        Mirrors: PHP actualizar action
        """
        # STEP 1: Validate numescritura exists
        kardex_obj = models.Kardex.objects.get(kardex=kardex)
        if not kardex_obj.numescritura:
            raise ValueError("Falta grabar el número de escritura. Por favor, complete el número de escritura antes de actualizar el documento.")
        
        # STEP 2: Get escrituracion data from database
        raw_data = self._consulta_escritura(kardex, "actualizar", None)
        data_escrituracion = self.formatter.format_escrituracion_data(raw_data)
        
        # STEP 3: Update document using generic method from utils
        output_buffer = self.template_manager.update_document_escrituracion(
            kardex, 
            data_escrituracion, 
            self.placeholder_processor
        )
        
        # STEP 4: Return HTTP response
        filename = f"__PROY__{kardex}.docx"
        return self._create_response_from_buffer(output_buffer, filename, kardex, mode)
    
    def _get_template_info(self, template_id):
        """
        Get template information from database
        """

        template = models.TplTemplate.objects.get(pktemplate=template_id)
        return {"filename": template.filename}

    def _create_response(self, doc, filename, kardex, mode):
        """
        Create HTTP response with the document
        """

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        if mode == "open":
            response = JsonResponse(
                {
                    "status": "success",
                    "mode": "open",
                    "filename": filename,
                    "kardex": kardex,
                    "message": "Document generated and ready to open in Word",
                }
            )
            response["Access-Control-Allow-Origin"] = "*"
            return response
        else:
            response = HttpResponse(
                buffer.read(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = f'inline; filename="{filename}"'
            response["Content-Length"] = str(buffer.getbuffer().nbytes)
            response["Access-Control-Allow-Origin"] = "*"
            return response

    def _create_response_from_buffer(self, buffer, filename, kardex, mode):
        """
        Create HTTP response from buffer
        """

        if mode == "open":
            response = JsonResponse(
                {
                    "status": "success",
                    "mode": "open",
                    "filename": filename,
                    "kardex": kardex,
                    "message": "Document generated and ready to open in Word",
                }
            )
            response["Access-Control-Allow-Origin"] = "*"
            return response
        else:
            response = HttpResponse(
                buffer.read(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = f'inline; filename="{filename}"'
            response["Content-Length"] = str(buffer.getbuffer().nbytes)
            response["Access-Control-Allow-Origin"] = "*"
            return response

    def _consulta_escritura(self, num_kardex, action, template_id):
        """
        MAIN SQL QUERY - Fetch ALL data from database in ONE query
        Mirrors: PHP consulta_escritura()

        TASK:
        - Execute raw SQL query with all LEFT JOINS
        - Get kardex, contratantes, vehiculos, pagos, empresas in ONE query
        - Return dictionary with all raw data

        TABLES NEEDED:
        - kardex (main table)
        - contratantesxacto, contratantes, cliente2 (contractors)
        - detallevehicular (vehicle data)
        - patrimonial, detallemediopago (payment data)
        - nacionalidades, tipodocumento, tipoestacivil, ubigeo (lookups)
        - tb_abogado, usuarios (notary/user data)

        SOLUTION: See app/ducumentation/services.py line 3032-3164
        """

        #########################################################
        """
        Get idtipoacto from kardex we need to use the id tipoacto from the parameters and do not query the database again
        """
        idtipoacto = models.Kardex.objects.get(kardex=num_kardex).idtipkar
        #########################################################

        query = """
            SELECT k.idkardex as id_kardex,
                k.kardex,
                k.numescritura as numero_escritura,
                k.fechaescritura as fecha_escritura,
                k.txa_minuta as registro_escritura,
                CURRENT_DATE() as fecha_generado,
                k.fechaconclusion as fecha_conclusion,
                k.numminuta as numero_minuta,
                k.kardexconexo as kardex_conexo,
                k.folioini as folio_inicial,
                k.foliofin as folio_final,
                k.papelini as papel_inicial,
                k.papelfin as papel_final,
                (SELECT desacto FROM tiposdeacto WHERE idtipoacto=%s) as acto,
                (SELECT fileName FROM tpl_template WHERE pkTemplate=%s) as plantilla,
                (SELECT urlTemplate FROM tpl_template WHERE pkTemplate=%s) as url_plantilla,
                k.fechaingreso as fecha_ingreso,
                k.responsable_new as usuario,
                abo.razonsocial as abogado,
                abo.matricula as matricula,
                abo.sede_colegio as sede_colegio,
                usu.dni as dni_usuario,
                GROUP_CONCAT(c2.idcliente) as id_cliente,
                GROUP_CONCAT(IF(c2.conyuge='','NO',c2.conyuge)) as id_conyuge,
                GROUP_CONCAT(cxa.idcontratante) as id_contratante,
                GROUP_CONCAT(TRIM(CONCAT(IFNULL(c2.prinom, ''), ' ', IFNULL(c2.segnom, ''), IF(c2.segnom='','',' ') ,IFNULL(c2.apepat, ''), ' ',IFNULL(c2.apemat, ''),
            IFNULL(c2.razonsocial, '')))) AS nombres,
                GROUP_CONCAT(cxa.uif) as uif,
                GROUP_CONCAT(ac.condicion) as condicion,
                GROUP_CONCAT(IF(n.descripcion IS NULL OR n.descripcion='','EMPRESA',n.descripcion)) as nacionalidad,
                GROUP_CONCAT(td.destipdoc) as tipo_documento,
                GROUP_CONCAT(c2.numdoc) AS numero_documento,
                GROUP_CONCAT(UPPER(c2.profesion_plantilla)) AS ocupacion,
                GROUP_CONCAT(IF(tec.desestcivil IS NULL OR tec.desestcivil='','EMPRESA',tec.desestcivil)) as estado_civil,
                GROUP_CONCAT(IF(c2.tipper='N',c2.direccion,c2.domfiscal) SEPARATOR ',,') as direccion,
                GROUP_CONCAT(IFNULL(u.codpto, '')) as codigo_departamento,
                GROUP_CONCAT(IFNULL(u.coddis, '')) as codigo_distrito,
                GROUP_CONCAT(IFNULL(u.codprov, '')) as codigo_provincia,
                GROUP_CONCAT(IFNULL(IF(SUBSTRING_INDEX(c2.ubigeo_plantilla, '/', -1)='',u.nomdis,SUBSTRING_INDEX(c2.ubigeo_plantilla, '/', -1)),(IFNULL(u.nomdis, '')))) AS distrito,
                GROUP_CONCAT(IFNULL(u.nomprov, '')) as provincia,
                GROUP_CONCAT(IFNULL(u.nomdpto, '')) as departamento,
                GROUP_CONCAT(c2.sexo) AS sexo,
                GROUP_CONCAT(c2.tipper) as tipo_persona,
                GROUP_CONCAT(IF(cn.firma = '0', 'NO', 'SI')) AS firma,
                GROUP_CONCAT(cn.firma) as n_firma,
                GROUP_CONCAT(cn.tiporepresentacion) AS tipo_representacion,
                dv.numplaca AS placa,
                dv.marca AS marca,
                dv.clase AS clase,
                dv.anofab AS anio,
                dv.numserie AS serie,
                dv.color AS color,
                dv.motor AS motor,
                dv.modelo AS modelo,
                dv.carroceria AS carroceria,
                dv.pregistral as partida,
                dv.fecinsc AS fecha_inscripcion,
                dv.combustible AS combustible,
                UPPER(sr.dessede) AS sede,
                UPPER(sr.num_zona) AS numero_zona,
                pat.importetrans AS precio ,
                pat.idmon AS moneda,
                pat.exhibiomp,
                pat.idoppago,
                uif.descripcion AS medio_pago,
                mon.simbolo as simbolo_moneda,
                mon.desmon as descripcion_moneda,
                mp.desmpagos as descripcion_medio_pago,
                mp.sunat as sunat_medio_pago,
                GROUP_CONCAT(cnr.idcontratanterp) as id_empresa,
                GROUP_CONCAT(TRIM(CONCAT(IFNULL(cr2.prinom, ''), ' ', IFNULL(cr2.segnom, ''), IF(cr2.segnom='','',' ') ,IFNULL(cr2.apepat, ''), ' ',IFNULL(cr2.apemat, ''),
            IFNULL(cr2.razonsocial, '')))) AS nombre_empresa,
                GROUP_CONCAT(cr2.tipper) as tipo_persona_empresa,
                GROUP_CONCAT(acr.condicion) as condicion_empresa,
                GROUP_CONCAT(tdr.destipdoc) as tipo_documento_empresa,
                GROUP_CONCAT(cr2.numdoc) AS numero_documento_empresa,
                GROUP_CONCAT(cr2.domfiscal) as domicilio_empresa,
                GROUP_CONCAT(ur.nomdis) as distrito_empresa,
                GROUP_CONCAT(ur.nomprov) as provincia_empresa,
                GROUP_CONCAT(ur.nomdpto) as departamento_empresa,
                GROUP_CONCAT(srr2.zona_depar SEPARATOR ',,') as oficina_registral,
                GROUP_CONCAT(cr2.numpartida) as numero_partida,
                dmp.foperacion as fecha_operacion,
                dmp.documentos as documentos,
                ban.desbanco as banco
            FROM kardex as k
            LEFT JOIN tb_abogado as abo on abo.idabogado=k.idabogado
            LEFT JOIN usuarios as usu on usu.idusuario=k.idusuario
            LEFT JOIN contratantesxacto as cxa on cxa.kardex=k.kardex
            LEFT JOIN actocondicion as ac ON cxa.idcondicion=ac.idcondicion
            LEFT JOIN contratantes cn ON cxa.idcontratante = cn.idcontratante
            LEFT JOIN cliente2 as c2 on c2.idcontratante=cxa.idcontratante
            LEFT JOIN nacionalidades as n on n.idnacionalidad=c2.nacionalidad
            LEFT JOIN tipodocumento as td ON td.idtipdoc = c2.idtipdoc
            LEFT OUTER JOIN tipoestacivil as tec ON tec.idestcivil = c2.idestcivil
            LEFT OUTER JOIN ubigeo as u ON u.coddis = c2.idubigeo
            LEFT JOIN detallevehicular as dv ON dv.kardex=k.kardex
            LEFT JOIN sedesregistrales as sr ON sr.idsedereg=dv.idsedereg
            LEFT JOIN patrimonial as pat ON pat.kardex=k.kardex
            LEFT JOIN fpago_uif as uif ON uif.id_fpago = pat.fpago
            LEFT JOIN monedas as mon ON mon.idmon = pat.idmon
            LEFT JOIN detallemediopago as dmp ON pat.kardex = dmp.kardex
            LEFT JOIN mediospago as mp ON dmp.codmepag = mp.codmepag
            LEFT JOIN contratantes as cnr ON cxa.idcontratante = cnr.idcontratante
            LEFT JOIN contratantesxacto as cxar on cxar.idcontratante=cnr.idcontratanterp
            LEFT JOIN actocondicion as acr ON acr.idcondicion=cxar.idcondicion
            LEFT JOIN cliente2 as cr2 on cr2.idcontratante=cnr.idcontratanterp
            left JOIN nacionalidades as nr on nr.idnacionalidad=cr2.nacionalidad
            LEFT JOIN sedesregistrales as srr2 on srr2.idsedereg=cr2.idsedereg
            LEFT JOIN tipodocumento as tdr ON tdr.idtipdoc = cr2.idtipdoc
            LEFT OUTER JOIN tipoestacivil as tecr ON tecr.idestcivil = cr2.idestcivil
            LEFT OUTER JOIN ubigeo as ur ON ur.coddis = cr2.idubigeo
            LEFT JOIN  bancos as ban ON ban.idbancos=dmp.idbancos
            WHERE k.kardex=%s and (c2.tipper='N')
            GROUP BY k.idkardex, dmp.detmp LIMIT 1
        """

        with connection.cursor() as cursor:
            cursor.execute(query, [idtipoacto, template_id, template_id, num_kardex])
            desc = cursor.description
            row = cursor.fetchone()
            if not row:
                return None
            result = dict(zip([col[0] for col in desc], row))
            
            # For constitution documents, also query company data (tipper='J')
            # This gets the company being constituted
            company_query = """
                SELECT 
                    c2.razonsocial as nombre_empresa_constitucion,
                    c2.domfiscal as domicilio_empresa_constitucion,
                    c2.tipper as tipo_persona_empresa_constitucion,
                    c2.numdoc as numero_documento_empresa_constitucion,
                    c2.numpartida as numero_partida_constitucion
                FROM contratantesxacto as cxa
                LEFT JOIN cliente2 as c2 on c2.idcontratante=cxa.idcontratante
                WHERE cxa.kardex=%s and c2.tipper='J'
                LIMIT 1
            """
            cursor.execute(company_query, [num_kardex])
            company_row = cursor.fetchone()
            
            if company_row and company_row[0]:  # If company data exists
                # Add company constitution data to result
                result['nombre_empresa_constitucion'] = company_row[0]
                result['domicilio_empresa_constitucion'] = company_row[1]
                result['tipo_persona_empresa_constitucion'] = company_row[2]
                result['numero_documento_empresa_constitucion'] = company_row[3]
                result['numero_partida_constitucion'] = company_row[4]
            
            return result


class EscriturasPublicasReportService:
    """Service for generating escrituras publicas reports matching PHP script format"""

    def _sanitize_cell_value(self, value):
        """Sanitize cell values to prevent Excel corruption"""
        if value is None:
            return ""

        # Convert to string and strip whitespace
        val_str = str(value).strip()

        # Remove control characters that can cause XML corruption
        import re

        val_str = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]", "", val_str)

        # Limit to Excel's cell character limit
        val_str = val_str[:32767]

        return val_str

    def _format_date_in_spanish(self, date_input):
        """Convert date to Spanish format like 'LUNES, 15 DE ENERO DEL 2025'"""
        try:
            # Handle both datetime objects and date strings
            if hasattr(date_input, "strftime"):
                date_obj = date_input
            else:
                from datetime import datetime

                # Try different formats
                if "-" in str(date_input) and len(str(date_input).split("-")[0]) == 4:
                    date_obj = datetime.strptime(str(date_input), "%Y-%m-%d")
                elif "/" in str(date_input):
                    date_obj = datetime.strptime(str(date_input), "%d/%m/%Y")
                else:
                    return str(date_input)

            # Spanish day names
            dias = ["LUNES", "MARTES", "MIÉRCOLES", "JUEVES", "VIERNES", "SÁBADO", "DOMINGO"]
            # Spanish month names
            meses = [
                "ENERO",
                "FEBRERO",
                "MARZO",
                "ABRIL",
                "MAYO",
                "JUNIO",
                "JULIO",
                "AGOSTO",
                "SEPTIEMBRE",
                "OCTUBRE",
                "NOVIEMBRE",
                "DICIEMBRE",
            ]

            dia_semana = dias[date_obj.weekday()]
            dia = date_obj.day
            mes = meses[date_obj.month - 1]
            anio = date_obj.year

            return f"{dia_semana}, {dia} DE {mes} DEL {anio}"
        except:
            return str(date_input)

    def _extract_year_from_date(self, date_input):
        """Extract year from date string DD/MM/YYYY or YYYY-MM-DD or datetime object"""
        try:
            # Handle datetime objects
            if hasattr(date_input, "year"):
                return str(date_input.year)

            # Handle strings
            date_str = str(date_input)
            # Try to parse as YYYY-MM-DD first
            if "-" in date_str and len(date_str.split("-")[0]) == 4:
                return date_str.split("-")[0]
            # Try to parse as DD/MM/YYYY
            elif "/" in date_str:
                return date_str.split("/")[-1]
            else:
                from datetime import datetime

                return str(datetime.now().year)
        except:
            from datetime import datetime

            return str(datetime.now().year)

    def _format_date_for_display(self, date_obj):
        """Format date like PHP script logic"""
        try:
            if hasattr(date_obj, "strftime"):
                return date_obj.strftime("%d/%m/%Y")
            return str(date_obj)
        except:
            return str(date_obj)

    def _get_notary_info(self):
        """Get notary configuration info from database"""
        from .utils import get_notary_config
        return get_notary_config()

    def _get_report_data(self, desde, hasta):
        """Fetch data for the report matching PHP query"""
        try:
            from django.db import connection
            from datetime import datetime

            # Convert dates to proper format for Django ORM
            if isinstance(desde, str):
                if "-" in desde and len(desde.split("-")[0]) == 4:
                    desde_dt = datetime.strptime(desde, "%Y-%m-%d")
                else:
                    desde_dt = datetime.strptime(desde, "%d/%m/%Y")
            else:
                desde_dt = desde

            if isinstance(hasta, str):
                if "-" in hasta and len(hasta.split("-")[0]) == 4:
                    hasta_dt = datetime.strptime(hasta, "%Y-%m-%d")
                else:
                    hasta_dt = datetime.strptime(hasta, "%d/%m/%Y")
            else:
                hasta_dt = hasta

            # First set group_concat_max_len to handle large strings
            with connection.cursor() as cursor:
                cursor.execute("SET SESSION group_concat_max_len = 1000000")

            query = """
                SELECT 
                    k.fechaescritura,
                    k.kardex,
                    k.contrato,
                    k.numescritura,
                    k.numminuta,
                    k.folioini,
                    CAST(k.numescritura AS SIGNED) AS numescritura2,
                    p.importetrans as precio,
                    m.simbolo as moneda
                FROM kardex as k 
                LEFT JOIN patrimonial as p ON p.kardex=k.kardex AND p.idtipoacto = k.codactos
                LEFT JOIN monedas as m ON m.idmon=p.idmon
                WHERE k.idtipkar='1' 
                    AND k.fechaescritura <> '' 
                    AND STR_TO_DATE(k.fechaescritura,'%%Y-%%m-%%d') >= STR_TO_DATE(%s,'%%Y-%%m-%%d')
                    AND STR_TO_DATE(k.fechaescritura,'%%Y-%%m-%%d') <= STR_TO_DATE(%s,'%%Y-%%m-%%d')
                ORDER BY numescritura2 ASC, fechaescritura ASC
            """

            with connection.cursor() as cursor:
                cursor.execute(
                    query, [desde_dt.strftime("%Y-%m-%d"), hasta_dt.strftime("%Y-%m-%d")]
                )
                escrituras = []
                rows = cursor.fetchall()

                # Get all kardex numbers
                kardex_list = [row[1] for row in rows]

                # Get all contractors in one query
                contractors_query = """
                    SELECT 
                        cxa.kardex,
                        c2.tipper,
                        UPPER(CONCAT(c2.apepat,' ',c2.apemat,' ',c2.prinom,' ',c2.segnom)) AS nombre,
                        cxa.idcontratante,
                        UPPER(c2.razonsocial) AS empresa,
                        cxa.parte,
                        cxa.uif,
                        (SELECT cxar.parte 
                            FROM contratantesxacto AS cxar
                            WHERE con.idcontratanterp = cxar.idcontratante 
                            AND cxar.kardex = cxa.kardex limit 1) as parte_representada
                    FROM contratantesxacto AS cxa
                    INNER JOIN contratantes AS con ON con.idcontratante=cxa.idcontratante
                    INNER JOIN cliente2 AS c2 ON c2.idcontratante=con.idcontratante
                    WHERE cxa.kardex IN %s
                    ORDER BY cxa.kardex, c2.tipper ASC
                """
                cursor.execute(contractors_query, [tuple(kardex_list)])
                all_contractors = cursor.fetchall()

                # Group contractors by kardex
                contractors_by_kardex = {}
                for contractor in all_contractors:
                    kardex = contractor[0]
                    if kardex not in contractors_by_kardex:
                        contractors_by_kardex[kardex] = []
                    contractors_by_kardex[kardex].append(contractor[1:])  # Skip kardex from tuple

                # Process each kardex
                for row in rows:
                    kardex = row[1]
                    otorgante = []
                    otorgado = []

                    # Process contractors if any exist for this kardex
                    for contractor in contractors_by_kardex.get(kardex, []):
                        tipper, nombre, idcontratante, empresa, parte, uif, parte_representada = (
                            contractor
                        )

                        # Process otorgante
                        if parte == 1 or parte_representada == 1 or uif == "O":
                            if not (uif == "O" and parte_representada == 2):
                                otorgante.append(empresa if tipper != "N" else nombre)

                        # Process otorgado
                        if parte == 2 or parte_representada == 2 or uif in ("B", "N"):
                            if not (uif == "B" and parte_representada == 1):
                                otorgado.append(empresa if tipper != "N" else nombre)

                    escrituras.append(
                        {
                            "numero_escritura": row[3],
                            "fecha": datetime.strptime(row[0], "%Y-%m-%d").strftime("%d/%m/%Y"),
                            "otorgante": (
                                "NO CORRE" if row[2] == "NO CORRE / " else ", ".join(otorgante)
                            ),
                            "otorgado": (
                                "NO CORRE" if row[2] == "NO CORRE / " else ", ".join(otorgado)
                            ),
                            "contrato": row[2].replace("/", "").upper() if row[2] else "",
                            "monto": f"{row[8]} {row[7]}" if row[7] else "",
                            "folio": row[5],
                        }
                    )

                return escrituras

        except Exception as e:
            import traceback

            traceback.print_exc()
            return []

    def generate_excel_report(self, desde, hasta):
        """Generate Excel report matching PHP script format"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            import io

            # Get data
            report_data = self._get_report_data(desde, hasta)

            notary_info = self._get_notary_info()

            anio = self._extract_year_from_date(hasta)

            # Create workbook and worksheet
            wb = Workbook()
            ws = wb.active
            ws.title = "ESCRITURAS PUBLICAS"

            # Styles
            title_font = Font(name="Arial", size=18.5, bold=True)
            header_font = Font(name="Arial", size=13.5, bold=True)
            data_font = Font(name="Arial", size=13.5)
            center_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            left_alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            right_alignment = Alignment(horizontal="right", vertical="top", wrap_text=True)

            # Borders
            thin_border = Border(
                left=Side(border_style="thin"),
                right=Side(border_style="thin"),
                top=Side(border_style="thin"),
                bottom=Side(border_style="thin"),
            )
            no_border = Border(
                left=Side(style=None),
                right=Side(style=None),
                top=Side(style=None),
                bottom=Side(style=None),
            )

            # Title section
            ws.merge_cells("A1:G1")
            ws["A1"] = "INDICE CRONOLOGICO - REGISTRO DE ESCRITURAS PUBLICAS"
            ws["A1"].font = title_font
            ws["A1"].alignment = center_alignment
            ws["A1"].border = no_border

            ws.merge_cells("A2:G2")
            ws["A2"] = f"AÑO {anio}"
            ws["A2"].font = title_font
            ws["A2"].alignment = center_alignment
            ws["A2"].border = no_border

            # Notary info section
            row = 4
            ws[f"A{row}"] = "NOTARIA"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"C{row}"] = f': {self._sanitize_cell_value(notary_info["nombre"])}'
            ws[f"C{row}"].font = data_font
            ws[f"C{row}"].border = no_border

            row += 1
            ws[f"A{row}"] = "DIRECCION"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"C{row}"] = f': {notary_info["direccion"]}'
            ws[f"C{row}"].font = data_font
            ws[f"C{row}"].border = no_border
            ws[f"F{row}"] = "TELEFONO"
            ws[f"F{row}"].font = header_font
            ws[f"F{row}"].border = no_border
            ws[f"H{row}"] = f': {notary_info["telefono"]}'
            ws[f"H{row}"].font = data_font
            ws[f"H{row}"].border = no_border

            row += 1
            ws[f"A{row}"] = "DEPARTAMENTO"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"C{row}"] = f': {notary_info["departamento"]}'
            ws[f"C{row}"].font = data_font
            ws[f"C{row}"].border = no_border
            ws[f"F{row}"] = "RUC"
            ws[f"F{row}"].font = header_font
            ws[f"F{row}"].border = no_border
            ws[f"H{row}"] = f': {notary_info["ruc"]}'
            ws[f"H{row}"].font = data_font
            ws[f"H{row}"].border = no_border

            row += 1
            ws[f"A{row}"] = "PROVINCIA"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"C{row}"] = f': {notary_info["provincia"]}'
            ws[f"C{row}"].font = data_font
            ws[f"C{row}"].border = no_border
            ws[f"F{row}"] = "DESDE"
            ws[f"F{row}"].font = header_font
            ws[f"F{row}"].border = no_border
            ws[f"H{row}"] = f": {self._format_date_in_spanish(desde).upper()}"
            ws[f"H{row}"].font = data_font
            ws[f"H{row}"].border = no_border

            row += 1
            ws[f"A{row}"] = "DISTRITO"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"C{row}"] = f': {notary_info["distrito"]}'
            ws[f"C{row}"].font = data_font
            ws[f"C{row}"].border = no_border
            ws[f"F{row}"] = "HASTA"
            ws[f"F{row}"].font = header_font
            ws[f"F{row}"].border = no_border
            ws[f"H{row}"] = f": {self._format_date_in_spanish(hasta).upper()}"
            ws[f"H{row}"].font = data_font
            ws[f"H{row}"].border = no_border

            # Data table headers
            row += 2
            headers = [
                "ESCR.",
                "FECH.ESCR.",
                "OTORGANTE",
                "A FAVOR",
                "ACTO JURIDICO",
                "PRECIO",
                "NUM.FOLIO",
            ]
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = header_font
                cell.alignment = center_alignment
                cell.border = thin_border

            # Data rows
            for i, data_row in enumerate(report_data, 1):
                row += 1

                # Row data
                row_data = [
                    data_row["numero_escritura"],
                    data_row["fecha"],
                    self._sanitize_cell_value(data_row["otorgante"]),
                    self._sanitize_cell_value(data_row["otorgado"]),
                    self._sanitize_cell_value(data_row["contrato"]),
                    data_row["monto"],
                    data_row["folio"],
                ]

                for col, value in enumerate(row_data, 1):
                    cell = ws.cell(row=row, column=col, value=value)
                    cell.font = data_font
                    cell.border = thin_border

                    # Alignment based on column
                    if col in [1, 2, 6, 7]:  # ESCR., FECHA, PRECIO, FOLIO
                        cell.alignment = right_alignment
                    else:
                        cell.alignment = left_alignment

            # Auto-adjust column widths
            for column in ws.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width

            # Create response
            from django.http import StreamingHttpResponse

            def file_iterator():
                buffer = io.BytesIO()
                wb.save(buffer)
                buffer.seek(0)
                yield buffer.getvalue()
                buffer.close()

            response = StreamingHttpResponse(
                file_iterator(),
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            response["Content-Disposition"] = (
                f"attachment; filename=INDICE_CRONOLOGICO_ESCRITURAS_PUBLICAS_{anio}.xlsx"
            )
            response["Access-Control-Allow-Origin"] = "*"

            return response

        except Exception as e:
            import traceback

            traceback.print_exc()
            from django.http import HttpResponse

            return HttpResponse(f"Error generating Excel report: {e}", status=500)

    def generate_word_report(self, desde, hasta):
        """Generate Word report matching PHP script format"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            import io

            # Get data
            report_data = self._get_report_data(desde, hasta)

            notary_info = self._get_notary_info()

            anio = self._extract_year_from_date(hasta)

            # Create document
            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            # Title
            title = doc.add_heading("INDICE CRONOLOGICO - REGISTRO DE ESCRITURAS PUBLICAS", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            subtitle = doc.add_heading(f"AÑO {anio}", 0)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add spacing
            doc.add_paragraph()

            # Notary info table
            info_table = doc.add_table(rows=5, cols=9)
            info_table.style = "Table Grid"

            # Row 1: NOTARIA
            row1 = info_table.rows[0]
            row1.cells[0].text = "NOTARIA"
            row1.cells[0].paragraphs[0].runs[0].bold = True
            row1.cells[2].text = f': {notary_info["nombre"]}'

            # Row 2: DIRECCION
            row2 = info_table.rows[1]
            row2.cells[0].text = "DIRECCION"
            row2.cells[0].paragraphs[0].runs[0].bold = True
            row2.cells[2].text = f': {notary_info["direccion"]}'
            row2.cells[4].text = "TELEFONO"
            row2.cells[4].paragraphs[0].runs[0].bold = True
            row2.cells[7].text = f': {notary_info["telefono"]}'

            # Row 3: DEPARTAMENTO
            row3 = info_table.rows[2]
            row3.cells[0].text = "DEPARTAMENTO"
            row3.cells[0].paragraphs[0].runs[0].bold = True
            row3.cells[2].text = f': {notary_info["departamento"]}'
            row3.cells[4].text = "RUC"
            row3.cells[4].paragraphs[0].runs[0].bold = True
            row3.cells[7].text = f': {notary_info["ruc"]}'

            # Row 4: PROVINCIA
            row4 = info_table.rows[3]
            row4.cells[0].text = "PROVINCIA"
            row4.cells[0].paragraphs[0].runs[0].bold = True
            row4.cells[2].text = f': {notary_info["provincia"]}'
            row4.cells[4].text = "DESDE"
            row4.cells[4].paragraphs[0].runs[0].bold = True
            row4.cells[7].text = f": {self._format_date_in_spanish(desde).upper()}"

            # Row 5: DISTRITO
            row5 = info_table.rows[4]
            row5.cells[0].text = "DISTRITO"
            row5.cells[0].paragraphs[0].runs[0].bold = True
            row5.cells[2].text = f': {notary_info["distrito"]}'
            row5.cells[4].text = "HASTA"
            row5.cells[4].paragraphs[0].runs[0].bold = True
            row5.cells[7].text = f": {self._format_date_in_spanish(hasta).upper()}"

            # Add spacing
            doc.add_paragraph()

            # Data table
            if report_data:
                data_table = doc.add_table(rows=1, cols=7)
                data_table.style = "Table Grid"

                # Headers
                header_row = data_table.rows[0]
                headers = [
                    "ESCR.",
                    "FECH.ESCR.",
                    "OTORGANTE",
                    "A FAVOR",
                    "ACTO JURIDICO",
                    "PRECIO",
                    "NUM.FOLIO",
                ]
                for i, header in enumerate(headers):
                    header_row.cells[i].text = header
                    header_row.cells[i].paragraphs[0].runs[0].bold = True

            # Data rows
            for i, data_row in enumerate(report_data, 1):
                row = data_table.add_row()
                row.cells[0].text = str(data_row["numero_escritura"])
                row.cells[1].text = data_row["fecha"]
                row.cells[2].text = self._sanitize_cell_value(data_row["otorgante"])
                row.cells[3].text = self._sanitize_cell_value(data_row["otorgado"])
                row.cells[4].text = self._sanitize_cell_value(data_row["contrato"])
                row.cells[5].text = data_row["monto"]
                row.cells[6].text = str(data_row["folio"])

            # Create response
            from django.http import StreamingHttpResponse

            def file_iterator():
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                yield buffer.getvalue()
                buffer.close()

            response = StreamingHttpResponse(
                file_iterator(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = (
                f"attachment; filename=INDICE_CRONOLOGICO_ESCRITURAS_PUBLICAS_{anio}.docx"
            )
            response["Access-Control-Allow-Origin"] = "*"

            return response

        except Exception as e:
            import traceback

            traceback.print_exc()
            from django.http import HttpResponse

            return HttpResponse(f"Error generating Word report: {e}", status=500)
