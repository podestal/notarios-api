from django.db import connection
from django.http import HttpResponse, JsonResponse
from docx import Document
import io
import re
from zipfile import ZipFile
from decimal import Decimal
from datetime import datetime
from rest_framework.response import Response
from notaria import models
from .utils import (
    NumberToLetterConverter,
    DocumentFormatter,
    PlaceholderProcessor,
    DocxTemplateProcessor,
    DataValidator,
    TemplateManager,
)


class TransferenciasVehicularesDocumentService:
    """
    Service to generate Transferencias Vehiculares documents
    Mirrors: PHP transferencia_vehicular() function
    Reuses all utilities from utils.py
    """

    def __init__(
        self,
        letras=None,
        formatter=None,
        placeholder_processor=None,
        docx_template_processor=None,
        data_validator=None,
        template_manager=None,
    ):
        """
        Initialize with dependency injection - SAME as other services
        All utilities are reusable!
        """
        self.letras = letras or NumberToLetterConverter()
        self.formatter = DocumentFormatter(self.letras)
        self.placeholder_processor = placeholder_processor or PlaceholderProcessor()
        self.docx_template_processor = docx_template_processor or DocxTemplateProcessor()
        self.data_validator = data_validator or DataValidator()
        self.template_manager = template_manager or TemplateManager()

    def generate_transferencias_document(self, template_id, kardex, action, mode):
        """
        Main entry point for generating Transferencias Vehiculares documents
        Mirrors: PHP transferencia_vehicular() function
        
        FLOW: (Same as Escrituras, No Contenciosos, Garantias, and Testamentos)
        1. Get template information
        2. Download template from R2 (or existing document for actualizar)
        3. Fetch data from database
        4. Process data into sections
        5. Replace placeholders
        6. Upload to R2
        7. Return HTTP response
        """
        
        # Handle "actualizar" action
        if action == "actualizar":
            return self._update_existing_document(kardex, mode)

        # STEP 1: Get template info
        template_info = self._get_template_info(template_id)

        # STEP 2: Download template from R2
        template_bytes = self.template_manager.get_template_from_r2(
            template_id, template_info["filename"]
        )

        # STEP 3: Fetch data from database
        raw_data = self._consulta_transferencias(kardex, action, template_id)

        # STEP 4: Format data using REUSABLE utilities
        data_documento = self.formatter.format_document_data(raw_data)
        data_vehiculos = self.formatter.format_vehicle_data(raw_data)
        data_pagos = self.formatter.format_payment_data(raw_data, kind="vehicular")
        data_escrituracion = self.formatter.format_escrituracion_data(raw_data)
        data_contratantes = self.formatter.format_contractor_data(raw_data)
        data_contratantes = self._normalize_transferencias_contractor_fields(
            data_contratantes
        )
        data_company = self.formatter.format_company_data(raw_data)

        # STEP 5: Combine all data
        final_data = self.formatter.combine_all_data(
            data_documento,
            data_vehiculos,
            data_pagos,
            data_escrituracion,
            data_contratantes,
            data_company,
        )

        # STEP 6: Replace placeholders (try docxtpl, fallback to PlaceholderProcessor)
        try:
            processed_bytes = self.docx_template_processor.replace_placeholders(
                template_bytes, final_data
            )
            if processed_bytes:
                final_bytes = processed_bytes
            else:
                raise Exception("DocxTemplateProcessor returned None")
        except Exception as e:
            print(f"DEBUG: Falling back to PlaceholderProcessor for Transferencias")
            buffer = io.BytesIO(template_bytes)
            doc = Document(buffer)

            self.placeholder_processor.replace_placeholders(doc, final_data)
            self.placeholder_processor.clean_unfilled_placeholders(doc)

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            final_bytes = buffer.getvalue()

        # Cleanup punctuation artifacts caused by optional placeholders in templates
        # (e.g. ", ,", ", ;") for Transferencias text blocks.
        final_bytes = self._cleanup_transferencias_punctuation_artifacts(final_bytes)
        buffer = io.BytesIO(final_bytes)
        buffer.seek(0)

        # STEP 7: Upload to R2
        self.template_manager.upload_document_to_r2(buffer, kardex)
        
        # STEP 8: Return HTTP response
        filename = f"__PROY__{kardex}.docx"
        return self._create_response_from_buffer(buffer, filename, kardex, mode)

    def _normalize_transferencias_contractor_fields(self, contractor_data):
        """
        Transferencias-only text normalization requested by business:
        - Contractor names end with comma.
        - Contractor addresses separate address and ubigeo with comma.

        Important: do not strip/normalize every contractor placeholder because many
        templates rely on intentional trailing spaces between placeholders.
        """
        normalized = {}

        contractor_comma_only_fields = re.compile(
            r"^(P|C)_(NOM(_\d+)?)$"
        )
        contractor_comma_space_fields = re.compile(
            r"^(P|C)_(NACIONALIDAD(_\d+)?|DOC(_\d+)?)$"
        )
        contractor_origen_fields = re.compile(r"^(P|C)_(ORIGEN_FONDO(_\d+)?)$")
        contractor_comma_space_estado_civil_fields = re.compile(
            r"^(P|C)_(ESTADO_CIVIL(_\d+)?)$"
        )
        contractor_no_trailing_space_fields = re.compile(r"^(P|C)_(OCUPACION(_\d+)?)$")
        contractor_empty_fields = re.compile(r"^(P|C)_(DOC_LETRAS(_\d+)?|IDE(_\d+)?)$")

        def _compact_spaces(text: str) -> str:
            # Includes non-breaking spaces from Word/html payloads.
            text = text.replace("\u00a0", " ")
            return re.sub(r"\s+", " ", text).strip()

        for key, value in contractor_data.items():
            if not isinstance(value, str):
                normalized[key] = value
                continue

            # Default: compact spaces and remove trailing whitespace.
            text = _compact_spaces(value)

            if "DOMICILIO" in key and value.strip():
                # "CON DOMICILIO EN <direccion> DEL DISTRITO..." -> "... <direccion>, DEL DISTRITO..."
                text = _compact_spaces(value)
                text = re.sub(
                    r"\b(CON DOMICILIO EN\s+)(.+?)\s+(DEL DISTRITO DE\b)",
                    lambda m: f"{m.group(1)}{m.group(2).strip()}, {m.group(3)}",
                    text,
                    count=1,
                )
                # Only normalize spaces around comma we introduced.
                text = re.sub(r"\s+,", ",", text)
                text = re.sub(r",\s*", ", ", text)
                # Fix legacy cases like "..., , DEL DISTRITO..."
                text = re.sub(r",\s*,\s*(DEL DISTRITO DE\b)", r", \1", text)
                # Add comma between distrito and provincia.
                # "... DEL DISTRITO DE AYAVIRI PROVINCIA DE MELGAR ..."
                # -> "... DEL DISTRITO DE AYAVIRI, PROVINCIA DE MELGAR ..."
                text = re.sub(
                    r"\b(DEL DISTRITO DE\s+.+?)\s+(PROVINCIA DE\b)",
                    r"\1, \2",
                    text,
                    count=1,
                )

            # Enforce clean separators for contractor placeholders:
            # - no double spaces
            # - NOMBRE ends in comma (template usually has next static space)
            # - core identity fields end in comma+space
            # - ocupacion/estado civil end in single space (template may add punctuation after ocupacion)
            if contractor_empty_fields.match(key):
                text = ""
            elif contractor_comma_only_fields.match(key) and text:
                text = text.rstrip(" ,;.") + ","
            elif contractor_comma_space_fields.match(key) and text:
                text = text.rstrip(" ,;.") + ", "
            elif contractor_origen_fields.match(key) and text:
                text = (
                    "DECLARA EL DINERO PARA LA ADQUISICION DEL PRESENTE BIEN MUEBLE "
                    f"ES PROVENIENTE DE {text.rstrip(' ,;.')}"
                )
            elif contractor_comma_space_estado_civil_fields.match(key) and text:
                text = text.rstrip(" ,;.") + ", "
            elif contractor_no_trailing_space_fields.match(key) and text:
                text = text.rstrip(" ,;.")
            elif re.match(r"^(P|C)_", key) or key in {"CALIDAD_P", "CALIDAD_C"}:
                text = _compact_spaces(text)

            normalized[key] = text
        return normalized

    def _cleanup_transferencias_punctuation_artifacts(self, docx_bytes: bytes) -> bytes:
        """
        Remove redundant punctuation left by optional placeholders in templates.
        Examples:
        - ", , " -> ", "
        - ", ;"  -> ";"
        """
        in_mem = io.BytesIO(docx_bytes)
        out_mem = io.BytesIO()
        with ZipFile(in_mem, "r") as zin, ZipFile(out_mem, "w") as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                    xml = data.decode("utf-8")
                    xml = re.sub(r",\s*,+", ", ", xml)
                    xml = re.sub(r",\s*;", ";", xml)
                    xml = re.sub(r";\s*,", "; ", xml)
                    data = xml.encode("utf-8")
                zout.writestr(item, data)
        return out_mem.getvalue()

    def _update_existing_document(self, kardex, mode):
        """
        Update existing Transferencias document with escrituracion data
        Mirrors: PHP actualizar action
        REUSES: TemplateManager.update_document_escrituracion()
        """
        # STEP 1: Validate numescritura exists
        kardex_obj = models.Kardex.objects.get(kardex=kardex)
        if not kardex_obj.numescritura:
            raise ValueError("Falta grabar el número de escritura. Por favor, complete el número de escritura antes de actualizar el documento.")
        
        # STEP 2: Get escrituracion data
        raw_data = self._consulta_transferencias(kardex, "actualizar", None)
        data_escrituracion = self.formatter.format_escrituracion_data(raw_data)
        
        # STEP 3: Update document using REUSABLE generic method
        output_buffer = self.template_manager.update_document_escrituracion(
            kardex, 
            data_escrituracion, 
            self.placeholder_processor
        )
        
        # STEP 4: Return HTTP response
        filename = f"__PROY__{kardex}.docx"
        return self._create_response_from_buffer(output_buffer, filename, kardex, mode)

    def _get_template_info(self, template_id):
        """Get template information from database"""
        template = models.TplTemplate.objects.get(pktemplate=template_id)
        return {"filename": template.filename}

    def _create_response_from_buffer(self, buffer, filename, kardex, mode):
        """Create HTTP response from buffer"""
        if mode == "open":
            response = JsonResponse(
                {
                    "status": "success",
                    "mode": "open",
                    "filename": filename,
                    "kardex": kardex,
                    "message": "Document generated and ready to open in Word",
                }
            )
            response["Access-Control-Allow-Origin"] = "*"
            return response
        else:
            response = HttpResponse(
                buffer.read(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = f'inline; filename="{filename}"'
            response["Content-Length"] = str(buffer.getbuffer().nbytes)
            response["Access-Control-Allow-Origin"] = "*"
            return response

    def _consulta_transferencias(self, num_kardex, action, template_id):
        """
        SQL Query for Transferencias Vehiculares documents
        Mirrors: PHP consulta_transferencia() in transferencia_vehicular.php
        
        NOTE: The query structure is IDENTICAL to other services
        P = TRANSFERENTE/VENDEDOR (parte=1 or uif='O')
        C = COMPRADOR/ADQUIRIENTE (parte=2 or uif='B')
        """
        
        # Get idtipoacto from kardex
        idtipoacto = models.Kardex.objects.get(kardex=num_kardex).idtipkar
        
        # Same query as PHP script
        query = """
            SELECT k.idkardex as id_kardex,
                k.kardex,
                k.numescritura as numero_escritura,
                k.fechaescritura as fecha_escritura,
                k.txa_minuta as registro_escritura,
                CURRENT_DATE() as fecha_generado,
                k.fechaconclusion as fecha_conclusion,
                k.numminuta as numero_minuta,
                k.kardexconexo as kardex_conexo,
                k.folioini as folio_inicial, 
                k.foliofin as folio_final, 
                k.papelini as papel_inicial, 
                k.papelfin as papel_final,
                (SELECT desacto FROM tiposdeacto WHERE idtipoacto=%s) as acto,
                (SELECT fileName FROM tpl_template WHERE pkTemplate=%s) as plantilla,
                (SELECT urlTemplate FROM tpl_template WHERE pkTemplate=%s) as url_plantilla,
                k.fechaingreso as fecha_ingreso,
                k.responsable_new as usuario,
                abo.razonsocial as abogado,
                abo.matricula as matricula,
                usu.dni as dni_usuario,
                GROUP_CONCAT(c2.idcliente) as id_cliente,
                GROUP_CONCAT(IF(c2.conyuge='','NO',c2.conyuge)) as id_conyuge,
                GROUP_CONCAT(cxa.idcontratante) as id_contratante,
                GROUP_CONCAT(TRIM(CONCAT(IFNULL(c2.prinom, ''), ' ', IFNULL(c2.segnom, ''), IF(c2.segnom='','',' ') ,IFNULL(c2.apepat, ''), ' ',IFNULL(c2.apemat, ''),
            IFNULL(c2.razonsocial, '')))) AS nombres,
                GROUP_CONCAT(cxa.uif) as uif,
                GROUP_CONCAT(ac.condicion) as condicion,
                GROUP_CONCAT(IF(n.descripcion IS NULL OR n.descripcion='','EMPRESA',n.descripcion)) as nacionalidad,
                GROUP_CONCAT(td.destipdoc) as tipo_documento,
                GROUP_CONCAT(c2.numdoc) AS numero_documento,
                GROUP_CONCAT(UPPER(c2.profesion_plantilla)) AS ocupacion,
                GROUP_CONCAT(IF(tec.desestcivil IS NULL OR tec.desestcivil='','EMPRESA',tec.desestcivil)) as estado_civil,
                GROUP_CONCAT(IF(c2.tipper='N',c2.direccion,c2.domfiscal) SEPARATOR ',,') as direccion,
                GROUP_CONCAT(IFNULL(u.codpto, '')) as codigo_departamento,
                GROUP_CONCAT(IFNULL(u.coddis, '')) as codigo_distrito,
                GROUP_CONCAT(IFNULL(u.codprov, '')) as codigo_provincia,
                GROUP_CONCAT(IFNULL(IF(SUBSTRING_INDEX(c2.ubigeo_plantilla, '/', -1)='',u.nomdis,SUBSTRING_INDEX(c2.ubigeo_plantilla, '/', -1)),(IFNULL(u.nomdis, '')))) AS distrito,
                GROUP_CONCAT(IFNULL(u.nomprov, '')) as provincia,
                GROUP_CONCAT(IFNULL(u.nomdpto, '')) as departamento,
                GROUP_CONCAT(c2.sexo) AS sexo,
                GROUP_CONCAT(c2.tipper) as tipo_persona,
                GROUP_CONCAT(IF(cn.firma = '0', 'NO', 'SI')) AS firma,
                GROUP_CONCAT(cn.firma) as n_firma,
                GROUP_CONCAT(cn.tiporepresentacion) AS tipo_representacion,
                GROUP_CONCAT(IFNULL(cxa.ofondo, '')) AS origen_fondo,
                dv.numplaca AS placa, 
                dv.marca AS marca, 
                dv.clase AS clase,
                dv.anofab AS anio,
                dv.numserie AS serie, 
                dv.color AS color,
                dv.motor AS motor, 
                dv.modelo AS modelo, 
                dv.carroceria AS carroceria,
                dv.pregistral as partida,
                dv.fecinsc AS fecha_inscripcion,
                dv.combustible AS combustible,
                UPPER(sr.dessede) AS sede,
                UPPER(sr.num_zona) AS numero_zona,
                pat.importetrans AS precio , 
                pat.idmon AS moneda,
                pat.exhibiomp, 
                pat.idoppago, 
                uif.descripcion AS medio_pago, 
                mon.simbolo as simbolo_moneda,
                mon.desmon as descripcion_moneda,
                mp.desmpagos as descripcion_medio_pago,
                mp.sunat as sunat_medio_pago,
                GROUP_CONCAT(IF(cnr.idcontratanterp='','0',cnr.idcontratanterp)) as id_empresa,
                GROUP_CONCAT(TRIM(CONCAT(IFNULL(cr2.prinom, ''), ' ', IFNULL(cr2.segnom, ''), IF(cr2.segnom='','',' ') ,IFNULL(cr2.apepat, ''), ' ',IFNULL(cr2.apemat, ''),
            IFNULL(cr2.razonsocial, '')))) AS nombre_empresa,
                GROUP_CONCAT(IF(cr2.tipper='J',cr2.tipper,'N')) as tipo_persona_empresa,
                GROUP_CONCAT(acr.condicion) as condicion_empresa,
                GROUP_CONCAT(tdr.destipdoc) as tipo_documento_empresa,
                GROUP_CONCAT(cr2.numdoc) AS numero_documento_empresa,
                GROUP_CONCAT(cr2.domfiscal) as domicilio_empresa,
                GROUP_CONCAT(ur.nomdis) as distrito_empresa,
                GROUP_CONCAT(ur.nomprov) as provincia_empresa,
                GROUP_CONCAT(ur.nomdpto) as departamento_empresa,
                GROUP_CONCAT(srr2.zona_depar SEPARATOR ',,') as oficina_registral,
                GROUP_CONCAT(cr2.numpartida) as numero_partida,
                dmp.foperacion as fecha_operacion,
                dmp.documentos as documentos,
                ban.desbanco as banco
            FROM kardex as k
            LEFT JOIN tb_abogado as abo on abo.idabogado=k.idabogado
            LEFT JOIN usuarios as usu on usu.idusuario=k.idusuario
            LEFT JOIN contratantesxacto as cxa on cxa.kardex=k.kardex
            LEFT JOIN actocondicion as ac ON cxa.idcondicion=ac.idcondicion
            LEFT JOIN contratantes cn ON cxa.idcontratante = cn.idcontratante
            LEFT JOIN cliente2 as c2 on c2.idcontratante=cxa.idcontratante
            LEFT JOIN nacionalidades as n on n.idnacionalidad=c2.nacionalidad
            LEFT JOIN tipodocumento as td ON td.idtipdoc = c2.idtipdoc
            LEFT OUTER JOIN tipoestacivil as tec ON tec.idestcivil = c2.idestcivil
            LEFT OUTER JOIN ubigeo as u ON u.coddis = c2.idubigeo
            LEFT JOIN detallevehicular as dv ON dv.kardex=k.kardex AND dv.idtipacto<>''
            LEFT JOIN sedesregistrales as sr ON sr.idsedereg=dv.idsedereg
            LEFT JOIN patrimonial as pat ON pat.kardex=k.kardex
            LEFT JOIN fpago_uif as uif ON uif.id_fpago = pat.fpago
            LEFT JOIN monedas as mon ON mon.idmon = pat.idmon
            LEFT JOIN detallemediopago as dmp ON pat.kardex = dmp.kardex
            LEFT JOIN mediospago as mp ON dmp.codmepag = mp.codmepag
            LEFT JOIN contratantes as cnr ON cxa.idcontratante = cnr.idcontratante
            LEFT JOIN contratantesxacto as cxar on cxar.idcontratante=cnr.idcontratanterp AND cxar.kardex = cxa.kardex
            LEFT JOIN actocondicion as acr ON acr.idcondicion=cxar.idcondicion
            LEFT JOIN cliente2 as cr2 on cr2.idcontratante=cnr.idcontratanterp
            left JOIN nacionalidades as nr on nr.idnacionalidad=cr2.nacionalidad
            LEFT JOIN sedesregistrales as srr2 on srr2.idsedereg=cr2.idsedereg
            LEFT JOIN tipodocumento as tdr ON tdr.idtipdoc = cr2.idtipdoc
            LEFT OUTER JOIN tipoestacivil as tecr ON tecr.idestcivil = cr2.idestcivil
            LEFT OUTER JOIN ubigeo as ur ON ur.coddis = cr2.idubigeo
            LEFT JOIN  bancos as ban ON ban.idbancos=dmp.idbancos
            WHERE k.kardex=%s and (c2.tipper IN ('N','J'))
            GROUP BY k.idkardex, dmp.detmp LIMIT 1
        """
        with connection.cursor() as cursor:
            cursor.execute(query, [idtipoacto, template_id, template_id, num_kardex])
            desc = cursor.description
            row = cursor.fetchone()
            if not row:
                return None
            result = dict(zip([col[0] for col in desc], row))
            
            # Also query company data if needed (for Juridical persons)
            company_query = """
                SELECT 
                    c2.razonsocial as nombre_empresa_constitucion,
                    c2.domfiscal as domicilio_empresa_constitucion,
                    c2.tipper as tipo_persona_empresa_constitucion,
                    c2.numdoc as numero_documento_empresa_constitucion,
                    c2.numpartida as numero_partida_constitucion
                FROM contratantesxacto as cxa
                LEFT JOIN cliente2 as c2 on c2.idcontratante=cxa.idcontratante
                WHERE cxa.kardex=%s and c2.tipper='J'
                LIMIT 1
            """
            cursor.execute(company_query, [num_kardex])
            company_row = cursor.fetchone()
            
            if company_row and company_row[0]:
                result['nombre_empresa_constitucion'] = company_row[0]
                result['domicilio_empresa_constitucion'] = company_row[1]
                result['tipo_persona_empresa_constitucion'] = company_row[2]
                result['numero_documento_empresa_constitucion'] = company_row[3]
                result['numero_partida_constitucion'] = company_row[4]

            # Debug dump: all contratantes rows for this kardex
            contratantes_debug_query = """
                SELECT
                    cxa.id AS cxa_id,
                    cxa.idcontratante,
                    IFNULL(ac.condicion, '') AS condicion,
                    IFNULL(cxa.parte, '') AS parte,
                    IFNULL(cxa.uif, '') AS uif,
                    IFNULL(c2.tipper, '') AS tipper,
                    TRIM(CONCAT(
                        IFNULL(c2.prinom, ''), ' ',
                        IFNULL(c2.segnom, ''), IF(IFNULL(c2.segnom,'')='','',' '),
                        IFNULL(c2.apepat, ''), ' ',
                        IFNULL(c2.apemat, '')
                    )) AS nombre_natural,
                    TRIM(IFNULL(c2.razonsocial, '')) AS razonsocial,
                    IFNULL(cn.idcontratanterp, '') AS idcontratanterp_raw
                FROM contratantesxacto cxa
                LEFT JOIN actocondicion ac ON ac.idcondicion = cxa.idcondicion
                LEFT JOIN contratantes cn ON cn.idcontratante = cxa.idcontratante
                LEFT JOIN cliente2 c2 ON c2.idcontratante = cxa.idcontratante
                WHERE cxa.kardex = %s
                ORDER BY cxa.id ASC
            """
            cursor.execute(contratantes_debug_query, [num_kardex])
            contratantes_debug_rows = cursor.fetchall()
            print(f"DEBUG: contratantes dump for {num_kardex} (rows={len(contratantes_debug_rows)})")
            for row in contratantes_debug_rows:
                print(
                    "DEBUG: "
                    f"cxa_id={row[0]} | idcontratante={row[1]} | condicion={row[2]} | "
                    f"parte={row[3]} | uif={row[4]} | tipper={row[5]} | "
                    f"nombre_natural='{row[6]}' | razonsocial='{row[7]}' | "
                    f"idcontratanterp_raw='{row[8]}'"
                )
            
            return result


class TransferenciasVehicularesReportService:
    """Service for generating transferencias vehiculares reports matching PHP script format"""

    def _sanitize_cell_value(self, value):
        """Sanitize cell values to prevent Excel corruption"""
        if value is None:
            return ""

        # Convert to string and strip whitespace
        val_str = str(value).strip()

        # Remove control characters that can cause XML corruption
        import re

        val_str = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]", "", val_str)

        # Limit to Excel's cell character limit
        val_str = val_str[:32767]

        return val_str

    def _format_date_in_spanish(self, date_input):
        """Convert date to Spanish format like 'LUNES, 15 DE ENERO DEL 2025'"""
        try:
            # Handle both datetime objects and date strings
            if hasattr(date_input, "strftime"):
                date_obj = date_input
            else:
                from datetime import datetime

                # Try different formats
                if "-" in str(date_input) and len(str(date_input).split("-")[0]) == 4:
                    date_obj = datetime.strptime(str(date_input), "%Y-%m-%d")
                elif "/" in str(date_input):
                    date_obj = datetime.strptime(str(date_input), "%d/%m/%Y")
                else:
                    return str(date_input)

            # Spanish day names
            dias = ["LUNES", "MARTES", "MIÉRCOLES", "JUEVES", "VIERNES", "SÁBADO", "DOMINGO"]
            # Spanish month names
            meses = [
                "ENERO",
                "FEBRERO",
                "MARZO",
                "ABRIL",
                "MAYO",
                "JUNIO",
                "JULIO",
                "AGOSTO",
                "SEPTIEMBRE",
                "OCTUBRE",
                "NOVIEMBRE",
                "DICIEMBRE",
            ]

            dia_semana = dias[date_obj.weekday()]
            dia = date_obj.day
            mes = meses[date_obj.month - 1]
            anio = date_obj.year

            return f"{dia_semana}, {dia} DE {mes} DEL {anio}"
        except:
            return str(date_input)

    def _extract_year_from_date(self, date_input):
        """Extract year from date string DD/MM/YYYY or YYYY-MM-DD or datetime object"""
        try:
            # Handle datetime objects
            if hasattr(date_input, "year"):
                return str(date_input.year)

            # Handle strings
            date_str = str(date_input)
            # Try to parse as YYYY-MM-DD first
            if "-" in date_str and len(date_str.split("-")[0]) == 4:
                return date_str.split("-")[0]
            # Try to parse as DD/MM/YYYY
            elif "/" in date_str:
                return date_str.split("/")[-1]
            else:
                from datetime import datetime

                return str(datetime.now().year)
        except:
            from datetime import datetime

            return str(datetime.now().year)

    def _format_date_for_display(self, date_obj):
        """Format date like PHP script logic"""
        try:
            if hasattr(date_obj, "strftime"):
                return date_obj.strftime("%d/%m/%Y")
            return str(date_obj)
        except:
            return str(date_obj)

    def _get_notary_info(self):
        """Get notary configuration info from database"""
        from .utils import get_notary_config
        return get_notary_config()

    def _get_report_data(self, desde, hasta):
        """Fetch data for the report matching PHP query"""
        try:
            from django.db import connection
            from datetime import datetime

            import time

            start_time = time.time()
            print(
                f"DEBUG: TransferenciasVehiculares - desde: {desde} (type: {type(desde)}), hasta: {hasta} (type: {type(hasta)})"
            )

            # Convert dates to proper format for Django ORM
            desde_dt = None
            hasta_dt = None

            try:
                if isinstance(desde, str):
                    if "-" in desde and len(desde.split("-")[0]) == 4:
                        desde_dt = datetime.strptime(desde, "%Y-%m-%d")
                    else:
                        desde_dt = datetime.strptime(desde, "%d/%m/%Y")
                else:
                    desde_dt = desde
            except (ValueError, TypeError) as e:
                print(f"DEBUG: Error parsing desde date '{desde}': {e}")
                return []

            try:
                if isinstance(hasta, str):
                    if "-" in hasta and len(hasta.split("-")[0]) == 4:
                        hasta_dt = datetime.strptime(hasta, "%Y-%m-%d")
                    else:
                        hasta_dt = datetime.strptime(hasta, "%d/%m/%Y")
                else:
                    hasta_dt = hasta
            except (ValueError, TypeError) as e:
                print(f"DEBUG: Error parsing hasta date '{hasta}': {e}")
                return []

            # Validate dates are not None
            if desde_dt is None or hasta_dt is None:
                print(f"DEBUG: Invalid dates - desde: {desde}, hasta: {hasta}")
                return []

            # Set session variables for optimization
            with connection.cursor() as cursor:
                cursor.execute("SET SESSION group_concat_max_len = 1000000")
                cursor.execute("SET SESSION sql_mode = 'NO_AUTO_VALUE_ON_ZERO'")
                cursor.execute("SET SESSION sort_buffer_size = 2097152")

                # Try to create indexes if they don't exist (for performance)
                try:
                    cursor.execute(
                        "CREATE INDEX IF NOT EXISTS idx_kardex_tipkar_fecha ON kardex(idtipkar, fechaescritura, nc)"
                    )
                    cursor.execute(
                        "CREATE INDEX IF NOT EXISTS idx_kardex_fecha ON kardex(fechaescritura)"
                    )
                    print("DEBUG: Database indexes created/verified")
                except Exception as e:
                    print(f"DEBUG: Index creation failed (may already exist): {e}")

            # Ultra-simple query to get basic data first
            query = """
                SELECT 
                    k.fechaescritura,
                    k.kardex,
                    k.contrato,
                    k.numescritura,
                    k.numminuta,
                    k.folioini,
                    k.numescritura as numescritura2,
                    '' as precio,
                    '' as moneda,
                    '' as placa
                FROM kardex as k 
                WHERE k.idtipkar='3' 
                    AND k.fechaescritura <> '' 
                    AND k.fechaescritura >= %s
                    AND k.fechaescritura <= %s
                ORDER BY k.numescritura ASC
            """

            with connection.cursor() as cursor:
                cursor.execute(
                    query, [desde_dt.strftime("%Y-%m-%d"), hasta_dt.strftime("%Y-%m-%d")]
                )
                transferencias = []
                rows = cursor.fetchall()
                print(
                    f"DEBUG: Main query completed in {time.time() - start_time:.2f}s, found {len(rows)} records"
                )

                # Get additional data for the kardex records we found
                if rows:
                    kardex_list = [row[1] for row in rows]

                    # Get additional data (precio, moneda, placa) in a separate optimized query
                    additional_data_query = """
                        SELECT 
                            k.kardex,
                            COALESCE(p.importetrans, '') as precio,
                            COALESCE(m.simbolo, '') as moneda,
                            COALESCE(dv.numplaca, '') as placa
                        FROM kardex as k 
                        LEFT JOIN patrimonial as p ON p.kardex=k.kardex AND p.idtipoacto = k.codactos
                        LEFT JOIN monedas as m ON m.idmon=p.idmon
                        LEFT JOIN detallevehicular as dv ON dv.kardex=k.kardex AND dv.idtipacto = k.codactos
                        WHERE k.kardex IN %s
                    """

                    additional_start = time.time()
                    cursor.execute(additional_data_query, [tuple(kardex_list)])
                    additional_data = cursor.fetchall()
                    print(
                        f"DEBUG: Additional data query completed in {time.time() - additional_start:.2f}s, found {len(additional_data)} records"
                    )

                    # Create a lookup dictionary for additional data
                    additional_lookup = {
                        row[0]: (row[1], row[2], row[3]) for row in additional_data
                    }
                else:
                    kardex_list = []
                    additional_lookup = {}

                # Get all contractors in one query with optimized subquery
                contractors_query = """
                    SELECT 
                        cxa.kardex,
                        c2.tipper,
                        UPPER(CONCAT(COALESCE(c2.apepat,''),' ',COALESCE(c2.apemat,''),' ',COALESCE(c2.prinom,''),' ',COALESCE(c2.segnom,''))) AS nombre,
                        cxa.idcontratante,
                        UPPER(COALESCE(c2.razonsocial,'')) AS empresa,
                        cxa.parte,
                        cxa.uif,
                        COALESCE(cxar.parte, 0) as parte_representada
                    FROM contratantesxacto AS cxa
                    INNER JOIN contratantes AS con ON con.idcontratante=cxa.idcontratante
                    INNER JOIN cliente2 AS c2 ON c2.idcontratante=con.idcontratante
                    LEFT JOIN contratantesxacto AS cxar ON con.idcontratanterp = cxar.idcontratante AND cxar.kardex = cxa.kardex
                    WHERE cxa.kardex IN %s
                    ORDER BY cxa.kardex, c2.tipper ASC
                """
                contractors_start = time.time()
                cursor.execute(contractors_query, [tuple(kardex_list)])
                all_contractors = cursor.fetchall()
                print(
                    f"DEBUG: Contractors query completed in {time.time() - contractors_start:.2f}s, found {len(all_contractors)} contractors"
                )

                # Group contractors by kardex
                contractors_by_kardex = {}
                for contractor in all_contractors:
                    kardex = contractor[0]
                    if kardex not in contractors_by_kardex:
                        contractors_by_kardex[kardex] = []
                    contractors_by_kardex[kardex].append(contractor[1:])  # Skip kardex from tuple

                # Process each kardex with optimized processing
                processing_start = time.time()
                transferencias = []
                for row in rows:
                    kardex = row[1]
                    otorgante = []
                    otorgado = []

                    # Process contractors if any exist for this kardex
                    for contractor in contractors_by_kardex.get(kardex, []):
                        tipper, nombre, idcontratante, empresa, parte, uif, parte_representada = (
                            contractor
                        )

                        # Process otorgante (vendedor) - parte=1 or parte_representada=1 or uif='O'
                        if parte == 1 or parte_representada == 1 or uif == "O":
                            otorgante.append(empresa if tipper != "N" else nombre)

                        # Process otorgado (comprador) - parte=2 or parte_representada=2 or uif='B'
                        if parte == 2 or parte_representada == 2 or uif == "B":
                            otorgado.append(empresa if tipper != "N" else nombre)

                    # Clean contract name like in PHP - optimized string operations
                    contrato_raw = row[2] or ""
                    contrato_clean = (
                        contrato_raw.replace("/", "").replace("DE VEHICULO AUTOMOTOR", "").upper()
                    )

                    # Optimized date formatting
                    fecha_str = row[0]
                    fecha_formatted = (
                        f"{fecha_str[8:10]}/{fecha_str[5:7]}/{fecha_str[0:4]}"
                        if len(fecha_str) >= 10
                        else fecha_str
                    )

                    # Get additional data from lookup
                    precio, moneda, placa = additional_lookup.get(kardex, ("", "", ""))

                    transferencias.append(
                        {
                            "numero_escritura": row[3],
                            "fecha": fecha_formatted,
                            "otorgante": (
                                "NO CORRE"
                                if contrato_raw == "NO CORRE / "
                                else ", ".join(otorgante)
                            ),
                            "otorgado": (
                                "NO CORRE" if contrato_raw == "NO CORRE / " else ", ".join(otorgado)
                            ),
                            "contrato": contrato_clean,
                            "placa": placa.upper(),
                            "folio": row[5],
                        }
                    )

                print(f"DEBUG: Data processing completed in {time.time() - processing_start:.2f}s")
                print(f"DEBUG: Total data fetching completed in {time.time() - start_time:.2f}s")
                return transferencias

        except Exception as e:
            import traceback

            traceback.print_exc()
            return []

    def generate_excel_report(self, desde, hasta):
        """Generate Excel report matching PHP script format"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            import io

            import time

            report_start = time.time()
            print(f"DEBUG: TransferenciasVehiculares Excel - desde: {desde}, hasta: {hasta}")

            # Get data
            report_data = self._get_report_data(desde, hasta)
            print(f"DEBUG: Data fetched in {time.time() - report_start:.2f}s, generating Excel...")

            notary_info = self._get_notary_info()

            anio = self._extract_year_from_date(hasta)

            # Create workbook and worksheet
            wb = Workbook()
            ws = wb.active
            ws.title = "TRANSFERENCIAS VEHICULARES"

            # Styles
            title_font = Font(name="Arial", size=18.5, bold=True)
            header_font = Font(name="Arial", size=13.5, bold=True)
            data_font = Font(name="Arial", size=13.5)
            center_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            left_alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            right_alignment = Alignment(horizontal="right", vertical="top", wrap_text=True)

            # Borders
            thin_border = Border(
                left=Side(border_style="thin"),
                right=Side(border_style="thin"),
                top=Side(border_style="thin"),
                bottom=Side(border_style="thin"),
            )
            no_border = Border(
                left=Side(style=None),
                right=Side(style=None),
                top=Side(style=None),
                bottom=Side(style=None),
            )

            # Title section
            ws.merge_cells("A1:G1")
            ws["A1"] = "INDICE CRONOLOGICO - REGISTRO DE TRANSFERENCIAS DE BIENES MUEBLES"
            ws["A1"].font = title_font
            ws["A1"].alignment = center_alignment
            ws["A1"].border = no_border

            ws.merge_cells("A2:G2")
            ws["A2"] = f"AÑO {anio}"
            ws["A2"].font = title_font
            ws["A2"].alignment = center_alignment
            ws["A2"].border = no_border

            # Notary info section
            row = 4
            ws[f"A{row}"] = "NOTARIA"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"C{row}"] = f': {self._sanitize_cell_value(notary_info["nombre"])}'
            ws[f"C{row}"].font = data_font
            ws[f"C{row}"].border = no_border

            row += 1
            ws[f"A{row}"] = "DIRECCION"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"C{row}"] = f': {notary_info["direccion"]}'
            ws[f"C{row}"].font = data_font
            ws[f"C{row}"].border = no_border
            ws[f"F{row}"] = "TELEFONO"
            ws[f"F{row}"].font = header_font
            ws[f"F{row}"].border = no_border
            ws[f"H{row}"] = f': {notary_info["telefono"]}'
            ws[f"H{row}"].font = data_font
            ws[f"H{row}"].border = no_border

            row += 1
            ws[f"A{row}"] = "DEPARTAMENTO"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"C{row}"] = f': {notary_info["departamento"]}'
            ws[f"C{row}"].font = data_font
            ws[f"C{row}"].border = no_border
            ws[f"F{row}"] = "RUC"
            ws[f"F{row}"].font = header_font
            ws[f"F{row}"].border = no_border
            ws[f"H{row}"] = f': {notary_info["ruc"]}'
            ws[f"H{row}"].font = data_font
            ws[f"H{row}"].border = no_border

            row += 1
            ws[f"A{row}"] = "PROVINCIA"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"C{row}"] = f': {notary_info["provincia"]}'
            ws[f"C{row}"].font = data_font
            ws[f"C{row}"].border = no_border
            ws[f"F{row}"] = "DESDE"
            ws[f"F{row}"].font = header_font
            ws[f"F{row}"].border = no_border
            ws[f"H{row}"] = f": {self._format_date_in_spanish(desde).upper()}"
            ws[f"H{row}"].font = data_font
            ws[f"H{row}"].border = no_border

            row += 1
            ws[f"A{row}"] = "DISTRITO"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"C{row}"] = f': {notary_info["distrito"]}'
            ws[f"C{row}"].font = data_font
            ws[f"C{row}"].border = no_border
            ws[f"F{row}"] = "HASTA"
            ws[f"F{row}"].font = header_font
            ws[f"F{row}"].border = no_border
            ws[f"H{row}"] = f": {self._format_date_in_spanish(hasta).upper()}"
            ws[f"H{row}"].font = data_font
            ws[f"H{row}"].border = no_border

            # Data table headers
            row += 2
            headers = [
                "ACTA",
                "FECHA",
                "VENDEDOR",
                "COMPRADOR",
                "ACTO JURIDICO",
                "PLACA",
                "NUM.FOLIO",
            ]
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = header_font
                cell.alignment = center_alignment
                cell.border = thin_border

            # Data rows
            for i, data_row in enumerate(report_data, 1):
                row += 1

                # Row data
                row_data = [
                    data_row["numero_escritura"],
                    data_row["fecha"],
                    self._sanitize_cell_value(data_row["otorgante"]),
                    self._sanitize_cell_value(data_row["otorgado"]),
                    self._sanitize_cell_value(data_row["contrato"]),
                    data_row["placa"],
                    data_row["folio"],
                ]

                for col, value in enumerate(row_data, 1):
                    cell = ws.cell(row=row, column=col, value=value)
                    cell.font = data_font
                    cell.border = thin_border

                    # Alignment based on column
                    if col in [1, 2, 6, 7]:  # ACTA, FECHA, PLACA, FOLIO
                        cell.alignment = right_alignment
                    else:
                        cell.alignment = left_alignment

            # Auto-adjust column widths
            for column in ws.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width

            # Create response
            from django.http import StreamingHttpResponse

            def file_iterator():
                buffer = io.BytesIO()
                wb.save(buffer)
                buffer.seek(0)
                yield buffer.getvalue()
                buffer.close()

            response = StreamingHttpResponse(
                file_iterator(),
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            response["Content-Disposition"] = (
                f"attachment; filename=INDICE_CRONOLOGICO_VEHICULAR_{anio}.xlsx"
            )
            response["Access-Control-Allow-Origin"] = "*"

            return response

        except Exception as e:
            import traceback

            traceback.print_exc()
            from django.http import HttpResponse

            return HttpResponse(f"Error generating Excel report: {e}", status=500)

    def generate_word_report(self, desde, hasta):
        """Generate Word report matching PHP script format"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            import io

            import time

            report_start = time.time()
            print(f"DEBUG: TransferenciasVehiculares Word - desde: {desde}, hasta: {hasta}")

            # Get data
            report_data = self._get_report_data(desde, hasta)
            print(f"DEBUG: Data fetched in {time.time() - report_start:.2f}s, generating Word...")

            notary_info = self._get_notary_info()

            anio = self._extract_year_from_date(hasta)

            # Create document
            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            # Title
            title = doc.add_heading(
                "INDICE CRONOLOGICO - REGISTRO DE TRANSFERENCIAS DE BIENES MUEBLES", 0
            )
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            subtitle = doc.add_heading(f"AÑO {anio}", 0)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add spacing
            doc.add_paragraph()

            # Notary info table
            info_table = doc.add_table(rows=5, cols=9)
            info_table.style = "Table Grid"

            # Row 1: NOTARIA
            row1 = info_table.rows[0]
            row1.cells[0].text = "NOTARIA"
            row1.cells[0].paragraphs[0].runs[0].bold = True
            row1.cells[2].text = f': {notary_info["nombre"]}'

            # Row 2: DIRECCION
            row2 = info_table.rows[1]
            row2.cells[0].text = "DIRECCION"
            row2.cells[0].paragraphs[0].runs[0].bold = True
            row2.cells[2].text = f': {notary_info["direccion"]}'
            row2.cells[4].text = "TELEFONO"
            row2.cells[4].paragraphs[0].runs[0].bold = True
            row2.cells[7].text = f': {notary_info["telefono"]}'

            # Row 3: DEPARTAMENTO
            row3 = info_table.rows[2]
            row3.cells[0].text = "DEPARTAMENTO"
            row3.cells[0].paragraphs[0].runs[0].bold = True
            row3.cells[2].text = f': {notary_info["departamento"]}'
            row3.cells[4].text = "RUC"
            row3.cells[4].paragraphs[0].runs[0].bold = True
            row3.cells[7].text = f': {notary_info["ruc"]}'

            # Row 4: PROVINCIA
            row4 = info_table.rows[3]
            row4.cells[0].text = "PROVINCIA"
            row4.cells[0].paragraphs[0].runs[0].bold = True
            row4.cells[2].text = f': {notary_info["provincia"]}'
            row4.cells[4].text = "DESDE"
            row4.cells[4].paragraphs[0].runs[0].bold = True
            row4.cells[7].text = f": {self._format_date_in_spanish(desde).upper()}"

            # Row 5: DISTRITO
            row5 = info_table.rows[4]
            row5.cells[0].text = "DISTRITO"
            row5.cells[0].paragraphs[0].runs[0].bold = True
            row5.cells[2].text = f': {notary_info["distrito"]}'
            row5.cells[4].text = "HASTA"
            row5.cells[4].paragraphs[0].runs[0].bold = True
            row5.cells[7].text = f": {self._format_date_in_spanish(hasta).upper()}"

            # Add spacing
            doc.add_paragraph()

            # Data table
            if report_data:
                data_table = doc.add_table(rows=1, cols=7)
                data_table.style = "Table Grid"

                # Headers
                header_row = data_table.rows[0]
                headers = [
                    "ACTA",
                    "FECHA",
                    "VENDEDOR",
                    "COMPRADOR",
                    "ACTO JURIDICO",
                    "PLACA",
                    "NUM.FOLIO",
                ]
                for i, header in enumerate(headers):
                    header_row.cells[i].text = header
                    header_row.cells[i].paragraphs[0].runs[0].bold = True

            # Data rows
            for i, data_row in enumerate(report_data, 1):
                row = data_table.add_row()
                row.cells[0].text = str(data_row["numero_escritura"])
                row.cells[1].text = data_row["fecha"]
                row.cells[2].text = self._sanitize_cell_value(data_row["otorgante"])
                row.cells[3].text = self._sanitize_cell_value(data_row["otorgado"])
                row.cells[4].text = self._sanitize_cell_value(data_row["contrato"])
                row.cells[5].text = data_row["placa"]
                row.cells[6].text = str(data_row["folio"])

            # Create response
            from django.http import StreamingHttpResponse

            def file_iterator():
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                yield buffer.getvalue()
                buffer.close()

            response = StreamingHttpResponse(
                file_iterator(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = (
                f"attachment; filename=INDICE_CRONOLOGICO_VEHICULAR_{anio}.docx"
            )
            response["Access-Control-Allow-Origin"] = "*"

            return response

        except Exception as e:
            import traceback

            traceback.print_exc()
            from django.http import HttpResponse

            return HttpResponse(f"Error generating Word report: {e}", status=500)
