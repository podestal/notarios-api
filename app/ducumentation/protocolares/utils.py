from datetime import datetime
from decimal import Decimal
import locale
import boto3
from botocore.config import Config
import os
import re
import io
from docx import Document
from docx.shared import RGBColor
from docxtpl import DocxTemplate
from ducumentation.shared.base_r2_documents import get_s3_client


def get_notary_config():
    """
    Get notary configuration from database (confinotario table).
    Returns a dict with notary information.
    
    This function queries the database directly and does NOT include
    hardcoded sensitive values as fallback.
    
    The ubigeo field format is: "distrito - provincia - departamento"
    Example: "Macusani - carabaya - Puno"
    
    If ubigeo column doesn't exist, falls back to individual distrito, provincia, departamento columns.
    
    Returns:
        dict: Notary configuration with keys: nombre, telefono, correo, ruc, 
              direccion, distrito, provincia, departamento, codnotario, codoficial, coduif
              Returns empty strings for missing values.
    
    Raises:
        Exception: If no configuration is found in database
    """
    from django.db import connection
    
    with connection.cursor() as cursor:
        # First check if ubigeo column exists
        cursor.execute("""
            SELECT COUNT(*) 
            FROM information_schema.columns 
            WHERE table_schema = DATABASE()
            AND table_name = 'confinotario' 
            AND column_name = 'ubigeo'
        """)
        has_ubigeo = cursor.fetchone()[0] > 0
        cursor.execute("""
            SELECT COUNT(*)
            FROM information_schema.columns
            WHERE table_schema = DATABASE()
            AND table_name = 'confinotario'
            AND column_name = 'notario'
        """)
        has_notario = cursor.fetchone()[0] > 0
        
        # Build query based on column availability
        if has_ubigeo:
            notario_select = "notario" if has_notario else "NULL as notario"
            cursor.execute(f"""
                SELECT nombre, apellido, telefono, correo, ruc, direccion, distrito,
                       codnotario, codoficial, coduif, {notario_select}, ubigeo, provincia, departamento
                FROM confinotario
                LIMIT 1
            """)
        else:
            notario_select = "notario" if has_notario else "NULL as notario"
            cursor.execute(f"""
                SELECT nombre, apellido, telefono, correo, ruc, direccion, distrito,
                       codnotario, codoficial, coduif, {notario_select}, NULL as ubigeo, provincia, departamento
                FROM confinotario
                LIMIT 1
            """)
        
        result = cursor.fetchone()
        
        if result:
            # Format phone with area code if needed
            telefono = result[2] or ""
            if telefono and not telefono.startswith("("):
                telefono = f"(051) {telefono}"
            
            # Parse ubigeo field if available: "distrito - provincia - departamento"
            ubigeo = result[11] or ""
            distrito = ""
            provincia = ""
            departamento = ""
            
            if ubigeo and has_ubigeo:
                parts = [part.strip() for part in ubigeo.split("-")]
                if len(parts) >= 1:
                    distrito = parts[0].upper()  # "MACUSANI"
                if len(parts) >= 2:
                    provincia = parts[1].upper()  # "CARABAYA"
                if len(parts) >= 3:
                    departamento = parts[2].upper()  # "PUNO"
            
            # Fallback to individual columns if ubigeo parsing fails or column doesn't exist
            if not distrito:
                distrito = (result[6] or "").upper()
            if not provincia:
                provincia = (result[12] or "").upper()
            if not departamento:
                departamento = (result[13] or "").upper()
            
            return {
                "nombre": result[10] or f"{result[0] or ''} {result[1] or ''}".strip(),  # Use 'notario' field or concatenate
                "telefono": telefono,
                "correo": result[3] or "",
                "ruc": result[4] or "",
                "direccion": result[5] or "",
                "distrito": distrito,
                "provincia": provincia,
                "departamento": departamento,
                "codnotario": result[7] or "",
                "codoficial": result[8] or "",
                "coduif": result[9] or "",
            }
        else:
            raise Exception("No se encontró configuración de notaría en la base de datos. Por favor, configure la tabla 'confinotario'.")


class NumberToLetterConverter:
    """
    Utility class to convert numbers and dates to letter format in Spanish
    """

    def __init__(self):
        # Set locale for Spanish date formatting
        try:
            locale.setlocale(locale.LC_TIME, "es_ES.UTF-8")
        except:
            try:
                locale.setlocale(locale.LC_TIME, "es_ES")
            except:
                print("WARNING: Spanish locale not available, falling back to default")

    def number_to_letters(self, number) -> str:
        """Convert number to words in Spanish"""
        try:
            number = int(number)
        except (ValueError, TypeError):
            print(f"WARNING: Invalid number value: {number}")
            return str(number)

        UNITS = ["", "UNO", "DOS", "TRES", "CUATRO", "CINCO", "SEIS", "SIETE", "OCHO", "NUEVE"]
        TENS = [
            "",
            "DIEZ",
            "VEINTE",
            "TREINTA",
            "CUARENTA",
            "CINCUENTA",
            "SESENTA",
            "SETENTA",
            "OCHENTA",
            "NOVENTA",
        ]
        TEENS = [
            "DIEZ",
            "ONCE",
            "DOCE",
            "TRECE",
            "CATORCE",
            "QUINCE",
            "DIECISEIS",
            "DIECISIETE",
            "DIECIOCHO",
            "DIECINUEVE",
        ]
        HUNDREDS = [
            "",
            "CIENTO",
            "DOSCIENTOS",
            "TRESCIENTOS",
            "CUATROCIENTOS",
            "QUINIENTOS",
            "SEISCIENTOS",
            "SETECIENTOS",
            "OCHOCIENTOS",
            "NOVECIENTOS",
        ]

        if number == 0:
            return "CERO"

        if number < 10:
            return UNITS[number]

        if number < 20:
            return TEENS[number - 10]

        if number < 100:
            tens = number // 10
            units = number % 10
            if units == 0:
                return TENS[tens]
            return f"{TENS[tens]} Y {UNITS[units]}"

        if number < 1000:
            hundreds = number // 100
            rest = number % 100
            if rest == 0:
                return HUNDREDS[hundreds]
            return f"{HUNDREDS[hundreds]} {self.number_to_letters(rest)}"

        if number < 1000000:
            thousands = number // 1000
            rest = number % 1000
            if thousands == 1:
                thousands_text = "MIL"
            else:
                thousands_text = self.number_to_letters(thousands) + " MIL"

            if rest == 0:
                return thousands_text
            return f"{thousands_text} {self.number_to_letters(rest)}"

        return str(number)

    def date_to_letters(self, date_str) -> str:
        """Convert date to words in Spanish"""
        if not date_str:
            print(f"WARNING: Empty date value")
            return ""

        try:
            if isinstance(date_str, str):
                date_obj = datetime.strptime(date_str, "%Y-%m-%d")
            else:
                date_obj = date_str

            day = self.number_to_letters(date_obj.day)

            # Manual Spanish month names
            spanish_months = {
                1: "ENERO",
                2: "FEBRERO",
                3: "MARZO",
                4: "ABRIL",
                5: "MAYO",
                6: "JUNIO",
                7: "JULIO",
                8: "AGOSTO",
                9: "SEPTIEMBRE",
                10: "OCTUBRE",
                11: "NOVIEMBRE",
                12: "DICIEMBRE",
            }

            month = spanish_months[date_obj.month]
            year = self.number_to_letters(date_obj.year)

            return f"{day} DE {month} DEL {year}"
        except Exception as e:
            print(f"ERROR: Failed to convert date to letters: {e}")
            return str(date_str)

    def money_to_letters(self, currency: str, amount: Decimal) -> str:
        """Convert money amount to words in Spanish"""
        try:
            integer_part = int(amount)
            decimal_part = int((amount % 1) * 100)

            integer_text = self.number_to_letters(integer_part)
            decimal_text = self.number_to_letters(decimal_part)

            if currency == "SOLES":
                return f"{integer_text} SOLES CON {decimal_text} CÉNTIMOS"
            elif currency == "DOLARES N.A.":
                return f"{integer_text} DÓLARES AMERICANOS CON {decimal_text} CENTAVOS"
            else:
                return f"{integer_text} {currency} CON {decimal_text} CENTAVOS"
        except Exception as e:
            print(f"ERROR: Failed to convert money to letters: {e}")
            return f"{amount} {currency}"


# In utils.py - Add this class
class DocumentFormatter:
    """
    Format document data for templates
    Handles complex formatting logic that was in PHP
    """

    def __init__(self, letras_converter):
        self.letras = letras_converter

    def _es_acto_donacion(self, raw_data: dict) -> bool:
        acto_u = (raw_data.get("acto") or "").upper()
        plantilla_u = (raw_data.get("plantilla") or "").upper()
        cond_part_u = (raw_data.get("condicion") or "").upper()
        cond_emp_u = (raw_data.get("condicion_empresa") or "").upper()
        return (
            "DONAC" in acto_u
            or "DONAC" in plantilla_u
            or "DONATARIO" in cond_emp_u
            or ("DONANTE" in cond_part_u and "REPRESENTANTE" in cond_part_u)
        )

    def _es_acto_constitucion(self, raw_data: dict) -> bool:
        acto_u = (raw_data.get("acto") or "").upper()
        plantilla_u = (raw_data.get("plantilla") or "").upper()
        return "CONSTITUC" in acto_u or "CONSTITUC" in plantilla_u

    def _infer_company_target_slot(self, raw_data: dict) -> int:
        """
        Decide whether current PJ should be exposed as slot _1 (P side)
        or slot _2 (C side), based on contractor role data.
        """
        condiciones = (raw_data.get("condicion") or "").split(",")
        tipos_persona = (raw_data.get("tipo_persona") or raw_data.get("tipper") or "").split(",")
        partes = (raw_data.get("parte") or "").split(",")
        uifs = (raw_data.get("uif") or "").split(",")

        transferor_roles = {
            "VENDEDOR",
            "DONANTE",
            "OTORGANTE",
            "TRANSFERENTE",
            "PROPIETARIO",
        }
        acquirer_roles = {
            "COMPRADOR",
            "DONATARIO",
            "ADQUIRIENTE",
            "OTORGADO",
            "ACREEDOR",
        }

        max_len = max(len(condiciones), len(tipos_persona), len(partes), len(uifs))
        for i in range(max_len):
            tip = (tipos_persona[i].strip().upper() if i < len(tipos_persona) else "")
            if tip != "J":
                continue

            cond = (condiciones[i].strip().upper() if i < len(condiciones) else "")
            parte = (partes[i].strip() if i < len(partes) else "")
            uif = (uifs[i].strip().upper() if i < len(uifs) else "")

            if cond in acquirer_roles or parte == "2" or uif in {"B", "N"}:
                return 2
            if cond in transferor_roles or parte == "1" or uif == "O":
                return 1

        return 1

    def format_document_data(self, raw_data):
        """
        Format basic document data
        """

        return {
            "K": raw_data.get("kardex", ""),
            "NUM_REG": "1",
            "USUARIO": raw_data.get("usuario", ""),
            "USUARIO_DNI": raw_data.get("dni_usuario", ""),
            "COMPROBANTE": "sin",
            "O_S": raw_data.get("kardex", ""),
            "ORDEN_SERVICIO": raw_data.get("kardex", ""),
            "DESCRIPCION_SELLO": f"{raw_data.get('abogado', '')} PUNO {raw_data.get('matricula', '')}",
        }

    def format_contractor_data(self, raw_data):
        """
        Format contractor data into template-ready format
        Mirrors: PHP get_data_contratantes() complex logic
        """

        # Define role classifications
        TRANSFEROR_ROLES = {
            "VENDEDOR",
            "DONANTE",
            "PODERDANTE",
            "OTORGANTE",
            "ANTICIPANTE",
            "ADJUDICANTE",
            "USUFRUCTUANTE",
            "TRANSFERENTE",
            "TITULAR",
            "MUTUANTE",
            "PROPIETARIO",
            "ASOCIANTE",
            "ASOCIADO",  # Added for constitution documents
            "TRANSFERENTE / PROPIETARIO (VENDEDOR)",
            "SOLICITANTE/BENEFICIARIO",  # No Contenciosos
        }

        ACQUIRER_ROLES = {
            "COMPRADOR",
            "DONATARIO",
            "APODERADO",
            "ANTICIPADO",
            "ADJUDICATARIO",
            "USUFRUCTUARIO",
            "TESTIGO A RUEGO",
            "ADQUIRIENTE",
            "ACREEDOR",
            "OTORGADO",
            "MUTUATARIO",
            "BENEFICIARIA",
            "ASOCIADO",
            "ADQUIRENTE / BENEFICIARIO (COMPRADOR)",
            "CAUSANTE",  # No Contenciosos
            "DEUDOR",  # Garantias Mobiliarias (goes to C_ side per PHP and template)
            "CONSTITUYENTE",  # Garantias Mobiliarias (goes to C_ side per template)
            "CONSTITUYENTE - DEUDOR",  # Garantias Mobiliarias (goes to C_ side per template)
        }

        REPRESENTATIVE_ROLES = {"APODERADO", "REPRESENTANTE"}

        is_donacion = self._es_acto_donacion(raw_data)

        # Parse contractor data from raw_data
        contractors = self._parse_contractor_data(raw_data)
        
        print(f"DEBUG: Parsed {len(contractors)} contractors")

        # Classify contractors
        transferors = []
        acquirers = []
        transferor_companies = []
        acquirer_companies = []

        for contractor in contractors:
            print(f"DEBUG: Classifying {contractor.get('nombres', 'NO_NAME')} - Role: {contractor.get('condicion_str', 'NO_ROLE')}")
            cond = (contractor.get("condicion_str") or "").strip()
            uif_c = (contractor.get("uif") or "").strip().upper()
            cond_emp = (contractor.get("condicion_empresa") or "").upper()
            parte_s = contractor.get("parte")
            try:
                parte_i = (
                    int(str(parte_s).strip())
                    if parte_s is not None and str(parte_s).strip() != ""
                    else 0
                )
            except (TypeError, ValueError):
                parte_i = 0

            # REPRESENTANTE / UIF R: lado P (donantes/otorgantes) vs C (donatarios) según la PJ
            # que representan (condicion_empresa desde actocondicion de la empresa en idcontratanterp),
            # parte y UIF. Sin esto, el gerente del donatario caía por defecto en P_.
            if cond == "REPRESENTANTE" or uif_c == "R":
                if "DONATARIO" in cond_emp:
                    acquirers.append(contractor)
                    print("DEBUG: -> ACQUIRER (representante de persona jurídica donataria)")
                    continue
                if "DONANTE" in cond_emp and "DONATARIO" not in cond_emp:
                    transferors.append(contractor)
                    print("DEBUG: -> TRANSFEROR (representante de persona jurídica donante)")
                    continue
                if parte_i == 2:
                    acquirers.append(contractor)
                    print("DEBUG: -> ACQUIRER (representante parte=2)")
                elif parte_i == 1:
                    transferors.append(contractor)
                    print("DEBUG: -> TRANSFEROR (representante parte=1)")
                elif uif_c in ("B", "N"):
                    acquirers.append(contractor)
                    print(f"DEBUG: -> ACQUIRER (representante uif={uif_c})")
                elif uif_c == "O":
                    transferors.append(contractor)
                    print("DEBUG: -> TRANSFEROR (representante uif=O)")
                elif is_donacion and (contractor.get("idContratanteRepresentado") or "").strip():
                    acquirers.append(contractor)
                    print("DEBUG: -> ACQUIRER (donación: representante con PJ, parte/uif ambiguos)")
                else:
                    transferors.append(contractor)
                    print("DEBUG: -> TRANSFEROR (representante default otorgante-side)")
                continue

            if contractor["tipper"] == "J":  # Company
                # Companies are classified based on their representatives' roles
                if contractor["condicion_str"] in TRANSFEROR_ROLES:
                    transferor_companies.append(contractor)
                    print(f"DEBUG: -> TRANSFEROR COMPANY")
                elif contractor["condicion_str"] in ACQUIRER_ROLES:
                    acquirer_companies.append(contractor)
                    print(f"DEBUG: -> ACQUIRER COMPANY")
            elif contractor["condicion_str"] in TRANSFEROR_ROLES:
                transferors.append(contractor)
                print(f"DEBUG: -> TRANSFEROR")
            elif contractor["condicion_str"] in ACQUIRER_ROLES:
                acquirers.append(contractor)
                print(f"DEBUG: -> ACQUIRER")
            else:
                print(f"DEBUG: -> UNCLASSIFIED!")
        
        print(f"DEBUG: Total transferors: {len(transferors)}, Total acquirers: {len(acquirers)}")

        # Generate contractor placeholders
        contractor_data = {}

        # Process transferors (P_ prefix)
        for idx, t in enumerate(transferors, 1):
            # Apply gender agreement for nationality and civil status
            nacionalidad = self._apply_gender_to_word(t["nacionalidad"], t["sexo"])
            estado_civil = self._apply_gender_to_word(t["estadoCivil"], t["sexo"])
            
            contractor_data[f"P_NOM_{idx}"] = t["nombres"] + " "
            contractor_data[f"P_NACIONALIDAD_{idx}"] = nacionalidad + " "
            contractor_data[f"P_DOC_{idx}"] = self._get_identification_phrase(
                t["sexo"], t["tipoDocumento"], t["numeroDocumento"]
            ) + " "
            contractor_data[f"P_IDE_{idx}"] = " "
            contractor_data[f"P_ORIGEN_FONDO_{idx}"] = (t.get("origen_fondo") or "").strip() + " "
            contractor_data[f"P_OCUPACION_{idx}"] = t["ocupacion"] + " "
            contractor_data[f"P_ESTADO_CIVIL_{idx}"] = estado_civil + " "
            contractor_data[f"P_DOMICILIO_{idx}"] = self._format_address_with_ubigeo(
                t["direccion"], t["distrito"], t["provincia"], t["departamento"]
            )

            # Add unnumbered versions for first person
            if idx == 1:
                contractor_data["P_NOM"] = t["nombres"] + " "
                contractor_data["P_NACIONALIDAD"] = nacionalidad + " "
                contractor_data["P_DOC"] = self._get_identification_phrase(
                    t["sexo"], t["tipoDocumento"], t["numeroDocumento"]
                ) + " "
                contractor_data["P_IDE"] = " "
                contractor_data["P_ORIGEN_FONDO"] = (t.get("origen_fondo") or "").strip() + " "
                contractor_data["P_OCUPACION"] = t["ocupacion"] + " "
                contractor_data["P_ESTADO_CIVIL"] = estado_civil + " "
                contractor_data["P_DOMICILIO"] = self._format_address_with_ubigeo(
                    t["direccion"], t["distrito"], t["provincia"], t["departamento"]
                )
                # CALIDAD_P will be set by _get_articles_and_grammar method

        # Process acquirers (C_ prefix)
        for idx, c in enumerate(acquirers, 1):
            # Apply gender agreement for nationality and civil status
            nacionalidad = self._apply_gender_to_word(c["nacionalidad"], c["sexo"])
            estado_civil = self._apply_gender_to_word(c["estadoCivil"], c["sexo"])
            
            contractor_data[f"C_NOM_{idx}"] = c["nombres"] + " "
            contractor_data[f"C_NACIONALIDAD_{idx}"] = nacionalidad + " "
            contractor_data[f"C_DOC_{idx}"] = self._get_identification_phrase(
                c["sexo"], c["tipoDocumento"], c["numeroDocumento"]
            ) + " "
            contractor_data[f"C_IDE_{idx}"] = " "
            contractor_data[f"C_ORIGEN_FONDO_{idx}"] = (c.get("origen_fondo") or "").strip() + " "
            contractor_data[f"C_OCUPACION_{idx}"] = c["ocupacion"] + " "
            contractor_data[f"C_ESTADO_CIVIL_{idx}"] = estado_civil + " "
            contractor_data[f"C_DOMICILIO_{idx}"] = self._format_address_with_ubigeo(
                c["direccion"], c["distrito"], c["provincia"], c["departamento"]
            )

            # Add unnumbered versions for first person
            if idx == 1:
                contractor_data["C_NOM"] = c["nombres"] + " "
                contractor_data["C_NACIONALIDAD"] = nacionalidad + " "
                contractor_data["C_DOC"] = self._get_identification_phrase(
                    c["sexo"], c["tipoDocumento"], c["numeroDocumento"]
                ) + " "
                contractor_data["C_IDE"] = " "
                contractor_data["C_ORIGEN_FONDO"] = (c.get("origen_fondo") or "").strip() + " "
                contractor_data["C_OCUPACION"] = c["ocupacion"] + " "
                contractor_data["C_ESTADO_CIVIL"] = estado_civil + " "
                contractor_data["C_DOMICILIO"] = self._format_address_with_ubigeo(
                    c["direccion"], c["distrito"], c["provincia"], c["departamento"]
                )
                # CALIDAD_C will be set by _get_articles_and_grammar method


        # Fill empty placeholders for unused slots
        self._fill_empty_contractor_placeholders(contractor_data, len(transferors), len(acquirers))
        self._fill_empty_company_placeholders(contractor_data, len(transferor_companies), len(acquirer_companies))

        # Add gender-based articles and grammar
        contractor_data.update(self._get_articles_and_grammar(transferors, "P"))
        contractor_data.update(self._get_articles_and_grammar(acquirers, "C"))

        # Backward-compatible aliases used by multiple templates
        # (many templates use P_CALIDAD/C_CALIDAD instead of CALIDAD_P/CALIDAD_C)
        if contractor_data.get("CALIDAD_P") and not contractor_data.get("P_CALIDAD"):
            contractor_data["P_CALIDAD"] = contractor_data["CALIDAD_P"]
        if contractor_data.get("CALIDAD_C") and not contractor_data.get("C_CALIDAD"):
            contractor_data["C_CALIDAD"] = contractor_data["CALIDAD_C"]
        
        # Add additional grammar placeholders
        contractor_data.update(self._get_additional_grammar(transferors, acquirers))

        if is_donacion:
            np, nq = len(transferors), len(acquirers)
            p_cal = "DONANTES" if np > 1 else "DONANTE"
            c_cal = "DONATARIOS" if nq > 1 else "DONATARIO"
            p_cal += " "
            c_cal += " "
            contractor_data["P_CALIDAD"] = p_cal
            contractor_data["C_CALIDAD"] = c_cal
            contractor_data["CALIDAD_P"] = p_cal
            contractor_data["CALIDAD_C"] = c_cal

        return contractor_data

    def _parse_contractor_data(self, raw_data):
        """
        Parse contractor data from raw SQL query result
        """

        # Split comma-separated values from raw_data
        def split_if_not_none(value, separator=","):
            return value.split(separator) if value else []

        # Extract arrays from raw_data
        condiciones = split_if_not_none(raw_data.get("condicion"))
        nombres = split_if_not_none(raw_data.get("nombres"))
        nacionalidades = split_if_not_none(raw_data.get("nacionalidad"))
        tipos_documento = split_if_not_none(raw_data.get("tipo_documento"))
        numeros_documento = split_if_not_none(raw_data.get("numero_documento"))
        ocupaciones = split_if_not_none(raw_data.get("ocupacion"))
        origenes_fondo = split_if_not_none(raw_data.get("origen_fondo"))
        estados_civil = split_if_not_none(raw_data.get("estado_civil"))
        direcciones = raw_data.get("direccion", "").split(",,") if raw_data.get("direccion") else []
        distritos = split_if_not_none(raw_data.get("distrito"))
        provincias = split_if_not_none(raw_data.get("provincia"))
        departamentos = split_if_not_none(raw_data.get("departamento"))
        sexos = split_if_not_none(raw_data.get("sexo"))
        id_clientes = split_if_not_none(raw_data.get("id_cliente"))
        id_conyuges = split_if_not_none(raw_data.get("id_conyuge"))
        uifs = split_if_not_none(raw_data.get("uif"))
        partes = split_if_not_none(raw_data.get("parte"))
        id_contratantes = split_if_not_none(raw_data.get("id_contratante"))
        # Per fila en contratantes: idcontratanterp → idcontratante de la persona jurídica representada (ej. empresa donataria)
        id_contratantes_representados = split_if_not_none(raw_data.get("id_empresa"))
        
        # Company fields - use separate company fields from query
        razones_sociales = split_if_not_none(raw_data.get("nombre_empresa"))
        domicilios_fiscales = split_if_not_none(raw_data.get("domicilio_empresa"))
        tipos_persona_empresa = split_if_not_none(raw_data.get("tipo_persona_empresa"))
        condiciones_empresa = split_if_not_none(raw_data.get("condicion_empresa"))
        # In SQL this field is exposed as "tipo_persona" (GROUP_CONCAT(c2.tipper)).
        # Keep fallback to "tipper" for compatibility with older query variants.
        tipos_persona = split_if_not_none(raw_data.get("tipo_persona") or raw_data.get("tipper"))

        contractors = []

        # Process each contractor
        for k, condicion in enumerate(condiciones):
            if not condicion:
                continue

            # Get data for this contractor
            nombre = nombres[k] if k < len(nombres) else ""
            nacionalidad = nacionalidades[k] if k < len(nacionalidades) else ""
            tipo_doc = tipos_documento[k] if k < len(tipos_documento) else ""
            num_doc = numeros_documento[k] if k < len(numeros_documento) else ""
            ocupacion = ocupaciones[k] if k < len(ocupaciones) else ""
            origen_fondo = origenes_fondo[k] if k < len(origenes_fondo) else ""
            estado_civil = estados_civil[k] if k < len(estados_civil) else ""
            direccion = direcciones[k] if k < len(direcciones) else ""
            distrito = distritos[k] if k < len(distritos) else ""
            provincia = provincias[k] if k < len(provincias) else ""
            departamento = departamentos[k] if k < len(departamentos) else ""
            sexo = sexos[k] if k < len(sexos) else "M"
            id_cliente = id_clientes[k] if k < len(id_clientes) else ""
            id_conyuge = id_conyuges[k] if k < len(id_conyuges) else "NO"
            uif_k = uifs[k] if k < len(uifs) else ""
            parte_k = partes[k] if k < len(partes) else ""
            id_cte = id_contratantes[k] if k < len(id_contratantes) else ""
            id_cte_representado = (
                id_contratantes_representados[k]
                if k < len(id_contratantes_representados)
                else ""
            )

            # Company data
            razon_social = razones_sociales[k] if k < len(razones_sociales) else ""
            domicilio_fiscal = domicilios_fiscales[k] if k < len(domicilios_fiscales) else ""
            tipper_empresa = tipos_persona_empresa[k] if k < len(tipos_persona_empresa) else ""
            condicion_empresa = condiciones_empresa[k] if k < len(condiciones_empresa) else ""
            tipper = (tipos_persona[k].strip().upper() if k < len(tipos_persona) else "N")
            if tipper not in {"N", "J"}:
                tipper = "N"

            contractor = {
                "condiciones": condicion,
                "condicion_str": condicion,  # For role checking
                "nombres": nombre,
                "nacionalidad": nacionalidad,
                "tipoDocumento": tipo_doc,
                "numeroDocumento": num_doc,
                "ocupacion": ocupacion,
                "origen_fondo": origen_fondo,
                "estadoCivil": estado_civil,
                "direccion": direccion,
                "distrito": distrito,
                "provincia": provincia,
                "departamento": departamento,
                "sexo": sexo,
                "idCliente": id_cliente,
                "idConyuge": id_conyuge,
                "uif": uif_k,
                "parte": parte_k,
                # contratantesxacto.idcontratante de ESTE participante (ej. el representante)
                "idContratante": id_cte,
                # contratantes.idcontratanterp: si firma por empresa, apunta al idcontratante de la persona jurídica
                "idContratanteRepresentado": id_cte_representado.strip()
                if id_cte_representado
                else "",
                "tipper": tipper,  # N = Natural, J = Juridical
                # Company fields
                "razonsocial": razon_social,
                "domfiscal": domicilio_fiscal,
                "tipper_empresa": tipper_empresa,
                "condicion_empresa": condicion_empresa,
                "numdoc_empresa": num_doc if tipper == "J" else "",
            }

            contractors.append(contractor)

        return contractors

    def _resolve_nombre_empresa_aggregate(
        self,
        nombre_raw: str,
        tipo_persona_empresa_raw: str = "",
        id_empresa_raw: str = "",
    ) -> str:
        """
        GROUP_CONCAT mezcla filas; el nombre PJ debe alinearse con idcontratanterp (id_empresa) o tipper J.
        """
        raw = (nombre_raw or "").strip().rstrip(",")
        if not raw:
            return ""
        segments = [s.strip() for s in raw.split(",") if s.strip()]
        if not segments:
            return ""
        ids_parts = (id_empresa_raw or "").split(",")
        if len(ids_parts) == len(segments):
            for i, eid in enumerate(ids_parts):
                eid = (eid or "").strip()
                if eid and i < len(segments) and (segments[i] or "").strip():
                    return segments[i]
        tipos = [t.strip() for t in (tipo_persona_empresa_raw or "").split(",") if t.strip() != ""]
        if tipos and len(tipos) == len(segments):
            for i, t in enumerate(tipos):
                if t == "J" and i < len(segments):
                    return segments[i]
        if len(segments) == 1:
            return segments[0]
        return max(segments, key=len)

    def format_company_data(self, raw_data):
        """
        Format company data separately from contractors
        Companies are represented by juridical persons in the database
        """
        company_data = {}
        
        # Get company data from raw_data (handle None values)
        # First try regular company data (from idcontratanterp)
        nombre_empresa = self._resolve_nombre_empresa_aggregate(
            raw_data.get("nombre_empresa") or "",
            raw_data.get("tipo_persona_empresa") or "",
            raw_data.get("id_empresa") or "",
        )
        domicilio_empresa = (raw_data.get("domicilio_empresa") or "").strip()
        tipo_persona_empresa = raw_data.get("tipo_persona_empresa") or ""
        condicion_empresa = (raw_data.get("condicion_empresa") or "").strip()
        numero_documento_empresa = (raw_data.get("numero_documento_empresa") or "").strip()
        numero_partida = (raw_data.get("numero_partida") or "").strip()
        distrito_empresa = (raw_data.get("distrito_empresa") or "").strip()
        provincia_empresa = (raw_data.get("provincia_empresa") or "").strip()
        departamento_empresa = (raw_data.get("departamento_empresa") or "").strip()
        oficina_registral = (raw_data.get("oficina_registral") or "").strip()
        
        # If no company data, try alternate company data from service query fallback.
        if not nombre_empresa:
            nombre_empresa = self._resolve_nombre_empresa_aggregate(
                raw_data.get("nombre_empresa_constitucion") or "",
                raw_data.get("tipo_persona_empresa_constitucion") or "",
                "",
            )
            domicilio_empresa = (raw_data.get("domicilio_empresa_constitucion") or "").strip()
            tipo_persona_empresa = raw_data.get("tipo_persona_empresa_constitucion") or ""
            numero_documento_empresa = (raw_data.get("numero_documento_empresa_constitucion") or "").strip()
            numero_partida = (raw_data.get("numero_partida_constitucion") or "").strip()
            # Mark as constitution only when the current act/template is really a constitution.
            if self._es_acto_constitucion(raw_data):
                condicion_empresa = "EMPRESA EN CONSTITUCION"
        
        
        # Process company data if it exists (joined cliente2 cr2 via idcontratanterp).
        # Do not require "J" in tipo_persona_empresa: MySQL GROUP_CONCAT skips NULL segments, so the
        # letter J can disappear from the aggregate even when nombre_empresa / RUC from cr2 are present.
        is_donacion = self._es_acto_donacion(raw_data)
        target_slot = self._infer_company_target_slot(raw_data)
        if (nombre_empresa or "").strip():
            ins_txt = (
                f" INSCRITA EN LA PARTIDA ELECTRONICA N° {numero_partida} DE LA OFICINA REGISTRAL {oficina_registral}"
                if numero_partida
                else ""
            )
            ruc_txt = f", CON RUC N° {numero_documento_empresa}, " if numero_documento_empresa else ""
            dom_txt = (
                f"CON DOMICILIO EN {domicilio_empresa} DEL DISTRITO DE {distrito_empresa} PROVINCIA DE {provincia_empresa} Y DEPARTAMENTO DE {departamento_empresa}"
                if domicilio_empresa
                else ""
            )
            cond_emp = condicion_empresa if condicion_empresa else ""
            # Determine which company slot to use based on condition
            if condicion_empresa in ['EMPRESA EN CONSTITUCION', 'ASOCIACION EN CONSTITUCION']:
                print(f"DEBUG: Setting NOMBRE_EMPRESA_2 (constitution)")
                company_data["NOMBRE_EMPRESA_2"] = nombre_empresa
                company_data["INS_EMPRESA_2"] = ins_txt
                company_data["RUC_2"] = ruc_txt
                company_data["DOMICILIO_EMPRESA_2"] = dom_txt
                company_data["CONDICION_EMPRESA_2"] = condicion_empresa
            elif is_donacion:
                # Donación: plantilla usa NOMBRE_EMPRESA_2 para el donatario PJ; algunas líneas usan INS/RUC_1
                print(f"DEBUG: Setting empresa donatario (donación) en _2 y duplicando legales en _1")
                company_data["NOMBRE_EMPRESA_2"] = nombre_empresa
                company_data["INS_EMPRESA_2"] = ins_txt
                company_data["RUC_2"] = ruc_txt
                company_data["DOMICILIO_EMPRESA_2"] = dom_txt
                company_data["CONDICION_EMPRESA_2"] = cond_emp
                company_data["NOMBRE_EMPRESA_1"] = nombre_empresa
                company_data["INS_EMPRESA_1"] = ins_txt
                company_data["RUC_1"] = ruc_txt
                company_data["DOMICILIO_EMPRESA_1"] = dom_txt
                company_data["CONDICION_EMPRESA_1"] = cond_emp
            else:
                if target_slot == 2:
                    print(f"DEBUG: Setting NOMBRE_EMPRESA_2 (normal C-side PJ)")
                    company_data["NOMBRE_EMPRESA_2"] = nombre_empresa
                    company_data["INS_EMPRESA_2"] = ins_txt
                    company_data["RUC_2"] = ruc_txt
                    company_data["DOMICILIO_EMPRESA_2"] = dom_txt
                    company_data["CONDICION_EMPRESA_2"] = cond_emp
                else:
                    print(f"DEBUG: Setting NOMBRE_EMPRESA_1 (normal P-side PJ)")
                    company_data["NOMBRE_EMPRESA_1"] = nombre_empresa
                    company_data["INS_EMPRESA_1"] = ins_txt
                    company_data["RUC_1"] = ruc_txt
                    company_data["DOMICILIO_EMPRESA_1"] = dom_txt
                    company_data["CONDICION_EMPRESA_1"] = cond_emp
        
        # Fill empty company placeholders
        if "NOMBRE_EMPRESA_1" not in company_data:
            company_data["NOMBRE_EMPRESA_1"] = ""
            company_data["INS_EMPRESA_1"] = ""
            company_data["RUC_1"] = ""
            company_data["DOMICILIO_EMPRESA_1"] = ""
            company_data["CONDICION_EMPRESA_1"] = ""
        
        if "NOMBRE_EMPRESA_2" not in company_data:
            company_data["NOMBRE_EMPRESA_2"] = ""
            company_data["INS_EMPRESA_2"] = ""
            company_data["RUC_2"] = ""
            company_data["DOMICILIO_EMPRESA_2"] = ""
            company_data["CONDICION_EMPRESA_2"] = ""
        
        return company_data

    def _apply_gender_to_word(self, word, gender):
        """
        Apply gender agreement to Spanish words
        Mirrors: PHP logic - changes last letter based on gender
        
        Examples:
        - PERUANO (M) -> PERUANO, PERUANO (F) -> PERUANA
        - SOLTERO (M) -> SOLTERO, SOLTERO (F) -> SOLTERA
        """
        if not word or not word.strip():
            return word
        
        word = word.strip()
        
        # For female, change last 'O' to 'A'
        if gender == "F":
            if word.endswith("O"):
                return word[:-1] + "A"
        # For male, ensure it ends with 'O' (if it ended with 'A', change it)
        elif gender == "M":
            if word.endswith("A"):
                return word[:-1] + "O"
        
        return word
    
    def _get_identification_phrase(self, gender, doc_type, doc_number):
        """
        Generate identification phrase based on gender and document type
        """
        if not doc_number:
            return ""

        if gender == "F":
            return f"IDENTIFICADA CON {doc_type} N° {doc_number}"
        else:
            return f"IDENTIFICADO CON {doc_type} N° {doc_number}"

    def _format_address_with_ubigeo(self, address, distrito, provincia, departamento):
        """
        Build address phrase always followed by ubigeo names.
        """
        address = (address or "").strip()
        if not address:
            return ""

        distrito = self._resolve_distrito_name((distrito or "").strip())
        provincia = (provincia or "").strip()
        departamento = (departamento or "").strip()

        return (
            f"CON DOMICILIO EN {address} DEL DISTRITO DE {distrito} "
            f"PROVINCIA DE {provincia} Y DEPARTAMENTO DE {departamento} "
        )

    def _resolve_distrito_name(self, distrito):
        """
        Some rows return coddis (e.g. 010102) instead of nomdis.
        Resolve the code to district name from ubigeo table.
        """
        distrito = (distrito or "").strip()
        if not distrito:
            return ""
        if not re.fullmatch(r"\d{6}", distrito):
            return distrito

        try:
            from django.db import connection

            with connection.cursor() as cursor:
                cursor.execute(
                    "SELECT IFNULL(nomdis, '') FROM ubigeo WHERE coddis = %s LIMIT 1",
                    [distrito],
                )
                row = cursor.fetchone()
                if row and (row[0] or "").strip():
                    return row[0].strip()
        except Exception:
            pass

        return distrito

    def _fill_empty_contractor_placeholders(self, contractor_data, num_transferors, num_acquirers):
        """
        Fill empty placeholders for unused contractor slots
        """
        # Fill empty transferor slots (P_ prefix) - use empty strings, not [E.PLACEHOLDER]
        for idx in range(num_transferors + 1, 11):
            contractor_data[f"P_NOM_{idx}"] = ""  # Empty string instead of [E.P_NOM_X]
            contractor_data[f"P_NACIONALIDAD_{idx}"] = ""
            contractor_data[f"P_DOC_{idx}"] = ""
            contractor_data[f"P_IDE_{idx}"] = ""
            contractor_data[f"P_ORIGEN_FONDO_{idx}"] = ""
            contractor_data[f"P_OCUPACION_{idx}"] = ""
            contractor_data[f"P_ESTADO_CIVIL_{idx}"] = ""
            contractor_data[f"P_DOMICILIO_{idx}"] = ""

        # Fill empty acquirer slots (C_ prefix) - use empty strings, not [E.PLACEHOLDER]
        for idx in range(num_acquirers + 1, 11):
            contractor_data[f"C_NOM_{idx}"] = ""  # Empty string instead of [E.C_NOM_X]
            contractor_data[f"C_NACIONALIDAD_{idx}"] = ""
            contractor_data[f"C_DOC_{idx}"] = ""
            contractor_data[f"C_IDE_{idx}"] = ""
            contractor_data[f"C_ORIGEN_FONDO_{idx}"] = ""
            contractor_data[f"C_OCUPACION_{idx}"] = ""
            contractor_data[f"C_ESTADO_CIVIL_{idx}"] = ""
            contractor_data[f"C_DOMICILIO_{idx}"] = ""

    def _fill_empty_company_placeholders(self, contractor_data, num_transferor_companies, num_acquirer_companies):
        """
        Fill empty placeholders for unused company slots
        """
        total_companies = num_transferor_companies + num_acquirer_companies
        
        # Fill empty company slots - use empty strings, not [E.PLACEHOLDER]
        for idx in range(total_companies + 1, 6):  # Support up to 5 companies
            contractor_data[f"NOMBRE_EMPRESA_{idx}"] = ""
            contractor_data[f"INS_EMPRESA_{idx}"] = ""
            contractor_data[f"RUC_{idx}"] = ""
            contractor_data[f"DOMICILIO_EMPRESA_{idx}"] = ""

    def _get_articles_and_grammar(self, contractors, role_prefix):
        """
        Generate gender-based articles and grammar
        Mirrors: PHP articulos_singular_plural() function
        """
        if not contractors:
            return {}
        
        # Determine gender and number
        genders = [c.get("sexo", "M") for c in contractors]
        num_contractors = len(contractors)
        
        # Determine if all women (for proper grammar)
        all_women = all(g == "F" for g in genders)
        
        # Generate articles based on gender and number
        articles = {}
        
        if num_contractors >= 2:
            # PLURAL - Based on PHP logic
            if all_women:
                articles[f"EL_{role_prefix}"] = "LAS"
                articles[f"INICIO_{role_prefix}"] = "SEÑORAS"
                # Set CALIDAD based on role and gender
                if role_prefix == "P":
                    articles[f"CALIDAD_{role_prefix}"] = "VENDEDORAS"
                else:
                    articles[f"CALIDAD_{role_prefix}"] = "COMPRADORAS"
            else:
                articles[f"EL_{role_prefix}"] = "LOS"
                articles[f"INICIO_{role_prefix}"] = "SEÑORES"
                # Set CALIDAD based on role and gender
                if role_prefix == "P":
                    articles[f"CALIDAD_{role_prefix}"] = "VENDEDORES"
                else:
                    articles[f"CALIDAD_{role_prefix}"] = "COMPRADORES"
            
            articles[f"ES_{role_prefix}"] = "ES"
            articles[f"S_{role_prefix}"] = "S"
            articles[f"ES_SON_{role_prefix}"] = "SON"
            articles[f"Y_CON_{role_prefix}"] = "Y"
            articles[f"N_{role_prefix}"] = "N"
            articles[f"Y_{role_prefix}"] = "Y"
            articles[f"L_{role_prefix}"] = "L"
            articles[f"O_A_{role_prefix}"] = "OS"
            articles[f"O_ERON_{role_prefix}"] = "ERON"
            articles[f"O_ARON_{role_prefix}"] = "ARON"
            articles[f"OR_{role_prefix}"] = "ES"  # For TESTADORES (plural)
            articles[f"{role_prefix}_FIRMA"] = "FIRMAN EN"
            
        else:
            # SINGULAR - Based on PHP logic
            if all_women:
                articles[f"EL_{role_prefix}"] = "LA"
                articles[f"INICIO_{role_prefix}"] = "SEÑORA"
                # Set CALIDAD based on role and gender (covers multiple document types)
                if role_prefix == "P":
                    articles[f"CALIDAD_{role_prefix}"] = "VENDEDORA"
                else:
                    articles[f"CALIDAD_{role_prefix}"] = "COMPRADORA"
                # Female singular
                articles[f"O_A_{role_prefix}"] = "A"
                articles[f"O_ERON_{role_prefix}"] = "A"
                articles[f"O_ARON_{role_prefix}"] = "A"
                articles[f"OR_{role_prefix}"] = "A"  # For TESTADORA (feminine)
            else:
                articles[f"EL_{role_prefix}"] = "EL"
                articles[f"INICIO_{role_prefix}"] = "SEÑOR"
                # Set CALIDAD based on role and gender (covers multiple document types)
                if role_prefix == "P":
                    articles[f"CALIDAD_{role_prefix}"] = "VENDEDOR"
                else:
                    articles[f"CALIDAD_{role_prefix}"] = "COMPRADOR"
                # Male singular
                articles[f"O_A_{role_prefix}"] = "O"
                articles[f"O_ERON_{role_prefix}"] = "O"
                articles[f"O_ARON_{role_prefix}"] = "O"
                articles[f"OR_{role_prefix}"] = "OR"  # For TESTADOR (masculine)
            
            articles[f"ES_{role_prefix}"] = ""
            articles[f"S_{role_prefix}"] = ""
            articles[f"ES_SON_{role_prefix}"] = "ES"
            articles[f"Y_CON_{role_prefix}"] = ""
            articles[f"N_{role_prefix}"] = ""
            articles[f"Y_{role_prefix}"] = ""
            articles[f"L_{role_prefix}"] = ""
            articles[f"{role_prefix}_FIRMA"] = "FIRMA EN"
            articles[f"{role_prefix}_AMBOS"] = " "
        
        # Add spaces for proper formatting
        for key, value in articles.items():
            articles[key] = value + " "
        
        return articles

    def _get_additional_grammar(self, transferors, acquirers):
        """
        Generate additional grammar placeholders commonly used in notary documents
        """
        grammar = {}
        
        # Determine if we have multiple parties
        has_multiple_transferors = len(transferors) > 1
        has_multiple_acquirers = len(acquirers) > 1
        
        # Articles for multiple parties
        if has_multiple_transferors:
            grammar["LOS_P"] = "LOS "
            grammar["LAS_P"] = "LAS "
        else:
            grammar["LOS_P"] = "EL "
            grammar["LAS_P"] = "LA "
        
        if has_multiple_acquirers:
            grammar["LOS_C"] = "LOS "
            grammar["LAS_C"] = "LAS "
        else:
            grammar["LOS_C"] = "EL "
            grammar["LAS_C"] = "LA "
        
        # Conjunctions
        grammar["Y"] = "Y "
        grammar["O"] = "O "
        
        # Prepositions
        grammar["DE"] = "DE "
        grammar["A"] = "A "
        grammar["EN"] = "EN "
        grammar["POR"] = "POR "
        grammar["CON"] = "CON "
        
        # Common phrases
        grammar["QUE"] = "QUE "
        grammar["QUIEN"] = "QUIEN "
        grammar["QUIENES"] = "QUIENES "
        grammar["CUYO"] = "CUYO "
        grammar["CUYA"] = "CUYA "
        grammar["CUYOS"] = "CUYOS "
        grammar["CUYAS"] = "CUYAS "
        
        return grammar

    def format_payment_data(self, raw_data):
        """
        Format payment data with sunat logic
        Mirrors: PHP get_data_pagos() switch statement
        """

        # Extract payment data
        precio = raw_data.get("precio") or 0
        moneda = raw_data.get("descripcion_moneda") or "SOLES"
        simbolo_moneda = raw_data.get("simbolo_moneda") or "S/."
        sunat_medio_pago = raw_data.get("sunat_medio_pago") or "008"

        # Convert amount to letters
        monto_letras = self.letras.money_to_letters(moneda, Decimal(str(precio)))

        # Handle sunat_medio_pago logic (switch statement from PHP)
        if sunat_medio_pago == "008":
            medio_pago = 'EL COMPRADOR DECLARA QUE HA PAGADO EL PRECIO DEL VEHICULO EN DINERO EN EFECTIVO. NO HABIENDO UTILIZADO NINGÚN MEDIO DE PAGO ESTABLECIDO EN LA LEY Nº 28194, PORQUE EL MONTO TOTAL NO ES IGUAL NI SUPERA LOS S/ 3,500.00 O US$ 1,000.00. EL TIPO Y CÓDIGO DEL MEDIO EMPLEADO ES: "EFECTIVO POR OPERACIONES EN LAS QUE NO EXISTE OBLIGACIÓN DE UTILIZAR MEDIOS DE PAGO-008". INAPLICABLE LA LEY 30730 POR SER EL PAGO DEL PRECIO INFERIOR A 3 UIT.'
            exhibio_medio_pago = "SE DEJA CONSTANCIA QUE PARA LA REALIZACIÓN DEL PRESENTE ACTO, LAS PARTES NO ME HAN EXHIBIDO NINGÚN MEDIO DE PAGO. DOY FE."
            fin_medio_pago = "EN DINERO EN EFECTIVO"
            forma_pago = "AL CONTADO CON DINERO EN EFECTIVO"
        elif sunat_medio_pago == "009":
            medio_pago = 'EL COMPRADOR DECLARA QUE HA PAGADO EL PRECIO DEL VEHICULO EN DINERO EN EFECTIVO Y CON ANTERIORIDAD A LA CELEBRACION DE LA PRESENTE ACTA DE TRANSFERENCIA. NO HABIENDO UTILIZADO NINGÚN MEDIO DE PAGO ESTABLECIDO EN LA LEY Nº 28194, EL TIPO Y CÓDIGO DEL MEDIO EMPLEADO ES: "EFECTIVO POR OPERACIONES EN LAS QUE NO EXISTE OBLIGACIÓN DE UTILIZAR MEDIOS DE PAGO-009". INAPLICABLE LA LEY 30730 POR SER EL PAGO DEL PRECIO INFERIOR A 3 UIT.'
            exhibio_medio_pago = "SE DEJA CONSTANCIA QUE PARA LA REALIZACIÓN DEL PRESENTE ACTO, LAS PARTES NO ME HAN EXHIBIDO NINGÚN MEDIO DE PAGO. DOY FE."
            fin_medio_pago = "EN DINERO EN EFECTIVO"
            forma_pago = "AL CONTADO CON DINERO EN EFECTIVO"
        else:
            medio_pago = 'EL COMPRADOR DECLARA QUE HA PAGADO EL PRECIO DEL VEHICULO CON CHEQUE DEL BANCO DE CREDITO DEL PERÚ N° 1111111 111111 1111, GIRADO POR: YYYYYYYYY A FAVOR DE: XXXXXXXXX POR LA SUMA DE S/ 15,000.00, JULIACA 16/08/2018 EL TIPO Y CÓDIGO DEL MEDIO EMPLEADO ES: "CHEQUE -001" '
            exhibio_medio_pago = "EN APLICACIÓN DE LA LEY 30730, SE DEJA CONSTANCIA QUE PARA LA REALIZACIÓN DEL PRESENTE ACTO, LAS PARTES ME HAN EXHIBIDO EL SIGUIENTE MEDIO DE PAGO: ……… CHEQUE DEL BANCO DE CREDITO DEL PERÚ N° 1111111 111111 1111, GIRADO POR: YYYYYYYYY A FAVOR DE: XXXXXXXXX POR LA SUMA DE S/ 15,000.00, JULIACA 16/08/2018. DOY FE."
            fin_medio_pago = "EN DINERO EN EFECTIVO"
            forma_pago = "AL CONTADO CON DINERO EN EFECTIVO"

        return {
            "MONTO": str(precio),
            "MON_VEHI": moneda,
            "MONTO_LETRAS": monto_letras,
            "MONEDA_C": f"{simbolo_moneda} ",
            "SUNAT_MED_PAGO": sunat_medio_pago,
            "DES_PRE_VEHI": monto_letras,
            "EXH_MED_PAGO": exhibio_medio_pago,
            "MED_PAGO": medio_pago,
            "FIN_MED_PAGO": fin_medio_pago,
            "FORMA_PAGO": forma_pago,
            "C_INICIO_MP": "",
            "TIPO_PAGO_E": "",
            "TIPO_PAGO_C": "",
            "MONTO_MP": "",
            "CONSTANCIA": "",
            "DETALLE_MP": "",
            "FORMA_PAGO_S": "",
            "MONEDA_C_MP": "",
            "MEDIO_PAGO_C": "",
            "MP_MEDIO_PAGO": "",
            "MP_COMPLETO": "",
            "USO": "",
        }

    def format_vehicle_data(self, raw_data):
        """
        Format vehicle data
        Mirrors: PHP get_data_vehiculos()
        """

        # Extract sede registral and parse it
        sede = raw_data.get("sede") or ""
        sede_parts = sede.split("-") if sede else ["", ""]
        sede_name = sede_parts[1].strip() if len(sede_parts) > 1 else ""

        return {
            "PLACA": str(raw_data.get("placa") or "").upper(),
            "CLASE": str(raw_data.get("clase") or "").upper(),
            "MARCA": str(raw_data.get("marca") or "").upper(),
            "MODELO": str(raw_data.get("modelo") or "").upper(),
            "AÑO_FABRICACION": str(raw_data.get("anio") or "").upper(),
            "CARROCERIA": str(raw_data.get("carroceria") or "").upper(),
            "COLOR": str(raw_data.get("color") or "").upper(),
            "NRO_MOTOR": str(raw_data.get("motor") or "").upper(),
            "NRO_SERIE": str(raw_data.get("serie") or "").upper(),
            "PARTIDA": str(raw_data.get("partida") or "").upper(),
            "FEC_INS": str(raw_data.get("fecha_inscripcion") or "").upper(),
            "FECHA_INSCRIPCION": str(raw_data.get("fecha_inscripcion") or "").upper(),
            "ZONA_REGISTRAL": str(sede).upper(),
            "NUM_ZONA_REG": str(raw_data.get("numero_zona") or "").upper(),
            "SEDE": str(sede_name).upper(),
            "ZONA_VEHICULO": str(sede_name).upper(),
            "INSTRUIDO": " ",
            "COMBUSTIBLE": str(raw_data.get("combustible") or "").upper(),
            "NRO_TARJETA": " ",
        }

    def format_escrituracion_data(self, raw_data):
        """
        Format escrituracion data (folio and papel numbers)
        """


        numero_escritura = raw_data.get("numero_escritura") or ""
        fecha_escritura = raw_data.get("fecha_escritura")
        numero_minuta = raw_data.get("numero_minuta") or ""

        numero_acta = (
            f"{numero_escritura}({self.letras.number_to_letters(numero_escritura)})"
            if numero_escritura
            else "{{NRO_ESC}}"
        )
        fecha_impresion = (
            self.letras.date_to_letters(fecha_escritura) if fecha_escritura else "{{F_IMPRESION}}"
        )
        fecha_acta = self.letras.date_to_letters(fecha_escritura) if fecha_escritura else "{{F}}"
        numero_minuta_formatted = numero_minuta if numero_minuta else "{{NRO_MIN}}"
        folio_inicial = raw_data.get("folio_inicial") or ""
        folio_final = raw_data.get("folio_final") or ""
        papel_inicial = raw_data.get("papel_inicial") or ""
        papel_final = raw_data.get("papel_final") or ""

        return {
            "NRO_ESC": numero_acta,
            "FEC_LET": self.letras.date_to_letters(fecha_escritura) if fecha_escritura else "",
            "F_IMPRESION": fecha_impresion,
            "NRO_MIN": numero_minuta_formatted,
            "F": fecha_acta,
            "FI": folio_inicial if folio_inicial else "{{FI}}",
            "S_IN": papel_inicial if papel_inicial else "{{S_IN}}",
            "FF": folio_final if folio_final else "{{FF}}",
            "S_FN": papel_final if papel_final else "{{S_FN}}",
        }

    def combine_all_data(
        self,
        data_documento,
        data_vehiculos,
        data_pagos,
        data_escrituracion,
        data_contratantes,
        data_company=None,
    ):
        """
        Combine all data dictionaries into one
        Mirrors: PHP array merging with +=
        """
        # Merge all dictionaries
        final_data = {
            **data_documento,
            **data_vehiculos,
            **data_pagos,
            **data_escrituracion,
            **data_contratantes,
        }
        
        # Add company data if provided
        if data_company:
            final_data.update(data_company)

        return final_data


# In utils.py - Add this class
class DocxTemplateProcessor:
    """
    Handle placeholder replacement using python-docx-template
    This library is specifically designed for template processing with formatting
    """
    
    def replace_placeholders(self, template_bytes, final_data):
        """
        Replace placeholders using python-docx-template
        This preserves formatting and handles colors better
        """
        try:
            # Load template from bytes
            template = DocxTemplate(io.BytesIO(template_bytes))
            
            # Render the template with data
            template.render(final_data)
            
            # Save to bytes
            output = io.BytesIO()
            template.save(output)
            output.seek(0)
            
            return output.getvalue()
        except Exception as e:
            print(f"ERROR: DocxTemplate processing failed: {e}")
            # Fallback to original method
            return None

# In utils.py - Add this class
class PlaceholderProcessor:
    """
    Handle placeholder replacement in Word documents
    Preserves formatting while replacing text
    """

    def replace_placeholders(self, doc, final_data):
        """
        Replace {{PLACEHOLDERS}} in Word document
        Processes: paragraphs, tables, headers, and footers
        """
        
        # Replace in main document paragraphs
        for paragraph in doc.paragraphs:
            if "{{" in paragraph.text and "}}" in paragraph.text:
                self._replace_in_paragraph(paragraph, final_data)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if "{{" in paragraph.text and "}}" in paragraph.text:
                            self._replace_in_paragraph(paragraph, final_data)
        
        # Replace in headers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                if "{{" in paragraph.text and "}}" in paragraph.text:
                    self._replace_in_paragraph(paragraph, final_data)
            
            # Replace in header tables
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if "{{" in paragraph.text and "}}" in paragraph.text:
                                self._replace_in_paragraph(paragraph, final_data)
        
        # Replace in footers
        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                if "{{" in paragraph.text and "}}" in paragraph.text:
                    self._replace_in_paragraph(paragraph, final_data)
            
            # Replace in footer tables
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if "{{" in paragraph.text and "}}" in paragraph.text:
                                self._replace_in_paragraph(paragraph, final_data)



    def debug_placeholder_formatting(self, doc):
        """
        Debug method to check if placeholders were properly formatted
        This will print information about the formatting of replaced text
        """
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.runs:
                for j, run in enumerate(paragraph.runs):
                    if run.text.strip():
                        is_bold = run.font.bold
                        
                        # Simple color detection
                        try:
                            color_rgb = run.font.color.rgb
                            if hasattr(color_rgb, 'rgb'):
                                color_rgb = color_rgb.rgb
                            elif isinstance(color_rgb, (tuple, list)) and len(color_rgb) >= 3:
                                color_rgb = tuple(color_rgb[:3])
                        except:
                            color_rgb = None
                        
                        is_red = (color_rgb == (255, 0, 0))
                        

    def _replace_in_paragraph(self, paragraph, final_data):
        """
        Replace placeholders in a paragraph while preserving formatting
        Only replaced values will be BOLD and RED, not the entire paragraph
        """
        full_text = paragraph.text

        # Quick check - if no placeholders, skip
        if "{{" not in full_text or "}}" not in full_text:
            return

        # Get original formatting from first run
        first_run_font = None
        if paragraph.runs:
            first_run_font = paragraph.runs[0].font

        # Find all placeholders in the text
        import re
        placeholder_pattern = r'\{\{([^}]+)\}\}'
        matches = list(re.finditer(placeholder_pattern, full_text))
        
        if not matches:
            return

        # Clear all runs
        for run in paragraph.runs:
            run.text = ""

        # Split text by placeholders and create separate runs
        last_end = 0
        
        # Process each placeholder match
        for match in matches:
            placeholder = match.group(0)  # Full placeholder like {{NOMBRE}}
            key = match.group(1)  # Just the key like NOMBRE
            start = match.start()
            end = match.end()

            # Add text before placeholder
            if start > last_end:
                text_before = full_text[last_end:start]
                if text_before:
                    new_run = paragraph.add_run(text_before)
                    if first_run_font:
                        try:
                            new_run.font.bold = first_run_font.bold
                            new_run.font.italic = first_run_font.italic
                            new_run.font.color.rgb = first_run_font.color.rgb
                            if first_run_font.size:
                                new_run.font.size = first_run_font.size
                            if first_run_font.name:
                                new_run.font.name = first_run_font.name
                        except:
                            pass

            # Add replaced placeholder with RED and BOLD formatting
            if key in final_data:
                value = str(final_data[key])
                new_run = paragraph.add_run(value)
                
                # Set bold first
                new_run.font.bold = True
                
                # Set RED color - use direct RGB
                new_run.font.color.rgb = RGBColor(255, 0, 0)
                
                # Preserve font size and name from original
                if first_run_font:
                    try:
                        if first_run_font.size:
                            new_run.font.size = first_run_font.size
                        if first_run_font.name:
                            new_run.font.name = first_run_font.name
                    except:
                        pass
            else:
                # Placeholder not found in data, keep original placeholder
                new_run = paragraph.add_run(placeholder)
                if first_run_font:
                    try:
                        new_run.font.bold = first_run_font.bold
                        new_run.font.italic = first_run_font.italic
                        new_run.font.color.rgb = first_run_font.color.rgb
                        if first_run_font.size:
                            new_run.font.size = first_run_font.size
                        if first_run_font.name:
                            new_run.font.name = first_run_font.name
                    except:
                        pass
            
            last_end = end

        # Add remaining text after last placeholder
        if last_end < len(full_text):
            text_after = full_text[last_end:]
            if text_after:
                new_run = paragraph.add_run(text_after)
                if first_run_font:
                    try:
                        new_run.font.bold = first_run_font.bold
                        new_run.font.italic = first_run_font.italic
                        new_run.font.color.rgb = first_run_font.color.rgb
                        if first_run_font.size:
                            new_run.font.size = first_run_font.size
                        if first_run_font.name:
                            new_run.font.name = first_run_font.name
                    except:
                        pass

    def clean_unfilled_placeholders(self, doc):
        """
        Remove or hide unfilled placeholders
        Mirrors: PHP placeholder cleanup
        """

        # Define placeholder categories
        escrituracion_placeholders = {
            "{{FI}}",
            "{{FF}}",
            "{{S_IN}}",
            "{{S_FN}}",
            "{{NRO_ESC}}",
            "{{F}}",
            "{{F_IMPRESION}}",
        }

        contractor_placeholders = {
            "{{P_NOM}}",
            "{{P_NOM_2}}",
            "{{P_NOM_3}}",
            "{{P_NOM_4}}",
            "{{P_NOM_5}}",
            "{{C_NOM}}",
            "{{C_NOM_2}}",
            "{{C_NOM_3}}",
            "{{C_NOM_4}}",
            "{{C_NOM_5}}",
            "{{P_DOC}}",
            "{{P_DOC_2}}",
            "{{P_DOC_3}}",
            "{{P_DOC_4}}",
            "{{P_DOC_5}}",
            "{{C_DOC}}",
            "{{C_DOC_2}}",
            "{{C_DOC_3}}",
            "{{C_DOC_4}}",
            "{{C_DOC_5}}",
            "{{P_NACIONALIDAD}}",
            "{{P_NACIONALIDAD_2}}",
            "{{P_NACIONALIDAD_3}}",
            "{{P_NACIONALIDAD_4}}",
            "{{P_NACIONALIDAD_5}}",
            "{{C_NACIONALIDAD}}",
            "{{C_NACIONALIDAD_2}}",
            "{{C_NACIONALIDAD_3}}",
            "{{C_NACIONALIDAD_4}}",
            "{{C_NACIONALIDAD_5}}",
            "{{P_OCUPACION}}",
            "{{P_OCUPACION_2}}",
            "{{P_OCUPACION_3}}",
            "{{P_OCUPACION_4}}",
            "{{P_OCUPACION_5}}",
            "{{C_OCUPACION}}",
            "{{C_OCUPACION_2}}",
            "{{C_OCUPACION_3}}",
            "{{C_OCUPACION_4}}",
            "{{C_OCUPACION_5}}",
            "{{P_ESTADO_CIVIL}}",
            "{{P_ESTADO_CIVIL_2}}",
            "{{P_ESTADO_CIVIL_3}}",
            "{{P_ESTADO_CIVIL_4}}",
            "{{P_ESTADO_CIVIL_5}}",
            "{{C_ESTADO_CIVIL}}",
            "{{C_ESTADO_CIVIL_2}}",
            "{{C_ESTADO_CIVIL_3}}",
            "{{C_ESTADO_CIVIL_4}}",
            "{{C_ESTADO_CIVIL_5}}",
            "{{P_DOMICILIO}}",
            "{{P_DOMICILIO_2}}",
            "{{P_DOMICILIO_3}}",
            "{{P_DOMICILIO_4}}",
            "{{P_DOMICILIO_5}}",
            "{{C_DOMICILIO}}",
            "{{C_DOMICILIO_2}}",
            "{{C_DOMICILIO_3}}",
            "{{C_DOMICILIO_4}}",
            "{{C_DOMICILIO_5}}",
            "{{P_IDE}}",
            "{{P_IDE_2}}",
            "{{P_IDE_3}}",
            "{{P_IDE_4}}",
            "{{P_IDE_5}}",
            "{{C_IDE}}",
            "{{C_IDE_2}}",
            "{{C_IDE_3}}",
            "{{C_IDE_4}}",
            "{{C_IDE_5}}",
        }

        # Clean paragraphs
        for paragraph in doc.paragraphs:
            self._clean_paragraph_placeholders(
                paragraph, escrituracion_placeholders, contractor_placeholders
            )

        # Clean tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._clean_paragraph_placeholders(
                            paragraph, escrituracion_placeholders, contractor_placeholders
                        )


    def _clean_paragraph_placeholders(
        self, paragraph, escrituracion_placeholders, contractor_placeholders
    ):
        """
        Clean placeholders in a single paragraph - SIMPLIFIED VERSION
        """
        full_text = paragraph.text

        # Quick check - if no placeholders, skip
        if "{{" not in full_text and "[" not in full_text:
            return


        # Get original formatting
        first_run_font = None
        if paragraph.runs:
            first_run_font = paragraph.runs[0].font

        # Clear all runs
        for run in paragraph.runs:
            run.text = ""

        # Simple string replacement approach
        new_text = full_text

        # Remove both {{PLACEHOLDER}} and [E.PLACEHOLDER] formats
        import re

        # Remove {{PLACEHOLDER}} format
        new_text = re.sub(r"\{\{[^}]+\}\}", "", new_text)

        # Remove [E.PLACEHOLDER] format
        new_text = re.sub(r"\[E\.[^\]]+\]", "", new_text)


        # Clean up extra spaces and punctuation
        new_text = self._clean_text_formatting(new_text)

        # Create new run with cleaned text
        if new_text.strip():
            new_run = paragraph.add_run(new_text)

            # Preserve original formatting
            if first_run_font:
                try:
                    new_run.font.bold = first_run_font.bold
                    new_run.font.italic = first_run_font.italic
                    new_run.font.color.rgb = first_run_font.color.rgb
                    new_run.font.size = first_run_font.size
                    new_run.font.name = first_run_font.name
                except:
                    pass

    def _clean_text_formatting(self, text):
        """
        Clean up text formatting after placeholder removal
        """
        import re

        # Remove extra spaces
        text = re.sub(r"\s+", " ", text)

        # Remove extra commas and semicolons
        text = re.sub(r",\s*,", ",", text)
        text = re.sub(r";\s*;", ";", text)

        # Remove trailing commas and semicolons
        text = re.sub(r",\s*$", "", text)
        text = re.sub(r";\s*$", "", text)

        # Remove extra spaces around punctuation
        text = re.sub(r"\s+([,;.])", r"\1", text)

        # Remove empty contractor slots (multiple commas in a row)
        text = re.sub(r",\s*,", ",", text)  # Remove double commas
        text = re.sub(r",\s*,", ",", text)  # Remove triple commas

        # Remove leading/trailing commas
        text = re.sub(r"^,\s*", "", text)
        text = re.sub(r",\s*$", "", text)

        return text.strip()


# In utils.py - Add this class
class DataValidator:
    """
    Validate data before processing
    Ensure data integrity and handle edge cases
    """

    def validate_raw_data(self, raw_data):
        """
        Validate raw data from database
        Check for required fields and data integrity

        SOLUTION:
        - Check required fields exist
        - Validate data types
        - Handle null/empty values
        - Return validation results
        """
        # TODO: Implement data validation
        # TODO: Check required fields
        # TODO: Validate data types
        # TODO: Handle edge cases
        pass

    def validate_template_data(self, final_data):
        """
        Validate final template data
        Ensure all placeholders have values

        SOLUTION:
        - Check for missing placeholders
        - Validate placeholder values
        - Handle empty values
        - Return validation results
        """
        # TODO: Implement template data validation
        # TODO: Check missing placeholders
        # TODO: Validate placeholder values
        # TODO: Handle empty values
        pass


# In utils.py - Add this class
class TemplateManager:
    """
    Handle template operations
    Download, process, and manage templates
    """

    def __init__(self):
        self.s3_client = None  # Will be set when needed

    def get_template_from_r2(self, template_id, filename):
        """
        Download template from R2 storage
        Mirrors: PHP template download
        """

        s3 = boto3.client(
            "s3",
            endpoint_url=os.environ.get("CLOUDFLARE_R2_ENDPOINT"),
            aws_access_key_id=os.environ.get("CLOUDFLARE_R2_ACCESS_KEY"),
            aws_secret_access_key=os.environ.get("CLOUDFLARE_R2_SECRET_KEY"),
            config=Config(signature_version="s3v4"),
            region_name="auto",
        )

        object_key = f"{os.environ.get('CLOUDFLARE_R2_MAIN_URL')}/plantillas/{filename}"

        try:
            response = s3.get_object(Bucket=os.environ.get("CLOUDFLARE_R2_BUCKET"), Key=object_key)
            template_bytes = response["Body"].read()
            return template_bytes
        except Exception as e:
            print(f"ERROR: Failed to download template from R2: {e}")
            raise

    def upload_document_to_r2(self, buffer, kardex):
        """
        Upload generated document to R2
        Mirrors: PHP document upload
        
        Args:
            buffer: BytesIO buffer containing the document
            kardex: Kardex number for filename
            
        Returns:
            bool: True if upload successful, False otherwise
        """
        try:
            # Reset buffer position
            buffer.seek(0)
            doc_content = buffer.read()
            buffer.seek(0)  # Reset for further use
            
            # Define object key for R2
            filename = f"__PROY__{kardex}.docx"
            print(f"DEBUG: Uploading document to R2: {filename}")
            object_key = f"{os.environ.get('CLOUDFLARE_R2_MAIN_URL')}/documentos/{filename}"
            
            # Get S3 client
            s3 = get_s3_client()
            
            # Upload to R2
            s3.upload_fileobj(
                io.BytesIO(doc_content),
                os.environ.get('CLOUDFLARE_R2_BUCKET'),
                object_key
            )
            
            print(f"DEBUG: Document uploaded to R2: {object_key}")
            return True
            
        except Exception as e:
            print(f"ERROR: No se pudo subir el documento a R2: {e}")
            return False
    
    def get_document_from_r2(self, kardex):
        """
        Download existing document from R2 for updating
        Mirrors: PHP actualizar action
        
        Args:
            kardex: Kardex number to identify the document
            
        Returns:
            bytes: Document content, or None if not found
        """
        try:
            filename = f"__PROY__{kardex}.docx"
            object_key = f"{os.environ.get('CLOUDFLARE_R2_MAIN_URL')}/documentos/{filename}"
            
            print(f"DEBUG: Downloading document from R2: {object_key}")
            
            # Get S3 client
            s3 = get_s3_client()
            
            # Download from R2
            response = s3.get_object(
                Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'),
                Key=object_key
            )
            
            doc_bytes = response['Body'].read()
            print(f"DEBUG: Downloaded document from R2: {len(doc_bytes)} bytes")
            
            return doc_bytes
            
        except Exception as e:
            print(f"ERROR: No se pudo descargar el documento de R2: {e}")
            return None
    
    def update_document_escrituracion(self, kardex, escrituracion_data, placeholder_processor):
        """
        Update existing document with escrituracion data only
        Mirrors: PHP actualizar action
        
        This is a generic method that all services can use to update documents
        
        Args:
            kardex: Kardex number
            escrituracion_data: Dictionary with escrituracion placeholders (NRO_ESC, F, FI, FF, etc.)
            placeholder_processor: PlaceholderProcessor instance to use for replacement
            
        Returns:
            BytesIO: Updated document buffer
        """
        # Download existing document from R2
        doc_bytes = self.get_document_from_r2(kardex)
        if not doc_bytes:
            raise ValueError(f"El documento __PROY__{kardex}.docx no existe. Debe generar el documento antes de actualizarlo.")
        
        # Load document
        buffer = io.BytesIO(doc_bytes)
        doc = Document(buffer)
        
        # Replace only escrituracion placeholders
        placeholder_processor.replace_placeholders(doc, escrituracion_data)
        
        # Save to buffer
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        
        # Upload updated document back to R2
        self.upload_document_to_r2(output_buffer, kardex)
        
        return output_buffer
