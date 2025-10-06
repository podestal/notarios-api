from django.db import connection
import io


class GarantiasReportService:
    """Service for generating garantias mobiliarias reports matching PHP script format"""

    def _sanitize_cell_value(self, value):
        """Sanitize cell values to prevent Excel corruption"""
        if value is None:
            return ""

        # Convert to string and strip whitespace
        val_str = str(value).strip()

        # Remove control characters that can cause XML corruption
        import re

        val_str = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]", "", val_str)

        # Limit to Excel's cell character limit
        val_str = val_str[:32767]

        return val_str

    def _format_date_in_spanish(self, date_input):
        """Convert date to Spanish format like 'LUNES, 15 DE ENERO DEL 2025'"""
        try:
            # Handle both datetime objects and date strings
            if hasattr(date_input, "strftime"):
                date_obj = date_input
            else:
                from datetime import datetime

                # Try different formats
                if "-" in str(date_input) and len(str(date_input).split("-")[0]) == 4:
                    date_obj = datetime.strptime(str(date_input), "%Y-%m-%d")
                elif "/" in str(date_input):
                    date_obj = datetime.strptime(str(date_input), "%d/%m/%Y")
                else:
                    return str(date_input)

            # Spanish day names
            dias = ["LUNES", "MARTES", "MIÉRCOLES", "JUEVES", "VIERNES", "SÁBADO", "DOMINGO"]
            # Spanish month names
            meses = [
                "ENERO",
                "FEBRERO",
                "MARZO",
                "ABRIL",
                "MAYO",
                "JUNIO",
                "JULIO",
                "AGOSTO",
                "SEPTIEMBRE",
                "OCTUBRE",
                "NOVIEMBRE",
                "DICIEMBRE",
            ]

            dia_semana = dias[date_obj.weekday()]
            dia = date_obj.day
            mes = meses[date_obj.month - 1]
            anio = date_obj.year

            return f"{dia_semana}, {dia} DE {mes} DEL {anio}"
        except:
            return str(date_input)

    def _extract_year_from_date(self, date_input):
        """Extract year from date string DD/MM/YYYY or YYYY-MM-DD or datetime object"""
        try:
            # Handle datetime objects
            if hasattr(date_input, "year"):
                return str(date_input.year)

            # Handle strings
            date_str = str(date_input)
            # Try to parse as YYYY-MM-DD first
            if "-" in date_str and len(date_str.split("-")[0]) == 4:
                return date_str.split("-")[0]
            # Try to parse as DD/MM/YYYY
            elif "/" in date_str:
                return date_str.split("/")[-1]
            else:
                from datetime import datetime

                return str(datetime.now().year)
        except:
            from datetime import datetime

            return str(datetime.now().year)

    def _format_date_for_display(self, date_obj):
        """Format date like PHP script logic"""
        try:
            if hasattr(date_obj, "strftime"):
                return date_obj.strftime("%d/%m/%Y")
            return str(date_obj)
        except:
            return str(date_obj)

    def _get_notary_info(self):
        """Get notary configuration info"""
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT nombre, apellido, telefono, correo, ruc, direccion, distrito FROM confinotario"
            )
            result = cursor.fetchone()
            if result:
                return {
                    "nombre": f"{result[0]} {result[1]}",
                    "telefono": result[2],
                    "correo": result[3],
                    "ruc": result[4],
                    "direccion": result[5],
                    "distrito": result[6],
                }
            return {
                "nombre": "NOTARIO",
                "telefono": "(051) 326609",
                "correo": "",
                "ruc": "10024231572",
                "direccion": "JR.BOLIVAR NRO. 340",
                "distrito": "JULIACA",
            }

    def _get_report_data(self, desde, hasta):
        """Fetch data for the report matching PHP query"""
        try:
            from django.db import connection
            from datetime import datetime

            import time

            start_time = time.time()
            print(
                f"DEBUG: Garantias - desde: {desde} (type: {type(desde)}), hasta: {hasta} (type: {type(hasta)})"
            )

            # Convert dates to proper format for Django ORM
            desde_dt = None
            hasta_dt = None

            try:
                if isinstance(desde, str):
                    if "-" in desde and len(desde.split("-")[0]) == 4:
                        desde_dt = datetime.strptime(desde, "%Y-%m-%d")
                    else:
                        desde_dt = datetime.strptime(desde, "%d/%m/%Y")
                else:
                    desde_dt = desde
            except (ValueError, TypeError) as e:
                print(f"DEBUG: Error parsing desde date '{desde}': {e}")
                return []

            try:
                if isinstance(hasta, str):
                    if "-" in hasta and len(hasta.split("-")[0]) == 4:
                        hasta_dt = datetime.strptime(hasta, "%Y-%m-%d")
                    else:
                        hasta_dt = datetime.strptime(hasta, "%d/%m/%Y")
                else:
                    hasta_dt = hasta
            except (ValueError, TypeError) as e:
                print(f"DEBUG: Error parsing hasta date '{hasta}': {e}")
                return []

            # Validate dates are not None
            if desde_dt is None or hasta_dt is None:
                print(f"DEBUG: Invalid dates - desde: {desde}, hasta: {hasta}")
                return []

            # Set session variables for optimization
            with connection.cursor() as cursor:
                cursor.execute("SET SESSION group_concat_max_len = 1000000")
                cursor.execute("SET SESSION sql_mode = 'NO_AUTO_VALUE_ON_ZERO'")
                cursor.execute("SET SESSION sort_buffer_size = 2097152")

                # Try to create indexes if they don't exist (for performance)
                try:
                    cursor.execute(
                        "CREATE INDEX IF NOT EXISTS idx_kardex_tipkar_fecha ON kardex(idtipkar, fechaescritura, nc)"
                    )
                    cursor.execute(
                        "CREATE INDEX IF NOT EXISTS idx_kardex_fecha ON kardex(fechaescritura)"
                    )
                    print("DEBUG: Database indexes created/verified")
                except Exception as e:
                    print(f"DEBUG: Index creation failed (may already exist): {e}")

            # Ultra-simple query to get basic data first
            query = """
                SELECT 
                    k.fechaescritura,
                    k.kardex,
                    k.contrato,
                    k.numescritura,
                    k.numminuta,
                    k.folioini,
                    k.numescritura as numescritura2,
                    '' as precio,
                    '' as moneda
                FROM kardex as k 
                WHERE k.idtipkar='4' 
                    AND k.fechaescritura <> '' 
                    AND k.fechaescritura >= %s
                    AND k.fechaescritura <= %s
                ORDER BY k.numescritura ASC
                LIMIT 2000
            """

            with connection.cursor() as cursor:
                cursor.execute(
                    query, [desde_dt.strftime("%Y-%m-%d"), hasta_dt.strftime("%Y-%m-%d")]
                )
                garantias = []
                rows = cursor.fetchall()
                print(
                    f"DEBUG: Main query completed in {time.time() - start_time:.2f}s, found {len(rows)} records"
                )

                # Get additional data for the kardex records we found
                if rows:
                    kardex_list = [row[1] for row in rows]

                    # Get additional data (precio, moneda) in a separate optimized query
                    additional_data_query = """
                        SELECT 
                            k.kardex,
                            COALESCE(p.importetrans, '') as precio,
                            COALESCE(m.simbolo, '') as moneda
                        FROM kardex as k 
                        LEFT JOIN patrimonial as p ON p.kardex=k.kardex
                        LEFT JOIN monedas as m ON m.idmon=p.idmon
                        WHERE k.kardex IN %s
                    """

                    additional_start = time.time()
                    cursor.execute(additional_data_query, [tuple(kardex_list)])
                    additional_data = cursor.fetchall()
                    print(
                        f"DEBUG: Additional data query completed in {time.time() - additional_start:.2f}s, found {len(additional_data)} records"
                    )

                    # Create a lookup dictionary for additional data
                    additional_lookup = {
                        row[0]: (row[1], row[2]) for row in additional_data
                    }
                else:
                    kardex_list = []
                    additional_lookup = {}

                # Get all contractors in one query with optimized subquery
                contractors_query = """
                    SELECT 
                        cxa.kardex,
                        c2.tipper,
                        UPPER(CONCAT(COALESCE(c2.apepat,''),' ',COALESCE(c2.apemat,''),' ',COALESCE(c2.prinom,''),' ',COALESCE(c2.segnom,''))) AS nombre,
                        cxa.idcontratante,
                        UPPER(COALESCE(c2.razonsocial,'')) AS empresa,
                        cxa.parte,
                        cxa.uif,
                        COALESCE(cxar.parte, 0) as parte_representada
                    FROM contratantesxacto AS cxa
                    INNER JOIN contratantes AS con ON con.idcontratante=cxa.idcontratante
                    INNER JOIN cliente2 AS c2 ON c2.idcontratante=con.idcontratante
                    LEFT JOIN contratantesxacto AS cxar ON con.idcontratanterp = cxar.idcontratante AND cxar.kardex = cxa.kardex
                    WHERE cxa.kardex IN %s
                    ORDER BY cxa.kardex, c2.tipper ASC
                """
                contractors_start = time.time()
                cursor.execute(contractors_query, [tuple(kardex_list)])
                all_contractors = cursor.fetchall()
                print(
                    f"DEBUG: Contractors query completed in {time.time() - contractors_start:.2f}s, found {len(all_contractors)} contractors"
                )

                # Group contractors by kardex
                contractors_by_kardex = {}
                for contractor in all_contractors:
                    kardex = contractor[0]
                    if kardex not in contractors_by_kardex:
                        contractors_by_kardex[kardex] = []
                    contractors_by_kardex[kardex].append(contractor[1:])  # Skip kardex from tuple

                # Process each kardex with optimized processing
                processing_start = time.time()
                garantias = []
                for row in rows:
                    kardex = row[1]
                    otorgante = []
                    otorgado = []

                    # Process contractors if any exist for this kardex
                    for contractor in contractors_by_kardex.get(kardex, []):
                        tipper, nombre, idcontratante, empresa, parte, uif, parte_representada = (
                            contractor
                        )

                        # Process otorgante - parte=1 OR parte_representada=1 OR uif='O'
                        # Special case: if uif='O' AND parte_representada=2, skip
                        if (parte == 1 or parte_representada == 1 or uif == "O"):
                            if not (uif == "O" and parte_representada == 2):
                                otorgante.append(empresa if tipper != "N" else nombre)

                        # Process otorgado - parte=2 OR parte_representada=2 OR uif='B' OR uif='N'
                        # Special case: if uif='B' AND parte_representada=1, skip
                        if (parte == 2 or parte_representada == 2 or uif == "B" or uif == "N"):
                            if not (uif == "B" and parte_representada == 1):
                                otorgado.append(empresa if tipper != "N" else nombre)

                    # Clean contract name like in PHP - optimized string operations
                    contrato_raw = row[2] or ""
                    contrato_clean = contrato_raw.replace("/", "").upper()

                    # Optimized date formatting
                    fecha_str = row[0]
                    fecha_formatted = (
                        f"{fecha_str[8:10]}/{fecha_str[5:7]}/{fecha_str[0:4]}"
                        if len(fecha_str) >= 10
                        else fecha_str
                    )

                    # Get additional data from lookup
                    precio, moneda = additional_lookup.get(kardex, ("", ""))

                    garantias.append(
                        {
                            "numero_escritura": row[3],
                            "fecha": fecha_formatted,
                            "otorgante": (
                                "NO CORRE"
                                if contrato_raw == "NO CORRE / "
                                else ", ".join(otorgante)
                            ),
                            "otorgado": (
                                "NO CORRE" if contrato_raw == "NO CORRE / " else ", ".join(otorgado)
                            ),
                            "contrato": contrato_clean,
                            "monto": f"{moneda} {precio}" if precio else "",
                            "folio": row[5],
                        }
                    )

                print(f"DEBUG: Data processing completed in {time.time() - processing_start:.2f}s")
                print(f"DEBUG: Total data fetching completed in {time.time() - start_time:.2f}s")
                return garantias

        except Exception as e:
            import traceback

            traceback.print_exc()
            return []

    def generate_excel_report(self, desde, hasta):
        """Generate Excel report matching PHP script format"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            import io

            import time

            report_start = time.time()
            print(f"DEBUG: Garantias Excel - desde: {desde}, hasta: {hasta}")

            # Get data
            report_data = self._get_report_data(desde, hasta)
            print(f"DEBUG: Data fetched in {time.time() - report_start:.2f}s, generating Excel...")

            notary_info = self._get_notary_info()

            anio = self._extract_year_from_date(hasta)

            # Create workbook and worksheet
            wb = Workbook()
            ws = wb.active
            ws.title = "GARANTIAS MOBILIARIAS"

            # Styles
            title_font = Font(name="Arial", size=18.5, bold=True)
            header_font = Font(name="Arial", size=13.5, bold=True)
            data_font = Font(name="Arial", size=13.5)
            center_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            left_alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            right_alignment = Alignment(horizontal="right", vertical="top", wrap_text=True)

            # Borders
            thin_border = Border(
                left=Side(border_style="thin"),
                right=Side(border_style="thin"),
                top=Side(border_style="thin"),
                bottom=Side(border_style="thin"),
            )
            no_border = Border(
                left=Side(style=None),
                right=Side(style=None),
                top=Side(style=None),
                bottom=Side(style=None),
            )

            # Title section
            ws.merge_cells("A1:G1")
            ws["A1"] = "INDICE CRONOLOGICO - CONSTITUCION DE GARANTIAS MOBILIARIAS"
            ws["A1"].font = title_font
            ws["A1"].alignment = center_alignment
            ws["A1"].border = no_border

            ws.merge_cells("A2:G2")
            ws["A2"] = f"AÑO {anio}"
            ws["A2"].font = title_font
            ws["A2"].alignment = center_alignment
            ws["A2"].border = no_border

            # Notary info section
            row = 4
            ws[f"A{row}"] = "NOTARIA"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"B{row}"] = f': {self._sanitize_cell_value(notary_info["nombre"])}'
            ws[f"B{row}"].font = data_font
            ws[f"B{row}"].border = no_border

            row += 1
            ws[f"A{row}"] = "DIRECCION"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"B{row}"] = f': {notary_info["direccion"]}'
            ws[f"B{row}"].font = data_font
            ws[f"B{row}"].border = no_border
            ws[f"C{row}"] = "TELEFONO"
            ws[f"C{row}"].font = header_font
            ws[f"C{row}"].border = no_border
            ws[f"D{row}"] = f': {notary_info["telefono"]}'
            ws[f"D{row}"].font = data_font
            ws[f"D{row}"].border = no_border

            row += 1
            ws[f"A{row}"] = "DEPARTAMENTO"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"B{row}"] = ": PUNO"
            ws[f"B{row}"].font = data_font
            ws[f"B{row}"].border = no_border
            ws[f"C{row}"] = "RUC"
            ws[f"C{row}"].font = header_font
            ws[f"C{row}"].border = no_border
            ws[f"D{row}"] = f': {notary_info["ruc"]}'
            ws[f"D{row}"].font = data_font
            ws[f"D{row}"].border = no_border

            row += 1
            ws[f"A{row}"] = "PROVINCIA"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"B{row}"] = ": SAN ROMAN"
            ws[f"B{row}"].font = data_font
            ws[f"B{row}"].border = no_border
            ws[f"C{row}"] = "DESDE"
            ws[f"C{row}"].font = header_font
            ws[f"C{row}"].border = no_border
            ws[f"D{row}"] = f": {self._format_date_in_spanish(desde).upper()}"
            ws[f"D{row}"].font = data_font
            ws[f"D{row}"].border = no_border

            row += 1
            ws[f"A{row}"] = "DISTRITO"
            ws[f"A{row}"].font = header_font
            ws[f"A{row}"].border = no_border
            ws[f"B{row}"] = f': {notary_info["distrito"]}'
            ws[f"B{row}"].font = data_font
            ws[f"B{row}"].border = no_border
            ws[f"C{row}"] = "HASTA"
            ws[f"C{row}"].font = header_font
            ws[f"C{row}"].border = no_border
            ws[f"D{row}"] = f": {self._format_date_in_spanish(hasta).upper()}"
            ws[f"D{row}"].font = data_font
            ws[f"D{row}"].border = no_border

            # Data table headers
            row += 2
            headers = [
                "ESCR.",
                "FECH.ESCR.",
                "OTORGANTE",
                "A FAVOR",
                "ACTO JURIDICO",
                "PRECIO",
                "NUM.FOLIO",
            ]
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = header_font
                cell.alignment = center_alignment
                cell.border = thin_border

            # Data rows
            for i, data_row in enumerate(report_data, 1):
                row += 1

                # Row data
                row_data = [
                    data_row["numero_escritura"],
                    data_row["fecha"],
                    self._sanitize_cell_value(data_row["otorgante"]),
                    self._sanitize_cell_value(data_row["otorgado"]),
                    self._sanitize_cell_value(data_row["contrato"]),
                    data_row["monto"],
                    data_row["folio"],
                ]

                for col, value in enumerate(row_data, 1):
                    cell = ws.cell(row=row, column=col, value=value)
                    cell.font = data_font
                    cell.border = thin_border

                    # Alignment based on column
                    if col in [1, 2, 6, 7]:  # ESCR., FECHA, PRECIO, FOLIO
                        cell.alignment = right_alignment
                    else:
                        cell.alignment = left_alignment

            # Auto-adjust column widths
            for column in ws.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width

            # Create response
            from django.http import StreamingHttpResponse

            def file_iterator():
                buffer = io.BytesIO()
                wb.save(buffer)
                buffer.seek(0)
                yield buffer.getvalue()
                buffer.close()

            response = StreamingHttpResponse(
                file_iterator(),
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            response["Content-Disposition"] = (
                f"attachment; filename=INDICE_CRONOLOGICO_GARANTIAS_MOBILIARIAS_{anio}.xlsx"
            )
            response["Access-Control-Allow-Origin"] = "*"

            return response

        except Exception as e:
            import traceback

            traceback.print_exc()
            from django.http import HttpResponse

            return HttpResponse(f"Error generating Excel report: {e}", status=500)

    def generate_word_report(self, desde, hasta):
        """Generate Word report matching PHP script format"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            import io

            import time

            report_start = time.time()
            print(f"DEBUG: Garantias Word - desde: {desde}, hasta: {hasta}")

            # Get data
            report_data = self._get_report_data(desde, hasta)
            print(f"DEBUG: Data fetched in {time.time() - report_start:.2f}s, generating Word...")

            notary_info = self._get_notary_info()

            anio = self._extract_year_from_date(hasta)

            # Create document
            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            # Title
            title = doc.add_heading(
                "INDICE CRONOLOGICO - CONSTITUCION DE GARANTIAS MOBILIARIAS", 0
            )
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            subtitle = doc.add_heading(f"AÑO {anio}", 0)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add spacing
            doc.add_paragraph()
            doc.add_paragraph()

            # Notary info table
            info_table = doc.add_table(rows=5, cols=4)
            info_table.style = "Table Grid"

            # Row 1: NOTARIA
            row1 = info_table.rows[0]
            row1.cells[0].text = "NOTARIA"
            row1.cells[0].paragraphs[0].runs[0].bold = True
            row1.cells[1].text = f': {notary_info["nombre"]}'

            # Row 2: DIRECCION
            row2 = info_table.rows[1]
            row2.cells[0].text = "DIRECCION"
            row2.cells[0].paragraphs[0].runs[0].bold = True
            row2.cells[1].text = f': {notary_info["direccion"]}'
            row2.cells[2].text = "TELEFONO"
            row2.cells[2].paragraphs[0].runs[0].bold = True
            row2.cells[3].text = f': {notary_info["telefono"]}'

            # Row 3: DEPARTAMENTO
            row3 = info_table.rows[2]
            row3.cells[0].text = "DEPARTAMENTO"
            row3.cells[0].paragraphs[0].runs[0].bold = True
            row3.cells[1].text = ": PUNO"
            row3.cells[2].text = "RUC"
            row3.cells[2].paragraphs[0].runs[0].bold = True
            row3.cells[3].text = f': {notary_info["ruc"]}'

            # Row 4: PROVINCIA
            row4 = info_table.rows[3]
            row4.cells[0].text = "PROVINCIA"
            row4.cells[0].paragraphs[0].runs[0].bold = True
            row4.cells[1].text = ": SAN ROMAN"
            row4.cells[2].text = "DESDE"
            row4.cells[2].paragraphs[0].runs[0].bold = True
            row4.cells[3].text = f": {self._format_date_in_spanish(desde).upper()}"

            # Row 5: DISTRITO
            row5 = info_table.rows[4]
            row5.cells[0].text = "DISTRITO"
            row5.cells[0].paragraphs[0].runs[0].bold = True
            row5.cells[1].text = f': {notary_info["distrito"]}'
            row5.cells[2].text = "HASTA"
            row5.cells[2].paragraphs[0].runs[0].bold = True
            row5.cells[3].text = f": {self._format_date_in_spanish(hasta).upper()}"

            # Add spacing
            doc.add_paragraph()

            # Data table
            if report_data:
                data_table = doc.add_table(rows=1, cols=7)
                data_table.style = "Table Grid"

                # Headers
                header_row = data_table.rows[0]
                headers = [
                    "ESCR.",
                    "FECH.ESCR.",
                    "OTORGANTE",
                    "A FAVOR",
                    "ACTO JURIDICO",
                    "PRECIO",
                    "NUM.FOLIO",
                ]
                for i, header in enumerate(headers):
                    header_row.cells[i].text = header
                    header_row.cells[i].paragraphs[0].runs[0].bold = True

            # Data rows
            for i, data_row in enumerate(report_data, 1):
                row = data_table.add_row()
                row.cells[0].text = str(data_row["numero_escritura"])
                row.cells[1].text = data_row["fecha"]
                row.cells[2].text = self._sanitize_cell_value(data_row["otorgante"])
                row.cells[3].text = self._sanitize_cell_value(data_row["otorgado"])
                row.cells[4].text = self._sanitize_cell_value(data_row["contrato"])
                row.cells[5].text = data_row["monto"]
                row.cells[6].text = str(data_row["folio"])

            # Create response
            from django.http import StreamingHttpResponse

            def file_iterator():
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                yield buffer.getvalue()
                buffer.close()

            response = StreamingHttpResponse(
                file_iterator(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = (
                f"attachment; filename=INDICE_CRONOLOGICO_GARANTIAS_MOBILIARIAS_{anio}.docx"
            )
            response["Access-Control-Allow-Origin"] = "*"

            return response

        except Exception as e:
            import traceback

            traceback.print_exc()
            from django.http import HttpResponse

            return HttpResponse(f"Error generating Word report: {e}", status=500)
