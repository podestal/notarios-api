import boto3
from botocore.client import Config
from django.conf import settings
import os
import io
from docx import Document
from docxtpl import DocxTemplate
from docxcompose.properties import CustomProperties
from docx.shared import RGBColor, Pt
from decimal import Decimal
from typing import Dict, Any
from .constants import ROLE_LABELS, TIPO_DOCUMENTO, CIVIL_STATUS
import re
from datetime import datetime
from django.http import HttpResponse, JsonResponse
from notaria.models import TplTemplate, Contratantesxacto, Detallevehicular, Patrimonial, Contratantes, Actocondicion, Cliente2, Nacionalidades, Kardex, Usuarios, Sedesregistrales, Ubigeo
from notaria.constants import MONEDAS, OPORTUNIDADES_PAGO, FORMAS_PAGO
from .utils import NumberToLetterConverter
import time
from django.db import connection

# Cached S3 client to avoid recreating on every request
_s3_client = None

def get_s3_client():
    """
    Get a cached S3 client for R2 operations
    """
    global _s3_client
    if _s3_client is None:
        _s3_client = boto3.client(
            's3',
            endpoint_url=os.environ.get('CLOUDFLARE_R2_ENDPOINT'),
            aws_access_key_id=os.environ.get('CLOUDFLARE_R2_ACCESS_KEY'),
            aws_secret_access_key=os.environ.get('CLOUDFLARE_R2_SECRET_KEY'),
            config=Config(signature_version='s3v4'),
            region_name='auto',
        )
    return _s3_client

class VehicleTransferDocumentService:
    """
    Django service to generate vehicle transfer documents based on the PHP logic
    """
    
    def __init__(self):
        self.letras = NumberToLetterConverter()
    
    def generate_vehicle_transfer_document(self, template_id: int, num_kardex: str, action: str = 'generate', mode: str = "download") -> HttpResponse:
        """
        Main method to generate vehicle transfer document
        """
        start_time = time.time()
        print(f"PERF: Starting vehicle document generation for kardex: {num_kardex}")
        
        try:
            # Step 1: Get template from R2
            template_start = time.time()
            template = self._get_template_from_r2(template_id)
            template_time = time.time() - template_start
            print(f"PERF: Template download took {template_time:.2f}s")
            
            # Step 2: Get document data
            data_start = time.time()
            document_data = self.get_document_data(num_kardex)
            data_time = time.time() - data_start
            print(f"PERF: Data retrieval took {data_time:.2f}s")
            
            # Step 3: Process document
            process_start = time.time()
            doc = self._process_document(template, document_data)
            process_time = time.time() - process_start
            print(f"PERF: Document processing took {process_time:.2f}s")
            
            # Step 4: Remove placeholders
            cleanup_start = time.time()
            self.remove_unfilled_placeholders(doc)
            cleanup_time = time.time() - cleanup_start
            print(f"PERF: Placeholder cleanup took {cleanup_time:.2f}s")
            
            # Step 5: Upload to R2
            upload_start = time.time()
            upload_success = self.create_documento_in_r2(doc, num_kardex)
            upload_time = time.time() - upload_start
            print(f"PERF: R2 upload took {upload_time:.2f}s")
            
            if not upload_success:
                print(f"WARNING: Failed to upload document to R2 for kardex: {num_kardex}")
            
            total_time = time.time() - start_time
            print(f"PERF: Total vehicle document generation took {total_time:.2f}s")
            
            return self._create_response(doc, f"__PROY__{num_kardex}.docx", num_kardex, mode)
        except FileNotFoundError as e:
            return HttpResponse(str(e), status=404)
        except Exception as e:
            total_time = time.time() - start_time
            print(f"PERF: Vehicle document generation failed after {total_time:.2f}s")
            return HttpResponse(f"Error generating document: {str(e)}", status=500)

    def create_documento_in_r2(self, doc, kardex):
        """
        Create a new document in R2 storage
        """
        try:
            print(f"DEBUG: Starting R2 upload for kardex: {kardex}")
            
            # Save the document to a bytes buffer
            from io import BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            doc_content = buffer.read()
            
            print(f"DEBUG: Document size: {len(doc_content)} bytes")
            
            # Define the object key for R2
            object_key = f"rodriguez-zea/documentos/__PROY__{kardex}.docx"
            
            # Check environment variables
            endpoint_url = os.environ.get('CLOUDFLARE_R2_ENDPOINT')
            access_key = os.environ.get('CLOUDFLARE_R2_ACCESS_KEY')
            secret_key = os.environ.get('CLOUDFLARE_R2_SECRET_KEY')
            bucket = os.environ.get('CLOUDFLARE_R2_BUCKET')
            
            print(f"DEBUG: R2 Configuration - Endpoint: {endpoint_url}")
            print(f"DEBUG: R2 Configuration - Access Key: {'SET' if access_key else 'NOT SET'}")
            print(f"DEBUG: R2 Configuration - Secret Key: {'SET' if secret_key else 'NOT SET'}")
            print(f"DEBUG: R2 Configuration - Bucket: {bucket}")
            
            # Upload to R2
            s3 = boto3.client(
                's3',
                endpoint_url=endpoint_url,
                aws_access_key_id=access_key,
                aws_secret_access_key=secret_key,
                config=Config(signature_version='s3v4'),
                region_name='auto',
            )
            
            print(f"DEBUG: S3 client created successfully")
            print(f"DEBUG: Uploading to bucket: {bucket}, key: {object_key}")
            
            s3.upload_fileobj(
                BytesIO(doc_content),
                bucket,
                object_key
            )
            
            print(f"DEBUG: Document uploaded to R2: {object_key}")
            return True
            
        except Exception as e:
            print(f"ERROR: Failed to upload document to R2: {e}")
            print(f"ERROR: Exception type: {type(e).__name__}")
            import traceback
            print(f"ERROR: Full traceback: {traceback.format_exc()}")
            return False


    def remove_unfilled_placeholders(self, doc):
        """
        Remove all [E.SOMETHING] placeholders and hide {{SOMETHING}} placeholders from user.
        Keep and hide: {{NRO_ESC}}, {{FI}}, {{FF}}, {{S_IN}}, {{S_FN}}, {{FECHA_ACT}}
        Remove all others: [E.P_NOM_2], [E.C_NOM_2], etc.
        """
        import re
        placeholder_pattern = re.compile(r'\[E\.[A-Z0-9_]+\]')
        curly_placeholder_pattern = re.compile(r'\{\{[A-Z0-9_]+\}\}')
        representation_pattern = re.compile(
            r'EN REPRESENTACION DE\s*Y[\s.·…‥⋯⋮⋱⋰⋯—–-]*', re.IGNORECASE
        )

        # List of placeholders to keep and hide
        keep_placeholders = ['{{NRO_ESC}}', '{{FI}}', '{{FF}}', '{{S_IN}}', '{{S_FN}}', '{{FECHA_ACT}}']

        def clean_runs(runs):
            # Remove [E.SOMETHING] placeholders and hide {{SOMETHING}} placeholders
            for run in runs:
                # Remove all [E.SOMETHING] placeholders
                run.text = placeholder_pattern.sub('', run.text)
                
                # Hide {{SOMETHING}} placeholders by making them white
                # BUT only if they are actually placeholders (not real values)
                if curly_placeholder_pattern.search(run.text):
                    # Check if this is actually a placeholder or a real value
                    should_hide = True
                    for placeholder in keep_placeholders:
                        if placeholder in run.text:
                            # If it's a placeholder we want to keep, check if it has a real value
                            # Real values would not be exactly the placeholder text
                            if run.text.strip() == placeholder:
                                # This is still a placeholder, hide it
                                run.font.color.rgb = RGBColor(255, 255, 255)  # White color
                                print(f"DEBUG: Hiding curly placeholder in run: '{run.text}'")
                            else:
                                # This has a real value, don't hide it
                                should_hide = False
                                print(f"DEBUG: Keeping visible value in run: '{run.text}'")
                            break
                    
                    # For other curly placeholders not in keep_placeholders, hide them
                    if should_hide and not any(placeholder in run.text for placeholder in keep_placeholders):
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White color
                        print(f"DEBUG: Hiding other curly placeholder in run: '{run.text}'")
                
                # Remove unwanted phrase
                run.text = representation_pattern.sub('', run.text)

        # Clean paragraphs
        for paragraph in doc.paragraphs:
            clean_runs(paragraph.runs)
        # Clean tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        clean_runs(paragraph.runs)
    
    def clean_text(self, text):
        # Remove duplicate commas, semicolons, and spaces
        text = re.sub(r'[;,]{2,}', lambda m: m.group(0)[0], text)  # Replace multiple , or ; with one
        text = re.sub(r'([;,])\s*([;,])', r'\1', text)  # Remove space between ,; or ;,
        text = re.sub(r'\s{2,}', ' ', text)  # Replace multiple spaces with one
        text = re.sub(r'\s+,', ',', text)  # Remove space before comma
        text = re.sub(r'\s+;', ';', text)  # Remove space before semicolon
        text = re.sub(r',\s*;', ';', text)  # Replace ', ;' with ';'
        text = re.sub(r';\s*,', ',', text)  # Replace '; ,' with ','
        text = re.sub(r'([;,]){2,}', r'\1', text)  # Remove any remaining double punctuation
        text = re.sub(r'\s{2,}', ' ', text)  # Again, just in case
        text = re.sub(r'\s*([;,])\s*', r'\1 ', text)  # Normalize space after punctuation
        text = re.sub(r'\s{2,}', ' ', text)  # Final pass for spaces
        text = re.sub(r'\s+\.', '.', text)  # Remove space before period
        text = re.sub(r'\s+\,', ',', text)  # Remove space before comma
        text = re.sub(r'\s+\;', ';', text)  # Remove space before semicolon
        text = re.sub(r'\s+\:', ':', text)  # Remove space before colon
        text = re.sub(r'\s+\?', '?', text)  # Remove space before question mark
        text = re.sub(r'\s+\!', '!', text)  # Remove space before exclamation mark
        # Remove leading/trailing punctuation and spaces
        text = re.sub(r'^[,;\s]+', '', text)
        text = re.sub(r'[,;\s]+$', '', text)
        return text.strip()
    
    def _get_template_from_r2(self, template_id: int) -> bytes:
        """
        Get template from R2 storage (placeholder for your existing logic)
        """
        template = TplTemplate.objects.get(pktemplate=template_id)
        s3 = boto3.client(
            's3',
            endpoint_url=os.environ.get('CLOUDFLARE_R2_ENDPOINT'),
            aws_access_key_id=os.environ.get('CLOUDFLARE_R2_ACCESS_KEY'),
            aws_secret_access_key=os.environ.get('CLOUDFLARE_R2_SECRET_KEY'),
            config=Config(signature_version='s3v4'),
            region_name='auto',
        )
        
        # Template path in simplified structure
        object_key = f"rodriguez-zea/plantillas/{template.filename}"
        
        try:
            s3_response = s3.get_object(
                Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'), 
                Key=object_key
            )
            return s3_response['Body'].read()
        except Exception as e:
            # Log the error and return a 404 or raise
            print(f'Template not found in R2: {e}')
            raise FileNotFoundError(f"Template not found in R2: {object_key}")
    
    def get_document_data(self, num_kardex: str) -> Dict[str, Any]:
        """
        Get dummy data for document placeholders
        """
        # Extract year from kardex (similar to PHP logic)
        arr_kardex = num_kardex.split('-')
        anio_kardex = arr_kardex[1] if len(arr_kardex) > 1 else "2024"
        
        # Document data
        document_data = self._get_document_data(num_kardex, anio_kardex)
        
        # Vehicle data
        vehicle_data = self._get_vehicle_data(num_kardex)
        
        # Payment data
        payment_data = self._get_payment_data(num_kardex)
        
        # Contractors data
        contractors_data = self._get_contractors_data(num_kardex)
        
        # Escrituración data 
        # escrituracion_data = self._get_escrituracion_data(num_kardex)

        # Merge all data
        final_data = {}
        final_data.update(document_data)
        final_data.update(vehicle_data)
        final_data.update(payment_data)
        final_data.update(contractors_data)
        # final_data.update(escrituracion_data)
        
        return final_data
    
    def _get_document_data(self, num_kardex: str, anio_kardex: str) -> Dict[str, str]:
        """
        Get basic document information
        """
        kardex = Kardex.objects.filter(kardex=num_kardex).first()
        if not kardex:
            raise ValueError(f"Kardex {num_kardex} not found")
        
        # Get user information
        usuario = kardex.responsable_new or ''
        usuario_dni = ''
        if kardex.idusuario:
            # idusuario is an integer ID, not a user object
            user = Usuarios.objects.filter(idusuario=kardex.idusuario).first()
            if user:
                usuario_dni = user.dni or ''
        
        # Get abogado information
        abogado = ''
        matricula = ''
        if kardex.idabogado:
            from notaria.models import TbAbogado
            abogado_obj = TbAbogado.objects.filter(idabogado=kardex.idabogado).first()
            if abogado_obj:
                abogado = abogado_obj.razonsocial or ''
                matricula = abogado_obj.matricula or ''
        
        numero_escritura = kardex.numescritura or ''
        fecha_escritura = kardex.fechaescritura or datetime.now()
        numero_minuta = kardex.numminuta or ''
        folioini = kardex.folioini or ''
        foliofin = kardex.foliofin or ''
        papelini = kardex.papelini or ''
        papelfin = kardex.papelfin or ''
        
        return {
            'K': num_kardex,
            'NRO_ESC': f"{numero_escritura}({self.letras.number_to_letters(numero_escritura)})" if numero_escritura else '{{NRO_ESC}}',
            'NUM_REG': '1',
            'FEC_LET': self.letras.date_to_letters(fecha_escritura) if fecha_escritura else '',
            'F_IMPRESION': self.letras.date_to_letters(fecha_escritura) if fecha_escritura else '{{F_IMPRESION}}',
            'USUARIO': usuario,
            'USUARIO_DNI': usuario_dni,
            'NRO_MIN': numero_minuta or '{{NRO_MIN}}',
            'COMPROBANTE': ' ',
            'O_S': ' ',
            'ORDEN_SERVICIO': ' ',
            'F': self.letras.date_to_letters(fecha_escritura) if fecha_escritura else '{{F}}',
            'DESCRIPCION_SELLO': f"{abogado} CAP. {matricula}",
            'FI': folioini or '{{FI}}',
            'FF': foliofin or '{{FF}}',
            'S_IN': papelini or '{{S_IN}}',
            'S_FN': papelfin or '{{S_FN}}',
        }
    
    def _get_vehicle_data(self, kardex) -> Dict[str, str]:
        vehicle = Detallevehicular.objects.filter(kardex=kardex).first()
        sede = ''
        num_zona = ''
        zona_registral = ''
        if vehicle and vehicle.idsedereg:
            sede_obj = Sedesregistrales.objects.filter(idsedereg=vehicle.idsedereg).first()
            if sede_obj:
                sede = sede_obj.dessede or ''
                num_zona = sede_obj.num_zona or ''
                zona_registral = sede_obj.dessede or ''
        return {
            'PLACA': vehicle.numplaca if vehicle else '',
            'CLASE': vehicle.clase if vehicle else '',
            'MARCA': vehicle.marca if vehicle else '',
            'MODELO': vehicle.modelo if vehicle else '',
            'AÑO_FABRICACION': vehicle.anofab if vehicle else '',
            'CARROCERIA': vehicle.carroceria if vehicle else '',
            'COLOR': vehicle.color if vehicle else '',
            'NRO_MOTOR': vehicle.motor if vehicle else '',
            'NRO_SERIE': vehicle.numserie if vehicle else '',
            'FEC_INS': vehicle.fecinsc if vehicle else '',
            'FECHA_INSCRIPCION': vehicle.fecinsc if vehicle else '',
            'ZONA_REGISTRAL': zona_registral,
            'NUM_ZONA_REG': num_zona,
            'SEDE': sede,
            'INSTRUIDO': 'INSTRUIDO',
            'COMBUSTIBLE': vehicle.combustible if vehicle else '',
            'NRO_TARJETA': '',
            'NRO_CILINDROS': vehicle.numcil if vehicle else '',
        }
    
    def _get_payment_data(self, kardex) -> Dict[str, str]:
        """
        Get payment information based on payment method
        """

        patrimonial = Patrimonial.objects.filter(
            kardex=kardex)
        if patrimonial.exists():
            patrimonial = patrimonial.first()

        sunat_medio_pago = "008"  # Cash payment
        precio = patrimonial.importetrans if patrimonial else '0.00'
        moneda = MONEDAS[patrimonial.idmon]['desmon'] if patrimonial else '' 
        simbolo_moneda = MONEDAS[patrimonial.idmon]['simbolo'] if patrimonial else ''
        forma_pago = FORMAS_PAGO[patrimonial.fpago]['descripcion'] if patrimonial else ''
        
        # Payment method logic (similar to PHP switch)
        if sunat_medio_pago == "008":
            medio_pago = f'EL COMPRADOR DECLARA QUE HA PAGADO EL PRECIO DEL VEHICULO EN DINERO EN EFECTIVO Y CON ANTERIORIDAD A LA CELEBRACION DE LA PRESENTE ACTA DE TRANSFERENCIA. NO HABIENDO UTILIZADO NINGÚN MEDIO DE PAGO ESTABLECIDO EN LA LEY Nº 28194, PORQUE EL MONTO TOTAL NO ES IGUAL NI SUPERA LOS S/ 2,000.00 O US$ 500.00. EL TIPO Y CÓDIGO DEL MEDIO EMPLEADO ES: "EFECTIVO POR OPERACIONES EN LAS QUE NO EXISTE OBLIGACIÓN DE UTILIZAR MEDIOS DE PAGO-008". POR SER EL PAGO DEL PRECIO INFERIOR A 1 UIT. MODIFICADO POR EL DECRETO LEGISLATIVO N° 1529.'
            exhibio_medio_pago = 'SE DEJA CONSTANCIA QUE PARA LA REALIZACIÓN DEL PRESENTE ACTO, LAS PARTES NO ME HAN EXHIBIDO NINGÚN MEDIO DE PAGO. DOY FE.'
            fin_medio_pago = 'EN DINERO EN EFECTIVO'
        else:
            # Default case
            medio_pago = 'EL COMPRADOR DECLARA QUE HA PAGADO EL PRECIO DEL VEHICULO EN DINERO EN EFECTIVO Y CON ANTERIORIDAD A LA CELEBRACION DE LA PRESENTE ACTA DE TRANSFERENCIA.'
            exhibio_medio_pago = 'SE DEJA CONSTANCIA QUE PARA LA REALIZACIÓN DEL PRESENTE ACTO, LAS PARTES NO ME HAN EXHIBIDO NINGÚN MEDIO DE PAGO. DOY FE.'
            fin_medio_pago = 'EN DINERO EN EFECTIVO'
        
        return {
            'MONTO': precio,
            'MON_VEHI': moneda,
            'MONTO_LETRAS': self.letras.money_to_letters(moneda, Decimal(precio)),
            'MONEDA_C': simbolo_moneda,
            'SUNAT_MED_PAGO': sunat_medio_pago,
            'DES_PRE_VEHI': self.letras.money_to_letters(moneda, Decimal(precio)),
            'EXH_MED_PAGO': exhibio_medio_pago,
            'MED_PAGO': medio_pago,
            'FIN_MED_PAGO': fin_medio_pago,
            'FORMA_PAGO': forma_pago,
            'C_INICIO_MP': '',
            'TIPO_PAGO_E': '',
            'TIPO_PAGO_C': '',
            'MONTO_MP': '',
            'CONSTANCIA': exhibio_medio_pago,
            'DETALLE_MP': '',
            'FORMA_PAGO_S': '',
            'MONEDA_C_MP': '',
            'MEDIO_PAGO_C': '',
            'MP_MEDIO_PAGO': '',
            'MP_COMPLETO': '',
            'USO': '',
        }

    def _get_contractors_data(self, kardex) -> Dict[str, str]:
        """
        Get contractors (transferor and acquirer) information, including both natural and legal persons.
        Robustly handles missing nacionalidad and fills all placeholders for both types.
        Handles EN REPRESENTACION DE for both persona natural and juridica.
        """
        TRANSFEROR_ROLES = {
            "VENDEDOR", "DONANTE", "PODERDANTE", "OTORGANTE", "REPRESENTANTE", "ANTICIPANTE",
            "ADJUDICANTE", "USUFRUCTUANTE", "TRANSFERENTE", "TITULAR", "MUTUANTE", "PROPIETARIO",
            "DEUDOR", "ASOCIANTE", "TRANSFERENTE / PROPIETARIO (VENDEDOR)", "APODERADO"
        }
        ACQUIRER_ROLES = {
            "COMPRADOR", "DONATARIO", "APODERADO", "ANTICIPADO", "ADJUDICATARIO", "USUFRUCTUARIO",
            "TESTIGO A RUEGO", "ADQUIRIENTE", "ACREEDOR", "OTORGADO", "MUTUATARIO", "BENEFICIARIA",
            "ASOCIADO", "ADQUIRENTE / BENEFICIARIO (COMPRADOR)", "REPRESENTANTE"
        }
        REPRESENTATIVE_ROLES = {"APODERADO", "REPRESENTANTE"}

        contratantes = list(Contratantes.objects.filter(kardex=kardex))
        id_to_contratante = {c.idcontratante: c for c in contratantes}
        contratantes_list = []

        for contratante in contratantes:
            condiciones = contratante.condicion.split('/')
            condiciones_list = []
            for condicion in condiciones:
                if condicion:
                    condicion_int = condicion.split('.')[0]
                    condicion_str = Actocondicion.objects.get(idcondicion=condicion_int).condicion
                    condiciones_list.append(condicion_str)
            cliente2 = Cliente2.objects.get(idcontratante=contratante.idcontratante)
            # Robust nacionalidad handling
            if cliente2.tipper == 'J':
                nacionalidad = 'EMPRESA'
                sexo = ''
                ocupacion = ''
                estado_civil = ''
                direccion = cliente2.domfiscal or ''
            else:
                nacionalidad_obj = Nacionalidades.objects.filter(idnacionalidad=cliente2.nacionalidad).first()
                nacionalidad = nacionalidad_obj.descripcion if nacionalidad_obj else ''
                sexo = cliente2.sexo or ''
                ocupacion = re.split(r'[/,;]', cliente2.detaprofesion)[0].strip() if cliente2.detaprofesion else ''
                estado_civil = self.get_civil_status_by_gender(CIVIL_STATUS[cliente2.idestcivil]['label'].upper(), sexo) if cliente2.idestcivil in CIVIL_STATUS else ''
                direccion = cliente2.direccion or ''
            contratante_obj = {
                'idcontratante': contratante.idcontratante,
                'sexo': sexo,
                'condiciones': (', ').join(condiciones_list),
                'condicion_str': condiciones_list[0] if condiciones_list else '',
                'nombres': f'{cliente2.prinom} {cliente2.segnom} {cliente2.apepat} {cliente2.apemat}' if cliente2.tipper == 'N' else cliente2.razonsocial or '',
                'nacionalidad': self.get_nationality_by_gender(nacionalidad, sexo) if cliente2.tipper == 'N' else nacionalidad,
                'tipoDocumento': TIPO_DOCUMENTO[cliente2.idtipdoc]['destipdoc'] if cliente2.idtipdoc in TIPO_DOCUMENTO else '',
                'numeroDocumento': cliente2.numdoc or '',
                'ocupacion': ocupacion,
                'estadoCivil': estado_civil,
                'direccion': direccion,
                'idcontratanterp': getattr(contratante, 'idcontratanterp', None),
                'tipper': cliente2.tipper,  # 'N' for natural, 'J' for juridica
                'razonsocial': cliente2.razonsocial or '',
                'domfiscal': cliente2.domfiscal or '',
                'numpartida': cliente2.numpartida or '',
                'idubigeo': cliente2.idubigeo or '',
                'numdoc_empresa': cliente2.numdoc if cliente2.tipper == 'J' else '',
            }
            contratantes_list.append(contratante_obj)

        naturals = [c for c in contratantes_list if c['tipper'] == 'N']
        companies = [c for c in contratantes_list if c['tipper'] == 'J']

        transferors = []
        transferor_companies = []
        for c in naturals:
            if c['condicion_str'] in REPRESENTATIVE_ROLES and c.get('idcontratanterp'):
                principal_contratante = id_to_contratante.get(c['idcontratanterp'])
                if principal_contratante:
                    principal_cliente = Cliente2.objects.get(idcontratante=principal_contratante.idcontratante)
                    if principal_cliente.tipper == 'J':
                        principal_name = principal_cliente.razonsocial or ''
                    else:
                        principal_name = f'{principal_cliente.prinom} {principal_cliente.segnom} {principal_cliente.apepat} {principal_cliente.apemat}'
                    c = c.copy()
                    c['nombres'] = f"{c['nombres'].strip()}, EN REPRESENTACION DE {principal_name.strip()}"
                transferors.append(c)
            elif c['condicion_str'] in TRANSFEROR_ROLES:
                represented = any(
                    a['condicion_str'] in REPRESENTATIVE_ROLES and a.get('idcontratanterp') == c.get('idcontratante')
                    for a in naturals
                )
                if not represented:
                    transferors.append(c)
        for c in companies:
            if c['condicion_str'] in TRANSFEROR_ROLES:
                transferor_companies.append(c)

        acquirers = []
        acquirer_companies = []
        for c in naturals:
            if c['condicion_str'] in REPRESENTATIVE_ROLES and c.get('idcontratanterp'):
                principal_contratante = id_to_contratante.get(c['idcontratanterp'])
                if principal_contratante:
                    principal_cliente = Cliente2.objects.get(idcontratante=principal_contratante.idcontratante)
                    if principal_cliente.tipper == 'J':
                        principal_name = principal_cliente.razonsocial or ''
                    else:
                        principal_name = f'{principal_cliente.prinom} {principal_cliente.segnom} {principal_cliente.apepat} {principal_cliente.apemat}'
                    principal_condiciones = principal_contratante.condicion.split('/')
                    principal_condicion_str = Actocondicion.objects.get(idcondicion=principal_condiciones[0].split('.')[0]).condicion if principal_condiciones and principal_condiciones[0] else ''
                    if principal_condicion_str in ACQUIRER_ROLES:
                        c = c.copy()
                        c['nombres'] = f"{c['nombres'].strip()}, EN REPRESENTACION DE {principal_name.strip()}"
                        acquirers.append(c)
            elif c['condicion_str'] in ACQUIRER_ROLES:
                represented = any(
                    a['condicion_str'] in REPRESENTATIVE_ROLES and a.get('idcontratanterp') == c.get('idcontratante')
                    for a in naturals
                )
                if not represented:
                    acquirers.append(c)
        for c in companies:
            if c['condicion_str'] in ACQUIRER_ROLES:
                acquirer_companies.append(c)

        contractors_data = {}

        for idx, t in enumerate(transferors, 1):
            contractors_data[f'P_NOM_{idx}'] = t['nombres'] + ', '
            contractors_data[f'P_NACIONALIDAD_{idx}'] = t['nacionalidad'] + ', '
            contractors_data[f'P_TIP_DOC_{idx}'] = t['tipoDocumento']
            contractors_data[f'P_DOC_{idx}'] = self.get_identification_phrase(t['sexo'], t['tipoDocumento'], t['numeroDocumento'])
            contractors_data[f'P_OCUPACION_{idx}'] = t['ocupacion']
            contractors_data[f'P_ESTADO_CIVIL_{idx}'] = t['estadoCivil']
            contractors_data[f'P_DOMICILIO_{idx}'] = 'CON DOMICILIO EN ' + t['direccion']
            contractors_data[f'P_IDE_{idx}'] = ' '
            contractors_data[f'SEXO_P_{idx}'] = t['sexo']
            contractors_data[f'P_FIRMAN_{idx}'] = 'FIRMA EN'
            contractors_data[f'P_IMPRIME_{idx}'] = 'IMPRIME'
            
            # Add unnumbered versions for first person (like PHP legacy code)
            if idx == 1:
                contractors_data['P_NOM'] = t['nombres'] + ', '
                contractors_data['P_NACIONALIDAD'] = t['nacionalidad'] + ', '
                contractors_data['P_TIP_DOC'] = t['tipoDocumento']
                contractors_data['P_DOC'] = self.get_identification_phrase(t['sexo'], t['tipoDocumento'], t['numeroDocumento'])
                contractors_data['P_OCUPACION'] = t['ocupacion']
                contractors_data['P_ESTADO_CIVIL'] = t['estadoCivil']
                contractors_data['P_DOMICILIO'] = 'CON DOMICILIO EN ' + t['direccion']
                contractors_data['P_IDE'] = ' '
                contractors_data['SEXO_P'] = t['sexo']
                contractors_data['P_FIRMAN'] = 'FIRMA EN'
                contractors_data['P_IMPRIME'] = 'IMPRIME'
        
        for idx, c in enumerate(acquirers, 1):
            contractors_data[f'C_NOM_{idx}'] = c['nombres'] + ', '
            contractors_data[f'C_NACIONALIDAD_{idx}'] = c['nacionalidad'] + ', '
            contractors_data[f'C_TIP_DOC_{idx}'] = c['tipoDocumento']
            contractors_data[f'C_DOC_{idx}'] = self.get_identification_phrase(c['sexo'], c['tipoDocumento'], c['numeroDocumento'])
            contractors_data[f'C_OCUPACION_{idx}'] = c['ocupacion']
            contractors_data[f'C_ESTADO_CIVIL_{idx}'] = c['estadoCivil']
            contractors_data[f'C_DOMICILIO_{idx}'] = 'CON DOMICILIO EN ' + c['direccion']
            contractors_data[f'C_IDE_{idx}'] = ' '
            contractors_data[f'SEXO_C_{idx}'] = c['sexo']
            contractors_data[f'C_FIRMAN_{idx}'] = 'FIRMA EN'
            contractors_data[f'C_IMPRIME_{idx}'] = 'IMPRIME'
            
            # Add unnumbered versions for first person (like PHP legacy code)
            if idx == 1:
                contractors_data['C_NOM'] = c['nombres'] + ', '
                contractors_data['C_NACIONALIDAD'] = c['nacionalidad'] + ', '
                contractors_data['C_TIP_DOC'] = c['tipoDocumento']
                contractors_data['C_DOC'] = self.get_identification_phrase(c['sexo'], c['tipoDocumento'], c['numeroDocumento'])
                contractors_data['C_OCUPACION'] = c['ocupacion']
                contractors_data['C_ESTADO_CIVIL'] = c['estadoCivil']
                contractors_data['C_DOMICILIO'] = 'CON DOMICILIO EN ' + c['direccion']
                contractors_data['C_IDE'] = c['numeroDocumento'] or ' '
                contractors_data['SEXO_C'] = c['sexo']
                contractors_data['C_FIRMAN'] = 'FIRMA EN'
                contractors_data['C_IMPRIME'] = 'IMPRIME'

        for idx, t in enumerate(transferor_companies, 1):
            contractors_data[f'NOMBRE_EMPRESA_{idx}'] = t['razonsocial']
            contractors_data[f'INS_EMPRESA_{idx}'] = f'INSCRITA EN LA PARTIDA ELECTRONICA N° {t["numpartida"]}'
            contractors_data[f'RUC_{idx}'] = t['numdoc_empresa']
            contractors_data[f'DOMICILIO_EMPRESA_{idx}'] = f'CON DOMICILIO EN {t["domfiscal"]}'

        for idx, c in enumerate(acquirer_companies, 1):
            contractors_data[f'NOMBRE_EMPRESA_{idx+len(transferor_companies)}'] = c['razonsocial']
            contractors_data[f'INS_EMPRESA_{idx+len(transferor_companies)}'] = f'INSCRITA EN LA PARTIDA ELECTRONICA N° {c["numpartida"]}'
            contractors_data[f'RUC_{idx+len(transferor_companies)}'] = c['numdoc_empresa']
            contractors_data[f'DOMICILIO_EMPRESA_{idx+len(transferor_companies)}'] = f'CON DOMICILIO EN {c["domfiscal"]}'

        for idx in range(len(transferors) + 1, 11):
            contractors_data[f'P_NOM_{idx}'] = f'[E.P_NOM_{idx}]'
            contractors_data[f'P_NACIONALIDAD_{idx}'] = f'[E.P_NACIONALIDAD_{idx}]'
            contractors_data[f'P_TIP_DOC_{idx}'] = f'[E.P_TIP_DOC_{idx}]'
            contractors_data[f'P_DOC_{idx}'] = f'[E.P_DOC_{idx}]'
            contractors_data[f'P_OCUPACION_{idx}'] = f'[E.P_OCUPACION_{idx}]'
            contractors_data[f'P_ESTADO_CIVIL_{idx}'] = f'[E.P_ESTADO_CIVIL_{idx}]'
            contractors_data[f'P_DOMICILIO_{idx}'] = f'[E.P_DOMICILIO_{idx}]'
            contractors_data[f'P_IDE_{idx}'] = f'[E.P_IDE_{idx}]'
            contractors_data[f'SEXO_P_{idx}'] = f'[E.SEXO_P_{idx}]'
            contractors_data[f'P_FIRMAN_{idx}'] = f'[E.P_FIRMAN_{idx}]'
            contractors_data[f'P_IMPRIME_{idx}'] = f'[E.P_IMPRIME_{idx}]'

        for idx in range(len(acquirers) + 1, 11):
            contractors_data[f'C_NOM_{idx}'] = f'[E.C_NOM_{idx}]'
            contractors_data[f'C_NACIONALIDAD_{idx}'] = f'[E.C_NACIONALIDAD_{idx}]'
            contractors_data[f'C_TIP_DOC_{idx}'] = f'[E.C_TIP_DOC_{idx}]'
            contractors_data[f'C_DOC_{idx}'] = f'[E.C_DOC_{idx}]'
            contractors_data[f'C_OCUPACION_{idx}'] = f'[E.C_OCUPACION_{idx}]'
            contractors_data[f'C_ESTADO_CIVIL_{idx}'] = f'[E.C_ESTADO_CIVIL_{idx}]'
            contractors_data[f'C_DOMICILIO_{idx}'] = f'[E.C_DOMICILIO_{idx}]'
            contractors_data[f'C_IDE_{idx}'] = f'[E.C_IDE_{idx}]'
            contractors_data[f'SEXO_C_{idx}'] = f'[E.C_SEXO_{idx}]'
            contractors_data[f'C_FIRMAN_{idx}'] = f'[E.C_FIRMAN_{idx}]'
            contractors_data[f'C_IMPRIME_{idx}'] = f'[E.C_IMPRIME_{idx}]'

        for idx in range(len(transferor_companies) + len(acquirer_companies) + 1, 6):
            contractors_data[f'NOMBRE_EMPRESA_{idx}'] = f'[E.NOMBRE_EMPRESA_{idx}]'
            contractors_data[f'INS_EMPRESA_{idx}'] = f'[E.INS_EMPRESA_{idx}]'
            contractors_data[f'RUC_{idx}'] = f'[E.RUC_{idx}]'
            contractors_data[f'DOMICILIO_EMPRESA_{idx}'] = f'[E.DOMICILIO_EMPRESA_{idx}]'

        if transferors:
            contractors_data['P_NOM'] = contractors_data['P_NOM_1']
            contractors_data['P_NACIONALIDAD'] = contractors_data['P_NACIONALIDAD_1']
            contractors_data['P_TIP_DOC'] = contractors_data['P_TIP_DOC_1']
            contractors_data['P_DOC'] = contractors_data['P_DOC_1']
            contractors_data['P_OCUPACION'] = contractors_data['P_OCUPACION_1']
            contractors_data['P_ESTADO_CIVIL'] = contractors_data['P_ESTADO_CIVIL_1']
            contractors_data['P_DOMICILIO'] = contractors_data['P_DOMICILIO_1']
            contractors_data['P_IDE'] = contractors_data['P_IDE_1']
            contractors_data['SEXO_P'] = contractors_data['SEXO_P_1']
            contractors_data['P_FIRMAN'] = contractors_data['P_FIRMAN_1']
            contractors_data['P_IMPRIME'] = contractors_data['P_IMPRIME_1']

        if acquirers:
            contractors_data['C_NOM'] = contractors_data['C_NOM_1']
            contractors_data['C_NACIONALIDAD'] = contractors_data['C_NACIONALIDAD_1']
            contractors_data['C_TIP_DOC'] = contractors_data['C_TIP_DOC_1']
            contractors_data['C_DOC'] = contractors_data['C_DOC_1']
            contractors_data['C_OCUPACION'] = contractors_data['C_OCUPACION_1']
            contractors_data['C_ESTADO_CIVIL'] = contractors_data['C_ESTADO_CIVIL_1']
            contractors_data['C_DOMICILIO'] = contractors_data['C_DOMICILIO_1']
            contractors_data['C_IDE'] = contractors_data['C_IDE_1']
            contractors_data['SEXO_C'] = contractors_data['SEXO_C_1']
            contractors_data['C_FIRMAN'] = contractors_data['C_FIRMAN_1']
            contractors_data['C_IMPRIME'] = contractors_data['C_IMPRIME_1']

        contractors_data.update(self.get_articles_and_grammar(transferors, 'P'))
        contractors_data.update(self.get_articles_and_grammar(acquirers, 'C'))

        return contractors_data


    def get_identification_phrase(self, gender, doc_type, doc_number):
        if gender == 'F':
            return f'IDENTIFICADA CON {doc_type} N° {doc_number}, '
        else:
            return f'IDENTIFICADO CON {doc_type} N° {doc_number}, '

    def get_civil_status_by_gender(self, civil_status, gender):
        if not civil_status:
            return ''
        if gender == 'F':
            return civil_status[:-1] + 'A, '
        else:
            return civil_status + ', '

    def classify_contratantes(self, contratantes):
        TRANSFEROR_ROLES = {'VENDEDOR', 'DONANTE', 'APODERADO', 'CEDENTE', 'ARRENDADOR', 'MUTUANTE', 'ADJUDICANTE'}
        ACQUIRER_ROLES = {'COMPRADOR', 'DONATARIO', 'CESIONARIO', 'ARRENDATARIO', 'MUTUARIO', 'ADJUDICATARIO'}

        transferors = [c for c in contratantes if c['condiciones'] in TRANSFEROR_ROLES]
        acquirers = [c for c in contratantes if c['condiciones'] in ACQUIRER_ROLES]
        return transferors, acquirers

    def get_nationality_by_gender(self, nationality, gender):
        if not nationality:
            return ''
        base = nationality[:-1]  # Remove last character
        if gender == 'F':
            return base + 'A'
        else:
            return base + 'O'
    def get_articles_and_grammar(self, people, role_prefix):
        count = len(people)
        if count == 0:
            return {
                f'EL_{role_prefix}': '',
                f'{role_prefix}_CALIDAD': '',
                f'{role_prefix}_INICIO': '',
                f'{role_prefix}_AMBOS': '',
            }
        all_female = all(p['sexo'] == 'F' for p in people)
        all_male = all(p['sexo'] == 'M' for p in people)
        ambos = ' AMBOS ' if count > 1 else ' '
        main_role = people[0]['condiciones'] if people else ''
        role_labels = ROLE_LABELS.get(main_role, {})
        if count > 1:
            calidad = role_labels.get('F_PL' if all_female else 'M_PL', main_role + 'S')
            inicio = ' SEÑORAS' if all_female else ' SEÑORES'
            el = 'LAS' if all_female else 'LOS'
        else:
            calidad = role_labels.get('F' if all_female else 'M', main_role)
            inicio = ' SEÑORA' if all_female else 'SEÑOR'
            el = 'LA' if all_female else 'EL'
        return {
            f'EL_{role_prefix}': el,
            f'{role_prefix}_CALIDAD': calidad,
            f'{role_prefix}_INICIO': inicio,
            f'{role_prefix}_AMBOS': ambos,
        }
    
    def _process_document(self, template_bytes: bytes, data: Dict[str, str]) -> Document:
        """
        Process the document template with data
        """
        buffer = io.BytesIO(template_bytes)
        doc = DocxTemplate(buffer)
        doc.render(data)
        return doc
    
    def _create_response(self, doc: Document, filename: str, kardex: str, mode: str = "download") -> HttpResponse:
        """
        Create HTTP response with the document
        """
        # Add the custom property
        custom_props = CustomProperties(doc)
        custom_props['documentoGeneradoId'] = kardex

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        if mode == "open":
            # Production mode: Return JSON with file info for Word opening
            response = JsonResponse({
                'status': 'success',
                'mode': 'open',
                'filename': filename,
                'kardex': kardex,
                'message': 'Document generated and ready to open in Word'
            })
            response['Access-Control-Allow-Origin'] = '*'
            return response
        else:
            # Testing mode: Download the document
            response = HttpResponse(
                buffer.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'inline; filename="{filename}"'
            response['Content-Length'] = str(buffer.getbuffer().nbytes)
            response['Access-Control-Allow-Origin'] = '*'
            return response


class NonContentiousDocumentService:
    """
    Django service to generate non-contentious documents based on the PHP logic
    """
    
    def __init__(self):
        self.letras = NumberToLetterConverter()
    
    def generate_non_contentious_document(self, template_id: int, num_kardex: str, idtipoacto: str, action: str = 'generate', mode: str = "download") -> HttpResponse:
        """
        Main method to generate non-contentious document
        """
        start_time = time.time()
        print(f"PERF: Starting non-contentious document generation for kardex: {num_kardex}")
        
        try:
            # Step 1: Get template from R2
            template_start = time.time()
            template = self._get_template_from_r2(template_id)
            template_time = time.time() - template_start
            print(f"PERF: Non-contentious template download took {template_time:.2f}s")
            
            # Step 2: Get document data
            data_start = time.time()
            document_data = self.get_document_data(num_kardex, idtipoacto)
            data_time = time.time() - data_start
            print(f"PERF: Non-contentious data retrieval took {data_time:.2f}s")
            
            # Step 3: Process document
            process_start = time.time()
            doc = self._process_document(template, document_data)
            process_time = time.time() - process_start
            print(f"PERF: Non-contentious document processing took {process_time:.2f}s")
            
            # Step 4: Remove placeholders
            cleanup_start = time.time()
            self.remove_unfilled_placeholders(doc)
            cleanup_time = time.time() - cleanup_start
            print(f"PERF: Non-contentious placeholder cleanup took {cleanup_time:.2f}s")
            
            # Step 5: Upload to R2
            upload_start = time.time()
            upload_success = self.create_documento_in_r2(doc, num_kardex)
            upload_time = time.time() - upload_start
            print(f"PERF: Non-contentious R2 upload took {upload_time:.2f}s")
            
            if not upload_success:
                print(f"WARNING: Failed to upload non-contentious document to R2 for kardex: {num_kardex}")
            
            total_time = time.time() - start_time
            print(f"PERF: Total non-contentious document generation took {total_time:.2f}s")
            
            return self._create_response(doc, f"__PROY__{num_kardex}.docx", num_kardex, mode)
        except FileNotFoundError as e:
            return HttpResponse(str(e), status=404)
        except Exception as e:
            total_time = time.time() - start_time
            print(f"PERF: Non-contentious document generation failed after {total_time:.2f}s")
            return HttpResponse(f"Error generating document: {str(e)}", status=500)

    def create_documento_in_r2(self, doc, kardex):
        """
        Create a new document in R2 storage
        """
        try:
            print(f"DEBUG: Starting R2 upload for non-contentious kardex: {kardex}")
            
            # Save the document to a bytes buffer
            from io import BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            doc_content = buffer.read()
            
            print(f"DEBUG: Non-contentious document size: {len(doc_content)} bytes")
            
            # Define the object key for R2 - simplified path structure
            object_key = f"rodriguez-zea/documentos/__PROY__{kardex}.docx"
            
            # Check environment variables
            endpoint_url = os.environ.get('CLOUDFLARE_R2_ENDPOINT')
            access_key = os.environ.get('CLOUDFLARE_R2_ACCESS_KEY')
            secret_key = os.environ.get('CLOUDFLARE_R2_SECRET_KEY')
            bucket = os.environ.get('CLOUDFLARE_R2_BUCKET')
            
            print(f"DEBUG: R2 Configuration - Endpoint: {endpoint_url}")
            print(f"DEBUG: R2 Configuration - Access Key: {'SET' if access_key else 'NOT SET'}")
            print(f"DEBUG: R2 Configuration - Secret Key: {'SET' if secret_key else 'NOT SET'}")
            print(f"DEBUG: R2 Configuration - Bucket: {bucket}")
            
            # Use cached S3 client
            s3 = get_s3_client()
            
            print(f"DEBUG: S3 client created successfully for non-contentious")
            print(f"DEBUG: Uploading non-contentious document to bucket: {bucket}, key: {object_key}")
            
            s3.upload_fileobj(
                BytesIO(doc_content),
                bucket,
                object_key
            )
            
            print(f"DEBUG: Non-contentious document uploaded to R2: {object_key}")
            return True
            
        except Exception as e:
            print(f"ERROR: Failed to upload non-contentious document to R2: {e}")
            print(f"ERROR: Exception type: {type(e).__name__}")
            import traceback
            print(f"ERROR: Full traceback: {traceback.format_exc()}")
            return False

    def remove_unfilled_placeholders(self, doc):
        """
        Remove all [E.SOMETHING] placeholders and hide {{SOMETHING}} placeholders from user.
        Keep and hide: {{NRO_ESC}}, {{FI}}, {{FF}}, {{S_IN}}, {{S_FN}}, {{FECHA_ACT}}
        Remove all others: [E.P_NOM_2], [E.C_NOM_2], etc.
        """
        import re
        placeholder_pattern = re.compile(r'\[E\.[A-Z0-9_]+\]')
        curly_placeholder_pattern = re.compile(r'\{\{[A-Z0-9_]+\}\}')
        representation_pattern = re.compile(
            r'EN REPRESENTACION DE\s*Y[\s.·…‥⋯⋮⋱⋰⋯—–-]*', re.IGNORECASE
        )

        # List of placeholders to keep and hide
        keep_placeholders = ['{{NRO_ESC}}', '{{FI}}', '{{FF}}', '{{S_IN}}', '{{S_FN}}', '{{FECHA_ACT}}']

        def clean_runs(runs):
            # Remove [E.SOMETHING] placeholders and hide {{SOMETHING}} placeholders
            for run in runs:
                # Remove all [E.SOMETHING] placeholders
                run.text = placeholder_pattern.sub('', run.text)
                
                # Hide {{SOMETHING}} placeholders by making them white
                # BUT only if they are actually placeholders (not real values)
                if curly_placeholder_pattern.search(run.text):
                    # Check if this is actually a placeholder or a real value
                    should_hide = True
                    for placeholder in keep_placeholders:
                        if placeholder in run.text:
                            # If it's a placeholder we want to keep, check if it has a real value
                            # Real values would not be exactly the placeholder text
                            if run.text.strip() == placeholder:
                                # This is still a placeholder, hide it
                                run.font.color.rgb = RGBColor(255, 255, 255)  # White color
                                print(f"DEBUG: Hiding curly placeholder in run: '{run.text}'")
                            else:
                                # This has a real value, don't hide it
                                should_hide = False
                                print(f"DEBUG: Keeping visible value in run: '{run.text}'")
                            break
                    
                    # For other curly placeholders not in keep_placeholders, hide them
                    if should_hide and not any(placeholder in run.text for placeholder in keep_placeholders):
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White color
                        print(f"DEBUG: Hiding other curly placeholder in run: '{run.text}'")
                
                # Remove unwanted phrase
                run.text = representation_pattern.sub('', run.text)

        # Clean paragraphs
        for paragraph in doc.paragraphs:
            clean_runs(paragraph.runs)
        # Clean tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        clean_runs(paragraph.runs)

    def _get_template_from_r2(self, template_id: int) -> bytes:
        """
        Get template from R2 storage for non-contentious documents
        """
        template = TplTemplate.objects.get(pktemplate=template_id)
        
        # Use cached S3 client
        s3 = get_s3_client()
        
        # Template path in simplified structure
        object_key = f"rodriguez-zea/plantillas/{template.filename}"
        
        try:
            s3_response = s3.get_object(
                Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'), 
                Key=object_key
            )
            return s3_response['Body'].read()
        except Exception as e:
            print(f"Error getting template from R2: {e}")
            raise FileNotFoundError(f"Template not found: {template.filename}")

    def get_document_data(self, num_kardex: str, idtipoacto: str) -> Dict[str, Any]:
        """
        Get all document data for non-contentious documents
        """
        try:
            # Get the year from kardex
            arr_kardex = num_kardex.split('-')
            anio_kardex = arr_kardex[1] if len(arr_kardex) > 1 else datetime.now().year
            
            # Get main document data
            document_data = self._get_document_data(num_kardex, anio_kardex)
            
            # Get contractors data (transferors and acquirers)
            contractors_data = self._get_contractors_data(num_kardex, idtipoacto)
            
            # Get payment data
            payment_data = self._get_payment_data(num_kardex)
            
            # Get escrituración data
            escrituracion_data = self._get_escrituracion_data(num_kardex)
            
            # Merge all data
            all_data = {}
            all_data.update(document_data)
            all_data.update(contractors_data)
            all_data.update(payment_data)
            all_data.update(escrituracion_data)
            
            # Debug: Print some key placeholders
            print(f"DEBUG: Generated data for kardex {num_kardex}:")
            print(f"DEBUG: C_NOM_1: {all_data.get('C_NOM_1', 'NOT FOUND')}")
            print(f"DEBUG: P_NOM_1: {all_data.get('P_NOM_1', 'NOT FOUND')}")
            print(f"DEBUG: C_NOM: {all_data.get('C_NOM', 'NOT FOUND')}")
            print(f"DEBUG: P_NOM: {all_data.get('P_NOM', 'NOT FOUND')}")
            print(f"DEBUG: C_DOMICILIO: {all_data.get('C_DOMICILIO', 'NOT FOUND')}")
            print(f"DEBUG: P_DOMICILIO: {all_data.get('P_DOMICILIO', 'NOT FOUND')}")
            print(f"DEBUG: C_IDE: {all_data.get('C_IDE', 'NOT FOUND')}")
            print(f"DEBUG: P_IDE: {all_data.get('P_IDE', 'NOT FOUND')}")
            print(f"DEBUG: All C_ keys: {[k for k in all_data.keys() if k.startswith('C_')]}")
            print(f"DEBUG: All P_ keys: {[k for k in all_data.keys() if k.startswith('P_')]}")
            print(f"DEBUG: Total placeholders: {len(all_data)}")
            
            return all_data
            
        except Exception as e:
            print(f"Error getting document data: {e}")
            raise

    def _get_document_data(self, num_kardex: str, anio_kardex: str) -> Dict[str, str]:
        """
        Get basic document information
        """
        kardex = Kardex.objects.filter(kardex=num_kardex).first()
        if not kardex:
            raise ValueError(f"Kardex {num_kardex} not found")
        
        # Get user information
        usuario = kardex.responsable_new or ''
        usuario_dni = ''
        if kardex.idusuario:
            # idusuario is an integer ID, not a user object
            user = Usuarios.objects.filter(idusuario=kardex.idusuario).first()
            if user:
                usuario_dni = user.dni or ''
        
        # Get abogado information
        abogado = ''
        matricula = ''
        if kardex.idabogado:
            from notaria.models import TbAbogado
            abogado_obj = TbAbogado.objects.filter(idabogado=kardex.idabogado).first()
            if abogado_obj:
                abogado = abogado_obj.razonsocial or ''
                matricula = abogado_obj.matricula or ''
        
        numero_escritura = kardex.numescritura or ''
        fecha_escritura = kardex.fechaescritura or datetime.now()
        numero_minuta = kardex.numminuta or ''
        folioini = kardex.folioini or ''
        foliofin = kardex.foliofin or ''
        papelini = kardex.papelini or ''
        papelfin = kardex.papelfin or ''
        
        return {
            'K': num_kardex,
            'NRO_ESC': f"{numero_escritura}({self.letras.number_to_letters(numero_escritura)})" if numero_escritura else '{{NRO_ESC}}',
            'NUM_REG': '1',
            'FEC_LET': self.letras.date_to_letters(fecha_escritura) if fecha_escritura else '',
            'F_IMPRESION': self.letras.date_to_letters(fecha_escritura) if fecha_escritura else '{{F_IMPRESION}}',
            'USUARIO': usuario,
            'USUARIO_DNI': usuario_dni,
            'NRO_MIN': numero_minuta or '{{NRO_MIN}}',
            'COMPROBANTE': ' ',
            'O_S': ' ',
            'ORDEN_SERVICIO': ' ',
            'F': self.letras.date_to_letters(fecha_escritura) if fecha_escritura else '{{F}}',
            'DESCRIPCION_SELLO': f"{abogado} CAP. {matricula}",
            'FI': folioini or '{{FI}}',
            'FF': foliofin or '{{FF}}',
            'S_IN': papelini or '{{S_IN}}',
            'S_FN': papelfin or '{{S_FN}}',
            'FECHA_ACT': self.letras.date_to_letters(fecha_escritura) if fecha_escritura else '{{FECHA_ACT}}',
        }

    def _get_contractors_data(self, num_kardex: str, idtipoacto: str) -> Dict[str, str]:
        """
        Get contractors (transferors and acquirers) information for non-contentious documents
        """
        # Get all contratantes for this kardex with optimized queries
        contratantes = Contratantesxacto.objects.filter(kardex=num_kardex)
        
        # Pre-fetch all related data to avoid N+1 queries
        contratante_ids = [cxa.idcontratante for cxa in contratantes]
        clientes = {c.idcontratante: c for c in Cliente2.objects.filter(idcontratante__in=contratante_ids)}
        condicion_ids = [cxa.idcondicion for cxa in contratantes if cxa.idcondicion]
        condiciones = {c.idcondicion: c for c in Actocondicion.objects.filter(idcondicion__in=condicion_ids)}
        
        # Pre-fetch nationality and ubigeo data
        nacionalidad_ids = [c.nacionalidad for c in clientes.values() if c.nacionalidad]
        nacionalidades = {n.idnacionalidad: n for n in Nacionalidades.objects.filter(idnacionalidad__in=nacionalidad_ids)}
        
        ubigeo_ids = [c.idubigeo for c in clientes.values() if c.idubigeo]
        ubigeos = {u.coddis: u for u in Ubigeo.objects.filter(coddis__in=ubigeo_ids)}
        
        transferors = []
        acquirers = []
        companies = []
        
        for cxa in contratantes:
            cliente = clientes.get(cxa.idcontratante)
            if not cliente:
                continue
                
            condicion = condiciones.get(cxa.idcondicion)
            condicion_str = condicion.condicion if condicion else ''
            
            # Build full name
            if cliente.tipper == 'J':  # Juridica
                nombres = cliente.razonsocial or ''
                companies.append({
                    'nombres': nombres,
                    'condicion_str': condicion_str,
                    'idcontratante': cxa.idcontratante,
                    'razonsocial': nombres,
                    'numdoc_empresa': cliente.numdoc or '',
                    'domfiscal': cliente.domfiscal or '',
                    'numpartida': cliente.numpartida or '',
                })
            else:  # Natural
                nombres = f"{cliente.prinom or ''} {cliente.segnom or ''} {cliente.apepat or ''} {cliente.apemat or ''}".strip()
                
                # Get nationality and civil status
                nacionalidad = ''
                if cliente.nacionalidad:
                    nac_obj = nacionalidades.get(cliente.nacionalidad)
                    if nac_obj:
                        nacionalidad = nac_obj.descripcion or ''
                
                estado_civil = ''
                if cliente.idestcivil:
                    from ducumentation.constants import CIVIL_STATUS
                    estado_civil = CIVIL_STATUS.get(cliente.idestcivil, {}).get('label', '')
                
                # Get document type
                tipo_documento = ''
                if cliente.idtipdoc:
                    from ducumentation.constants import TIPO_DOCUMENTO
                    tipo_documento = TIPO_DOCUMENTO.get(cliente.idtipdoc, {}).get('destipdoc', '')
                
                # Get ubigeo
                direccion = ''
                if cliente.idubigeo:
                    ubigeo_obj = ubigeos.get(cliente.idubigeo)
                    if ubigeo_obj:
                        direccion = f"{cliente.direccion or ''} DEL DISTRITO DE {ubigeo_obj.nomdis or ''} PROVINCIA DE {ubigeo_obj.nomprov or ''} Y DEPARTAMENTO DE {ubigeo_obj.nomdpto or ''}"
                
                # If no ubigeo or empty address, use a default
                if not direccion.strip():
                    direccion = "DIRECCIÓN NO ESPECIFICADA"
                
                person_data = {
                    'nombres': nombres,
                    'condicion_str': condicion_str,
                    'idcontratante': cxa.idcontratante,
                    'nacionalidad': self.get_nationality_by_gender(nacionalidad, cliente.sexo),
                    'tipoDocumento': tipo_documento,
                    'numeroDocumento': cliente.numdoc or '',
                    'ocupacion': cliente.profesion_plantilla or '',
                    'estadoCivil': self.get_civil_status_by_gender(estado_civil, cliente.sexo),
                    'direccion': direccion,
                    'sexo': cliente.sexo or 'M',
                }
                
                # Classify as transferor or acquirer
                if condicion_str in ['VENDEDOR', 'DONANTE', 'PODERDANTE', 'OTORGANTE', 'REPRESENTANTE', 'ANTICIPANTE', 'ADJUDICANTE', 'USUFRUCTUANTE', 'TRANSFERENTE', 'DEUDOR', 'SOLICITANTE/BENEFICIARIO']:
                    transferors.append(person_data)
                elif condicion_str in ['COMPRADOR', 'APODERADO', 'ANTICIPADO', 'ADJUDICATARIO', 'DONATARIO', 'USUFRUCTUARIO', 'TESTIGO A RUEGO', 'ADQUIRIENTE', 'ACREEDOR', 'CAUSANTE']:
                    acquirers.append(person_data)
        
        # Build contractors data
        contractors_data = {}
        
        # Add transferors
        for idx, t in enumerate(transferors, 1):
            contractors_data[f'P_NOM_{idx}'] = self.clean_commas(t['nombres'] + ', ')
            contractors_data[f'P_NACIONALIDAD_{idx}'] = self.clean_commas(t['nacionalidad'] + ', ')
            contractors_data[f'P_TIP_DOC_{idx}'] = t['tipoDocumento']
            contractors_data[f'P_DOC_{idx}'] = self.get_identification_phrase(t['sexo'], t['tipoDocumento'], t['numeroDocumento'])
            contractors_data[f'P_OCUPACION_{idx}'] = t['ocupacion']
            contractors_data[f'P_ESTADO_CIVIL_{idx}'] = self.clean_commas(t['estadoCivil'] + ', ')
            contractors_data[f'P_DOMICILIO_{idx}'] = 'CON DOMICILIO EN ' + t['direccion']
            contractors_data[f'P_IDE_{idx}'] = t['numeroDocumento'] or ' '
            contractors_data[f'SEXO_P_{idx}'] = t['sexo']
            
            # Add unnumbered versions for first person (like PHP legacy code)
            if idx == 1:
                contractors_data['P_NOM'] = self.clean_commas(t['nombres'] + ', ')
                contractors_data['P_NACIONALIDAD'] = self.clean_commas(t['nacionalidad'] + ', ')
                contractors_data['P_TIP_DOC'] = t['tipoDocumento']
                contractors_data['P_DOC'] = self.get_identification_phrase(t['sexo'], t['tipoDocumento'], t['numeroDocumento'])
                contractors_data['P_OCUPACION'] = t['ocupacion']
                contractors_data['P_ESTADO_CIVIL'] = self.clean_commas(t['estadoCivil'] + ', ')
                contractors_data['P_DOMICILIO'] = 'CON DOMICILIO EN ' + t['direccion']
                contractors_data['P_IDE'] = t['numeroDocumento'] or ' '
                contractors_data['SEXO_P'] = t['sexo']
        
        # Add acquirers
        for idx, c in enumerate(acquirers, 1):
            contractors_data[f'C_NOM_{idx}'] = self.clean_commas(c['nombres'] + ', ')
            contractors_data[f'C_NACIONALIDAD_{idx}'] = self.clean_commas(c['nacionalidad'] + ', ')
            contractors_data[f'C_TIP_DOC_{idx}'] = c['tipoDocumento']
            contractors_data[f'C_DOC_{idx}'] = self.get_identification_phrase(c['sexo'], c['tipoDocumento'], c['numeroDocumento'])
            contractors_data[f'C_OCUPACION_{idx}'] = c['ocupacion']
            contractors_data[f'C_ESTADO_CIVIL_{idx}'] = self.clean_commas(c['estadoCivil'] + ', ')
            contractors_data[f'C_DOMICILIO_{idx}'] = 'CON DOMICILIO EN ' + c['direccion']
            contractors_data[f'C_IDE_{idx}'] = c['numeroDocumento'] or ' '
            contractors_data[f'SEXO_C_{idx}'] = c['sexo']
            
            # Add unnumbered versions for first person (like PHP legacy code)
            if idx == 1:
                contractors_data['C_NOM'] = self.clean_commas(c['nombres'] + ', ')
                contractors_data['C_NACIONALIDAD'] = self.clean_commas(c['nacionalidad'] + ', ')
                contractors_data['C_TIP_DOC'] = c['tipoDocumento']
                contractors_data['C_DOC'] = self.get_identification_phrase(c['sexo'], c['tipoDocumento'], c['numeroDocumento'])
                contractors_data['C_OCUPACION'] = c['ocupacion']
                contractors_data['C_ESTADO_CIVIL'] = self.clean_commas(c['estadoCivil'] + ', ')
                contractors_data['C_DOMICILIO'] = 'CON DOMICILIO EN ' + c['direccion']
                contractors_data['C_IDE'] = c['numeroDocumento'] or ' '
                contractors_data['SEXO_C'] = c['sexo']
        
        # Add companies
        for idx, comp in enumerate(companies, 1):
            contractors_data[f'NOMBRE_EMPRESA_{idx}'] = comp['razonsocial']
            contractors_data[f'INS_EMPRESA_{idx}'] = ' '
            contractors_data[f'RUC_{idx}'] = f', CON RUC N° {comp["numdoc_empresa"]}, '
            contractors_data[f'DOMICILIO_EMPRESA_{idx}'] = f'CON DOMICILIO EN {comp["domfiscal"]}'
        
        # Fill empty placeholders
        for idx in range(len(transferors) + 1, 11):
            contractors_data[f'P_NOM_{idx}'] = f'[E.P_NOM_{idx}]'
            contractors_data[f'P_NACIONALIDAD_{idx}'] = f'[E.P_NACIONALIDAD_{idx}]'
            contractors_data[f'P_TIP_DOC_{idx}'] = f'[E.P_TIP_DOC_{idx}]'
            contractors_data[f'P_DOC_{idx}'] = f'[E.P_DOC_{idx}]'
            contractors_data[f'P_OCUPACION_{idx}'] = f'[E.P_OCUPACION_{idx}]'
            contractors_data[f'P_ESTADO_CIVIL_{idx}'] = f'[E.P_ESTADO_CIVIL_{idx}]'
            contractors_data[f'P_DOMICILIO_{idx}'] = f'[E.P_DOMICILIO_{idx}]'
            contractors_data[f'P_IDE_{idx}'] = f'[E.P_IDE_{idx}]'
            contractors_data[f'SEXO_P_{idx}'] = f'[E.SEXO_P_{idx}]'

        for idx in range(len(acquirers) + 1, 11):
            contractors_data[f'C_NOM_{idx}'] = f'[E.C_NOM_{idx}]'
            contractors_data[f'C_NACIONALIDAD_{idx}'] = f'[E.C_NACIONALIDAD_{idx}]'
            contractors_data[f'C_TIP_DOC_{idx}'] = f'[E.C_TIP_DOC_{idx}]'
            contractors_data[f'C_DOC_{idx}'] = f'[E.C_DOC_{idx}]'
            contractors_data[f'C_OCUPACION_{idx}'] = f'[E.C_OCUPACION_{idx}]'
            contractors_data[f'C_ESTADO_CIVIL_{idx}'] = f'[E.C_ESTADO_CIVIL_{idx}]'
            contractors_data[f'C_DOMICILIO_{idx}'] = f'[E.C_DOMICILIO_{idx}]'
            contractors_data[f'C_IDE_{idx}'] = f'[E.C_IDE_{idx}]'
            contractors_data[f'SEXO_C_{idx}'] = f'[E.C_SEXO_{idx}]'

        for idx in range(len(companies) + 1, 6):
            contractors_data[f'NOMBRE_EMPRESA_{idx}'] = f'[E.NOMBRE_EMPRESA_{idx}]'
            contractors_data[f'INS_EMPRESA_{idx}'] = f'[E.INS_EMPRESA_{idx}]'
            contractors_data[f'RUC_{idx}'] = f'[E.RUC_{idx}]'
            contractors_data[f'DOMICILIO_EMPRESA_{idx}'] = f'[E.DOMICILIO_EMPRESA_{idx}]'
        
        # Add unnumbered empty placeholders if no data exists
        if len(transferors) == 0:
            contractors_data['P_NOM'] = '[E.P_NOM]'
            contractors_data['P_NACIONALIDAD'] = '[E.P_NACIONALIDAD]'
            contractors_data['P_TIP_DOC'] = '[E.P_TIP_DOC]'
            contractors_data['P_DOC'] = '[E.P_DOC]'
            contractors_data['P_OCUPACION'] = '[E.P_OCUPACION]'
            contractors_data['P_ESTADO_CIVIL'] = '[E.P_ESTADO_CIVIL]'
            contractors_data['P_DOMICILIO'] = '[E.P_DOMICILIO]'
            contractors_data['P_IDE'] = '[E.P_IDE]'
            contractors_data['SEXO_P'] = '[E.SEXO_P]'
        
        if len(acquirers) == 0:
            contractors_data['C_NOM'] = '[E.C_NOM]'
            contractors_data['C_NACIONALIDAD'] = '[E.C_NACIONALIDAD]'
            contractors_data['C_TIP_DOC'] = '[E.C_TIP_DOC]'
            contractors_data['C_DOC'] = '[E.C_DOC]'
            contractors_data['C_OCUPACION'] = '[E.C_OCUPACION]'
            contractors_data['C_ESTADO_CIVIL'] = '[E.C_ESTADO_CIVIL]'
            contractors_data['C_DOMICILIO'] = '[E.C_DOMICILIO]'
            contractors_data['C_IDE'] = '[E.C_IDE]'
            contractors_data['SEXO_C'] = '[E.SEXO_C]'
        
        # Always ensure C_NOM exists (for the TERCERO section)
        if 'C_NOM' not in contractors_data:
            # Use the first acquirer's name or a default
            if acquirers:
                contractors_data['C_NOM'] = self.clean_commas(acquirers[0]['nombres'] + ', ')
            else:
                contractors_data['C_NOM'] = '[E.C_NOM]'

        # Add grammar and articles
        contractors_data.update(self.get_articles_and_grammar(transferors, 'P'))
        contractors_data.update(self.get_articles_and_grammar(acquirers, 'C'))

        return contractors_data

    def _get_payment_data(self, kardex: str) -> Dict[str, str]:
        """
        Get payment information for non-contentious documents
        """
        patrimonial = Patrimonial.objects.filter(kardex=kardex).first()
        
        if not patrimonial:
            return {
                'MONTO': '',
                'MON_VEHI': '',
                'MONTO_LETRAS': '',
                'MONEDA_C': '',
                'SUNAT_MED_PAGO': '',
                'DES_PRE_VEHI': '',
                'EXH_MED_PAGO': '',
                'MED_PAGO': '',
                'FIN_MED_PAGO': '',
                'FORMA_PAGO': '',
            }
        
        precio = patrimonial.importetrans or 0
        moneda = MONEDAS[patrimonial.idmon]['desmon'] if patrimonial else '' 
        simbolo_moneda = MONEDAS[patrimonial.idmon]['simbolo'] if patrimonial else ''
        forma_pago = FORMAS_PAGO[patrimonial.fpago]['descripcion'] if patrimonial else ''
        
        # Get medio de pago
        medio_pago_obj = None
        sunat_medio_pago = ''
        if patrimonial.fpago:
            # For now, use a default value since FpagoUif is not available
            sunat_medio_pago = "008"  # Default to cash payment
        
        # Payment method logic (similar to PHP switch)
        if sunat_medio_pago == "008":
            medio_pago = f'EL COMPRADOR DECLARA QUE HA PAGADO EL PRECIO DEL VEHICULO EN DINERO EN EFECTIVO. NO HABIENDO UTILIZADO NINGÚN MEDIO DE PAGO ESTABLECIDO EN LA LEY Nº 28194, PORQUE EL MONTO TOTAL NO ES IGUAL NI SUPERA LOS S/ 3,500.00 O US$ 1,000.00. EL TIPO Y CÓDIGO DEL MEDIO EMPLEADO ES: "EFECTIVO POR OPERACIONES EN LAS QUE NO EXISTE OBLIGACIÓN DE UTILIZAR MEDIOS DE PAGO-008". INAPLICABLE LA LEY 30730 POR SER EL PAGO DEL PRECIO INFERIOR A 3 UIT.'
            exhibio_medio_pago = 'SE DEJA CONSTANCIA QUE PARA LA REALIZACIÓN DEL PRESENTE ACTO, LAS PARTES NO ME HAN EXHIBIDO NINGÚN MEDIO DE PAGO. DOY FE.'
            fin_medio_pago = 'EN DINERO EN EFECTIVO'
            forma_pago = 'AL CONTADO CON DINERO EN EFECTIVO'
        elif sunat_medio_pago == "009":
            medio_pago = f'EL COMPRADOR DECLARA QUE HA PAGADO EL PRECIO DEL VEHICULO EN DINERO EN EFECTIVO Y CON ANTERIORIDAD A LA CELEBRACION DE LA PRESENTE ACTA DE TRANSFERENCIA. NO HABIENDO UTILIZADO NINGÚN MEDIO DE PAGO ESTABLECIDO EN LA LEY Nº 28194, EL TIPO Y CÓDIGO DEL MEDIO EMPLEADO ES: "EFECTIVO POR OPERACIONES EN LAS QUE NO EXISTE OBLIGACIÓN DE UTILIZAR MEDIOS DE PAGO-009". INAPLICABLE LA LEY 30730 POR SER EL PAGO DEL PRECIO INFERIOR A 3 UIT.'
            exhibio_medio_pago = 'SE DEJA CONSTANCIA QUE PARA LA REALIZACIÓN DEL PRESENTE ACTO, LAS PARTES NO ME HAN EXHIBIDO NINGÚN MEDIO DE PAGO. DOY FE.'
            fin_medio_pago = 'EN DINERO EN EFECTIVO'
            forma_pago = 'AL CONTADO CON DINERO EN EFECTIVO'
        else:
            # Default case
            medio_pago = 'EL COMPRADOR DECLARA QUE HA PAGADO EL PRECIO DEL VEHICULO CON CHEQUE DEL BANCO DE CREDITO DEL PERÚ N° 1111111 111111 1111, GIRADO POR: YYYYYYYYY A FAVOR DE: XXXXXXXXX POR LA SUMA DE S/ 15,000.00, JULIACA 16/08/2018 EL TIPO Y CÓDIGO DEL MEDIO EMPLEADO ES: "CHEQUE -001" '
            exhibio_medio_pago = 'EN APLICACIÓN DE LA LEY 30730, SE DEJA CONSTANCIA QUE PARA LA REALIZACIÓN DEL PRESENTE ACTO, LAS PARTES ME HAN EXHIBIDO EL SIGUIENTE MEDIO DE PAGO: ……… CHEQUE DEL BANCO DE CREDITO DEL PERÚ N° 1111111 111111 1111, GIRADO POR: YYYYYYYYY A FAVOR DE: XXXXXXXXX POR LA SUMA DE S/ 15,000.00, JULIACA 16/08/2018. DOY FE.'
            fin_medio_pago = 'EN DINERO EN EFECTIVO'
            forma_pago = 'AL CONTADO CON DINERO EN EFECTIVO'
        
        return {
            'MONTO': precio,
            'MON_VEHI': moneda,
            'MONTO_LETRAS': self.letras.money_to_letters(moneda, Decimal(precio)),
            'MONEDA_C': simbolo_moneda + ' ',
            'SUNAT_MED_PAGO': sunat_medio_pago,
            'DES_PRE_VEHI': self.letras.money_to_letters(moneda, Decimal(precio)),
            'EXH_MED_PAGO': exhibio_medio_pago,
            'MED_PAGO': medio_pago,
            'FIN_MED_PAGO': fin_medio_pago,
            'FORMA_PAGO': forma_pago,
            'C_INICIO_MP': '',
            'TIPO_PAGO_E': '',
            'TIPO_PAGO_C': '',
            'MONTO_MP': '',
            'CONSTANCIA': '',
            'DETALLE_MP': '',
            'FORMA_PAGO_S': '',
            'MONEDA_C_MP': '',
            'MEDIO_PAGO_C': '',
            'MP_MEDIO_PAGO': '',
            'MP_COMPLETO': '',
            'USO': '',
        }

    def _get_escrituracion_data(self, num_kardex: str) -> Dict[str, str]:
        """
        Get escrituración (folios, papeles) data from kardex.
        """
        kardex = Kardex.objects.filter(kardex=num_kardex).first()
        if not kardex:
            return {
                'FI': '{{FI}}',
                'FF': '{{FF}}',
                'S_IN': '{{S_IN}}',
                'S_FN': '{{S_FN}}',
            }
        
        folioini = kardex.folioini or ''
        foliofin = kardex.foliofin or ''
        papelini = kardex.papelini or ''
        papelfin = kardex.papelfin or ''
        
        return {
            'FI': folioini if folioini else '{{FI}}',
            'FF': foliofin if foliofin else '{{FF}}',
            'S_IN': papelini if papelini else '{{S_IN}}',
            'S_FN': papelfin if papelfin else '{{S_FN}}',
        }

    def _process_document(self, template_bytes: bytes, data: Dict[str, str]) -> Document:
        """
        Process the document using the same approach as other services.
        """
        doc = Document(io.BytesIO(template_bytes))
        
        # Process each paragraph
        for paragraph in doc.paragraphs:
            self._replace_placeholders_in_paragraph(paragraph, data)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_placeholders_in_paragraph(paragraph, data)
        
        return doc

    def _replace_placeholders_in_paragraph(self, paragraph, data: Dict[str, str]):
        """
        Replace placeholders in a paragraph, handling cases where placeholders span multiple runs
        """
        # Get the full text of the paragraph
        full_text = paragraph.text
        modified_text = full_text
        
        # Debug: Check if this paragraph contains placeholders
        if '{{' in full_text and '}}' in full_text:
            print(f"DEBUG: Found placeholders in paragraph: '{full_text[:100]}...'")
        
        # Replace all placeholders in the text
        for key, value in data.items():
            placeholder = f'{{{{{key}}}}}'
            if placeholder in modified_text:
                print(f"DEBUG: Replacing {placeholder} with '{value}'")
                modified_text = modified_text.replace(placeholder, str(value))
        
        # If the text changed, rebuild the paragraph with proper coloring
        if modified_text != full_text:
            print(f"DEBUG: Text changed from '{full_text[:50]}...' to '{modified_text[:50]}...'")
            # Clear all runs
            paragraph.clear()
            
            # Split the text by all possible placeholders to identify what was replaced
            import re
            placeholder_pattern = re.compile(r'\{\{[A-Z0-9_]+\}\}')
            
            # Find all placeholders in the original text
            original_placeholders = placeholder_pattern.findall(full_text)
            
            # Build the new text with coloring
            current_pos = 0
            for match in placeholder_pattern.finditer(full_text):
                placeholder = match.group()
                start_pos = match.start()
                end_pos = match.end()
                
                # Add text before this placeholder
                if start_pos > current_pos:
                    text_before = full_text[current_pos:start_pos]
                    if text_before:
                        paragraph.add_run(text_before)
                
                # Find the replacement value
                replacement = None
                for key, value in data.items():
                    if f'{{{{{key}}}}}' == placeholder:
                        replacement = value
                        break
                
                if replacement:
                    # Add replacement in red
                    red_run = paragraph.add_run(str(replacement))
                    red_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                else:
                    # Keep original placeholder if no replacement found
                    paragraph.add_run(placeholder)
                
                current_pos = end_pos
            
            # Add remaining text after the last placeholder
            if current_pos < len(full_text):
                remaining_text = full_text[current_pos:]
                if remaining_text:
                    paragraph.add_run(remaining_text)
        else:
            print(f"DEBUG: No changes detected in paragraph: '{full_text[:50]}...'")
            # If no changes in full text, try replacing in individual runs
            for run in paragraph.runs:
                run_text = run.text
                modified_run_text = run_text
                
                for key, value in data.items():
                    placeholder = f'{{{{{key}}}}}'
                    if placeholder in modified_run_text:
                        print(f"DEBUG: Replacing in run: {placeholder} -> {value}")
                        modified_run_text = modified_run_text.replace(placeholder, str(value))
                
                if modified_run_text != run_text:
                    run.text = modified_run_text
                    # Color the entire run red since we can't easily identify just the replacement
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red color

    def _create_response(self, doc: Document, filename: str, kardex: str, mode: str = "download") -> HttpResponse:
        """
        Create HTTP response for the generated document
        """
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        if mode == "download":
            response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
        else:
            return HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document') 

    def clean_commas(self, text: str) -> str:
        """
        Clean up repeated commas and normalize comma usage
        """
        import re
        # Remove multiple consecutive commas
        text = re.sub(r',{2,}', ',', text)
        # Remove comma followed by space and comma
        text = re.sub(r',\s*,', ',', text)
        # Remove leading/trailing commas
        text = re.sub(r'^,+|,+$', '', text)
        # Remove comma before period
        text = re.sub(r',\s*\.', '.', text)
        # Remove comma before semicolon
        text = re.sub(r',\s*;', ';', text)
        # Normalize spaces around commas
        text = re.sub(r'\s*,\s*', ', ', text)
        return text.strip()

    def _get_contractors_data(self, num_kardex: str, idtipoacto: str) -> Dict[str, str]:
        """
        Get contractors (transferors and acquirers) information for non-contentious documents
        """
        # Get all contratantes for this kardex with optimized queries
        contratantes = Contratantesxacto.objects.filter(kardex=num_kardex)
        
        # Pre-fetch all related data to avoid N+1 queries
        contratante_ids = [cxa.idcontratante for cxa in contratantes]
        clientes = {c.idcontratante: c for c in Cliente2.objects.filter(idcontratante__in=contratante_ids)}
        condicion_ids = [cxa.idcondicion for cxa in contratantes if cxa.idcondicion]
        condiciones = {c.idcondicion: c for c in Actocondicion.objects.filter(idcondicion__in=condicion_ids)}
        
        # Pre-fetch nationality and ubigeo data
        nacionalidad_ids = [c.nacionalidad for c in clientes.values() if c.nacionalidad]
        nacionalidades = {n.idnacionalidad: n for n in Nacionalidades.objects.filter(idnacionalidad__in=nacionalidad_ids)}
        
        ubigeo_ids = [c.idubigeo for c in clientes.values() if c.idubigeo]
        ubigeos = {u.coddis: u for u in Ubigeo.objects.filter(coddis__in=ubigeo_ids)}
        
        transferors = []
        acquirers = []
        companies = []
        
        for cxa in contratantes:
            cliente = clientes.get(cxa.idcontratante)
            if not cliente:
                continue
                
            condicion = condiciones.get(cxa.idcondicion)
            condicion_str = condicion.condicion if condicion else ''
            
            # Build full name
            if cliente.tipper == 'J':  # Juridica
                nombres = cliente.razonsocial or ''
                companies.append({
                    'nombres': nombres,
                    'condicion_str': condicion_str,
                    'idcontratante': cxa.idcontratante,
                    'razonsocial': nombres,
                    'numdoc_empresa': cliente.numdoc or '',
                    'domfiscal': cliente.domfiscal or '',
                    'numpartida': cliente.numpartida or '',
                })
            else:  # Natural
                nombres = f"{cliente.prinom or ''} {cliente.segnom or ''} {cliente.apepat or ''} {cliente.apemat or ''}".strip()
                
                # Get nationality and civil status
                nacionalidad = ''
                if cliente.nacionalidad:
                    nac_obj = nacionalidades.get(cliente.nacionalidad)
                    if nac_obj:
                        nacionalidad = nac_obj.descripcion or ''
                
                estado_civil = ''
                if cliente.idestcivil:
                    from ducumentation.constants import CIVIL_STATUS
                    estado_civil = CIVIL_STATUS.get(cliente.idestcivil, {}).get('label', '')
                
                # Get document type
                tipo_documento = ''
                if cliente.idtipdoc:
                    from ducumentation.constants import TIPO_DOCUMENTO
                    tipo_documento = TIPO_DOCUMENTO.get(cliente.idtipdoc, {}).get('destipdoc', '')
                
                # Get ubigeo
                direccion = ''
                if cliente.idubigeo:
                    ubigeo_obj = ubigeos.get(cliente.idubigeo)
                    if ubigeo_obj:
                        direccion = f"{cliente.direccion or ''} DEL DISTRITO DE {ubigeo_obj.nomdis or ''} PROVINCIA DE {ubigeo_obj.nomprov or ''} Y DEPARTAMENTO DE {ubigeo_obj.nomdpto or ''}"
                
                # If no ubigeo or empty address, use a default
                if not direccion.strip():
                    direccion = "DIRECCIÓN NO ESPECIFICADA"
                
                person_data = {
                    'nombres': nombres,
                    'condicion_str': condicion_str,
                    'idcontratante': cxa.idcontratante,
                    'nacionalidad': self.get_nationality_by_gender(nacionalidad, cliente.sexo),
                    'tipoDocumento': tipo_documento,
                    'numeroDocumento': cliente.numdoc or '',
                    'ocupacion': cliente.profesion_plantilla or '',
                    'estadoCivil': self.get_civil_status_by_gender(estado_civil, cliente.sexo),
                    'direccion': direccion,
                    'sexo': cliente.sexo or 'M',
                }
                
                # Classify as transferor or acquirer
                if condicion_str in ['VENDEDOR', 'DONANTE', 'PODERDANTE', 'OTORGANTE', 'REPRESENTANTE', 'ANTICIPANTE', 'ADJUDICANTE', 'USUFRUCTUANTE', 'TRANSFERENTE', 'DEUDOR', 'SOLICITANTE/BENEFICIARIO']:
                    transferors.append(person_data)
                elif condicion_str in ['COMPRADOR', 'APODERADO', 'ANTICIPADO', 'ADJUDICATARIO', 'DONATARIO', 'USUFRUCTUARIO', 'TESTIGO A RUEGO', 'ADQUIRIENTE', 'ACREEDOR', 'CAUSANTE']:
                    acquirers.append(person_data)
        
        # Build contractors data
        contractors_data = {}
        
        # Add transferors
        for idx, t in enumerate(transferors, 1):
            contractors_data[f'P_NOM_{idx}'] = self.clean_commas(t['nombres'] + ', ')
            contractors_data[f'P_NACIONALIDAD_{idx}'] = self.clean_commas(t['nacionalidad'] + ', ')
            contractors_data[f'P_TIP_DOC_{idx}'] = t['tipoDocumento']
            contractors_data[f'P_DOC_{idx}'] = self.get_identification_phrase(t['sexo'], t['tipoDocumento'], t['numeroDocumento'])
            contractors_data[f'P_OCUPACION_{idx}'] = t['ocupacion']
            contractors_data[f'P_ESTADO_CIVIL_{idx}'] = self.clean_commas(t['estadoCivil'] + ', ')
            contractors_data[f'P_DOMICILIO_{idx}'] = 'CON DOMICILIO EN ' + t['direccion']
            contractors_data[f'P_IDE_{idx}'] = t['numeroDocumento'] or ' '
            contractors_data[f'SEXO_P_{idx}'] = t['sexo']
            
            # Add unnumbered versions for first person (like PHP legacy code)
            if idx == 1:
                contractors_data['P_NOM'] = self.clean_commas(t['nombres'] + ', ')
                contractors_data['P_NACIONALIDAD'] = self.clean_commas(t['nacionalidad'] + ', ')
                contractors_data['P_TIP_DOC'] = t['tipoDocumento']
                contractors_data['P_DOC'] = self.get_identification_phrase(t['sexo'], t['tipoDocumento'], t['numeroDocumento'])
                contractors_data['P_OCUPACION'] = t['ocupacion']
                contractors_data['P_ESTADO_CIVIL'] = self.clean_commas(t['estadoCivil'] + ', ')
                contractors_data['P_DOMICILIO'] = 'CON DOMICILIO EN ' + t['direccion']
                contractors_data['P_IDE'] = t['numeroDocumento'] or ' '
                contractors_data['SEXO_P'] = t['sexo']
        
        # Add acquirers
        for idx, c in enumerate(acquirers, 1):
            contractors_data[f'C_NOM_{idx}'] = self.clean_commas(c['nombres'] + ', ')
            contractors_data[f'C_NACIONALIDAD_{idx}'] = self.clean_commas(c['nacionalidad'] + ', ')
            contractors_data[f'C_TIP_DOC_{idx}'] = c['tipoDocumento']
            contractors_data[f'C_DOC_{idx}'] = self.get_identification_phrase(c['sexo'], c['tipoDocumento'], c['numeroDocumento'])
            contractors_data[f'C_OCUPACION_{idx}'] = c['ocupacion']
            contractors_data[f'C_ESTADO_CIVIL_{idx}'] = self.clean_commas(c['estadoCivil'] + ', ')
            contractors_data[f'C_DOMICILIO_{idx}'] = 'CON DOMICILIO EN ' + c['direccion']
            contractors_data[f'C_IDE_{idx}'] = c['numeroDocumento'] or ' '
            contractors_data[f'SEXO_C_{idx}'] = c['sexo']
            
            # Add unnumbered versions for first person (like PHP legacy code)
            if idx == 1:
                contractors_data['C_NOM'] = self.clean_commas(c['nombres'] + ', ')
                contractors_data['C_NACIONALIDAD'] = self.clean_commas(c['nacionalidad'] + ', ')
                contractors_data['C_TIP_DOC'] = c['tipoDocumento']
                contractors_data['C_DOC'] = self.get_identification_phrase(c['sexo'], c['tipoDocumento'], c['numeroDocumento'])
                contractors_data['C_OCUPACION'] = c['ocupacion']
                contractors_data['C_ESTADO_CIVIL'] = self.clean_commas(c['estadoCivil'] + ', ')
                contractors_data['C_DOMICILIO'] = 'CON DOMICILIO EN ' + c['direccion']
                contractors_data['C_IDE'] = c['numeroDocumento'] or ' '
                contractors_data['SEXO_C'] = c['sexo']
        
        # Add companies
        for idx, comp in enumerate(companies, 1):
            contractors_data[f'NOMBRE_EMPRESA_{idx}'] = comp['razonsocial']
            contractors_data[f'INS_EMPRESA_{idx}'] = ' '
            contractors_data[f'RUC_{idx}'] = f', CON RUC N° {comp["numdoc_empresa"]}, '
            contractors_data[f'DOMICILIO_EMPRESA_{idx}'] = f'CON DOMICILIO EN {comp["domfiscal"]}'
        
        # Fill empty placeholders
        for idx in range(len(transferors) + 1, 11):
            contractors_data[f'P_NOM_{idx}'] = f'[E.P_NOM_{idx}]'
            contractors_data[f'P_NACIONALIDAD_{idx}'] = f'[E.P_NACIONALIDAD_{idx}]'
            contractors_data[f'P_TIP_DOC_{idx}'] = f'[E.P_TIP_DOC_{idx}]'
            contractors_data[f'P_DOC_{idx}'] = f'[E.P_DOC_{idx}]'
            contractors_data[f'P_OCUPACION_{idx}'] = f'[E.P_OCUPACION_{idx}]'
            contractors_data[f'P_ESTADO_CIVIL_{idx}'] = f'[E.P_ESTADO_CIVIL_{idx}]'
            contractors_data[f'P_DOMICILIO_{idx}'] = f'[E.P_DOMICILIO_{idx}]'
            contractors_data[f'P_IDE_{idx}'] = f'[E.P_IDE_{idx}]'
            contractors_data[f'SEXO_P_{idx}'] = f'[E.SEXO_P_{idx}]'

        for idx in range(len(acquirers) + 1, 11):
            contractors_data[f'C_NOM_{idx}'] = f'[E.C_NOM_{idx}]'
            contractors_data[f'C_NACIONALIDAD_{idx}'] = f'[E.C_NACIONALIDAD_{idx}]'
            contractors_data[f'C_TIP_DOC_{idx}'] = f'[E.C_TIP_DOC_{idx}]'
            contractors_data[f'C_DOC_{idx}'] = f'[E.C_DOC_{idx}]'
            contractors_data[f'C_OCUPACION_{idx}'] = f'[E.C_OCUPACION_{idx}]'
            contractors_data[f'C_ESTADO_CIVIL_{idx}'] = f'[E.C_ESTADO_CIVIL_{idx}]'
            contractors_data[f'C_DOMICILIO_{idx}'] = f'[E.C_DOMICILIO_{idx}]'
            contractors_data[f'C_IDE_{idx}'] = f'[E.C_IDE_{idx}]'
            contractors_data[f'SEXO_C_{idx}'] = f'[E.C_SEXO_{idx}]'

        for idx in range(len(companies) + 1, 6):
            contractors_data[f'NOMBRE_EMPRESA_{idx}'] = f'[E.NOMBRE_EMPRESA_{idx}]'
            contractors_data[f'INS_EMPRESA_{idx}'] = f'[E.INS_EMPRESA_{idx}]'
            contractors_data[f'RUC_{idx}'] = f'[E.RUC_{idx}]'
            contractors_data[f'DOMICILIO_EMPRESA_{idx}'] = f'[E.DOMICILIO_EMPRESA_{idx}]'
        
        # Add unnumbered empty placeholders if no data exists
        if len(transferors) == 0:
            contractors_data['P_NOM'] = '[E.P_NOM]'
            contractors_data['P_NACIONALIDAD'] = '[E.P_NACIONALIDAD]'
            contractors_data['P_TIP_DOC'] = '[E.P_TIP_DOC]'
            contractors_data['P_DOC'] = '[E.P_DOC]'
            contractors_data['P_OCUPACION'] = '[E.P_OCUPACION]'
            contractors_data['P_ESTADO_CIVIL'] = '[E.P_ESTADO_CIVIL]'
            contractors_data['P_DOMICILIO'] = '[E.P_DOMICILIO]'
            contractors_data['P_IDE'] = '[E.P_IDE]'
            contractors_data['SEXO_P'] = '[E.SEXO_P]'
        
        if len(acquirers) == 0:
            contractors_data['C_NOM'] = '[E.C_NOM]'
            contractors_data['C_NACIONALIDAD'] = '[E.C_NACIONALIDAD]'
            contractors_data['C_TIP_DOC'] = '[E.C_TIP_DOC]'
            contractors_data['C_DOC'] = '[E.C_DOC]'
            contractors_data['C_OCUPACION'] = '[E.C_OCUPACION]'
            contractors_data['C_ESTADO_CIVIL'] = '[E.C_ESTADO_CIVIL]'
            contractors_data['C_DOMICILIO'] = '[E.C_DOMICILIO]'
            contractors_data['C_IDE'] = '[E.C_IDE]'
            contractors_data['SEXO_C'] = '[E.SEXO_C]'
        
        # Always ensure C_NOM exists (for the TERCERO section)
        if 'C_NOM' not in contractors_data:
            # Use the first acquirer's name or a default
            if acquirers:
                contractors_data['C_NOM'] = self.clean_commas(acquirers[0]['nombres'] + ', ')
            else:
                contractors_data['C_NOM'] = '[E.C_NOM]'

        # Add grammar and articles
        contractors_data.update(self.get_articles_and_grammar(transferors, 'P'))
        contractors_data.update(self.get_articles_and_grammar(acquirers, 'C'))

        return contractors_data

class TestamentoDocumentService:
    """
    Django service to generate testamento documents by replicating the legacy PHP logic.
    """
    
    def __init__(self):
        self.letras = NumberToLetterConverter()

    def generate_testamento_document(self, template_id: int, num_kardex: str, idtipoacto: str, action: str = 'generate', mode: str = "download") -> HttpResponse:
        """
        Main method to generate testamento document.
        """
        try:
            # Step 1: Fetch all data using a raw SQL query mirroring the legacy script
            raw_data = self._fetch_all_data_raw(num_kardex)
            if not raw_data:
                raise ValueError(f"No data found for kardex {num_kardex}")

            # Step 2: Get template from R2
            template_bytes = self._get_template_from_r2(template_id)
            
            # Step 3: Process data into the required format for the template
            document_data = self._get_document_data(raw_data)
            contractors_data = self._get_contractors_data(raw_data)
            
            # Combine all data sources
            final_data = {**document_data, **contractors_data}
            
            # Step 4: Process the docx template
            doc = self._process_document(template_bytes, final_data)
            
            # Step 5: Remove placeholders
            self.remove_unfilled_placeholders(doc)
            
            # Step 6: Upload to R2 (optional, can be disabled if not needed)
            self.create_documento_in_r2(doc, num_kardex)
            
            # Step 7: Create and return the HTTP response
            filename = f"testamento_{num_kardex}.docx"
            return self._create_response(doc, filename, num_kardex, mode)
            
        except Exception as e:
            print(f"ERROR: Failed to generate testamento document: {str(e)}")
            import traceback
            traceback.print_exc()
            return JsonResponse({'error': f'Failed to generate testamento document: {str(e)}'}, status=500)

    def _fetch_all_data_raw(self, num_kardex: str) -> dict:
        """
        Executes a raw SQL query to fetch all data in a single row, mimicking the PHP script.
        """
        query = """
           SELECT
                k.idkardex as id_kardex,
                k.kardex,
                k.numescritura as numero_escritura,
                k.fechaescritura as fecha_escritura,
                k.txa_minuta as registro_escritura,
                CURRENT_DATE() as fecha_generado,
                k.fechaconclusion as fecha_conclusion,
                k.numminuta as numero_minuta,
                k.kardexconexo as kardex_conexo,
                k.folioini as folio_inicial,
                k.foliofin as folio_final,
                k.papelini as papel_inicial,
                k.papelfin as papel_final,
                k.fechaingreso as fecha_ingreso,
                k.responsable_new as usuario,
                abo.razonsocial as abogado,
                abo.matricula as matricula,
                usu.dni as dni_usuario,
                GROUP_CONCAT(ac.condicion SEPARATOR '|') as condiciones,
                GROUP_CONCAT(
                    TRIM(CONCAT_WS(' ', c2.prinom, c2.segnom, c2.apepat, c2.apemat))
                    SEPARATOR '|'
                ) as nombres,
                GROUP_CONCAT(IFNULL(n.descripcion, '') SEPARATOR '|') as nacionalidades,
                GROUP_CONCAT(IFNULL(td.destipdoc, '') SEPARATOR '|') as tipos_documento,
                GROUP_CONCAT(IFNULL(c2.numdoc, '') SEPARATOR '|') as numeros_documento,
                GROUP_CONCAT(IFNULL(c2.profesion_plantilla, '') SEPARATOR '|') as ocupaciones,
                GROUP_CONCAT(IFNULL(tec.desestcivil, '') SEPARATOR '|') as estados_civil,
                GROUP_CONCAT(c2.sexo SEPARATOR '|') as sexos,
                GROUP_CONCAT(IFNULL(c2.direccion, '') SEPARATOR '|') as direcciones,
                GROUP_CONCAT(IFNULL(u.nomdis, '') SEPARATOR '|') as distritos,
                GROUP_CONCAT(IFNULL(u.nomprov, '') SEPARATOR '|') as provincias,
                GROUP_CONCAT(IFNULL(u.nomdpto, '') SEPARATOR '|') as departamentos
            FROM kardex k
            LEFT JOIN contratantesxacto cxa ON k.kardex = cxa.kardex
            LEFT JOIN cliente2 c2 ON cxa.idcontratante = c2.idcontratante
            LEFT JOIN actocondicion ac ON cxa.idcondicion = ac.idcondicion
            LEFT JOIN usuarios usu ON k.idusuario = usu.idusuario
            LEFT JOIN tb_abogado abo ON k.idabogado = abo.idabogado
            LEFT JOIN nacionalidades n ON c2.nacionalidad = n.idnacionalidad
            LEFT JOIN tipodocumento td ON c2.idtipdoc = td.idtipdoc
            LEFT JOIN tipoestacivil tec ON c2.idestcivil = tec.idestcivil
            LEFT JOIN ubigeo u ON c2.idubigeo = u.coddis
            WHERE k.kardex = %s AND c2.tipper = 'N'
            GROUP BY k.idkardex
        """
        with connection.cursor() as cursor:
            cursor.execute(query, [num_kardex])
            desc = cursor.description
            row = cursor.fetchone()
            if not row:
                return None
            return dict(zip([col[0] for col in desc], row))

    def _get_document_data(self, raw_data: dict) -> Dict[str, str]:
        """
        Get basic document information from the raw query data.
        """
        numero_escritura = raw_data.get('numero_escritura') or ''
        fecha_escritura = raw_data.get('fecha_escritura')
        
        return {
            'K': raw_data.get('kardex', ''),
            'NRO_ESC': f"{numero_escritura}({self.letras.number_to_letters(numero_escritura)})" if numero_escritura else '{{NRO_ESC}}',
            'F': self.letras.date_to_letters(fecha_escritura) if fecha_escritura else '{{F}}',
            'FI': raw_data.get('folio_inicial') or '{{FI}}',
            'FF': raw_data.get('folio_final') or '{{FF}}',
            'S_IN': raw_data.get('papel_inicial') or '{{S_IN}}',
            'S_FN': raw_data.get('papel_final') or '{{S_FN}}',
        }

    def _get_contractors_data(self, raw_data: dict) -> dict:
        """
        Process the GROUP_CONCAT fields from the raw query to build contractor data.
        """
        data = {}

        def split_if_not_none(value, separator='|'):
            return value.split(separator) if value else []

        # Extract data from the raw query result
        condiciones = split_if_not_none(raw_data.get('condiciones'))
        nombres = split_if_not_none(raw_data.get('nombres'))
        nacionalidades = split_if_not_none(raw_data.get('nacionalidades'))
        tipos_documento = split_if_not_none(raw_data.get('tipos_documento'))
        numeros_documento = split_if_not_none(raw_data.get('numeros_documento'))
        ocupaciones = split_if_not_none(raw_data.get('ocupaciones'))
        estados_civil = split_if_not_none(raw_data.get('estados_civil'))
        sexos = split_if_not_none(raw_data.get('sexos'))
        direcciones = split_if_not_none(raw_data.get('direcciones'))
        distritos = split_if_not_none(raw_data.get('distritos'))
        provincias = split_if_not_none(raw_data.get('provincias'))
        departamentos = split_if_not_none(raw_data.get('departamentos'))

        # Create a list of participant dictionaries
        participants = []
        for i, cond in enumerate(condiciones):
            participants.append({
                'condicion': cond,
                'nombre_completo': nombres[i] if i < len(nombres) else '',
                'nacionalidad': nacionalidades[i] if i < len(nacionalidades) else '',
                'tipo_doc': tipos_documento[i] if i < len(tipos_documento) else '',
                'num_doc': numeros_documento[i] if i < len(numeros_documento) else '',
                'profesion': ocupaciones[i] if i < len(ocupaciones) else '',
                'estado_civil': estados_civil[i] if i < len(estados_civil) else '',
                'sexo': sexos[i] if i < len(sexos) else 'M',
                'direccion': direcciones[i] if i < len(direcciones) else '',
                'distrito': distritos[i] if i < len(distritos) else '',
                'provincia': provincias[i] if i < len(provincias) else '',
                'departamento': departamentos[i] if i < len(departamentos) else '',
            })
            
        testadores = [p for p in participants if p['condicion'].upper() in ["OTORGANTE", "TESTADOR"]]
        testigos = [p for p in participants if p['condicion'].upper() in ["TESTIGO", "TESTIGO A RUEGO"]]

        print(f"DEBUG: Found {len(testadores)} testadores and {len(testigos)} testigos")
        print(f"DEBUG: All participants: {[(p['condicion'], p['nombre_completo']) for p in participants]}")
        
        # Process Testadores (P)
        if testadores:
            data.update(self.get_articles_and_grammar(testadores, "P"))
            person = testadores[0]
            
            p_nom = person.get('nombre_completo', '')
            p_nac = self.get_nationality_by_gender(person.get('nacionalidad', ''), person.get('sexo'))
            p_doc = self.get_identification_phrase(person.get('sexo'), person.get('tipo_doc'), person.get('num_doc'))
            p_ocup = person.get('profesion', '')
            p_est_civ = self.get_civil_status_by_gender(person.get('estado_civil', ''), person.get('sexo'))

            data['P_NOM'] = f"{p_nom}, "
            data['P_NACIONALIDAD'] = f"{p_nac}, " if p_nac else ""
            data['P_DOC'] = f"{p_doc}, " if p_doc else ""
            data['P_OCUPACION'] = f"{p_ocup}" if p_ocup else ""
            data['P_ESTADO_CIVIL'] = f", QUIEN DECLARA SER {p_est_civ}, " if p_est_civ else ""

            domicilio = ""
            if person.get('direccion'):
                domicilio = f"con domicilio en {person.get('direccion')}"
                if person.get('distrito'):
                    domicilio += f" del distrito de {person.get('distrito')} provincia de {person.get('provincia')} y departamento de {person.get('departamento')}"
            data['P_DOMICILIO'] = domicilio
            data['P_DOC_LETRAS'] = ''
            data['P_IDE'] = ''
        
        # Process Testigos (C)
        if testigos:
            data.update(self.get_articles_and_grammar(testigos, "C"))
            for i, person in enumerate(testigos):
                if i >= 2: break
                suffix = '' if i == 0 else '_2'
                
                c_nom = person.get('nombre_completo', '')
                c_nac = self.get_nationality_by_gender(person.get('nacionalidad', ''), person.get('sexo'))
                c_doc = self.get_identification_phrase(person.get('sexo'), person.get('tipo_doc'), person.get('num_doc'))
                c_ocup = person.get('profesion', '')
                c_est_civ = self.get_civil_status_by_gender(person.get('estado_civil', ''), person.get('sexo'))

                data[f'C_NOM{suffix}'] = f"{c_nom}, "
                data[f'C_NACIONALIDAD{suffix}'] = f"{c_nac}, " if c_nac else ""
                data[f'C_DOC{suffix}'] = f"{c_doc}, " if c_doc else ""
                data[f'C_OCUPACION{suffix}'] = f"{c_ocup}" if c_ocup else ""
                data[f'C_ESTADO_CIVIL{suffix}'] = f", QUIEN DECLARA SER {c_est_civ}, " if c_est_civ else ""
                
                domicilio = ""
                if person.get('direccion'):
                    domicilio = f"con domicilio en {person.get('direccion')}"
                    if person.get('distrito'):
                        domicilio += f" del distrito de {person.get('distrito')} provincia de {person.get('provincia')} y departamento de {person.get('departamento')}"
                data[f'C_DOMICILIO{suffix}'] = domicilio
                data[f'C_DOC_LETRAS{suffix}'] = ''
                data[f'C_IDE{suffix}'] = ''

        # Clean up unused placeholders and set connectives
        self._cleanup_placeholders(data, len(testadores), len(testigos))
        return data

    def _cleanup_placeholders(self, data, num_testadores, num_testigos):
        """Fills unused placeholders with empty strings."""
        if num_testadores == 0:
            for key in ['P_NOM', 'P_NACIONALIDAD', 'P_DOC', 'P_OCUPACION', 'P_ESTADO_CIVIL', 'P_DOMICILIO', 'P_DOC_LETRAS', 'P_IDE', 'EL_P', 'INICIO_P', 'OR_P']:
                data[key] = ''
        if num_testigos < 2:
            for key in ['C_NOM_2', 'C_NACIONALIDAD_2', 'C_DOC_2', 'C_OCUPACION_2', 'C_ESTADO_CIVIL_2', 'C_DOMICILIO_2', 'C_DOC_LETRAS_2', 'C_IDE_2']:
                data[key] = ''
        if num_testigos < 1:
            for key in ['C_NOM', 'C_NACIONALIDAD', 'C_DOC', 'C_OCUPACION', 'C_ESTADO_CIVIL', 'C_DOMICILIO', 'C_DOC_LETRAS', 'C_IDE', 'EL_C', 'INICIO_C', 'C_AMBOS']:
                data[key] = ''
        if num_testigos > 1:
            data['Y_C'] = 'y'
            data['Y_CON_C'] = 'y'
        else:
            data['Y_C'] = ''
            data['Y_CON_C'] = ''
    
    def _get_payment_data(self, num_kardex: str) -> Dict[str, str]:
        # This service does not seem to have payment data according to the PHP legacy code
        return {}

    def _get_template_from_r2(self, template_id: int) -> bytes:
        """
        Get template from R2 storage - same as VehicleTransferDocumentService
        """
        template = TplTemplate.objects.get(pktemplate=template_id)
        s3 = get_s3_client()
        
        # Template path in simplified structure
        object_key = f"rodriguez-zea/plantillas/{template.filename}"
        
        try:
            response = s3.get_object(Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'), Key=object_key)
            return response['Body'].read()
        except Exception as e:
            print(f"Error downloading template from R2: {e}")
            raise

    def _process_document(self, template_bytes: bytes, data: Dict[str, str]) -> Document:
        """
        Process the document template with data - SAME AS VehicleTransferDocumentService
        """
        buffer = io.BytesIO(template_bytes)
        doc = DocxTemplate(buffer)
        doc.render(data)
        return doc

    def remove_unfilled_placeholders(self, doc):
        """
        Remove all [E.SOMETHING] placeholders and hide {{SOMETHING}} placeholders from user.
        SAME AS VehicleTransferDocumentService
        """
        import re
        placeholder_pattern = re.compile(r'\[E\.[A-Z0-9_]+\]')
        curly_placeholder_pattern = re.compile(r'\{\{[A-Z0-9_]+\}\}')
        representation_pattern = re.compile(
            r'EN REPRESENTACION DE\s*Y[\s.·…‥⋯⋮⋱⋰⋯—–-]*', re.IGNORECASE
        )

        # List of placeholders to keep and hide
        keep_placeholders = ['{{NRO_ESC}}', '{{FI}}', '{{FF}}', '{{S_IN}}', '{{S_FN}}', '{{FECHA_ACT}}']

        def clean_runs(runs):
            # Remove [E.SOMETHING] placeholders and hide {{SOMETHING}} placeholders
            for run in runs:
                # Remove all [E.SOMETHING] placeholders
                run.text = placeholder_pattern.sub('', run.text)
                
                # Hide {{SOMETHING}} placeholders by making them white
                # BUT only if they are actually placeholders (not real values)
                if curly_placeholder_pattern.search(run.text):
                    # Check if this is actually a placeholder or a real value
                    should_hide = True
                    for placeholder in keep_placeholders:
                        if placeholder in run.text:
                            # If it's a placeholder we want to keep, check if it has a real value
                            # Real values would not be exactly the placeholder text
                            if run.text.strip() == placeholder:
                                # This is still a placeholder, hide it
                                run.font.color.rgb = RGBColor(255, 255, 255)  # White color
                                print(f"DEBUG: Hiding curly placeholder in run: '{run.text}'")
                            else:
                                # This has a real value, don't hide it
                                should_hide = False
                                print(f"DEBUG: Keeping visible value in run: '{run.text}'")
                            break
                    
                    # For other curly placeholders not in keep_placeholders, hide them
                    if should_hide and not any(placeholder in run.text for placeholder in keep_placeholders):
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White color
                        print(f"DEBUG: Hiding other curly placeholder in run: '{run.text}'")
                
                # Remove unwanted phrase
                run.text = representation_pattern.sub('', run.text)

        # Clean paragraphs
        for paragraph in doc.paragraphs:
            clean_runs(paragraph.runs)
        # Clean tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        clean_runs(paragraph.runs)

    def create_documento_in_r2(self, doc, kardex):
        """
        Upload the generated document to R2 - SAME AS VehicleTransferDocumentService
        """
        try:
            from io import BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            object_key = f"rodriguez-zea/documentos/__PROY__{kardex}.docx"
            s3 = get_s3_client()
            s3.upload_fileobj(
                buffer,
                os.environ.get('CLOUDFLARE_R2_BUCKET'),
                object_key
            )
            return True
        except Exception as e:
            print(f"Error uploading testamento document to R2: {e}")
            return False

    def _create_response(self, doc, filename: str, kardex: str, mode: str = "download") -> HttpResponse:
        """
        Create HTTP response with the document - SAME AS VehicleTransferDocumentService
        """
        # Add the custom property
        custom_props = CustomProperties(doc)
        custom_props['documentoGeneradoId'] = kardex

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        if mode == "open":
            # Production mode: Return JSON with file info for Word opening
            response = JsonResponse({
                'status': 'success',
                'mode': 'open',
                'filename': filename,
                'kardex': kardex,
                'message': 'Document generated and ready to open in Word'
            })
            response['Access-Control-Allow-Origin'] = '*'
            return response
        else:
            # Testing mode: Download the document
            response = HttpResponse(
                buffer.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'inline; filename="{filename}"'
            response['Content-Length'] = str(buffer.getbuffer().nbytes)
            response['Access-Control-Allow-Origin'] = '*'
            return response

    def get_identification_phrase(self, gender, doc_type, doc_number):
        if not doc_type or not doc_number:
            return ""
        ident_str = 'IDENTIFICADA' if gender == 'F' else 'IDENTIFICADO'
        return f'{ident_str} CON {doc_type} N° {doc_number}'

    def get_civil_status_by_gender(self, civil_status, gender):
        if not civil_status:
            return ''
        civil_status = civil_status.strip().upper()
        if gender == 'F':
            if civil_status.endswith('O'):
                return civil_status[:-1] + 'A'
        return civil_status

    def get_nationality_by_gender(self, nationality, gender):
        if not nationality:
            return ''
        nationality = nationality.strip().upper()
        
        if gender == 'F':
            if nationality.endswith('O'):
                return nationality[:-1] + 'A'
            elif not nationality.endswith('A'):
                 return nationality + 'A'
        else: # Male
            if nationality.endswith('A'):
                return nationality[:-1] + 'O'
            elif not nationality.endswith('O') and nationality:
                return nationality + 'O'
        
        return nationality

    def get_articles_and_grammar(self, people, role_prefix):
        count = len(people)
        data = {
            f'EL_{role_prefix}': '',
            f'INICIO_{role_prefix}': '',
        }
        if role_prefix == 'P':
            data['OR_P'] = ''
        if role_prefix == 'C':
            data['C_AMBOS'] = ''

        if count == 0:
            return data

        first_person = people[0]
        # The 'people' list now contains dicts, not objects
        all_female = all(p['sexo'] == 'F' for p in people)

        if count > 1:
            data[f'EL_{role_prefix}'] = 'LAS' if all_female else 'LOS'
            data[f'INICIO_{role_prefix}'] = ' SEÑORAS' if all_female else ' SEÑORES'
            if role_prefix == 'P':
                data['OR_P'] = 'ORAS' if all_female else 'ORES'
            if role_prefix == 'C':
                data['C_AMBOS'] = 'AMBAS' if all_female else 'AMBOS'
        else:
            data[f'EL_{role_prefix}'] = 'LA' if first_person['sexo'] == 'F' else 'EL'
            data[f'INICIO_{role_prefix}'] = ' SEÑORA' if first_person['sexo'] == 'F' else ' SEÑOR'
            if role_prefix == 'P':
                data['OR_P'] = 'ORA' if first_person['sexo'] == 'F' else 'OR'
        
        return data


class GarantiasMobiliariasDocumentService:
    """
    Django service to generate garantias mobiliarias documents based on the PHP legacy script
    """
    
    def __init__(self):
        self.letras = NumberToLetterConverter()
    
    def generate_garantias_mobiliarias_document(self, template_id: int, num_kardex: str, idtipoacto: str, action: str = 'generate', mode: str = "download") -> HttpResponse:
        """
        Main method to generate garantias mobiliarias document
        """
        try:
            # Step 1: Fetch all data using raw SQL query mirroring the PHP script
            raw_data = self._consulta_transferencia(num_kardex, idtipoacto, template_id)
            if not raw_data:
                raise ValueError(f"No data found for kardex {num_kardex}")

            # Step 2: Get template from R2
            template_bytes = self._get_template_from_r2(template_id)
            
            # Step 3: Process data into the required format for the template
            document_data = self._get_data_documento(raw_data)
            vehiculos_data = self._get_data_vehiculos(raw_data)
            pagos_data = self._get_data_pagos(raw_data)
            contratantes_data = self._get_data_contratantes(raw_data)
            escrituracion_data = self._get_data_escrituracion(raw_data)
            
            # Step 4: Process contratantes data and add empty placeholders
            contratantes_processed = self._process_contratantes_data(contratantes_data)
            articulos_contratantes = self._get_articulos_contratantes(contratantes_data)
            
            # Step 5: Combine all data sources
            final_data = {}
            final_data.update(document_data)
            final_data.update(vehiculos_data)
            final_data.update(pagos_data)
            final_data.update(contratantes_processed)
            final_data.update(articulos_contratantes)
            final_data.update(escrituracion_data)
            
            # Step 6: Process the docx template
            doc = self._process_document(template_bytes, final_data)
            
            # Step 7: Remove placeholders
            self.remove_unfilled_placeholders(doc)
            
            # Step 8: Upload to R2 (optional)
            self.create_documento_in_r2(doc, num_kardex)
            
            # Step 9: Create and return the HTTP response
            filename = f"garantias_mobiliarias_{num_kardex}.docx"
            return self._create_response(doc, filename, num_kardex, mode)
            
        except Exception as e:
            print(f"ERROR: Failed to generate garantias mobiliarias document: {str(e)}")
            import traceback
            traceback.print_exc()
            return JsonResponse({'error': f'Failed to generate garantias mobiliarias document: {str(e)}'}, status=500)

    def _consulta_transferencia(self, num_kardex: str, idtipoacto: str, template_id: int) -> dict:
        """
        Raw SQL query that mirrors the PHP consulta_transferencia function
        """
        query = """
            SELECT k.idkardex as id_kardex,
                k.kardex,
                k.numescritura as numero_escritura,
                k.fechaescritura as fecha_escritura,
                k.txa_minuta as registro_escritura,
                CURRENT_DATE() as fecha_generado,
                k.fechaconclusion as fecha_conclusion,
                k.numminuta as numero_minuta,
                k.kardexconexo as kardex_conexo,
                k.folioini as folio_inicial, 
                k.foliofin as folio_final, 
                k.papelini as papel_inicial, 
                k.papelfin as papel_final,
                (SELECT desacto FROM tiposdeacto WHERE idtipoacto=%s) as acto,
                (SELECT fileName FROM tpl_template WHERE pkTemplate=%s) as plantilla,
                (SELECT urlTemplate FROM tpl_template WHERE pkTemplate=%s) as url_plantilla,
                k.fechaingreso as fecha_ingreso,
                k.responsable_new as usuario,
                abo.razonsocial as abogado,
                abo.matricula as matricula,
                usu.dni as dni_usuario,
                GROUP_CONCAT(cxa.idcontratante) as id_contratante,
                GROUP_CONCAT(TRIM(CONCAT(IFNULL(c2.prinom, ''), ' ', IFNULL(c2.segnom, ''), ' ',IFNULL(c2.apepat, ''), ' ',IFNULL(c2.apemat, ''),
            IFNULL(c2.razonsocial, '')))) AS nombres,
                GROUP_CONCAT(cxa.uif) as uif,
                GROUP_CONCAT(ac.condicion) as condicion,
                GROUP_CONCAT(IF(n.descripcion IS NULL OR n.descripcion='','EMPRESA',n.descripcion)) as nacionalidad,
                GROUP_CONCAT(td.destipdoc) as tipo_documento,
                GROUP_CONCAT(c2.numdoc) AS numero_documento,
                GROUP_CONCAT(UPPER(c2.profesion_plantilla)) AS ocupacion,
                GROUP_CONCAT(IF(tec.desestcivil IS NULL OR tec.desestcivil='','EMPRESA',tec.desestcivil)) as estado_civil,
                GROUP_CONCAT(IF(c2.tipper='N',c2.direccion,c2.domfiscal)) as direccion,
                GROUP_CONCAT(IFNULL(u.codpto, '')) as codigo_departamento,
                GROUP_CONCAT(IFNULL(u.coddis, '')) as codigo_distrito,
                GROUP_CONCAT(IFNULL(u.codprov, '')) as codigo_provincia,
                GROUP_CONCAT(IFNULL(IF(SUBSTRING_INDEX(c2.ubigeo_plantilla, '/', -1)='',u.nomdis,SUBSTRING_INDEX(c2.ubigeo_plantilla, '/', -1)),(IFNULL(u.nomdis, '')))) AS distrito,
                GROUP_CONCAT(IFNULL(u.nomprov, '')) as provincia,
                GROUP_CONCAT(IFNULL(u.nomdpto, '')) as departamento,
                GROUP_CONCAT(c2.sexo) AS sexo,
                GROUP_CONCAT(c2.tipper) as tipo_persona,
                GROUP_CONCAT(IF(cn.firma = '0', 'NO', 'SI')) AS firma,
                GROUP_CONCAT(cn.firma) as n_firma,
                GROUP_CONCAT(cn.tiporepresentacion) AS tipo_representacion,
                dv.numplaca AS placa, 
                dv.marca AS marca, 
                dv.clase AS clase,
                dv.anofab AS anio,
                dv.numserie AS serie, 
                dv.color AS color,
                dv.motor AS motor, 
                dv.modelo AS modelo, 
                dv.carroceria AS carroceria,
                dv.pregistral as partida,
                dv.fecinsc AS fecha_inscripcion,
                dv.combustible AS combustible,
                UPPER(sr.dessede) AS sede,
                UPPER(sr.num_zona) AS numero_zona,
                pat.importetrans AS precio , 
                pat.idmon AS moneda,
                pat.exhibiomp, 
                pat.idoppago, 
                uif.descripcion AS medio_pago, 
                mon.simbolo as simbolo_moneda,
                mon.desmon as descripcion_moneda,
                mp.desmpagos as descripcion_medio_pago,
                mp.sunat as sunat_medio_pago,
                GROUP_CONCAT(cnr.idcontratanterp) as id_empresa,
                GROUP_CONCAT(TRIM(CONCAT(IFNULL(cr2.prinom, ''), ' ', IFNULL(cr2.segnom, ''), ' ',IFNULL(cr2.apepat, ''), ' ',IFNULL(cr2.apemat, ''),
            IFNULL(cr2.razonsocial, '')))) AS nombre_empresa,
                GROUP_CONCAT(cr2.tipper) as tipo_persona_empresa,
                GROUP_CONCAT(acr.condicion) as condicion_empresa,
                GROUP_CONCAT(tdr.destipdoc) as tipo_documento_empresa,
                GROUP_CONCAT(cr2.numdoc) AS numero_documento_empresa,
                GROUP_CONCAT(cr2.domfiscal) as domicilio_empresa,
                GROUP_CONCAT(ur.nomdis) as distrito_empresa,
                GROUP_CONCAT(ur.nomprov) as provincia_empresa,
                GROUP_CONCAT(ur.nomdpto) as departamento_empresa
            FROM kardex as k
            LEFT JOIN tb_abogado as abo on abo.idabogado=k.idabogado
            LEFT JOIN usuarios as usu on usu.idusuario=k.idusuario
            LEFT JOIN contratantesxacto as cxa on cxa.kardex=k.kardex
            LEFT JOIN actocondicion as ac ON cxa.idcondicion=ac.idcondicion
            LEFT JOIN contratantes cn ON cxa.idcontratante = cn.idcontratante
            LEFT JOIN cliente2 as c2 on c2.idcontratante=cxa.idcontratante
            LEFT JOIN nacionalidades as n on n.idnacionalidad=c2.nacionalidad
            LEFT JOIN tipodocumento as td ON td.idtipdoc = c2.idtipdoc
            LEFT OUTER JOIN tipoestacivil as tec ON tec.idestcivil = c2.idestcivil
            LEFT OUTER JOIN ubigeo as u ON u.coddis = c2.idubigeo
            LEFT JOIN detallevehicular as dv ON dv.kardex=k.kardex
            LEFT JOIN sedesregistrales as sr ON sr.idsedereg=dv.idsedereg
            LEFT JOIN patrimonial as pat ON pat.kardex=k.kardex
            LEFT JOIN fpago_uif as uif ON uif.id_fpago = pat.fpago
            LEFT JOIN monedas as mon ON mon.idmon = pat.idmon
            LEFT JOIN detallemediopago as dmp ON pat.kardex = dmp.kardex
            LEFT JOIN mediospago as mp ON dmp.codmepag = mp.codmepag
            LEFT JOIN contratantes as cnr ON cxa.idcontratante = cnr.idcontratante
            LEFT JOIN contratantesxacto as cxar on cxar.idcontratante=cnr.idcontratanterp
            LEFT JOIN actocondicion as acr ON acr.idcondicion=cxar.idcondicion
            LEFT JOIN cliente2 as cr2 on cr2.idcontratante=cnr.idcontratanterp
            left JOIN nacionalidades as nr on nr.idnacionalidad=cr2.nacionalidad
            LEFT JOIN tipodocumento as tdr ON tdr.idtipdoc = cr2.idtipdoc
            LEFT OUTER JOIN tipoestacivil as tecr ON tecr.idestcivil = cr2.idestcivil
            LEFT OUTER JOIN ubigeo as ur ON ur.coddis = cr2.idubigeo
            WHERE k.kardex=%s and (c2.tipper='N')
            GROUP BY k.idkardex
        """
        
        with connection.cursor() as cursor:
            cursor.execute(query, [idtipoacto, template_id, template_id, num_kardex])
            desc = cursor.description
            row = cursor.fetchone()
            if not row:
                return None
            return dict(zip([col[0] for col in desc], row))

    def _get_data_documento(self, raw_data: dict) -> Dict[str, str]:
        """
        Get basic document information - mirrors get_data_documento PHP function
        """
        numero_escritura = raw_data.get('numero_escritura') or ''
        fecha_escritura = raw_data.get('fecha_escritura')
        numero_minuta = raw_data.get('numero_minuta') or ''
        
        numero_acta = f"{numero_escritura}({self.letras.number_to_letters(numero_escritura)})" if numero_escritura else '{{NRO_ESC}}'
        fecha_impresion = self.letras.date_to_letters(fecha_escritura) if fecha_escritura else '{{F_IMPRESION}}'
        fecha_acta = self.letras.date_to_letters(fecha_escritura) if fecha_escritura else '{{F}}'
        numero_minuta_formatted = numero_minuta if numero_minuta else '{{NRO_MIN}}'
        
        return {
            'NRO_ESC': numero_acta,
            'K': raw_data.get('kardex', ''),
            'NUM_REG': '1',
            'FEC_LET': self.letras.date_to_letters(fecha_escritura) if fecha_escritura else '',
            'F_IMPRESION': fecha_impresion,
            'USUARIO': raw_data.get('usuario', ''),
            'USUARIO_DNI': raw_data.get('dni_usuario', ''),
            'NRO_MIN': numero_minuta_formatted,
            'COMPROBANTE': ' ',
            'O_S': raw_data.get('kardex', ''),  # ORDEN DE SERVICIO
            'ORDEN_SERVICIO': ' ',
            'F': fecha_acta,
            'DESCRIPCION_SELLO': f"{raw_data.get('abogado', '')} CAP. {raw_data.get('matricula', '')}",
        }

    def _get_data_vehiculos(self, raw_data: dict) -> Dict[str, str]:
        """
        Get vehicle data - mirrors get_data_vehiculos PHP function
        """
        sede = raw_data.get('sede') or ''
        sede_parts = sede.split('-') if sede else ['', '']
        sede_name = sede_parts[1].strip() if len(sede_parts) > 1 else ''
        
        return {
            'PLACA': str(raw_data.get('placa') or '').upper(),
            'CLASE': str(raw_data.get('clase') or '').upper(),
            'MARCA': str(raw_data.get('marca') or '').upper(),
            'MODELO': str(raw_data.get('modelo') or '').upper(),
            'AÑO_FABRICACION': str(raw_data.get('anio') or '').upper(),
            'CARROCERIA': str(raw_data.get('carroceria') or '').upper(),
            'COLOR': str(raw_data.get('color') or '').upper(),
            'NRO_MOTOR': str(raw_data.get('motor') or '').upper(),
            'NRO_SERIE': str(raw_data.get('serie') or '').upper(),
            'FEC_INS': str(raw_data.get('fecha_inscripcion') or '').upper(),
            'FECHA_INSCRIPCION': str(raw_data.get('fecha_inscripcion') or '').upper(),
            'ZONA_REGISTRAL': str(sede).upper(),
            'NUM_ZONA_REG': str(raw_data.get('numero_zona') or '').upper(),
            'SEDE': str(sede_name).upper(),
            'INSTRUIDO': ' ',
            'COMBUSTIBLE': ' ',
            'NRO_TARJETA': ' ',
        }

    def _get_data_pagos(self, raw_data: dict) -> Dict[str, str]:
        """
        Get payment data - mirrors get_data_pagos PHP function
        """
        sunat_medio_pago = raw_data.get('sunat_medio_pago', '008')
        precio = raw_data.get('precio', 0)
        moneda = raw_data.get('moneda', 1)
        simbolo_moneda = raw_data.get('simbolo_moneda', '')
        
        # Payment method logic matching PHP switch statement
        if sunat_medio_pago == '008':
            medio_pago = 'EL COMPRADOR DECLARA QUE HA PAGADO EL PRECIO DEL VEHICULO EN DINERO EN EFECTIVO. NO HABIENDO UTILIZADO NINGÚN MEDIO DE PAGO ESTABLECIDO EN LA LEY Nº 28194, PORQUE EL MONTO TOTAL NO ES IGUAL NI SUPERA LOS S/ 3,500.00 O US$ 1,000.00. EL TIPO Y CÓDIGO DEL MEDIO EMPLEADO ES: "EFECTIVO POR OPERACIONES EN LAS QUE NO EXISTE OBLIGACIÓN DE UTILIZAR MEDIOS DE PAGO-008". INAPLICABLE LA LEY 30730 POR SER EL PAGO DEL PRECIO INFERIOR A 3 UIT.'
            exhibio_medio_pago = 'SE DEJA CONSTANCIA QUE PARA LA REALIZACIÓN DEL PRESENTE ACTO, LAS PARTES NO ME HAN EXHIBIDO NINGÚN MEDIO DE PAGO. DOY FE.'
            fin_medio_pago = 'EN DINERO EN EFECTIVO'
            forma_pago = 'AL CONTADO CON DINERO EN EFECTIVO'
        elif sunat_medio_pago == '009':
            medio_pago = 'EL COMPRADOR DECLARA QUE HA PAGADO EL PRECIO DEL VEHICULO EN DINERO EN EFECTIVO Y CON ANTERIORIDAD A LA CELEBRACION DE LA PRESENTE ACTA DE TRANSFERENCIA. NO HABIENDO UTILIZADO NINGÚN MEDIO DE PAGO ESTABLECIDO EN LA LEY Nº 28194, EL TIPO Y CÓDIGO DEL MEDIO EMPLEADO ES: "EFECTIVO POR OPERACIONES EN LAS QUE NO EXISTE OBLIGACIÓN DE UTILIZAR MEDIOS DE PAGO-009". INAPLICABLE LA LEY 30730 POR SER EL PAGO DEL PRECIO INFERIOR A 3 UIT.'
            exhibio_medio_pago = 'SE DEJA CONSTANCIA QUE PARA LA REALIZACIÓN DEL PRESENTE ACTO, LAS PARTES NO ME HAN EXHIBIDO NINGÚN MEDIO DE PAGO. DOY FE.'
            fin_medio_pago = 'EN DINERO EN EFECTIVO'
            forma_pago = 'AL CONTADO CON DINERO EN EFECTIVO'
        else:
            medio_pago = 'EL COMPRADOR DECLARA QUE HA PAGADO EL PRECIO DEL VEHICULO CON CHEQUE DEL BANCO DE CREDITO DEL PERÚ N° 1111111 111111 1111, GIRADO POR: YYYYYYYYY A FAVOR DE: XXXXXXXXX POR LA SUMA DE S/ 15,000.00, JULIACA 16/08/2018 EL TIPO Y CÓDIGO DEL MEDIO EMPLEADO ES: "CHEQUE -001" '
            exhibio_medio_pago = 'EN APLICACIÓN DE LA LEY 30730, SE DEJA CONSTANCIA QUE PARA LA REALIZACIÓN DEL PRESENTE ACTO, LAS PARTES ME HAN EXHIBIDO EL SIGUIENTE MEDIO DE PAGO: ……… CHEQUE DEL BANCO DE CREDITO DEL PERÚ N° 1111111 111111 1111, GIRADO POR: YYYYYYYYY A FAVOR DE: XXXXXXXXX POR LA SUMA DE S/ 15,000.00, JULIACA 16/08/2018. DOY FE.'
            fin_medio_pago = 'EN DINERO EN EFECTIVO'
            forma_pago = 'AL CONTADO CON DINERO EN EFECTIVO'
        
        # Get moneda description from constants
        descripcion_moneda = MONEDAS.get(moneda, {}).get('desmon', '') if moneda else ''
        
        return {
            'MONTO': str(precio),
            'MON_VEHI': descripcion_moneda,
            'MONTO_LETRAS': self.letras.money_to_letters(descripcion_moneda, Decimal(str(precio))),
            'MONEDA_C': simbolo_moneda + ' ' if simbolo_moneda else '',
            'SUNAT_MED_PAGO': sunat_medio_pago,
            'DES_PRE_VEHI': self.letras.money_to_letters(descripcion_moneda, Decimal(str(precio))),
            'EXH_MED_PAGO': exhibio_medio_pago,
            'MED_PAGO': medio_pago,
            'FIN_MED_PAGO': fin_medio_pago,
            'FORMA_PAGO': forma_pago,
            'C_INICIO_MP': '',
            'TIPO_PAGO_E': '',
            'TIPO_PAGO_C': '',
            'MONTO_MP': '',
            'CONSTANCIA': '',
            'DETALLE_MP': '',
            'FORMA_PAGO_S': '',
            'MONEDA_C_MP': '',
            'MEDIO_PAGO_C': '',
            'MP_MEDIO_PAGO': '',
            'MP_COMPLETO': '',
            'USO': '',
        }

    def _get_data_contratantes(self, raw_data: dict) -> Dict[str, list]:
        """
        Get contractors data - mirrors get_data_contratantes PHP function
        """
        def split_if_not_none(value, separator=','):
            return value.split(separator) if value else []

        # Extract data from raw query result
        condiciones = split_if_not_none(raw_data.get('condicion'))
        nombres = split_if_not_none(raw_data.get('nombres'))
        nacionalidades = split_if_not_none(raw_data.get('nacionalidad'))
        tipos_documento = split_if_not_none(raw_data.get('tipo_documento'))
        numeros_documento = split_if_not_none(raw_data.get('numero_documento'))
        ocupaciones = split_if_not_none(raw_data.get('ocupacion'))
        estados_civil = split_if_not_none(raw_data.get('estado_civil'))
        direcciones = split_if_not_none(raw_data.get('direccion'))
        distritos = split_if_not_none(raw_data.get('distrito'))
        provincias = split_if_not_none(raw_data.get('provincia'))
        departamentos = split_if_not_none(raw_data.get('departamento'))
        sexos = split_if_not_none(raw_data.get('sexo'))
        
        # Company data
        nombres_empresa = split_if_not_none(raw_data.get('nombre_empresa'))
        numeros_documento_empresa = split_if_not_none(raw_data.get('numero_documento_empresa'))
        direcciones_empresa = split_if_not_none(raw_data.get('domicilio_empresa'))
        distritos_empresa = split_if_not_none(raw_data.get('distrito_empresa'))
        provincias_empresa = split_if_not_none(raw_data.get('provincia_empresa'))
        departamentos_empresa = split_if_not_none(raw_data.get('departamento_empresa'))
        condiciones_empresa = split_if_not_none(raw_data.get('condicion_empresa'))
        
        transferentes = []
        adquirientes = []
        empresas = []
        
        contador_vendedor = 1
        contador_adquiriente = 1
        contador_empresa = 1
        
        # Process each contractor
        for k, condicion in enumerate(condiciones):
            if not condicion:
                continue
                
            # TRANSFERENTES (VENDEDOR, PODERDANTE, etc.)
            if condicion in ['VENDEDOR', 'PODERDANTE', 'OTORGANTE', 'REPRESENTANTE', 'ANTICIPANTE', 'ADJUDICANTE', 'DONANTE', 'USUFRUCTUANTE', 'TRANSFERENTE']:
                agregado = '' if contador_vendedor == 1 else f'_{contador_vendedor}'
                
                sexo = sexos[k] if k < len(sexos) else 'M'
                nombre = nombres[k] if k < len(nombres) else ''
                nacionalidad = nacionalidades[k] if k < len(nacionalidades) else ''
                tipo_doc = tipos_documento[k] if k < len(tipos_documento) else ''
                num_doc = numeros_documento[k] if k < len(numeros_documento) else ''
                ocupacion = ocupaciones[k] if k < len(ocupaciones) else ''
                estado_civil = estados_civil[k] if k < len(estados_civil) else ''
                direccion = direcciones[k] if k < len(direcciones) else ''
                distrito = distritos[k] if k < len(distritos) else ''
                provincia = provincias[k] if k < len(provincias) else ''
                departamento = departamentos[k] if k < len(departamentos) else ''
                
                # Gender-specific formatting
                if sexo == 'F':
                    documento_texto = f'IDENTIFICADA CON DNI N° {num_doc}, '
                    nacionalidad_texto = nacionalidad[:-1] + 'A, ' if nacionalidad.endswith('O') else nacionalidad + ', '
                    estado_civil_texto = estado_civil[:-1] + 'A, ' if estado_civil.endswith('O') else estado_civil + ', '
                else:
                    documento_texto = f'IDENTIFICADO CON DNI N° {num_doc}, '
                    nacionalidad_texto = nacionalidad[:-1] + 'O, ' if nacionalidad.endswith('A') else nacionalidad + ', '
                    estado_civil_texto = estado_civil + ', '
                
                transferentes.append({
                    f'P_NOM{agregado}': nombre + ', ',
                    f'P_NACIONALIDAD{agregado}': nacionalidad_texto,
                    f'P_TIP_DOC{agregado}': tipo_doc,
                    f'P_DOC{agregado}': documento_texto,
                    f'P_DOC_LETRAS{agregado}': documento_texto,
                    f'P_OCUPACION{agregado}': ocupacion,
                    f'P_ESTADO_CIVIL{agregado}': estado_civil_texto,
                    f'P_DOMICILIO{agregado}': f'CON DOMICILIO EN {direccion} DEL DISTRITO DE {distrito} PROVINCIA DE {provincia} Y DEPARTAMENTO DE {departamento}',
                    f'P_IDE{agregado}': ' ',
                    f'SEXO_P{agregado}': sexo,
                })
                contador_vendedor += 1
                
            # ADQUIRIENTES (COMPRADOR, APODERADO, etc.)
            elif condicion in ['COMPRADOR', 'APODERADO', 'ANTICIPADO', 'ADJUDICATARIO', 'DONATARIO', 'USUFRUCTUARIO', 'TESTIGO A RUEGO', 'ADQUIRIENTE', 'DEUDOR']:
                agregado = '' if contador_adquiriente == 1 else f'_{contador_adquiriente}'
                
                sexo = sexos[k] if k < len(sexos) else 'M'
                nombre = nombres[k] if k < len(nombres) else ''
                nacionalidad = nacionalidades[k] if k < len(nacionalidades) else ''
                tipo_doc = tipos_documento[k] if k < len(tipos_documento) else ''
                num_doc = numeros_documento[k] if k < len(numeros_documento) else ''
                ocupacion = ocupaciones[k] if k < len(ocupaciones) else ''
                estado_civil = estados_civil[k] if k < len(estados_civil) else ''
                direccion = direcciones[k] if k < len(direcciones) else ''
                distrito = distritos[k] if k < len(distritos) else ''
                provincia = provincias[k] if k < len(provincias) else ''
                departamento = departamentos[k] if k < len(departamentos) else ''
                
                # Gender-specific formatting
                if sexo == 'F':
                    documento_texto = f'IDENTIFICADA CON DNI N° {num_doc}, '
                    nacionalidad_texto = nacionalidad[:-1] + 'A, ' if nacionalidad.endswith('O') else nacionalidad + ', '
                    estado_civil_texto = estado_civil[:-1] + 'A, ' if estado_civil.endswith('O') else estado_civil + ', '
                else:
                    documento_texto = f'IDENTIFICADO CON DNI N° {num_doc}, '
                    nacionalidad_texto = nacionalidad[:-1] + 'O, ' if nacionalidad.endswith('A') else nacionalidad + ', '
                    estado_civil_texto = estado_civil + ', '
                
                adquirientes.append({
                    f'C_NOM{agregado}': nombre + ', ',
                    f'C_NACIONALIDAD{agregado}': nacionalidad_texto,
                    f'C_TIP_DOC{agregado}': tipo_doc,
                    f'C_DOC{agregado}': documento_texto,
                    f'C_DOC_LETRAS{agregado}': documento_texto,
                    f'C_OCUPACION{agregado}': ocupacion,
                    f'C_ESTADO_CIVIL{agregado}': estado_civil_texto,
                    f'C_DOMICILIO{agregado}': f'CON DOMICILIO EN {direccion} DEL DISTRITO DE {distrito} PROVINCIA DE {provincia} Y DEPARTAMENTO DE {departamento}',
                    f'C_IDE{agregado}': ' ',
                    f'SEXO_C{agregado}': sexo,
                })
                contador_adquiriente += 1
        
        # Process companies
        for k, condicion_emp in enumerate(condiciones_empresa):
            if not condicion_emp:
                continue
                
            nombre_emp = nombres_empresa[k] if k < len(nombres_empresa) else ''
            num_doc_emp = numeros_documento_empresa[k] if k < len(numeros_documento_empresa) else ''
            direccion_emp = direcciones_empresa[k] if k < len(direcciones_empresa) else ''
            distrito_emp = distritos_empresa[k] if k < len(distritos_empresa) else ''
            provincia_emp = provincias_empresa[k] if k < len(provincias_empresa) else ''
            departamento_emp = departamentos_empresa[k] if k < len(departamentos_empresa) else ''
            
            if condicion_emp == 'EMPRESA EN CONSTITUCION':
                empresas.append({
                    'NOMBRE_EMPRESA_2': nombre_emp,
                    'INS_EMPRESA_2': ' ',
                    'RUC_2': f', CON RUC N° {num_doc_emp}, ',
                    'DOMICILIO_EMPRESA_2': f'CON DOMICILIO EN {direccion_emp} DEL DISTRITO DE {distrito_emp} PROVINCIA DE {provincia_emp} Y DEPARTAMENTO DE {departamento_emp}',
                })
            else:
                empresas.append({
                    f'NOMBRE_EMPRESA_{contador_empresa}': nombre_emp,
                    f'INS_EMPRESA_{contador_empresa}': ' ',
                    f'RUC_{contador_empresa}': f', CON RUC N° {num_doc_emp}, ',
                    f'DOMICILIO_EMPRESA_{contador_empresa}': f'CON DOMICILIO EN {direccion_emp} DEL DISTRITO DE {distrito_emp} PROVINCIA DE {provincia_emp} Y DEPARTAMENTO DE {departamento_emp}',
                    f'CONDICION_EMPRESA_{contador_empresa}': condicion_emp
                })
                contador_empresa += 1
        
        return {
            'transferentes': transferentes,
            'adquirientes': adquirientes,
            'empresas': empresas
        }

    def _process_contratantes_data(self, contratantes_data: Dict[str, list]) -> Dict[str, str]:
        """
        Process contractors data and add empty placeholders - mirrors PHP processing
        """
        final_data = {}
        
        # Add transferentes data
        for transferente in contratantes_data['transferentes']:
            final_data.update(transferente)
        
        # Add adquirientes data
        for adquiriente in contratantes_data['adquirientes']:
            final_data.update(adquiriente)
        
        # Add empresas data
        for empresa in contratantes_data['empresas']:
            final_data.update(empresa)
        
        # Add empty placeholders for transferentes (P)
        total_transferentes = len(contratantes_data['transferentes'])
        for i in range(total_transferentes + 1, 11):
            final_data.update({
                f'P_NOM_{i}': '{{P_NOM_' + str(i) + '}}',
                f'P_NACIONALIDAD_{i}': '{{P_NACIONALIDAD_' + str(i) + '}}',
                f'P_TIP_DOC_{i}': '{{P_TIP_DOC_' + str(i) + '}}',
                f'P_DOC_{i}': '{{P_DOC_' + str(i) + '}}',
                f'P_DOC_LETRAS_{i}': '{{P_DOC_LETRAS_' + str(i) + '}}',
                f'P_OCUPACION_{i}': '{{P_OCUPACION_' + str(i) + '}}',
                f'P_ESTADO_CIVIL_{i}': '{{P_ESTADO_CIVIL_' + str(i) + '}}',
                f'P_DOMICILIO_{i}': '{{P_DOMICILIO_' + str(i) + '}}',
                f'P_IDE_{i}': '{{P_IDE_' + str(i) + '}}',
                f'P_FIRMA_{i}': '{{P_FIRMA_' + str(i) + '}}',
                f'P_AMBOS_{i}': '{{P_AMBOS_' + str(i) + '}}',
                f'SEXO_P_{i}': '{{SEXO_P_' + str(i) + '}}',
            })
        
        # Add empty placeholders for adquirientes (C)
        total_adquirientes = len(contratantes_data['adquirientes'])
        for i in range(total_adquirientes + 1, 11):
            final_data.update({
                f'C_NOM_{i}': '{{C_NOM_' + str(i) + '}}',
                f'C_NACIONALIDAD_{i}': '{{C_NACIONALIDAD_' + str(i) + '}}',
                f'C_TIP_DOC_{i}': '{{C_TIP_DOC_' + str(i) + '}}',
                f'C_DOC_{i}': '{{C_DOC_' + str(i) + '}}',
                f'C_DOC_LETRAS_{i}': '{{C_DOC_LETRAS_' + str(i) + '}}',
                f'C_OCUPACION_{i}': '{{C_OCUPACION_' + str(i) + '}}',
                f'C_ESTADO_CIVIL_{i}': '{{C_ESTADO_CIVIL_' + str(i) + '}}',
                f'C_DOMICILIO_{i}': '{{C_DOMICILIO_' + str(i) + '}}',
                f'C_IDE_{i}': '{{C_IDE_' + str(i) + '}}',
                f'C_FIRMA_{i}': '{{C_FIRMA_' + str(i) + '}}',
                f'C_AMBOS_{i}': '{{C_AMBOS_' + str(i) + '}}',
                f'SEXO_C_{i}': '{{SEXO_C_' + str(i) + '}}',
            })
        
        # Add empty placeholders for empresas
        total_empresas = len(contratantes_data['empresas'])
        for i in range(total_empresas + 1, 6):
            final_data.update({
                f'NOMBRE_EMPRESA_{i}': '{{NOMBRE_EMPRESA_' + str(i) + '}}',
                f'INS_EMPRESA_{i}': '{{INS_EMPRESA_' + str(i) + '}}',
                f'RUC_{i}': '{{RUC_' + str(i) + '}}',
                f'DOMICILIO_EMPRESA_{i}': '{{DOMICILIO_EMPRESA_' + str(i) + '}}',
            })
        
        return final_data

    def _get_articulos_contratantes(self, contratantes_data: Dict[str, list]) -> Dict[str, str]:
        """
        Get grammatical articles for contractors - mirrors PHP articulosContratantes logic
        """
        return {
            'EL_P': 'EL ',  # Will be determined by actual data processing
            'EL_C': 'EL ',  # Will be determined by actual data processing
            'CALIDAD_P': '',  # Will be determined by actual data processing
            'CALIDAD_C': '',  # Will be determined by actual data processing
            'Y_P': ' y ',
            'AMBOS': ' AMBOS ',
            'S_P': '  ',
            'ES_P': '  ',
            'INICIO_C': ' SEÑOR',
            'ES_C': '',
            'S_C': '',
            'ES_SON_C': 'ES',
            'Y_CON_C': '',
            'N_C': '',
            'Y_C': '',
            'L_C': '',
            'O_A_C': 'O',
            'O_ERON_C': '',
            'C_FIRMA': 'FIRMA EN',
            'C_AMBOS': '',
            'INICIO_P': ' SEÑOR',
            'ES_P': '',
            'S_P': '',
            'ES_SON_P': 'ES',
            'Y_CON_P': '',
            'N_P': '',
            'Y_P': '',
            'L_P': '',
            'O_A_P': 'O',
            'O_ERON_P': '',
            'P_FIRMA': 'FIRMA EN',
            'P_AMBOS': '',
        }

    def _get_data_escrituracion(self, raw_data: dict) -> Dict[str, str]:
        """
        Get escrituracion data - mirrors get_data_escrituracion PHP function
        """
        folio_inicial = raw_data.get('folio_inicial') or '{{FI}}'
        folio_final = raw_data.get('folio_final') or '{{FF}}'
        papel_inicial = raw_data.get('papel_inicial') or '{{S_IN}}'
        papel_final = raw_data.get('papel_final') or '{{S_FN}}'
        
        return {
            'FI': folio_inicial,
            'FF': folio_final,
            'S_IN': papel_inicial,
            'S_FN': papel_final,
        }

    def _get_template_from_r2(self, template_id: int) -> bytes:
        """
        Get template from R2 storage - same as other services
        """
        template = TplTemplate.objects.get(pktemplate=template_id)
        s3 = get_s3_client()
        
        object_key = f"rodriguez-zea/plantillas/{template.filename}"
        
        try:
            response = s3.get_object(Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'), Key=object_key)
            return response['Body'].read()
        except Exception as e:
            print(f"Error downloading template from R2: {e}")
            raise

    def _process_document(self, template_bytes: bytes, data: Dict[str, str]) -> Document:
        """
        Process the document template with data - same as other services
        """
        buffer = io.BytesIO(template_bytes)
        doc = DocxTemplate(buffer)
        doc.render(data)
        return doc

    def remove_unfilled_placeholders(self, doc):
        """
        Remove unfilled {{SOMETHING}} placeholders - adapted for garantias mobiliarias
        """
        import re
        curly_placeholder_pattern = re.compile(r'\{\{[A-Z0-9_]+\}\}')

        def clean_runs(runs):
            for run in runs:
                # Hide {{SOMETHING}} placeholders by making them white
                if curly_placeholder_pattern.search(run.text):
                    run.font.color.rgb = RGBColor(255, 255, 255)  # White color

        # Clean paragraphs
        for paragraph in doc.paragraphs:
            clean_runs(paragraph.runs)
        # Clean tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        clean_runs(paragraph.runs)

    def create_documento_in_r2(self, doc, kardex):
        """
        Upload the generated document to R2 - same as other services
        """
        try:
            from io import BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            object_key = f"rodriguez-zea/documentos/__PROY__{kardex}.docx"
            s3 = get_s3_client()
            s3.upload_fileobj(
                buffer,
                os.environ.get('CLOUDFLARE_R2_BUCKET'),
                object_key
            )
            return True
        except Exception as e:
            print(f"Error uploading garantias mobiliarias document to R2: {e}")
            return False

    def _create_response(self, doc, filename: str, kardex: str, mode: str = "download") -> HttpResponse:
        """
        Create HTTP response with the document - same as other services
        """
        from docxcompose.properties import CustomProperties
        
        # Add the custom property
        custom_props = CustomProperties(doc)
        custom_props['documentoGeneradoId'] = kardex

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        if mode == "open":
            # Production mode: Return JSON with file info for Word opening
            response = JsonResponse({
                'status': 'success',
                'mode': 'open',
                'filename': filename,
                'kardex': kardex,
                'message': 'Document generated and ready to open in Word'
            })
            response['Access-Control-Allow-Origin'] = '*'
            return response
        else:
            # Testing mode: Download the document
            response = HttpResponse(
                buffer.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'inline; filename="{filename}"'
            response['Content-Length'] = str(buffer.getbuffer().nbytes)
            response['Access-Control-Allow-Origin'] = '*'
            return response