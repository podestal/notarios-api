import os
import io
import traceback
from typing import Dict, Any, Optional

from django.db import connection
from django.http import HttpResponse, JsonResponse
from docxtpl import DocxTemplate

from ..shared.base_r2_documents import get_s3_client, BaseR2DocumentService
from ..utils import NumberToLetterConverter


class LibrosDocumentService(BaseR2DocumentService):
    """
    Service for Certificación de Apertura de Libros.

    - Templates expected in R2:
      - 'CERTIFICACION APERTURA DE LIBRO HORIZONTAL.docx'
      - 'CERTIFICACION APERTURA DE LIBRO VERTICAL.docx'
    - Output filename: '__LIBRO__{num_libro}-{anio_libro}.docx'
    """

    H_TEMPLATE = "CERTIFICACION APERTURA DE LIBRO HORIZONTAL.docx"
    V_TEMPLATE = "CERTIFICACION APERTURA DE LIBRO VERTICAL.docx"

    def __init__(self) -> None:
        self.letras = NumberToLetterConverter()
        # Default; can be overridden per-call based on orientation
        self.template_filename = self.V_TEMPLATE

    def retrieve_libro_document(self, num_libro: str, anio_libro: str, mode: str = "download") -> HttpResponse:
        try:
            if not num_libro or not anio_libro:
                return HttpResponse("Error: num_libro and anio_libro are required", status=400)

            filename = f"__LIBRO__{num_libro}-{anio_libro}.docx"

            if mode == "open":
                return self._create_response(None, filename, f"{num_libro}-{anio_libro}", mode)

            s3 = get_s3_client()
            object_key = self._object_key_for_document(filename)
            response = s3.get_object(Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'), Key=object_key)
            buffer = io.BytesIO(response['Body'].read())
            return self._create_response(buffer, filename, f"{num_libro}-{anio_libro}", mode)
        except Exception as e:
            if hasattr(e, 'response') and isinstance(getattr(e, 'response'), dict):
                if e.response.get('Error', {}).get('Code') == 'NoSuchKey':
                    return HttpResponse(f"Error: Document '{filename}' not found in R2.", status=404)
            traceback.print_exc()
            return HttpResponse(f"Error retrieving document: {e}", status=500)

    def generate_libro_document(self, num_libro: str, anio_libro: str, orientation: str = "V", mode: str = "download") -> HttpResponse:
        try:
            if not num_libro or not anio_libro:
                return self.json_error(400, "num_libro and anio_libro are required")

            filename = f"__LIBRO__{num_libro}-{anio_libro}.docx"
            # if self._document_exists_in_r2(filename):
            #     return self.json_error(409, "Document already exists. Use action=retrieve to fetch it.", {
            #         'num_libro': num_libro,
            #         'anio_libro': anio_libro,
            #         'filename': filename,
            #     })

            # Select template by orientation
            orientation_upper = (orientation or "V").upper()
            self.template_filename = self.H_TEMPLATE if orientation_upper.startswith("H") else self.V_TEMPLATE

            template_bytes = self._get_template_from_r2()
            if template_bytes is None:
                return self.json_error(404, f"Template '{self.template_filename}' not found in 'rodriguez-zea/plantillas/'.")

            libro_data = self._get_libro_data(num_libro, anio_libro)
            if not libro_data:
                return self.json_error(404, f"libros record {num_libro}-{anio_libro} not found")

            context: Dict[str, Any] = {}
            context.update(self._get_notary_data())
            context.update(libro_data)

            doc = DocxTemplate(io.BytesIO(template_bytes))
            doc.render(context)

            buffer = io.BytesIO()
            doc.save(buffer)
            self._save_document_to_r2(buffer, filename)
            return self._create_response(buffer, filename, f"{num_libro}-{anio_libro}", mode)
        except Exception as e:
            traceback.print_exc()
            return self.json_error(500, f"Error generating document: {e}")

    def _get_template_from_r2(self) -> Optional[bytes]:
        s3 = get_s3_client()
        object_key = self._object_key_for_template(self.template_filename)
        try:
            response = s3.get_object(Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'), Key=object_key)
            return response['Body'].read()
        except Exception:
            return None

    def _save_document_to_r2(self, buffer: io.BytesIO, filename: str) -> None:
        s3 = get_s3_client()
        object_key = self._object_key_for_document(filename)
        buffer.seek(0)
        s3.put_object(
            Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'),
            Key=object_key,
            Body=buffer.read(),
        )
        buffer.seek(0)

    def _create_response(self, buffer: Optional[io.BytesIO], filename: str, key_id: str, mode: str = "download") -> HttpResponse:
        if mode == "open":
            s3 = get_s3_client()
            object_key = self._object_key_for_document(filename)
            try:
                url = s3.generate_presigned_url(
                    'get_object',
                    Params={'Bucket': os.environ.get('CLOUDFLARE_R2_BUCKET'), 'Key': object_key},
                    ExpiresIn=3600,
                )
                response = JsonResponse({
                    'status': 'success', 'mode': 'open', 'url': url,
                    'filename': filename, 'libro': key_id,
                    'message': 'Document is ready to be opened.'
                })
                response['Access-Control-Allow-Origin'] = '*'
                return response
            except Exception as e:
                return HttpResponse(f"Error generating pre-signed URL: {e}", status=500)
        if buffer is None:
            return HttpResponse("Error: Document buffer is missing for download mode.", status=500)
        buffer.seek(0)
        response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'inline; filename="{filename}"'
        response['Content-Length'] = str(buffer.getbuffer().nbytes)
        response['Access-Control-Allow-Origin'] = '*'
        return response

    def _get_notary_data(self) -> Dict[str, str]:
        with connection.cursor() as cursor:
            cursor.execute("SELECT CONCAT(nombre, ' ', apellido) AS notario, direccion, distrito FROM confinotario")
            row = cursor.fetchone()
            if row:
                return {
                    'NOTARIO': str(row[0]).upper() if row[0] else '',
                    'DIRECCION_NOTARIO': str(row[1]).upper() if row[1] else '',
                    'DISTRITO_NOTARIO': str(row[2]).upper() if row[2] else '',
                }
        return {'NOTARIO': '', 'DIRECCION_NOTARIO': '', 'DISTRITO_NOTARIO': ''}

    def _get_libro_data(self, num_libro: str, anio_libro: str) -> Dict[str, Any]:
        d: Dict[str, Any] = {}
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT 
                    l.numlibro,
                    UPPER(l.descritiplib) AS des_libro,
                    UPPER(l.empresa) AS nom_empresa,
                    l.ruc AS num_ruc,
                    l.folio AS num_hojas2,
                    tf.destipfol AS tip_folio,
                    l.numlibro AS num_crono,
                    CONCAT(l.apepat,' ',l.apemat,', ',l.prinom,' ', l.segnom) AS nombre_persona,
                    (CASE WHEN l.tipper='N' THEN l.ruc ELSE l.ruc END) AS documento,
                    (CASE WHEN l.tipper='N' THEN 'D.N.I.' ELSE 'R.U.C.' END) AS tipo_documento,
                    nl.idnlibro AS nro_libro,
                    nl.desnlibro AS des_nro_libro,
                    UPPER(l.solicitante) AS nombre_solici,
                    l.dni AS dni_solici_raw,
                    UPPER(l.comentario) AS comentario,
                    UPPER(l.comentario2) AS comentario2,
                    l.tipper,
                    l.fecing AS fecha_ingreso_libro,
                    tl.deslegal AS tipolegalizacion,
                    u.nomdis AS distrito,
                    u.nomprov AS provincia,
                    u.nomdpto AS departamento,
                    l.domfiscal,
                    l.ano AS anio_crono,
                    l.detalle,
                    l.numdoc_plantilla
                FROM libros l
                INNER JOIN tipofolio tf ON tf.idtipfol = l.idtipfol
                INNER JOIN nlibro nl ON nl.idnlibro = l.idnlibro
                INNER JOIN tipolegal tl ON tl.idlegal = l.idlegal
                LEFT JOIN cliente c ON c.idcliente = l.codclie
                LEFT JOIN ubigeo u ON u.coddis = c.idubigeo
                WHERE l.numlibro = %s AND l.ano = %s
                LIMIT 1
                """,
                [num_libro, anio_libro],
            )
            row = cursor.fetchone()
            if not row:
                return {}
            cols = [col[0] for col in cursor.description]
            d = dict(zip(cols, row))

        # Derived values
        d['des_nro_libro'] = 'APERTURA' if (d.get('des_nro_libro') or '').upper() == 'PRIMERO' else (d.get('des_nro_libro') or '')

        nombre_persona = (d.get('nombre_persona') or '').upper()
        nom_empresa = (d.get('nom_empresa') or '').upper()
        d['eval_persona'] = nom_empresa if nom_empresa else nombre_persona

        documento = d.get('documento') or ''
        if documento:
            d['num_doc'] = documento
            d['domicilio_fiscal'] = f"CON DOMICILIO EN {d.get('domfiscal','')}".strip()
            d['ubigeo'] = f"DEL DISTRITO DE {d.get('distrito','')} PROVINCIA DE {d.get('provincia','')} Y DEPARTAMENTO DE {d.get('departamento','')}"
        else:
            d['num_doc'] = ''
            d['domicilio_fiscal'] = ''
            d['ubigeo'] = ''

        dni_solici = d.get('dni_solici_raw') or ''
        d['dni_solici'] = f"IDENTIFICADO CON DNI N°{dni_solici}" if dni_solici else ''

        # Date strings
        fecha_ing = d.get('fecha_ingreso_libro')
        if fecha_ing:
            try:
                d['fec_letras_completa'] = self.letras.date_to_letters(fecha_ing).upper()
                d['fec_completa'] = (fecha_ing.strftime('%d/%m/%Y') if hasattr(fecha_ing, 'strftime') else str(fecha_ing))
            except Exception:
                d['fec_letras_completa'] = ''
                d['fec_completa'] = ''
        else:
            d['fec_letras_completa'] = ''
            d['fec_completa'] = ''

        # Aliases and computed text for template
        # ano_crono alias expected by template (without i)
        d['ano_crono'] = d.get('anio_crono')

        # Libro text blocks depending on type
        des_libro_upper = (d.get('des_libro') or '').upper()
        detalle = d.get('detalle') or ''
        nombre_solici = (d.get('nombre_solici') or '').upper()
        d['titulo_libro'] = 'Apertura de Cuaderno de Obra' if des_libro_upper == 'CUADERNO DE OBRA' else 'Apertura de Libro'
        d['tipo_libro'] = 'CUADERNO' if des_libro_upper == 'CUADERNO DE OBRA' else 'LIBRO'
        d['solicitud_libro'] = '' if des_libro_upper == 'CUADERNO DE OBRA' else 'a solicitud de '
        d['solicitud_obra'] = 'SOLICITADO POR' if des_libro_upper == 'CUADERNO DE OBRA' else ''
        d['obra'] = f'OBRA: "{detalle}"' if des_libro_upper == 'CUADERNO DE OBRA' and detalle else ''
        d['solicitante_libro'] = '' if des_libro_upper == 'CUADERNO DE OBRA' else (f"{nombre_solici}, {d.get('dni_solici','')}" if nombre_solici else '')
        d['solicitante_obra'] = (f"{nombre_solici}, {d.get('dni_solici','')}" if des_libro_upper == 'CUADERNO DE OBRA' and nombre_solici else '')
        d['calidad_obra'] = 'EN SU CALIDAD DE RESIDENTE' if des_libro_upper == 'CUADERNO DE OBRA' else ''

        # Tipo persona block
        tipper = (d.get('tipper') or '').upper()
        if tipper == 'J':
            # persona jurídica
            d['tipo_persona'] = 'CON RUC NUMERO'
        else:
            d['tipo_persona'] = 'CON DNI NUMERO'

        # Defaults for footer variables if template includes them
        d['usuario'] = d.get('usuario', '') or ''
        d['usuario_dni'] = d.get('usuario_dni', '') or ''

        # Composed blocks to avoid dangling commas in templates
        d['ubigeo_block'] = ' '.join([part for part in [d.get('domicilio_fiscal', ''), d.get('ubigeo', '')] if part]).strip()
        if des_libro_upper == 'CUADERNO DE OBRA':
            d['solicitud_libro_line'] = ''
        else:
            # For libro variant: "A SOLICITUD DE {nombre}, {dni}"
            prefix = (d.get('solicitud_libro') or '').strip()
            who = (f"{nombre_solici}, {d.get('dni_solici','')}" if nombre_solici else '').strip()
            d['solicitud_libro_line'] = f"{prefix} {who}".strip()

        return {k: (v.upper() if isinstance(v, str) else v) for k, v in d.items()} 

class LibrosReportService:
    """Service for generating libros reports matching PHP script format"""
    
    def _sanitize_cell_value(self, value):
        """Sanitize cell values to prevent Excel corruption"""
        if value is None:
            return ""
        
        # Convert to string and strip whitespace
        val_str = str(value).strip()
        
        # Remove control characters that can cause XML corruption
        import re
        val_str = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', val_str)
        
        # Limit to Excel's cell character limit
        val_str = val_str[:32767]
        
        return val_str
    
    def _get_report_data(self, desde, hasta):
        """Fetch data for the report matching PHP query"""
        with connection.cursor() as cursor:
            # Convert YYYY-MM-DD back to DD/MM/YYYY for the SQL query
            desde_formatted = desde.strftime('%d/%m/%Y') if hasattr(desde, 'strftime') else desde
            hasta_formatted = hasta.strftime('%d/%m/%Y') if hasattr(hasta, 'strftime') else hasta
            
            # Print the dates for debugging
            print(f"Searching for records between {desde_formatted} and {hasta_formatted}")
            
            # First check if numdoc_plantilla column exists
            # Check if numdoc_plantilla column exists
            cursor.execute("""
                SELECT COUNT(*) 
                FROM information_schema.columns 
                WHERE table_schema = DATABASE()
                AND table_name = 'libros' 
                AND column_name = 'numdoc_plantilla'
            """)
            has_numdoc_plantilla = cursor.fetchone()[0] > 0
            
            # Build query based on available columns
            query = """
                SELECT
                    concat(libros.numlibro) as num_crono,
                    libros.fecing as fecha,
                    concat(libros.apepat,' ',libros.apemat,' ',libros.prinom,' ',libros.segnom) as cliente,
                    libros.empresa as empresa,
                    libros.descritiplib as tip_lib,
                    nlibro.desnlibro as n_lib,
                    libros.folio as folio,
                    tipofolio.destipfol as tip_fol,
                    libros.ruc as ruc,
                    libros.dni as dni,
                    libros.descritiplib as deslibro,
                    libros.solicitante as solicitante,
                    {numdoc_plantilla_col}
                FROM
                    libros
                    LEFT JOIN nlibro ON libros.idnlibro = nlibro.idnlibro
                    LEFT JOIN tipofolio ON libros.idtipfol = tipofolio.idtipfol
                    LEFT JOIN tipolibro ON libros.idtiplib = tipolibro.idtiplib
                WHERE STR_TO_DATE(fecing, '%%Y-%%m-%%d') BETWEEN STR_TO_DATE(%s, '%%d/%%m/%%Y') AND STR_TO_DATE(%s, '%%d/%%m/%%Y')
                ORDER BY num_crono
            """
            
            # Replace numdoc_plantilla column with NULL if it doesn't exist
            query = query.format(
                numdoc_plantilla_col='libros.numdoc_plantilla as ruc_plantilla' if has_numdoc_plantilla else 'NULL as ruc_plantilla'
            )
            
            cursor.execute(query, [desde_formatted, hasta_formatted])
            
            result = cursor.fetchall()
            print(f"Found {len(result) if result else 0} records")
            
            if result:
                return result
            return []
    
    def _get_notary_info(self):
        """Get notary configuration info"""
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    CONCAT(nombre, ' ', apellido) as nombre_completo,
                    direccion,
                    telefono,
                    departamento,
                    ruc,
                    provincia,
                    distrito
                FROM confinotario
                LIMIT 1
            """)
            result = cursor.fetchone()
            if result:
                return {
                    'nombre': result[0],
                    'direccion': result[1] or '',
                    'telefono': result[2] or '',
                    'departamento': result[3] or '',
                    'ruc': result[4] or '',
                    'provincia': result[5] or '',
                    'distrito': result[6] or ''
                }
            return {
                'nombre': 'NOTARIO',
                'direccion': '',
                'telefono': '',
                'departamento': '',
                'ruc': '', 
                'provincia': '',
                'distrito': ''
            }
    
    def _format_date_in_spanish(self, date_str):
        """Convert date to Spanish format like 'LUNES, 15 DE ENERO DEL 2025'"""
        try:
            # Parse the date string (assuming YYYY-MM-DD format)
            from datetime import datetime
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            
            # Spanish day names
            dias = ['LUNES', 'MARTES', 'MIÉRCOLES', 'JUEVES', 'VIERNES', 'SÁBADO', 'DOMINGO']
            # Spanish month names
            meses = ['ENERO', 'FEBRERO', 'MARZO', 'ABRIL', 'MAYO', 'JUNIO', 
                    'JULIO', 'AGOSTO', 'SEPTIEMBRE', 'OCTUBRE', 'NOVIEMBRE', 'DICIEMBRE']
            
            dia_semana = dias[date_obj.weekday()]
            dia = date_obj.day
            mes = meses[date_obj.month - 1]
            anio = date_obj.year
            
            return f"{dia_semana}, {dia} DE {mes} DEL {anio}"
        except:
            return date_str
    
    def _extract_year_from_date(self, date_str):
        """Extract year from date string DD/MM/YYYY or YYYY-MM-DD"""
        try:
            # Try to parse as YYYY-MM-DD first
            if '-' in date_str and len(date_str.split('-')[0]) == 4:
                return date_str.split('-')[0]
            # Try to parse as DD/MM/YYYY
            elif '/' in date_str:
                return date_str.split('/')[-1]
            else:
                from datetime import datetime
                return str(datetime.now().year)
        except:
            from datetime import datetime
            return str(datetime.now().year)
    
    def _format_date_for_display(self, date_obj):
        """Format date like PHP fechabd_an function"""
        try:
            if hasattr(date_obj, 'strftime'):
                return date_obj.strftime('%d/%m/%Y')
            return str(date_obj)
        except:
            return str(date_obj)
    
    def _format_dni_ruc(self, dni, ruc, ruc_plantilla):
        """Format DNI/RUC like PHP script logic"""
        if ruc_plantilla and 'CODJU' in str(ruc_plantilla):
            return ruc_plantilla
        else:
            if dni:
                return f"DNI: {dni}"
            elif ruc:
                return f"RUC: {ruc}"
            else:
                return ""
    
    def generate_excel_report(self, desde, hasta, orientation='horizontal'):
        """Generate Excel report matching PHP script format
        
        Args:
            desde: Start date
            hasta: End date
            orientation: 'horizontal' or 'vertical' (default: 'horizontal')
        """
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            import json
            
            # Get data
            report_data = self._get_report_data(desde, hasta)
            notary_info = self._get_notary_info()
            anio = self._extract_year_from_date(hasta)
            
            # Create workbook and worksheet
            wb = Workbook()
            ws = wb.active
            ws.title = "LIBROS CONTABLES"
            
            # Styles matching PHP script
            title_font = Font(name='Arial', size=18.5, bold=True)
            header_font = Font(name='Arial', size=13.5, bold=True)
            data_font = Font(name='Arial', size=13.5)
            center_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            left_alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
            
            # Borders for data table
            thin_border = Border(
                left=Side(border_style="thin"),
                right=Side(border_style="thin"),
                top=Side(border_style="thin"),
                bottom=Side(border_style="thin")
            )
            
            # No borders for header section
            no_border = Border(
                left=Side(style=None),
                right=Side(style=None),
                top=Side(style=None),
                bottom=Side(style=None)
            )
            
            # Title section
            ws.merge_cells('A1:G1')
            ws['A1'] = 'INDICE CRONOLOGICO - LEGALIZACION DE APERTURA DE LIBROS'
            ws['A1'].font = title_font
            ws['A1'].alignment = center_alignment
            ws['A1'].border = no_border
            
            ws.merge_cells('A2:G2')
            ws['A2'] = f'AÑO {anio}'
            ws['A2'].font = title_font
            ws['A2'].alignment = center_alignment
            ws['A2'].border = no_border
            
            # Notary info section
            row = 4
            ws[f'A{row}'] = 'NOTARIA'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].border = no_border
            ws[f'B{row}'] = f': {self._sanitize_cell_value(notary_info["nombre"])}'
            ws[f'B{row}'].font = data_font
            ws[f'B{row}'].border = no_border
            
            row += 1
            ws[f'A{row}'] = 'DIRECCION'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].border = no_border
            ws[f'B{row}'] = f': {self._sanitize_cell_value(notary_info["direccion"])}'
            ws[f'B{row}'].font = data_font
            ws[f'B{row}'].border = no_border
            ws[f'E{row}'] = 'TELEFONO'
            ws[f'E{row}'].font = header_font
            ws[f'E{row}'].border = no_border
            ws[f'F{row}'] = f': {self._sanitize_cell_value(notary_info["telefono"])}'
            ws[f'F{row}'].font = data_font
            ws[f'F{row}'].border = no_border
            
            row += 1
            ws[f'A{row}'] = 'DEPARTAMENTO'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].border = no_border
            ws[f'B{row}'] = f': {self._sanitize_cell_value(notary_info["departamento"])}'
            ws[f'B{row}'].font = data_font
            ws[f'B{row}'].border = no_border
            ws[f'E{row}'] = 'RUC'
            ws[f'E{row}'].font = header_font
            ws[f'E{row}'].border = no_border
            ws[f'F{row}'] = f': {self._sanitize_cell_value(notary_info["ruc"])}'
            ws[f'F{row}'].font = data_font
            ws[f'F{row}'].border = no_border
            
            row += 1
            ws[f'A{row}'] = 'PROVINCIA'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].border = no_border
            ws[f'B{row}'] = f': {self._sanitize_cell_value(notary_info["provincia"])}'
            ws[f'B{row}'].font = data_font
            ws[f'B{row}'].border = no_border
            ws[f'E{row}'] = 'DESDE'
            ws[f'E{row}'].font = header_font
            ws[f'E{row}'].border = no_border
            ws[f'F{row}'] = f': {self._format_date_in_spanish(desde)}'
            ws[f'F{row}'].font = data_font
            ws[f'F{row}'].border = no_border
            
            row += 1
            ws[f'A{row}'] = 'DISTRITO'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].border = no_border
            ws[f'B{row}'] = f': {self._sanitize_cell_value(notary_info["distrito"])}'
            ws[f'B{row}'].font = data_font
            ws[f'B{row}'].border = no_border
            ws[f'E{row}'] = 'HASTA'
            ws[f'E{row}'].font = header_font
            ws[f'E{row}'].border = no_border
            ws[f'F{row}'] = f': {self._format_date_in_spanish(hasta)}'
            ws[f'F{row}'].font = data_font
            ws[f'F{row}'].border = no_border
            
            # Table headers with borders
            row += 2
            if orientation == 'vertical':
                headers = ['NRO.', 'INGRESO LIBRO', 'PERTENECE A', '', 'NRO. LIBRO', 'NRO. FOLIOS', 'TIPO FOL.']
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=row, column=col, value=header)
                    cell.font = header_font
                    cell.alignment = center_alignment
                    cell.border = thin_border
                
                # Set column widths for vertical layout
                ws.column_dimensions['A'].width = 15  # NRO
                ws.column_dimensions['B'].width = 35  # INGRESO LIBRO
                ws.column_dimensions['C'].width = 20  # PERTENECE A
                ws.column_dimensions['D'].width = 20  # Empty column
                ws.column_dimensions['E'].width = 20  # NRO. LIBRO
                ws.column_dimensions['F'].width = 20  # NRO. FOLIOS
                ws.column_dimensions['G'].width = 20  # TIPO FOL.
            else:  # horizontal layout
                headers = ['NRO.', 'FECHA', 'TIPO LIBRO', 'SOLICITANTE', 'DNI', 'PROPIETARIO', 'RUC', 'NRO. LIBRO', 'NRO. FOLIOS', 'TIPO FOL.']
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=row, column=col, value=header)
                    cell.font = header_font
                    cell.alignment = center_alignment
                    cell.border = thin_border
                
                # Set minimum column widths for horizontal layout
                min_widths = {
                    'A': 10,  # NRO
                    'B': 15,  # FECHA
                    'C': 25,  # TIPO LIBRO
                    'D': 35,  # SOLICITANTE
                    'E': 15,  # DNI
                    'F': 35,  # PROPIETARIO
                    'G': 15,  # RUC
                    'H': 15,  # NRO. LIBRO
                    'I': 15,  # NRO. FOLIOS
                    'J': 15,  # TIPO FOL.
                }
                
                # Function to calculate text width
                def get_text_width(text):
                    if not text:
                        return 0
                    return len(str(text)) + 2  # +2 for padding
                
                # Track maximum width needed for each column
                max_widths = {col: width for col, width in min_widths.items()}
            
            # Data rows
            row += 1
            if report_data:
                for data_row in report_data:
                    try:
                        # Extract and sanitize data
                        num_crono = self._sanitize_cell_value(data_row[0] if len(data_row) > 0 else '')
                        fecha = data_row[1] if len(data_row) > 1 else ''
                        tip_lib = self._sanitize_cell_value(data_row[4] if len(data_row) > 4 else '')
                        solicitante = self._sanitize_cell_value(data_row[11] if len(data_row) > 11 else '')
                        cliente = self._sanitize_cell_value(data_row[2] if len(data_row) > 2 else '')
                        empresa = self._sanitize_cell_value(data_row[3] if len(data_row) > 3 else '')
                        n_lib = self._sanitize_cell_value(data_row[5] if len(data_row) > 5 else '')
                        folio = self._sanitize_cell_value(data_row[6] if len(data_row) > 6 else '')
                        tip_fol = self._sanitize_cell_value(data_row[7] if len(data_row) > 7 else '')
                        ruc = self._sanitize_cell_value(data_row[8] if len(data_row) > 8 else '')
                        dni = self._sanitize_cell_value(data_row[9] if len(data_row) > 9 else '')
                        ruc_plantilla = self._sanitize_cell_value(data_row[12] if len(data_row) > 12 else '')
                        
                        # Format DNI/RUC like PHP script
                        dni_ruc_formatted = self._format_dni_ruc(dni, ruc, ruc_plantilla)
                        
                        if orientation == 'vertical':
                            # Main data row - Vertical layout
                            ws.cell(row=row, column=1, value=int(num_crono) if num_crono.isdigit() else num_crono).font = data_font
                            ws.cell(row=row, column=2, value=f"{self._format_date_for_display(fecha)}     {tip_lib}").font = data_font
                            ws.cell(row=row, column=3, value="SOLICITANTE:\nPROPIETARIO:").font = data_font
                            ws.cell(row=row, column=4, value=f"{dni_ruc_formatted}").font = data_font
                            ws.cell(row=row, column=5, value=n_lib).font = data_font
                            ws.cell(row=row, column=6, value=folio).font = data_font
                            ws.cell(row=row, column=7, value=tip_fol).font = data_font
                            
                            # Apply borders and alignment
                            for col in range(1, 8):
                                cell = ws.cell(row=row, column=col)
                                cell.border = thin_border
                                if col in [1, 5, 6, 7]:  # Center align NRO, NRO. LIBRO, NRO. FOLIOS, TIPO FOL.
                                    cell.alignment = center_alignment
                                else:  # Left align other columns
                                    cell.alignment = left_alignment
                            
                            # Set row height
                            ws.row_dimensions[row].height = 45
                            
                            # Additional info row
                            row += 1
                            ws.cell(row=row, column=1, value='').font = data_font
                            ws.cell(row=row, column=2, value=f"{solicitante}\n{cliente}{empresa}").font = data_font
                            ws.cell(row=row, column=3, value='').font = data_font
                            ws.cell(row=row, column=4, value='').font = data_font
                            ws.cell(row=row, column=5, value='').font = data_font
                            ws.cell(row=row, column=6, value='').font = data_font
                            ws.cell(row=row, column=7, value='').font = data_font
                            
                            # Apply borders to info row
                            for col in range(1, 8):
                                cell = ws.cell(row=row, column=col)
                                cell.border = thin_border
                                if col == 2:  # Main content column
                                    cell.alignment = left_alignment
                                else:
                                    cell.alignment = center_alignment
                            
                            ws.row_dimensions[row].height = 45
                            row += 1
                        else:
                            # Single row - Horizontal layout
                            ws.cell(row=row, column=1, value=int(num_crono) if num_crono.isdigit() else num_crono).font = data_font
                            ws.cell(row=row, column=2, value=self._format_date_for_display(fecha)).font = data_font
                            ws.cell(row=row, column=3, value=tip_lib).font = data_font
                            ws.cell(row=row, column=4, value=solicitante).font = data_font
                            ws.cell(row=row, column=5, value=dni if dni else '').font = data_font
                            ws.cell(row=row, column=6, value=f"{cliente}{empresa}").font = data_font
                            ws.cell(row=row, column=7, value=ruc if ruc else '').font = data_font
                            ws.cell(row=row, column=8, value=n_lib).font = data_font
                            ws.cell(row=row, column=9, value=folio).font = data_font
                            ws.cell(row=row, column=10, value=tip_fol).font = data_font
                            
                            # Apply borders and alignment
                            for col in range(1, 11):  # Updated to include new column
                                cell = ws.cell(row=row, column=col)
                                cell.border = thin_border
                                if col in [1, 2, 7, 8, 9]:  # Center align numeric and date columns
                                    cell.alignment = center_alignment
                                else:  # Left align text columns
                                    cell.alignment = left_alignment
                            
                            # Set row height for single row
                            ws.row_dimensions[row].height = 30
                            row += 1
                        
                    except Exception as e:
                        # Error handling
                        ws.cell(row=row, column=1, value="ERROR").font = data_font
                        ws.cell(row=row, column=2, value=f"Error: {str(e)[:50]}").font = data_font
                        for col in range(1, 8):
                            ws.cell(row=row, column=col).border = thin_border
                        row += 1
            else:
                # No data message
                ws.cell(row=row, column=1, value='No se encontraron registros').font = data_font
                for col in range(1, 8):
                    ws.cell(row=row, column=col).border = thin_border
            
            # Update column widths based on content
            for row in ws.rows:
                for cell in row:
                    col_letter = get_column_letter(cell.column)
                    if col_letter in max_widths:
                        width = get_text_width(cell.value)
                        max_widths[col_letter] = max(max_widths[col_letter], width)

            # Apply the calculated widths, but not less than minimum widths
            for col_letter, width in max_widths.items():
                ws.column_dimensions[col_letter].width = max(width, min_widths[col_letter])

            # Save to BytesIO
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            
            # Create HTTP response
            filename = f"INDICE_CRONOLOGICO_LIBROS_CONTABLES_{anio}.xlsx"
            response = HttpResponse(
                output.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
            
        except Exception as e:
            return HttpResponse(
                json.dumps({'error': f'Error generating Excel report: {str(e)}'}),
                content_type='application/json',
                status=500
            )
    
    def generate_word_report(self, desde, hasta, orientation='horizontal'):
        """Generate Word report matching PHP script format
        
        Args:
            desde: Start date
            hasta: End date
            orientation: 'horizontal' or 'vertical' (default: 'horizontal')
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_ORIENT
            
            # Get data
            report_data = self._get_report_data(desde, hasta)
            notary_info = self._get_notary_info()
            anio = self._extract_year_from_date(hasta)
            
            # Create a new document
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            # Title - matching PHP font-size:18.5px
            title = doc.add_heading('INDICE CRONOLOGICO - LEGALIZACION DE APERTURA DE LIBROS', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(18.5)
                run.font.bold = True
            
            # Year
            year_heading = doc.add_heading(f'AÑO {anio}', 0)
            year_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in year_heading.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(18.5)
                run.font.bold = True
            
            # Add spacing
            doc.add_paragraph()
            doc.add_paragraph()
            
            # Notary info table - NO BORDERS
            info_table = doc.add_table(rows=5, cols=7)
            # No table style = no borders
            
            # Row 1: NOTARIA
            row = info_table.rows[0]
            row.cells[0].merge(row.cells[2])
            row.cells[0].text = 'NOTARIA'
            row.cells[3].text = f': {notary_info["nombre"]}'
            
            # Row 2: DIRECCION
            row = info_table.rows[1]
            row.cells[0].merge(row.cells[2])
            row.cells[0].text = 'DIRECCION'
            row.cells[3].text = f': {notary_info["direccion"]}'
            row.cells[4].text = 'TELEFONO'
            row.cells[5].merge(row.cells[6])
            row.cells[5].text = f': {notary_info["telefono"]}'
            
            # Row 3: DEPARTAMENTO
            row = info_table.rows[2]
            row.cells[0].merge(row.cells[2])
            row.cells[0].text = 'DEPARTAMENTO'
            row.cells[3].text = f': {notary_info["departamento"]}'
            row.cells[4].text = 'RUC'
            row.cells[5].merge(row.cells[6])
            row.cells[5].text = f': {notary_info["ruc"]}'
            
            # Row 4: PROVINCIA
            row = info_table.rows[3]
            row.cells[0].merge(row.cells[2])
            row.cells[0].text = 'PROVINCIA'
            row.cells[3].text = f': {notary_info["provincia"]}'
            row.cells[4].text = 'DESDE'
            row.cells[5].merge(row.cells[6])
            row.cells[5].text = f': {self._format_date_in_spanish(desde)}'
            
            # Row 5: DISTRITO
            row = info_table.rows[4]
            row.cells[0].merge(row.cells[2])
            row.cells[0].text = 'DISTRITO'
            row.cells[3].text = f': {notary_info["distrito"]}'
            row.cells[4].text = 'HASTA'
            row.cells[5].merge(row.cells[6])
            row.cells[5].text = f': {self._format_date_in_spanish(hasta)}'
            
            # Style the info table - Simple styling without borders
            for row in info_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if ':' in run.text:
                                run.font.bold = True
                            run.font.size = Pt(10)  # Set font size to 10pt
            
            # Add spacing
            doc.add_paragraph()
            
            # Main data table - ALWAYS CREATE, even if empty
            try:
                # Create table with headers based on orientation
                if orientation == 'vertical':
                    headers = ['NRO.', 'INGRESO LIBRO', 'PERTENECE A', '', 'NRO. LIBRO', 'NRO. FOLIOS', 'TIPO FOL.']
                    data_table = doc.add_table(rows=1, cols=7)
                else:  # horizontal layout
                    headers = ['NRO.', 'FECHA', 'TIPO LIBRO', 'SOLICITANTE', 'DNI', 'PROPIETARIO', 'RUC', 'NRO. LIBRO', 'NRO. FOLIOS', 'TIPO FOL.']
                    data_table = doc.add_table(rows=1, cols=10)
                    # Set page orientation to landscape for horizontal layout
                    section = doc.sections[-1]
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width = Inches(11.69)  # A4 height
                    section.page_height = Inches(8.27)  # A4 width
                
                data_table.style = 'Table Grid'
                
                # Set column widths for better content fit with smaller font
                for i, width in enumerate([
                    0.8,   # NRO
                    1.0,   # FECHA
                    1.6,   # TIPO LIBRO
                    2.0,   # SOLICITANTE
                    2.0,   # PROPIETARIO
                    1.2,   # DNI
                    1.2,   # RUC
                    1.0,   # NRO. LIBRO
                    1.0,   # NRO. FOLIOS
                    1.0,   # TIPO FOL.
                ]):
                    for cell in data_table.columns[i].cells:
                        cell.width = Inches(width)
            except Exception as e:
                self.logger.error(f"Error creating table: {str(e)}")
                raise
            
            try:
                # Add headers
                header_row = data_table.rows[0]
                for i, header in enumerate(headers):
                    cell = header_row.cells[i]
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run(header)
                    run.font.bold = True
                    run.font.size = Pt(12)  # Headers stay at 12pt
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                self.logger.error(f"Error adding headers: {str(e)}")
                raise
            
            # Add data rows if data exists
            if report_data and len(report_data) > 0:
                for data_row in report_data:
                    try:
                        num_crono = data_row[0] if len(data_row) > 0 else ''
                        fecha = data_row[1] if len(data_row) > 1 else ''
                        tip_lib = data_row[4] if len(data_row) > 4 else ''
                        solicitante = data_row[11] if len(data_row) > 11 else ''
                        cliente = data_row[2] if len(data_row) > 2 else ''
                        empresa = data_row[3] if len(data_row) > 3 else ''
                        n_lib = data_row[5] if len(data_row) > 5 else ''
                        folio = data_row[6] if len(data_row) > 6 else ''
                        tip_fol = data_row[7] if len(data_row) > 7 else ''
                        ruc = data_row[8] if len(data_row) > 8 else ''
                        dni = data_row[9] if len(data_row) > 9 else ''
                        ruc_plantilla = data_row[12] if len(data_row) > 12 else ''
                        
                        if orientation == 'vertical':
                            # Row 1: Main data - Vertical layout
                            row = data_table.add_row()
                            cells_data = [
                                str(num_crono),
                                f"{self._format_date_for_display(fecha)}     {tip_lib}",
                                "SOLICITANTE:\nPROPIETARIO:",
                                dni_ruc_formatted,
                                str(n_lib),
                                str(folio),
                                str(tip_fol)
                            ]
                            for idx, data in enumerate(cells_data):
                                cell = row.cells[idx]
                                paragraph = cell.paragraphs[0]
                                run = paragraph.add_run(data)
                                run.font.size = Pt(10)  # Data rows at 10pt
                                if idx in [0, 4, 5, 6]:  # Center align NRO, NRO. LIBRO, NRO. FOLIOS, TIPO FOL.
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                else:  # Left align other columns
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            
                            # Row 2: Additional info
                            row = data_table.add_row()
                            cells_data = [
                                '',
                                f"{solicitante}\n{cliente}{empresa}",
                                '',
                                '',
                                '',
                                '',
                                ''
                            ]
                            for idx, data in enumerate(cells_data):
                                cell = row.cells[idx]
                                paragraph = cell.paragraphs[0]
                                run = paragraph.add_run(data)
                                run.font.size = Pt(10)  # Data rows at 10pt
                                if idx == 1:  # Main content column
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                else:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            # Single row - Horizontal layout
                            row = data_table.add_row()
                            cells_data = [
                                str(num_crono),
                                self._format_date_for_display(fecha),
                                tip_lib,
                                solicitante,
                                dni if dni else '',
                                f"{cliente}{empresa}",
                                ruc if ruc else '',
                                str(n_lib),
                                str(folio),
                                str(tip_fol)
                            ]
                            for idx, data in enumerate(cells_data):
                                cell = row.cells[idx]
                                paragraph = cell.paragraphs[0]
                                run = paragraph.add_run(data)
                                run.font.size = Pt(10)  # Data rows at 10pt
                                if idx in [0, 1, 5, 6, 7, 8, 9]:  # Center align numeric, date, and document columns
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                else:  # Left align text columns
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    except Exception as e:
                        self.logger.error(f"Error processing row: {str(e)}")
                        continue
            else:
                # No data message
                no_data_row = data_table.add_row()
                no_data_cell = no_data_row.cells[0]
                no_data_cell.merge(no_data_row.cells[-1])  # Merge all cells in the row
                paragraph = no_data_cell.paragraphs[0]
                run = paragraph.add_run('No se encontraron registros para el período especificado')
                run.font.size = Pt(10)  # Set font size to 10pt
                run.font.italic = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            try:
                # Save to BytesIO
                output = io.BytesIO()
                doc.save(output)
                output.seek(0)
                
                # Create HTTP response
                filename = f"INDICE_CRONOLOGICO_LIBROS_CONTABLES_{anio}.docx"
                response = HttpResponse(
                    output.getvalue(),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = f'attachment; filename="{filename}"'
                
                return response
            except Exception as e:
                self.logger.error(f"Error saving document: {str(e)}")
                raise
            
        except Exception as e:
            from rest_framework.response import Response
            from rest_framework import status
            return Response(
                {'error': f'Error generating Word report: {str(e)}'}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            ) 