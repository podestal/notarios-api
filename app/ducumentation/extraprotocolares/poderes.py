import boto3
from botocore.client import Config
from botocore.exceptions import ClientError
import os
import io
import logging
from typing import Dict, Any, List, Optional, Tuple
from django.http import HttpResponse, JsonResponse
from django.db import connection
from docxtpl import DocxTemplate, RichText
import traceback
from ..shared.base_r2_documents import get_s3_client, BaseR2DocumentService
from ..utils import NumberToLetterConverter
from datetime import datetime
import re

logger = logging.getLogger(__name__)


class BasePoderDocumentService(BaseR2DocumentService):
    """
    Base service with common logic for generating Poder documents.
    - Loads templates from R2 under rodriguez-zea/plantillas/
    - Saves generated docs to R2 under rodriguez-zea/documentos/
    - Supports 'generate' and 'retrieve' workflows
    """
    def __init__(self) -> None:
        self.letras = NumberToLetterConverter()
        self.template_filename: Optional[str] = None

    def retrieve_document(self, id_poder: int, filename: str, mode: str = "download") -> HttpResponse:
        try:
            if not filename:
                return self.json_error(400, "filename is required to retrieve document")

            # Use the provided filename only (legacy-specific per endpoint)
            s3 = get_s3_client()
            bucket = os.environ.get('CLOUDFLARE_R2_BUCKET')
            object_key = f"rodriguez-zea/documentos/{filename}"
            if mode == "open":
                url = s3.generate_presigned_url(
                    'get_object',
                    Params={'Bucket': bucket, 'Key': object_key},
                    ExpiresIn=3600,
                )
                response = JsonResponse({
                    'status': 'success', 'mode': 'open', 'url': url,
                    'filename': filename, 'id_poder': id_poder,
                    'message': 'Document is ready to be opened.'
                })
                response['Access-Control-Allow-Origin'] = '*'
                return response

            response = s3.get_object(Bucket=bucket, Key=object_key)
            buffer = io.BytesIO(response['Body'].read())
            return self._create_response(buffer, filename, id_poder, mode)
        except ClientError as e:
            code = e.response.get('Error', {}).get('Code')
            if code == 'NoSuchKey':
                return self.json_error(404, f"Document not found in R2. Generate it first.", {
                    'id_poder': id_poder,
                    'filename': filename,
                })
            traceback.print_exc()
            return self.json_error(500, f"Error retrieving document: {e}")
        except Exception as e:
            traceback.print_exc()
            return self.json_error(500, f"Error retrieving document: {e}")

    def _get_template_from_r2(self) -> Optional[bytes]:
        if not self.template_filename:
            raise ValueError("template_filename must be set in child class")
        s3 = get_s3_client()
        object_key = f"rodriguez-zea/plantillas/{self.template_filename}"
        try:
            response = s3.get_object(Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'), Key=object_key)
            return response['Body'].read()
        except ClientError as e:
            if e.response['Error']['Code'] == 'NoSuchKey':
                return None
            raise

    def _render_with_coloring(self, template_bytes: bytes, context: Dict[str, Any]) -> DocxTemplate:
        doc = DocxTemplate(io.BytesIO(template_bytes))
        colored: Dict[str, Any] = {}
        for key, value in context.items():
            if isinstance(value, list):
                new_list = []
                for item in value:
                    if isinstance(item, dict):
                        new_item = {k: RichText(str(v) if v is not None else '', color='#FF0000') for k, v in item.items()}
                        new_list.append(new_item)
                    else:
                        new_list.append(RichText(str(item) if item is not None else '', color='#FF0000'))
                colored[key] = new_list
            elif isinstance(value, dict):
                colored[key] = {k: RichText(str(v) if v is not None else '', color='#FF0000') for k, v in value.items()}
            else:
                colored[key] = RichText(str(value) if value is not None else '', color='#FF0000')
        doc.render(colored)
        return doc

    def _create_response(self, buffer: Optional[io.BytesIO], filename: str, id_poder: int, mode: str = "download") -> HttpResponse:
        if mode == "open":
            s3 = get_s3_client()
            object_key = f"rodriguez-zea/documentos/{filename}"
            try:
                url = s3.generate_presigned_url(
                    'get_object',
                    Params={'Bucket': os.environ.get('CLOUDFLARE_R2_BUCKET'), 'Key': object_key},
                    ExpiresIn=3600,
                )
                response = JsonResponse({
                    'status': 'success', 'mode': 'open', 'url': url,
                    'filename': filename, 'id_poder': id_poder,
                    'message': 'Document is ready to be opened.'
                })
                response['Access-Control-Allow-Origin'] = '*'
                return response
            except Exception as e:
                return HttpResponse(f"Error generating pre-signed URL: {e}", status=500)
        if buffer is None:
            return HttpResponse("Error: Document buffer is missing for download mode.", status=500)
        buffer.seek(0)
        response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'inline; filename="{filename}"'
        response['Content-Length'] = str(buffer.getbuffer().nbytes)
        response['Access-Control-Allow-Origin'] = '*'
        return response

    def _save_document_to_r2(self, buffer: io.BytesIO, filename: str) -> None:
        s3 = get_s3_client()
        object_key = f"rodriguez-zea/documentos/{filename}"
        buffer.seek(0)
        s3.put_object(
            Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'),
            Key=object_key,
            Body=buffer.read(),
        )
        buffer.seek(0)

    def _get_notary_data(self) -> Dict[str, str]:
        with connection.cursor() as cursor:
            cursor.execute("SELECT CONCAT(nombre, ' ', apellido) AS notario, direccion, distrito AS distrito_notario FROM confinotario")
            notary_data = cursor.fetchone()
            if notary_data:
                columns = [col[0] for col in cursor.description]
                return {k.upper(): v for k, v in dict(zip(columns, notary_data)).items()}
        return {}

    def _get_user_data(self, usuario_imprime: Optional[str]) -> Dict[str, str]:
        if not usuario_imprime:
            return {'usuario': '?', 'dni_usuario': '?'}
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT loginusuario, dni FROM usuarios
                WHERE CONCAT(apepat,' ',prinom) = %s
                """,
                [usuario_imprime],
            )
            row = cursor.fetchone()
            if row:
                return {'usuario': row[0] or '?', 'dni_usuario': row[1] or '?'}
            return {'usuario': '?', 'dni_usuario': '?'}


class PoderFueraDeRegistroDocumentService(BasePoderDocumentService):
    """
    Poder Fuera de Registro document service.
    Template expected name (in R2): 'PODER FUERA DE REGISTRO BASE.docx'
    Output filename format: '__PROY__{num_kardex}.docx'
    """
    def __init__(self) -> None:
        super().__init__()
        self.template_filename = "PODER FUERA DE REGISTRO BASE.docx"

    def generate_poder_fuera_registro_document(self, id_poder: int, mode: str = "download") -> HttpResponse:
        try:
            poder_data = self._get_poder_data(id_poder)
            if not poder_data.get('NUM_KARDEX'):
                return HttpResponse(f"Error: num_kardex is empty for id_poder {id_poder}", status=400)

            anio_kardex = poder_data.get('anio_crono', '')
            filename = f"__PODER__{id_poder}-{anio_kardex}.docx"

            # Prevent regenerating if already exists
            # if self._document_exists_in_r2(filename):
            #     return self.json_error(409, "Document already exists. Use action=retrieve to fetch it.", {
            #         'id_poder': id_poder,
            #         'filename': filename,
            #     })

            template_bytes = self._get_template_from_r2()
            if template_bytes is None:
                return HttpResponse(
                    f"Error: Template '{self.template_filename}' not found in 'rodriguez-zea/plantillas/'.",
                    status=404,
                )

            context = self._build_context(id_poder, poder_data)
            doc = self._render_with_coloring(template_bytes, context)
            buffer = io.BytesIO()
            doc.save(buffer)
            self._save_document_to_r2(buffer, filename)
            return self._create_response(buffer, filename, id_poder, mode)
        except Exception as e:
            traceback.print_exc()
            return HttpResponse(f"Error generating document: {e}", status=500)

    def _get_poder_data(self, id_poder: int) -> Dict[str, Any]:
        data: Dict[str, Any] = {}
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT 
                    id_poder,
                    num_kardex,
                    fec_ingreso,
                    fec_crono,
                    num_formu
                FROM ingreso_poderes
                WHERE id_poder = %s
                """,
                [id_poder],
            )
            row = cursor.fetchone()
            if row:
                data['ID_PODER'] = str(row[0])
                data['NUM_KARDEX'] = row[1]
                data['FEC_INGRESO'] = row[2]
                data['FEC_CRONO'] = row[3]
                data['NUM_FORMU'] = row[4]
                if data['NUM_KARDEX'] and len(data['NUM_KARDEX']) >= 5:
                    data['anio_crono'] = data['NUM_KARDEX'][:4]
                    data['num_crono3'] = data['NUM_KARDEX'][4:]
                else:
                    data['anio_crono'] = ''
                    data['num_crono3'] = ''
                if data['FEC_INGRESO']:
                    data['fecha_letras_viaext'] = self.letras.date_to_letters(data['FEC_INGRESO'])
                else:
                    data['fecha_letras_viaext'] = ''
            # poderes_fuerareg
            cursor.execute(
                """
                SELECT f_fecotor, f_fecvcto, f_plazopoder, f_solicita
                FROM poderes_fuerareg
                WHERE id_poder = %s
                """,
                [id_poder],
            )
            row2 = cursor.fetchone()
            if row2:
                data['FEC_OTORGAMIENTO'] = row2[0]
                data['FEC_VENCIMIENTO'] = row2[1]
                data['PLAZO_PODER'] = row2[2]
                data['solicita'] = row2[3] or ''
        return data

    def _get_participants_data(self, id_poder: int) -> Dict[str, Any]:
        """
        Build the 'E' structure consumed by the Jinja2 template.
        This version fetches principals first, then runs a separate query for each principal
        to find their linked witness, mimicking the imperative style of the legacy PHP.
        """
        E: Dict[str, Any] = {}
        with connection.cursor() as cursor:
            # 1. Fetch up to two principals
            cursor.execute(
                """
                SELECT 
                    UPPER(CONCAT_WS(' ', c.prinom, c.segnom, c.apepat, c.apemat)) AS p_nom,
                    UPPER(n.descripcion) AS p_nacionalidad,
                    CASE WHEN c.sexo='F' THEN 'IDENTIFICADA CON' ELSE 'IDENTIFICADO CON' END AS p_ide,
                    UPPER(td.td_abrev) AS p_tipo_doc,
                    UPPER(c.numdoc) AS p_numdoc,
                    UPPER(te.desestcivil) AS p_estado_civil,
                    UPPER(COALESCE(c.detaprofesion, '')) AS p_ocupacion,
                    CONCAT('CON DOMICILIO EN ', UPPER(c.direccion), ' ',
                           UPPER(COALESCE(CONCAT('DEL DISTRITO DE ', u.nomdis, ', PROVINCIA DE ', u.nomprov, ', DEPARTAMENTO DE ', u.nomdpto), ''))) AS p_domicilio,
                    c.sexo AS p_sexo
                FROM poderes_contratantes pc
                JOIN cliente c ON c.numdoc = pc.c_codcontrat
                JOIN tipodocumento td ON td.idtipdoc = c.idtipdoc
                JOIN tipoestacivil te ON te.idestcivil = c.idestcivil
                LEFT JOIN nacionalidades n ON n.idnacionalidad = c.nacionalidad
                LEFT JOIN ubigeo u ON u.coddis = c.idubigeo
                WHERE pc.id_poder = %s AND pc.c_condicontrat IN ('007','011','009')
                ORDER BY pc.id_poder, pc.c_codcontrat
                LIMIT 2
                """,
                [id_poder],
            )
            principal_cols = [col[0] for col in cursor.description]
            principal_rows: List[Dict[str, Any]] = [dict(zip(principal_cols, row)) for row in cursor.fetchall()]

            # 2. For each principal, find their linked witness (if any)
            for idx, p_row in enumerate(principal_rows, start=1):
                suffix = '' if idx == 1 else '_2'

                # Populate principal data
                nat = (p_row.get('p_nacionalidad') or '')
                sex = (p_row.get('p_sexo') or 'M')
                if nat:
                    if sex == 'M' and nat.endswith('A'): nat = nat[:-1] + 'O'
                    elif sex == 'F' and nat.endswith('O'): nat = nat[:-1] + 'A'

                E[f'P_NOM{suffix}'] = p_row.get('p_nom') or ''
                E[f'P_NACIONALIDAD{suffix}'] = nat
                E[f'P_IDE{suffix}'] = p_row.get('p_ide') or 'IDENTIFICADO CON'
                E[f'P_DOC{suffix}'] = p_row.get('p_numdoc') or ''
                E[f'P_TIPO_DOC{suffix}'] = p_row.get('p_tipo_doc') or ''
                E[f'P_ESTADO_CIVIL{suffix}'] = p_row.get('p_estado_civil') or ''
                E[f'P_OCUPACION{suffix}'] = p_row.get('p_ocupacion') or ''
                E[f'P_DOMICILIO{suffix}'] = p_row.get('p_domicilio') or ''
                E[f'P_SEXO{suffix}'] = sex

                # Find linked witness for this principal
                principal_doc_num = p_row.get('p_numdoc')
                if principal_doc_num:
                    logger.info(f"Searching for witness linked to principal DNI: {principal_doc_num} for id_poder: {id_poder}")
                    cursor.execute(
                        """
                        SELECT
                            UPPER(CONCAT_WS(' ', c.prinom, c.segnom, c.apepat, c.apemat)) AS t_nom,
                            UPPER(n.descripcion) AS t_nacionalidad,
                            CASE WHEN c.sexo='F' THEN 'IDENTIFICADA CON' ELSE 'IDENTIFICADO CON' END AS t_ide,
                            UPPER(td.td_abrev) AS t_tipo_doc,
                            UPPER(c.numdoc) AS t_numdoc,
                            UPPER(te.desestcivil) AS t_estado_civil,
                            UPPER(COALESCE(c.detaprofesion, '')) AS t_ocupacion,
                            CONCAT('CON DOMICILIO EN ', UPPER(c.direccion), ' ',
                                   UPPER(COALESCE(CONCAT('DEL DISTRITO DE ', u.nomdis, ', PROVINCIA DE ', u.nomprov, ', DEPARTAMENTO DE ', u.nomdpto), ''))) AS t_domicilio,
                            c.sexo AS t_sexo
                        FROM poderes_contratantes pc
                        JOIN cliente c ON c.numdoc = pc.c_codcontrat
                        JOIN tipodocumento td ON td.idtipdoc = c.idtipdoc
                        JOIN tipoestacivil te ON te.idestcivil = c.idestcivil
                        LEFT JOIN nacionalidades n ON n.idnacionalidad = c.nacionalidad
                        LEFT JOIN ubigeo u ON u.coddis = c.idubigeo
                        WHERE pc.id_poder = %s AND pc.codi_testigo = %s AND pc.c_condicontrat = '008'
                        LIMIT 1
                        """,
                        [id_poder, principal_doc_num]
                    )
                    witness_cols = [col[0] for col in cursor.description]
                    w_row = cursor.fetchone()

                    if w_row:
                        logger.info(f"Found witness data: {w_row}")
                        witness_data = dict(zip(witness_cols, w_row))
                        natw = (witness_data.get('t_nacionalidad') or '')
                        sexw = (witness_data.get('t_sexo') or 'M')
                        if natw:
                            if sexw == 'M' and natw.endswith('A'): natw = natw[:-1] + 'O'
                            elif sexw == 'F' and natw.endswith('O'): natw = natw[:-1] + 'A'

                        E[f'T_INTERVIENE{suffix}'] = 'INTERVIENE:'
                        E[f'T_NOM{suffix}'] = witness_data.get('t_nom') or ''
                        E[f'T_NACIONALIDAD{suffix}'] = natw
                        E[f'T_IDE{suffix}'] = witness_data.get('t_ide') or 'IDENTIFICADO CON'
                        E[f'T_DOC{suffix}'] = witness_data.get('t_numdoc') or ''
                        E[f'T_TIPO_DOC{suffix}'] = witness_data.get('t_tipo_doc') or ''
                        E[f'T_ESTADO_CIVIL{suffix}'] = witness_data.get('t_estado_civil') or ''
                        E[f'T_OCUPACION{suffix}'] = witness_data.get('t_ocupacion') or ''
                        E[f'T_DOMICILIO{suffix}'] = witness_data.get('t_domicilio') or ''
                        E[f'T_CALIDAD{suffix}'] = 'QUIEN INTERVIENE EN CALIDAD DE TESTIGO A RUEGO'
                    else:
                        logger.info("No witness found for this principal.")

            # 3. Fetch Apoderado (representative) summary line used in text
            cursor.execute(
                """
                SELECT UPPER(CONCAT_WS(' ', c.prinom, c.segnom, c.apepat, c.apemat)) AS nom,
                       UPPER(td.td_abrev) AS tipo_doc,
                       UPPER(c.numdoc) AS numdoc,
                       c.sexo
                FROM poderes_contratantes pc
                JOIN cliente c ON c.numdoc = pc.c_codcontrat
                JOIN tipodocumento td ON td.idtipdoc = c.idtipdoc
                WHERE pc.id_poder = %s AND pc.c_condicontrat = '006'
                ORDER BY pc.id_poder, pc.c_codcontrat
                LIMIT 1
                """,
                [id_poder],
            )
            ap = cursor.fetchone()
            if ap:
                ap_text = f"{ap[0]} {'IDENTIFICADA' if (ap[3] or 'M')=='F' else 'IDENTIFICADO'} CON {ap[1]} N° {ap[2]}"
                E['apoderadoPoder'] = ap_text
            else:
                E['apoderadoPoder'] = ''

        # Pluralization helpers derived from number of principals
        num_principals = 0
        if 'P_NOM' in E and E['P_NOM']:
            num_principals += 1
        if 'P_NOM_2' in E and E['P_NOM_2']:
            num_principals += 1
        if num_principals > 1:
            return {
                **E,
                'P_EL': 'Los', 'P_S': 's', 'P_ES_SON': 'son', 'P_ES': 'es',
                'P_O_ARON': 'aron', 'P_N': 'n', 'P_DEL_LOS': 'de los', 'P_A_LOS': 'a los'
            }
        else:
            # Use first principal's sex if available; fallback by inferring from P_IDE
            sex = E.get('P_SEXO') or ('F' if str(E.get('P_IDE','')).strip().endswith('A') else 'M')
            if sex == 'F':
                return {
                    **E,
                    'P_EL': 'La', 'P_S': '', 'P_ES_SON': 'es', 'P_ES': '',
                    'P_O_ARON': 'ó', 'P_N': '', 'P_DEL_LOS': 'de la', 'P_A_LOS': 'a la'
                }
            return {
                **E,
                'P_EL': 'El', 'P_S': '', 'P_ES_SON': '', 'P_ES': '',
                'P_O_ARON': 'ó', 'P_N': '', 'P_DEL_LOS': 'del', 'P_A_LOS': 'al'
            }

    def _get_user_data(self, usuario_imprime: Optional[str]) -> Dict[str, str]:
        if not usuario_imprime:
            return {'usuario': '?', 'dni_usuario': '?'}
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT loginusuario, dni FROM usuarios
                WHERE CONCAT(apepat,' ',prinom) = %s
                """,
                [usuario_imprime],
            )
            row = cursor.fetchone()
            if row:
                return {'usuario': row[0] or '?', 'dni_usuario': row[1] or '?'}
            return {'usuario': '?', 'dni_usuario': '?'}

    def _build_context(self, id_poder: int, poder_data: Dict[str, Any]) -> Dict[str, Any]:
        context: Dict[str, Any] = {}
        # Map top-level vars expected by the template
        context.update({
            'aniocrono': poder_data.get('anio_crono', ''),
            'numcrono3': poder_data.get('num_crono3', ''),
            'fecha_letras_viaext': poder_data.get('fecha_letras_viaext', ''),
            'solicita': poder_data.get('solicita', ''),
        })
        # Pull user data if available (optional: you can adapt who is usuario_imprime)
        # Here we attempt to get from last editor; if not applicable, leave blanks
        context.update(self._get_user_data(None))
        # Participants block "E"
        E = self._get_participants_data(id_poder)
        context['E'] = E
        # Pronoun helpers also placed at top level for convenience
        for k in ['P_EL', 'P_S', 'P_ES_SON', 'P_ES', 'P_O_ARON', 'P_N', 'P_DEL_LOS', 'P_A_LOS', 'apoderadoPoder']:
            context[k] = E.get(k, '')
        # Flatten E so template can use top-level placeholders like {{P_NOM}}, {{T_DOC}}, etc.
        context.update(E)
        
        # Compose safe participant lines to avoid dangling commas
        def join_principal(idx: int) -> str:
            suf = '' if idx == 1 else '_2'
            nom = E.get(f'P_NOM{suf}', '')
            if not nom:
                return ''
            parts = [
                E.get(f'P_NOM{suf}', ''),
                E.get(f'P_NACIONALIDAD{suf}', ''),
                f"{E.get(f'P_IDE{suf}', '')} {E.get(f'P_TIPO_DOC{suf}', '')} N° {E.get(f'P_DOC{suf}', '')}".strip(),
                E.get(f'P_ESTADO_CIVIL{suf}', ''),
                E.get(f'P_OCUPACION{suf}', ''),
                E.get(f'P_DOMICILIO{suf}', ''),
            ]
            # Filter empties, then join with ', ' and end with period
            return (', '.join([p for p in parts if p]).replace('  ', ' ') + '.').strip()
        
        def join_testigo(idx: int) -> str:
            suf = '' if idx == 1 else '_2'
            nom = E.get(f'T_NOM{suf}', '')
            if not nom:
                return ''
            parts = [
                E.get(f'T_INTERVIENE{suf}', ''),
                E.get(f'T_NOM{suf}', ''),
                E.get(f'T_NACIONALIDAD{suf}', ''),
                f"{E.get(f'T_IDE{suf}', '')} {E.get(f'T_TIPO_DOC{suf}', '')} N° {E.get(f'T_DOC{suf}', '')}".strip(),
                E.get(f'T_ESTADO_CIVIL{suf}', ''),
                E.get(f'T_OCUPACION{suf}', ''),
                E.get(f'T_DOMICILIO{suf}', ''),
                E.get(f'T_CALIDAD{suf}', ''),
            ]
            return (', '.join([p for p in parts if p]).replace('  ', ' ') + '.').strip()
        
        context['PRINCIPAL_1_TEXT'] = join_principal(1)
        context['PRINCIPAL_2_TEXT'] = join_principal(2)
        context['TESTIGO_1_TEXT'] = join_testigo(1)
        context['TESTIGO_2_TEXT'] = join_testigo(2)
        
        return context


class PoderONPDocumentService(BasePoderDocumentService):
    def __init__(self) -> None:
        super().__init__()
        self.template_filename = "PODER ONP.docx"


class PoderEssaludDocumentService(BasePoderDocumentService):
    def __init__(self) -> None:
        super().__init__()
        self.template_filename = "plantilla_poder_essalud.docx"

    def generate_poder_essalud_document(self, id_poder: int, mode: str = "download") -> HttpResponse:
        try:
            poder_data = self._get_poder_data(id_poder)
            if not poder_data.get('NUM_KARDEX'):
                return HttpResponse(f"Error: num_kardex is empty for id_poder {id_poder}", status=400)

            anio_kardex = (poder_data.get('NUM_KARDEX') or '')[:4]
            filename = f"__PODER__{id_poder}-{anio_kardex}.docx"

            # Prevent regenerating if already exists
            # if self._document_exists_in_r2(filename):
            #     return self.json_error(409, "Document already exists. Use action=retrieve to fetch it.", {
            #         'id_poder': id_poder,
            #         'filename': filename,
            #     })

            template_bytes = self._get_template_from_r2()
            if template_bytes is None:
                return HttpResponse(
                    f"Error: Template '{self.template_filename}' not found in 'rodriguez-zea/plantillas/'.",
                    status=404,
                )

            context = self._build_context(id_poder, poder_data)
            # Since we are not coloring this document, we use a simpler render method
            doc = DocxTemplate(io.BytesIO(template_bytes))
            doc.render(context)
            
            buffer = io.BytesIO()
            doc.save(buffer)
            self._save_document_to_r2(buffer, filename)
            return self._create_response(buffer, filename, id_poder, mode)
        except Exception as e:
            traceback.print_exc()
            return HttpResponse(f"Error generating document: {e}", status=500)

    def _get_poder_data(self, id_poder: int) -> Dict[str, Any]:
        data: Dict[str, Any] = {}
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT num_kardex, fec_ingreso FROM ingreso_poderes WHERE id_poder = %s",
                [id_poder],
            )
            row = cursor.fetchone()
            if row:
                data['NUM_KARDEX'] = row[0]
                fec_ingreso = row[1]
                if fec_ingreso:
                    data['fecha_letras_viaext'] = self.letras.date_to_letters(fec_ingreso).upper()
                else:
                    data['fecha_letras_viaext'] = ''
                if data.get('NUM_KARDEX') and len(data['NUM_KARDEX']) >= 5:
                    data['numcrono2'] = f"{data['NUM_KARDEX'][4:]}-{data['NUM_KARDEX'][:4]}"
                else:
                    data['numcrono2'] = ''

            cursor.execute(
                """
                SELECT e_fecotor, e_fecvcto, e_montosep, e_montolact, e_montomater, e_plazopoder 
                FROM poderes_essalud WHERE id_poder = %s
                """,
                [id_poder],
            )
            row2 = cursor.fetchone()
            if row2:
                data['emision'] = self.letras.date_to_letters(row2[0]).upper() if row2[0] else ''
                data['vigencia_fin'] = row2[1] or ''
                data['sepelio'] = row2[2] or ''
                data['lactancia'] = row2[3] or ''
                data['maternidad'] = row2[4] or ''
                data['plazo'] = (row2[5] or '').upper()
        return data

    def _get_participants_data(self, id_poder: int) -> Dict[str, Any]:
        context: Dict[str, Any] = {}
        with connection.cursor() as cursor:
            # Poderdante (Grantor, role '007')
            cursor.execute("""
                SELECT
                    UPPER(CONCAT_WS(' ', c.prinom, c.segnom, c.apepat, c.apemat)) AS poderdante,
                    UPPER(td.destipdoc) AS tip_doc,
                    UPPER(c.numdoc) AS num_doc,
                    UPPER(tec.desestcivil) AS est_civil,
                    UPPER(n.descripcion) AS nacionalidad,
                    UPPER(c.direccion) AS direccion,
                    IF(u.coddis='070101','DISTRITO DE CALLAO , PROVINCIA CONSTITUCIONAL DEL CALLAO',CONCAT('DISTRITO DE ',u.nomdis, ', PROVINCIA DE ', u.nomprov,', DEPARTAMENTO DE ',u.nomdpto )) AS ubigeo,
                    pc.codi_asegurado AS seguro
                FROM poderes_contratantes pc
                JOIN cliente c ON c.numdoc = pc.c_codcontrat
                JOIN tipodocumento td ON td.idtipdoc = c.idtipdoc
                JOIN tipoestacivil tec ON tec.idestcivil = c.idestcivil
                JOIN nacionalidades n ON n.idnacionalidad = c.nacionalidad
                LEFT JOIN ubigeo u ON u.coddis = c.idubigeo
                WHERE pc.c_condicontrat = '007' AND pc.id_poder = %s
                LIMIT 1
            """, [id_poder])
            poderdante_row = cursor.fetchone()
            if poderdante_row:
                cols = [col[0] for col in cursor.description]
                p_data = dict(zip(cols, poderdante_row))
                p_data['domicilio'] = f"{p_data.get('direccion', '')} {p_data.get('ubigeo', '')}".strip()
                del p_data['direccion']
                del p_data['ubigeo']
                context.update(p_data)

            # Apoderado (Representative, role '006')
            cursor.execute("""
                SELECT
                    UPPER(CONCAT_WS(' ', c.prinom, c.segnom, c.apepat, c.apemat)) AS apoderado,
                    UPPER(td.destipdoc) AS tdoc_apoderado,
                    UPPER(c.numdoc) AS doc_apoderado,
                    UPPER(c.direccion) AS direccion,
                    IF(u.coddis='070101','DISTRITO DE CALLAO , PROVINCIA CONSTITUCIONAL DEL CALLAO',CONCAT('DISTRITO DE ',u.nomdis, ', PROVINCIA DE ', u.nomprov,', DEPARTAMENTO DE ',u.nomdpto )) AS ubigeo
                FROM poderes_contratantes pc
                JOIN cliente c ON c.numdoc = pc.c_codcontrat
                JOIN tipodocumento td ON td.idtipdoc = c.idtipdoc
                LEFT JOIN ubigeo u ON u.coddis = c.idubigeo
                WHERE pc.c_condicontrat = '006' AND pc.id_poder = %s
                LIMIT 1
            """, [id_poder])
            apoderado_row = cursor.fetchone()
            if apoderado_row:
                cols = [col[0] for col in cursor.description]
                ap_data = dict(zip(cols, apoderado_row))
                ap_data['domi_apoderado'] = f"{ap_data.get('direccion', '')} {ap_data.get('ubigeo', '')}".strip()
                del ap_data['direccion']
                del ap_data['ubigeo']
                context.update(ap_data)

        # Ensure all keys from the template exist to prevent errors
        expected_keys = [
            'poderdante', 'tip_doc', 'num_doc', 'est_civil', 'nacionalidad', 
            'domicilio', 'seguro', 'apoderado', 'tdoc_apoderado', 'doc_apoderado', 'domi_apoderado'
        ]
        for key in expected_keys:
            if key not in context:
                context[key] = ''
        return context

    def _build_context(self, id_poder: int, poder_data: Dict[str, Any]) -> Dict[str, Any]:
        context: Dict[str, Any] = {}
        context.update(poder_data)
        context.update(self._get_notary_data())
        context.update(self._get_participants_data(id_poder))
        return context 


class PoderPensionDocumentService(BasePoderDocumentService):
    def __init__(self) -> None:
        super().__init__()
        # Template expected in R2 at rodriguez-zea/plantillas/
        self.template_filename = "COBRO DE PENSION ONP.docx"

    def generate_poder_pension_document(self, id_poder: int, mode: str = "download") -> HttpResponse:
        try:
            poder_data = self._get_poder_data(id_poder)
            if not poder_data.get('NUM_KARDEX'):
                return HttpResponse(f"Error: num_kardex is empty for id_poder {id_poder}", status=400)

            anio_kardex = (poder_data.get('NUM_KARDEX') or '')[:4]
            filename = f"__PODER__{id_poder}-{anio_kardex}.docx"

            # Prevent regenerating if already exists
            # if self._document_exists_in_r2(filename):
            #     return self.json_error(409, "Document already exists. Use action=retrieve to fetch it.", {
            #         'id_poder': id_poder,
            #         'filename': filename,
            #     })

            template_bytes = self._get_template_from_r2()
            if template_bytes is None:
                return HttpResponse(
                    f"Error: Template '{self.template_filename}' not found in 'rodriguez-zea/plantillas/'.",
                    status=404,
                )

            context = self._build_context(id_poder, poder_data)
            doc = DocxTemplate(io.BytesIO(template_bytes))
            doc.render(context)

            buffer = io.BytesIO()
            doc.save(buffer)
            self._save_document_to_r2(buffer, filename)
            return self._create_response(buffer, filename, id_poder, mode)
        except Exception as e:
            traceback.print_exc()
            return HttpResponse(f"Error generating document: {e}", status=500)

    def _get_poder_data(self, id_poder: int) -> Dict[str, Any]:
        data: Dict[str, Any] = {}
        with connection.cursor() as cursor:
            # Basic ingreso data
            cursor.execute(
                "SELECT num_kardex, fec_ingreso FROM ingreso_poderes WHERE id_poder = %s",
                [id_poder],
            )
            row = cursor.fetchone()
            if row:
                data['NUM_KARDEX'] = row[0]
                fec_ingreso = row[1]
                # Build numcrono2 as NN...-YYYY if kardex is valid
                if data.get('NUM_KARDEX') and len(data['NUM_KARDEX']) >= 5:
                    data['numcrono2'] = f"{data['NUM_KARDEX'][4:]}-{data['NUM_KARDEX'][:4]}"
                else:
                    data['numcrono2'] = ''

            # Poderes Pension data
            cursor.execute(
                """
                SELECT p_fecotor, p_fecvcto, p_pension, p_presauto
                FROM poderes_pension
                WHERE id_poder = %s
                """,
                [id_poder],
            )
            row2 = cursor.fetchone()
            if row2:
                # Emission date in letters
                data['emision'] = self.letras.date_to_letters(row2[0]).upper() if row2[0] else ''
                data['vigencia_fin'] = row2[1] or ''
                data['prestacion'] = (row2[2] or '')
                data['prestacion_autorizada'] = (row2[3] or '')
        return data

    def _get_participants_data(self, id_poder: int) -> Dict[str, Any]:
        """
        Reuse the participants structure from Poder Fuera de Registro, producing P_*, T_* and pronoun helpers.
        """
        E: Dict[str, Any] = {}
        with connection.cursor() as cursor:
            # Principals
            cursor.execute(
                """
                SELECT 
                    UPPER(CONCAT_WS(' ', c.prinom, c.segnom, c.apepat, c.apemat)) AS p_nom,
                    UPPER(n.descripcion) AS p_nacionalidad,
                    CASE WHEN c.sexo='F' THEN 'IDENTIFICADA CON' ELSE 'IDENTIFICADO CON' END AS p_ide,
                    UPPER(td.td_abrev) AS p_tipo_doc,
                    UPPER(c.numdoc) AS p_numdoc,
                    UPPER(te.desestcivil) AS p_estado_civil,
                    UPPER(COALESCE(c.detaprofesion, '')) AS p_ocupacion,
                    CONCAT('CON DOMICILIO EN ', UPPER(c.direccion), ' ',
                           UPPER(COALESCE(CONCAT('DEL DISTRITO DE ', u.nomdis, ', PROVINCIA DE ', u.nomprov, ', DEPARTAMENTO DE ', u.nomdpto), ''))) AS p_domicilio,
                    c.sexo AS p_sexo
                FROM poderes_contratantes pc
                JOIN cliente c ON c.numdoc = pc.c_codcontrat
                JOIN tipodocumento td ON td.idtipdoc = c.idtipdoc
                JOIN tipoestacivil te ON te.idestcivil = c.idestcivil
                LEFT JOIN nacionalidades n ON n.idnacionalidad = c.nacionalidad
                LEFT JOIN ubigeo u ON u.coddis = c.idubigeo
                WHERE pc.id_poder = %s AND pc.c_condicontrat IN ('007','011','009')
                ORDER BY pc.id_poder, pc.c_codcontrat
                LIMIT 2
                """,
                [id_poder],
            )
            principal_cols = [col[0] for col in cursor.description]
            principal_rows: List[Dict[str, Any]] = [dict(zip(principal_cols, row)) for row in cursor.fetchall()]

            # For each principal, find linked witness
            for idx, p_row in enumerate(principal_rows, start=1):
                suffix = '' if idx == 1 else '_2'

                nat = (p_row.get('p_nacionalidad') or '')
                sex = (p_row.get('p_sexo') or 'M')
                if nat:
                    if sex == 'M' and nat.endswith('A'): nat = nat[:-1] + 'O'
                    elif sex == 'F' and nat.endswith('O'): nat = nat[:-1] + 'A'

                E[f'P_NOM{suffix}'] = p_row.get('p_nom') or ''
                E[f'P_NACIONALIDAD{suffix}'] = nat
                E[f'P_IDE{suffix}'] = p_row.get('p_ide') or 'IDENTIFICADO CON'
                E[f'P_DOC{suffix}'] = p_row.get('p_numdoc') or ''
                E[f'P_TIPO_DOC{suffix}'] = p_row.get('p_tipo_doc') or ''
                E[f'P_ESTADO_CIVIL{suffix}'] = p_row.get('p_estado_civil') or ''
                E[f'P_OCUPACION{suffix}'] = p_row.get('p_ocupacion') or ''
                E[f'P_DOMICILIO{suffix}'] = p_row.get('p_domicilio') or ''
                E[f'P_SEXO{suffix}'] = sex

                principal_doc_num = p_row.get('p_numdoc')
                if principal_doc_num:
                    logger.info(f"Searching for witness linked to principal DNI: {principal_doc_num} for id_poder: {id_poder}")
                    cursor.execute(
                        """
                        SELECT
                            UPPER(CONCAT_WS(' ', c.prinom, c.segnom, c.apepat, c.apemat)) AS t_nom,
                            UPPER(n.descripcion) AS t_nacionalidad,
                            CASE WHEN c.sexo='F' THEN 'IDENTIFICADA CON' ELSE 'IDENTIFICADO CON' END AS t_ide,
                            UPPER(td.td_abrev) AS t_tipo_doc,
                            UPPER(c.numdoc) AS t_numdoc,
                            UPPER(te.desestcivil) AS t_estado_civil,
                            UPPER(COALESCE(c.detaprofesion, '')) AS t_ocupacion,
                            CONCAT('CON DOMICILIO EN ', UPPER(c.direccion), ' ',
                                   UPPER(COALESCE(CONCAT('DEL DISTRITO DE ', u.nomdis, ', PROVINCIA DE ', u.nomprov, ', DEPARTAMENTO DE ', u.nomdpto), ''))) AS t_domicilio,
                            c.sexo AS t_sexo
                        FROM poderes_contratantes pc
                        JOIN cliente c ON c.numdoc = pc.c_codcontrat
                        JOIN tipodocumento td ON td.idtipdoc = c.idtipdoc
                        JOIN tipoestacivil te ON te.idestcivil = c.idestcivil
                        LEFT JOIN nacionalidades n ON n.idnacionalidad = c.nacionalidad
                        LEFT JOIN ubigeo u ON u.coddis = c.idubigeo
                        WHERE pc.id_poder = %s AND pc.codi_testigo = %s AND pc.c_condicontrat = '008'
                        LIMIT 1
                        """,
                        [id_poder, principal_doc_num]
                    )
                    witness_cols = [col[0] for col in cursor.description]
                    w_row = cursor.fetchone()

                    if w_row:
                        logger.info(f"Found witness data: {w_row}")
                        witness_data = dict(zip(witness_cols, w_row))
                        natw = (witness_data.get('t_nacionalidad') or '')
                        sexw = (witness_data.get('t_sexo') or 'M')
                        if natw:
                            if sexw == 'M' and natw.endswith('A'): natw = natw[:-1] + 'O'
                            elif sexw == 'F' and natw.endswith('O'): natw = natw[:-1] + 'A'

                        E[f'T_INTERVIENE{suffix}'] = 'INTERVIENE:'
                        E[f'T_NOM{suffix}'] = witness_data.get('t_nom') or ''
                        E[f'T_NACIONALIDAD{suffix}'] = natw
                        E[f'T_IDE{suffix}'] = witness_data.get('t_ide') or 'IDENTIFICADO CON'
                        E[f'T_DOC{suffix}'] = witness_data.get('t_numdoc') or ''
                        E[f'T_TIPO_DOC{suffix}'] = witness_data.get('t_tipo_doc') or ''
                        E[f'T_ESTADO_CIVIL{suffix}'] = witness_data.get('t_estado_civil') or ''
                        E[f'T_OCUPACION{suffix}'] = witness_data.get('t_ocupacion') or ''
                        E[f'T_DOMICILIO{suffix}'] = witness_data.get('t_domicilio') or ''
                        E[f'T_CALIDAD{suffix}'] = 'QUIEN INTERVIENE EN CALIDAD DE TESTIGO A RUEGO'

            # Pluralization helpers
            num_principals = 0
            if 'P_NOM' in E and E['P_NOM']:
                num_principals += 1
            if 'P_NOM_2' in E and E['P_NOM_2']:
                num_principals += 1
            if num_principals > 1:
                return {
                    **E,
                    'P_EL': 'Los', 'P_S': 's', 'P_ES_SON': 'son', 'P_ES': 'es',
                    'P_O_ARON': 'aron', 'P_N': 'n', 'P_DEL_LOS': 'de los', 'P_A_LOS': 'a los'
                }
            else:
                sex = E.get('P_SEXO') or ('F' if str(E.get('P_IDE','')).strip().endswith('A') else 'M')
                if sex == 'F':
                    return {
                        **E,
                        'P_EL': 'La', 'P_S': '', 'P_ES_SON': 'es', 'P_ES': '',
                        'P_O_ARON': 'ó', 'P_N': '', 'P_DEL_LOS': 'de la', 'P_A_LOS': 'a la'
                    }
                return {
                    **E,
                    'P_EL': 'El', 'P_S': '', 'P_ES_SON': '', 'P_ES': '',
                    'P_O_ARON': 'ó', 'P_N': '', 'P_DEL_LOS': 'del', 'P_A_LOS': 'al'
                }

    def _get_apoderado_fragment(self, id_poder: int) -> Dict[str, str]:
        """
        Build C_* variables used by the template to reference the apoderado in a natural sentence.
        - C_NOM: "{APODERADO}, "
        - C_DOC: "IDENTIFICADO(A) CON {TDOC} N° {NUM}, "
        - C_IDE: (kept empty for compatibility with legacy template)
        - C_O_A: 'a' if female, else 'o' (used in 'apoderad{{C_O_A}}')
        """
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT UPPER(CONCAT_WS(' ', c.prinom, c.segnom, c.apepat, c.apemat)) AS nom,
                       UPPER(td.td_abrev) AS tipo_doc,
                       UPPER(c.numdoc) AS numdoc,
                       c.sexo
                FROM poderes_contratantes pc
                JOIN cliente c ON c.numdoc = pc.c_codcontrat
                JOIN tipodocumento td ON td.idtipdoc = c.idtipdoc
                WHERE pc.id_poder = %s AND pc.c_condicontrat = '006'
                ORDER BY pc.id_poder, pc.c_codcontrat
                LIMIT 1
                """,
                [id_poder],
            )
            row = cursor.fetchone()
            if not row:
                return {'C_NOM': '', 'C_DOC': '', 'C_IDE': '', 'C_O_A': 'o'}
            nom, tipo_doc, numdoc, sexo = row
            is_female = (sexo or 'M') == 'F'
            c_nom = f"{(nom or '').upper()}, "
            c_doc = f"{'IDENTIFICADA' if is_female else 'IDENTIFICADO'} CON {(tipo_doc or '').upper()} N° {(numdoc or '').upper()}, "
            return {
                'C_NOM': c_nom,
                'C_DOC': c_doc,
                'C_IDE': '',
                'C_O_A': 'a' if is_female else 'o',
            }

    def _build_context(self, id_poder: int, poder_data: Dict[str, Any]) -> Dict[str, Any]:
        context: Dict[str, Any] = {}
        context.update(poder_data)
        # Optional user data (kept blank if unavailable)
        context.update(self._get_user_data(None))
        # Participants and pronouns
        E = self._get_participants_data(id_poder)
        context['E'] = E
        for k in ['P_EL', 'P_S', 'P_ES_SON', 'P_ES', 'P_O_ARON', 'P_N', 'P_DEL_LOS', 'P_A_LOS']:
            context[k] = E.get(k, '')
        context.update(E)
        # Apoderado fragment used in the prose
        context.update(self._get_apoderado_fragment(id_poder))
        return context 


class PoderesReportService:
    """Service for generating poderes reports"""
    
    def _sanitize_cell_value(self, value):
        """Sanitize cell values to prevent Excel corruption"""
        if value is None:
            return ""
        
        # Convert to string and strip whitespace
        val_str = str(value).strip()
        
        # Remove control characters that can cause XML corruption
        val_str = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', val_str)
        
        # Limit to Excel's cell character limit
        val_str = val_str[:32767]
        
        return val_str
    
    def _get_report_data(self, desde, hasta):
        """Fetch data for the poderes report"""
        with connection.cursor() as cursor:
            # Convert YYYY-MM-DD back to DD/MM/YYYY for the SQL query if needed
            try:
                desde_formatted = datetime.strptime(desde, '%Y-%m-%d').strftime('%Y-%m-%d')
                hasta_formatted = datetime.strptime(hasta, '%Y-%m-%d').strftime('%Y-%m-%d')
            except:
                desde_formatted = desde
                hasta_formatted = hasta
            
            cursor.execute("""
                SELECT
                    ingreso_poderes.id_poder as id_poder,
                    ingreso_poderes.num_kardex as kardex,
                    poderes_asunto.des_asunto as tip_poder,
                    ingreso_poderes.fec_crono as fec_crono,
                    ingreso_poderes.referencia as referencia,
                    ingreso_poderes.fec_ingreso as fec_ingreso,
                    ingreso_poderes.swt_est as estado,
                    ingreso_poderes.num_formu as formulario
                FROM ingreso_poderes 
                INNER JOIN poderes_asunto ON poderes_asunto.id_asunto = ingreso_poderes.id_asunto
                WHERE STR_TO_DATE(ingreso_poderes.fec_crono,'%%Y-%%m-%%d') >= STR_TO_DATE(%s,'%%Y-%%m-%%d') 
                AND STR_TO_DATE(ingreso_poderes.fec_crono,'%%Y-%%m-%%d') <= STR_TO_DATE(%s,'%%Y-%%m-%%d') 
                ORDER BY kardex
            """, [desde_formatted, hasta_formatted])
            
            result = cursor.fetchall()
            return result if result else []
    
    def _get_notary_info(self):
        """Get notary configuration info"""
        with connection.cursor() as cursor:
            cursor.execute("SELECT CONCAT(nombre, ' ', apellido) as notario FROM confinotario LIMIT 1")
            result = cursor.fetchone()
            return result[0] if result else "NOTARIO"
    
    def _get_contratantes_for_poder(self, id_poder):
        """Fetch contratantes for a specific poder"""
        with connection.cursor() as cursor:
            query = """
                SELECT
                    poderes_contratantes.id_contrata as id,
                    poderes_contratantes.c_descontrat as nombres,
                    poderes_contratantes.c_fircontrat as firma,
                    c_condiciones.des_condicion as condicion,
                    poderes_contratantes.c_codcontrat as dni_contratantes
                FROM poderes_contratantes 
                INNER JOIN c_condiciones ON c_condiciones.id_condicion = poderes_contratantes.c_condicontrat
                WHERE poderes_contratantes.id_poder = %s
                ORDER BY poderes_contratantes.id_contrata
            """
            
            try:
                cursor.execute(query, [id_poder])
                columns = [col[0] for col in cursor.description]
                return [dict(zip(columns, row)) for row in cursor.fetchall()]
            except Exception as e:
                print(f"DEBUG: Error getting contratantes for poder {id_poder}: {e}")
                return []
    
    def _get_contratantes_batch(self, poder_ids):
        """Fetch contratantes for multiple poderes in a single query"""
        if not poder_ids:
            return {}
            
        with connection.cursor() as cursor:
            placeholders = ','.join(['%s'] * len(poder_ids))
            
            query = f"""
                SELECT
                    poderes_contratantes.id_poder,
                    poderes_contratantes.id_contrata as id,
                    poderes_contratantes.c_descontrat as nombres,
                    poderes_contratantes.c_fircontrat as firma,
                    c_condiciones.des_condicion as condicion,
                    poderes_contratantes.c_codcontrat as dni_contratantes
                FROM poderes_contratantes 
                INNER JOIN c_condiciones ON c_condiciones.id_condicion = poderes_contratantes.c_condicontrat
                WHERE poderes_contratantes.id_poder IN ({placeholders})
                ORDER BY poderes_contratantes.id_poder, poderes_contratantes.id_contrata
                LIMIT 1000
            """
            
            try:
                cursor.execute(query, poder_ids)
                columns = [col[0] for col in cursor.description]
                results = cursor.fetchall()
                
                # Group contratantes by poder_id
                contratantes_by_poder = {}
                for row in results:
                    poder_id = row[0]
                    if poder_id not in contratantes_by_poder:
                        contratantes_by_poder[poder_id] = []
                    
                    contratante = dict(zip(columns, row))
                    contratantes_by_poder[poder_id].append(contratante)
                
                return contratantes_by_poder
                
            except Exception as e:
                print(f"DEBUG: Error in batch contratantes query: {e}")
                return {}
    
    def _format_date_in_spanish(self, date_str):
        """Convert date to Spanish format like 'LUNES, 15 DE ENERO DEL 2025'"""
        try:
            # Parse the date string (assuming YYYY-MM-DD format)
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            
            # Spanish day names
            dias = ['LUNES', 'MARTES', 'MIÉRCOLES', 'JUEVES', 'VIERNES', 'SÁBADO', 'DOMINGO']
            # Spanish month names
            meses = ['ENERO', 'FEBRERO', 'MARZO', 'ABRIL', 'MAYO', 'JUNIO', 
                    'JULIO', 'AGOSTO', 'SEPTIEMBRE', 'OCTUBRE', 'NOVIEMBRE', 'DICIEMBRE']
            
            dia_semana = dias[date_obj.weekday()]
            dia = date_obj.day
            mes = meses[date_obj.month - 1]
            anio = date_obj.year
            
            return f"{dia_semana}, {dia} DE {mes} DEL {anio}"
        except:
            return date_str
    
    def _extract_year_from_date(self, date_str):
        """Extract year from date string DD/MM/YYYY or YYYY-MM-DD"""
        try:
            # Try to parse as YYYY-MM-DD first
            if '-' in date_str and len(date_str.split('-')[0]) == 4:
                return date_str.split('-')[0]
            # Try to parse as DD/MM/YYYY
            elif '/' in date_str:
                return date_str.split('/')[-1]
            else:
                return str(datetime.now().year)
        except:
            return str(datetime.now().year)
    
    def generate_excel_report(self, desde, hasta):
        """Generate Excel report matching PHP script format"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            import json
            
            # Get data
            report_data = self._get_report_data(desde, hasta)
            notary_name = self._get_notary_info()
            anio = self._extract_year_from_date(hasta)
            
            # Create workbook and worksheet
            wb = Workbook()
            ws = wb.active
            ws.title = "PODERES FUERA DE REGISTRO"
            
            # Styles matching PHP script
            title_font = Font(name='Arial', size=18, bold=True)  # PHP: font-size:18.5px
            header_font = Font(name='Arial', size=13, bold=True)  # PHP: font-size: 13.5px
            data_font = Font(name='Arial', size=13)  # PHP: font-size: 13.5px
            center_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            left_alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
            
            # Borders for data table only
            thin_border = Border(
                left=Side(border_style="thin", color="333333"),
                right=Side(border_style="thin", color="333333"),
                top=Side(border_style="thin", color="333333"),
                bottom=Side(border_style="thin", color="333333")
            )
            
            # No borders for header section
            no_border = Border(
                left=Side(style=None),
                right=Side(style=None),
                top=Side(style=None),
                bottom=Side(style=None)
            )
            
            # Title section - spans across 9 columns (like PHP colspan="9")
            ws.merge_cells('A1:I1')
            ws['A1'] = 'INDICE CRONOLOGICO - PODERES FUERA DE REGISTRO'
            ws['A1'].font = title_font
            ws['A1'].alignment = center_alignment
            ws['A1'].border = no_border
            
            ws.merge_cells('A2:I2')
            ws['A2'] = f'AÑO {anio}'
            ws['A2'].font = title_font
            ws['A2'].alignment = center_alignment
            ws['A2'].border = no_border
            
            # Notary info section - matching PHP table structure
            row = 4
            # Row 1: NOTARIA (colspan="2")
            ws.merge_cells(f'A{row}:B{row}')
            ws[f'A{row}'] = 'NOTARIA'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].alignment = left_alignment
            ws[f'A{row}'].border = no_border
            ws.merge_cells(f'C{row}:D{row}')
            ws[f'C{row}'] = f': {self._sanitize_cell_value(notary_name)}'
            ws[f'C{row}'].font = data_font
            ws[f'C{row}'].alignment = left_alignment
            ws[f'C{row}'].border = no_border
            
            row += 1
            # Row 2: DIRECCION
            ws.merge_cells(f'A{row}:B{row}')
            ws[f'A{row}'] = 'DIRECCION'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].alignment = left_alignment
            ws[f'A{row}'].border = no_border
            ws.merge_cells(f'C{row}:D{row}')
            ws[f'C{row}'] = ': JR.BOLIVAR NRO. 340'
            ws[f'C{row}'].font = data_font
            ws[f'C{row}'].alignment = left_alignment
            ws[f'C{row}'].border = no_border
            ws[f'E{row}'] = 'TELEFONO'
            ws[f'E{row}'].font = header_font
            ws[f'E{row}'].alignment = left_alignment
            ws[f'E{row}'].border = no_border
            ws.merge_cells(f'F{row}:I{row}')
            ws[f'F{row}'] = ': (051) 326609'
            ws[f'F{row}'].font = data_font
            ws[f'F{row}'].alignment = left_alignment
            ws[f'F{row}'].border = no_border
            
            row += 1
            # Row 3: DEPARTAMENTO
            ws.merge_cells(f'A{row}:B{row}')
            ws[f'A{row}'] = 'DEPARTAMENTO'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].alignment = left_alignment
            ws[f'A{row}'].border = no_border
            ws.merge_cells(f'C{row}:D{row}')
            ws[f'C{row}'] = ': PUNO'
            ws[f'C{row}'].font = data_font
            ws[f'C{row}'].alignment = left_alignment
            ws[f'C{row}'].border = no_border
            ws[f'E{row}'] = 'RUC'
            ws[f'E{row}'].font = header_font
            ws[f'E{row}'].alignment = left_alignment
            ws[f'E{row}'].border = no_border
            ws.merge_cells(f'F{row}:I{row}')
            ws[f'F{row}'] = ': 10024231572'
            ws[f'F{row}'].font = data_font
            ws[f'F{row}'].alignment = left_alignment
            ws[f'F{row}'].border = no_border
            
            row += 1
            # Row 4: PROVINCIA
            ws.merge_cells(f'A{row}:B{row}')
            ws[f'A{row}'] = 'PROVINCIA'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].alignment = left_alignment
            ws[f'A{row}'].border = no_border
            ws.merge_cells(f'C{row}:D{row}')
            ws[f'C{row}'] = ': SAN ROMAN'
            ws[f'C{row}'].font = data_font
            ws[f'C{row}'].alignment = left_alignment
            ws[f'C{row}'].border = no_border
            ws[f'E{row}'] = 'DESDE'
            ws[f'E{row}'].font = header_font
            ws[f'E{row}'].alignment = left_alignment
            ws[f'E{row}'].border = no_border
            ws.merge_cells(f'F{row}:I{row}')
            ws[f'F{row}'] = f': {self._format_date_in_spanish(desde)}'
            ws[f'F{row}'].font = data_font
            ws[f'F{row}'].alignment = left_alignment
            ws[f'F{row}'].border = no_border
            
            row += 1
            # Row 5: DISTRITO
            ws.merge_cells(f'A{row}:B{row}')
            ws[f'A{row}'] = 'DISTRITO'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].alignment = left_alignment
            ws[f'A{row}'].border = no_border
            ws.merge_cells(f'C{row}:D{row}')
            ws[f'C{row}'] = ': JULIACA'
            ws[f'C{row}'].font = data_font
            ws[f'C{row}'].alignment = left_alignment
            ws[f'C{row}'].border = no_border
            ws[f'E{row}'] = 'HASTA'
            ws[f'E{row}'].font = header_font
            ws[f'E{row}'].alignment = left_alignment
            ws[f'E{row}'].border = no_border
            ws.merge_cells(f'F{row}:I{row}')
            ws[f'F{row}'] = f': {self._format_date_in_spanish(hasta)}'
            ws[f'F{row}'].font = data_font
            ws[f'F{row}'].alignment = left_alignment
            ws[f'F{row}'].border = no_border
            
            # Table headers with borders (matching PHP structure)
            row += 2
            headers = ['FECHA', 'N°', 'PODERDANTE', 'N° DNI', 'INTERVINIENTE', 'N° DNI', 'APODERADO', 'N° DNI', 'MOTIVO']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = header_font
                cell.alignment = center_alignment
                cell.border = thin_border
            
            # Set column widths to accommodate all data
            ws.column_dimensions['A'].width = 12  # FECHA
            ws.column_dimensions['B'].width = 8   # N°
            ws.column_dimensions['C'].width = 25  # PODERDANTE
            ws.column_dimensions['D'].width = 12  # N° DNI
            ws.column_dimensions['E'].width = 25  # INTERVINIENTE
            ws.column_dimensions['F'].width = 12  # N° DNI
            ws.column_dimensions['G'].width = 25  # APODERADO
            ws.column_dimensions['H'].width = 12  # N° DNI
            ws.column_dimensions['I'].width = 30  # MOTIVO
            
            # Data rows
            row += 1
            if report_data:
                # Batch query all contratantes for all poderes at once
                poder_ids = [data_row[0] for data_row in report_data if len(data_row) > 0]
                all_contratantes = self._get_contratantes_batch(poder_ids)
                
                for data_row in report_data:
                    try:
                        # Extract main poder data
                        id_poder = data_row[0] if len(data_row) > 0 else 0
                        kardex = self._sanitize_cell_value(data_row[1] if len(data_row) > 1 else '')
                        tip_poder = self._sanitize_cell_value(data_row[2] if len(data_row) > 2 else '')
                        fec_crono = self._sanitize_cell_value(data_row[3] if len(data_row) > 3 else '')
                        
                        # Extract correlative number like PHP: (int)substr($arr_poder[$j][1], 4,8)
                        correlativo = ""
                        if kardex and len(str(kardex)) > 4:
                            try:
                                correlativo = str(int(str(kardex)[4:]))  # Skip first 4 chars, convert to int
                            except:
                                correlativo = str(kardex)[4:] if len(str(kardex)) > 4 else str(kardex)
                        
                        # Get contratantes for this poder from batch results
                        contratantes = all_contratantes.get(id_poder, [])
                        
                        # Separate contratantes by type (like PHP)
                        poderdantes = []
                        testigos = []
                        apoderados = []
                        
                        for contratante in contratantes:
                            condicion = contratante.get('condicion', '').upper()
                            if condicion == 'PODERDANTE':
                                poderdantes.append(contratante)
                            elif condicion == 'TESTIGO A RUEGO':
                                testigos.append(contratante)
                            elif condicion == 'APODERADO':
                                apoderados.append(contratante)
                        
                        # Format participant texts
                        poderdante_text = '\n'.join([self._sanitize_cell_value(p.get('nombres', '')).upper() for p in poderdantes])
                        poderdante_dni = '\n'.join([self._sanitize_cell_value(p.get('dni_contratantes', '')).upper() for p in poderdantes])
                        
                        testigo_text = '\n'.join([self._sanitize_cell_value(t.get('nombres', '')).upper() for t in testigos])
                        testigo_dni = '\n'.join([self._sanitize_cell_value(t.get('dni_contratantes', '')).upper() for t in testigos])
                        
                        apoderado_text = '\n'.join([self._sanitize_cell_value(a.get('nombres', '')).upper() for a in apoderados])
                        apoderado_dni = '\n'.join([self._sanitize_cell_value(a.get('dni_contratantes', '')).upper() for a in apoderados])
                        
                        # Main data row
                        ws.cell(row=row, column=1, value=fec_crono).font = data_font
                        ws.cell(row=row, column=2, value=correlativo).font = data_font
                        ws.cell(row=row, column=3, value=poderdante_text).font = data_font
                        ws.cell(row=row, column=4, value=poderdante_dni).font = data_font
                        ws.cell(row=row, column=5, value=testigo_text).font = data_font
                        ws.cell(row=row, column=6, value=testigo_dni).font = data_font
                        ws.cell(row=row, column=7, value=apoderado_text).font = data_font
                        ws.cell(row=row, column=8, value=apoderado_dni).font = data_font
                        ws.cell(row=row, column=9, value=tip_poder.upper()).font = data_font
                        
                        # Apply borders and alignment
                        for col in range(1, 10):  # Columns A-I
                            cell = ws.cell(row=row, column=col)
                            cell.border = thin_border
                            if col in [1, 2, 9]:  # FECHA, N°, MOTIVO - center align
                                cell.alignment = center_alignment
                            else:  # Names and DNI - left align
                                cell.alignment = left_alignment
                        
                        # Calculate row height based on content
                        max_lines = 1
                        for text in [poderdante_text, testigo_text, apoderado_text]:
                            if text:
                                lines = len(text.split('\n'))
                                max_lines = max(max_lines, lines)
                        
                        # Set dynamic row height
                        calculated_height = max(30, max_lines * 20)
                        ws.row_dimensions[row].height = calculated_height
                        
                        row += 1
                        
                    except Exception as e:
                        # Error handling - add error row instead of crashing
                        ws.cell(row=row, column=1, value="ERROR").font = data_font
                        ws.cell(row=row, column=2, value=f"Error: {str(e)[:50]}").font = data_font
                        for col in range(1, 10):
                            ws.cell(row=row, column=col).border = thin_border
                        row += 1
            else:
                # No data message
                ws.merge_cells(f'A{row}:I{row}')
                ws.cell(row=row, column=1, value='No se encontraron registros para el período especificado').font = data_font
                ws.cell(row=row, column=1).alignment = center_alignment
                ws.cell(row=row, column=1).border = thin_border
            
            # Save to BytesIO
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            
            # Create HTTP response
            filename = f"INDICE_CRONOLOGICO_PODERES_{anio}.xlsx"
            response = HttpResponse(
                output.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
            
        except Exception as e:
            return HttpResponse(
                json.dumps({'error': f'Error generating Excel report: {str(e)}'}),
                content_type='application/json',
                status=500
            )
    
    def generate_word_report(self, desde, hasta):
        """Generate Word report matching PHP script format"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import json
            
            report_data = self._get_report_data(desde, hasta)
            notary_name = self._get_notary_info()
            anio = self._extract_year_from_date(hasta)
            
            # Create a new document
            doc = Document()
            
            # Title
            title = doc.add_heading('INDICE CRONOLOGICO - PODERES FUERA DE REGISTRO', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading(f'AÑO {anio}', 0)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add spacing
            doc.add_paragraph()
            
            # Notary info table - NO BORDERS
            info_table = doc.add_table(rows=5, cols=6)
            # No table style = no borders
            
            # Row 1: NOTARIA
            row = info_table.rows[0]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = 'NOTARIA'
            row.cells[2].merge(row.cells[3])
            row.cells[2].text = f': {notary_name}'
            
            # Row 2: DIRECCION
            row = info_table.rows[1]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = 'DIRECCION'
            row.cells[2].merge(row.cells[3])
            row.cells[2].text = ': JR.BOLIVAR NRO. 340'
            row.cells[4].text = 'TELEFONO'
            row.cells[5].text = ': (051) 326609'
            
            # Row 3: DEPARTAMENTO
            row = info_table.rows[2]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = 'DEPARTAMENTO'
            row.cells[2].merge(row.cells[3])
            row.cells[2].text = ': PUNO'
            row.cells[4].text = 'RUC'
            row.cells[5].text = ': 10024231572'
            
            # Row 4: PROVINCIA
            row = info_table.rows[3]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = 'PROVINCIA'
            row.cells[2].merge(row.cells[3])
            row.cells[2].text = ': SAN ROMAN'
            row.cells[4].text = 'DESDE'
            row.cells[5].text = f': {self._format_date_in_spanish(desde)}'
            
            # Row 5: DISTRITO
            row = info_table.rows[4]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = 'DISTRITO'
            row.cells[2].merge(row.cells[3])
            row.cells[2].text = ': JULIACA'
            row.cells[4].text = 'HASTA'
            row.cells[5].text = f': {self._format_date_in_spanish(hasta)}'
            
            # Style the info table
            for row in info_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if ':' in run.text:
                                run.font.bold = True
                            run.font.size = Pt(12)
            
            # Add spacing
            doc.add_paragraph()
            
            # Main data table
            headers = ['FECHA', 'N°', 'PODERDANTE', 'N° DNI', 'INTERVINIENTE', 'N° DNI', 'APODERADO', 'N° DNI', 'MOTIVO']
            data_table = doc.add_table(rows=1, cols=9)
            data_table.style = 'Table Grid'
            
            # Add headers
            header_row = data_table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(12)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add data rows if data exists
            if report_data and len(report_data) > 0:
                # Batch query all contratantes for all poderes at once
                poder_ids = [data_row[0] for data_row in report_data if len(data_row) > 0]
                all_contratantes = self._get_contratantes_batch(poder_ids)
                
                for data_row in report_data:
                    try:
                        # Extract main poder data
                        id_poder = data_row[0] if len(data_row) > 0 else 0
                        kardex = self._sanitize_cell_value(data_row[1] if len(data_row) > 1 else '')
                        tip_poder = self._sanitize_cell_value(data_row[2] if len(data_row) > 2 else '')
                        fec_crono = self._sanitize_cell_value(data_row[3] if len(data_row) > 3 else '')
                        
                        # Extract correlative number
                        correlativo = ""
                        if kardex and len(str(kardex)) > 4:
                            try:
                                correlativo = str(int(str(kardex)[4:]))
                            except:
                                correlativo = str(kardex)[4:] if len(str(kardex)) > 4 else str(kardex)
                        
                        # Get contratantes for this poder
                        contratantes = all_contratantes.get(id_poder, [])
                        
                        # Separate contratantes by type
                        poderdantes = []
                        testigos = []
                        apoderados = []
                        
                        for contratante in contratantes:
                            condicion = contratante.get('condicion', '').upper()
                            if condicion == 'PODERDANTE':
                                poderdantes.append(contratante)
                            elif condicion == 'TESTIGO A RUEGO':
                                testigos.append(contratante)
                            elif condicion == 'APODERADO':
                                apoderados.append(contratante)
                        
                        # Format participant texts
                        poderdante_text = '\n'.join([self._sanitize_cell_value(p.get('nombres', '')).upper() for p in poderdantes])
                        poderdante_dni = '\n'.join([self._sanitize_cell_value(p.get('dni_contratantes', '')).upper() for p in poderdantes])
                        
                        testigo_text = '\n'.join([self._sanitize_cell_value(t.get('nombres', '')).upper() for t in testigos])
                        testigo_dni = '\n'.join([self._sanitize_cell_value(t.get('dni_contratantes', '')).upper() for t in testigos])
                        
                        apoderado_text = '\n'.join([self._sanitize_cell_value(a.get('nombres', '')).upper() for a in apoderados])
                        apoderado_dni = '\n'.join([self._sanitize_cell_value(a.get('dni_contratantes', '')).upper() for a in apoderados])
                        
                        # Add data row
                        row = data_table.add_row()
                        row.cells[0].text = str(fec_crono)
                        row.cells[1].text = str(correlativo)
                        row.cells[2].text = poderdante_text
                        row.cells[3].text = poderdante_dni
                        row.cells[4].text = testigo_text
                        row.cells[5].text = testigo_dni
                        row.cells[6].text = apoderado_text
                        row.cells[7].text = apoderado_dni
                        row.cells[8].text = tip_poder.upper()
                        
                        # Center align FECHA, N°, MOTIVO
                        for i in [0, 1, 8]:
                            for paragraph in row.cells[i].paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Left align names and DNI
                        for i in [2, 3, 4, 5, 6, 7]:
                            for paragraph in row.cells[i].paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                    except Exception as e:
                        # Add error row instead of crashing
                        row = data_table.add_row()
                        row.cells[0].text = "ERROR"
                        row.cells[1].text = f"Error: {str(e)[:50]}"
                        for i in range(2, 9):
                            row.cells[i].text = ""
            else:
                # Add "No se encontraron registros" message
                row = data_table.add_row()
                row.cells[0].merge(row.cells[8])
                row.cells[0].text = "No se encontraron registros para el período especificado"
                for paragraph in row.cells[0].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save to buffer
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            # Create response
            filename = f"INDICE_CRONOLOGICO_PODERES_{anio}.docx"
            response = HttpResponse(
                output.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
            
        except Exception as e:
            return HttpResponse(
                json.dumps({'error': f'Error generating Word report: {str(e)}'}),
                content_type='application/json',
                status=500
            ) 