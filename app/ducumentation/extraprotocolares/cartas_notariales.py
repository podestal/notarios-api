import os
import io
import traceback
from typing import Dict, Any, Optional
from datetime import datetime
from io import BytesIO

from django.db import connection
from django.http import HttpResponse, JsonResponse
from rest_framework.response import Response
from rest_framework import status

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from ..shared.base_r2_documents import get_s3_client, BaseR2DocumentService
from ..utils import NumberToLetterConverter
from ..protocolares.utils import get_notary_config


class CartasNotarialesDocumentService(BaseR2DocumentService):
    """
    Service for generating and retrieving Certificación de Entrega de Carta Notarial documents.

    - Template expected in R2: 'CERTIFICACION ENTREGA DE CARTA NOTARIAL.docx'
    - Output filename: '__CARTA__{num_carta}.docx'
    - Stores under: {os.environ.get('CLOUDFLARE_R2_MAIN_URL')}/documentos/
    """

    def __init__(self) -> None:
        self.letras = NumberToLetterConverter()
        self.template_filename = "CERTIFICACION ENTREGA DE CARTA NOTARIAL.docx"

    def retrieve_carta_document(self, num_carta: str, mode: str = "download") -> HttpResponse:
        try:
            if not num_carta:
                return self.json_error(400, "num_carta is required to retrieve document")

            formatted_num_carta = self._format_num_carta(num_carta)
            filename = f"__CARTA__{formatted_num_carta}.docx"

            if mode == "open":
                return self._create_response(None, filename, num_carta, mode)

            s3 = get_s3_client()
            object_key = self._object_key_for_document(filename)
            response = s3.get_object(Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'), Key=object_key)
            buffer = io.BytesIO(response['Body'].read())
            return self._create_response(buffer, filename, num_carta, mode)
        except Exception as e:
            # Map not-found to 404 JSON
            if hasattr(e, 'response') and isinstance(getattr(e, 'response'), dict):
                code = e.response.get('Error', {}).get('Code')
                if code == 'NoSuchKey':
                    return self.json_error(404, "Document not found in R2. Generate it first.", {
                        'num_carta': num_carta,
                        'filename': f"__CARTA__{num_carta}.docx",
                    })
            traceback.print_exc()
            return self.json_error(500, f"Error retrieving document: {e}")

    def generate_carta_document(self, num_carta: str, mode: str = "download") -> HttpResponse:
        try:
            if not num_carta:
                return self.json_error(400, "num_carta is required to generate document")

            formatted_num_carta = self._format_num_carta(num_carta)
            filename = f"__CARTA__{formatted_num_carta}.docx"
            # if self._document_exists_in_r2(filename):
            #     return self.json_error(409, "Document already exists. Use action=retrieve to fetch it.", {
            #         'num_carta': formatted_num_carta,
            #         'filename': filename,
            #     })

            template_bytes = self._get_template_from_r2()
            if template_bytes is None:
                return self.json_error(404, f"Template '{self.template_filename}' not found in '{os.environ.get('CLOUDFLARE_R2_MAIN_URL')}/plantillas/'.")

            carta_data = self._get_carta_data(num_carta)
            if not carta_data:
                return self.json_error(404, f"ingreso_cartas record with num_carta {num_carta} not found")

            context: Dict[str, Any] = {}
            context.update(carta_data)
            context.update(self._get_user_data(carta_data.get('USUARIO_IMPRIME')))
            context.update(self._get_notary_data())

            # Aliases to match template variable names (docxtpl is case-sensitive)
            context['contenido_carta'] = str(
                context.get('CONTENIDO_CARTA')
                or context.get('conte_carta')
                or context.get('CONTENIDO')
                or ''
            ).strip()
            context['fec_ingreso'] = context.get('FECHA_INGRESO_LETRAS', '')
            context['num_carta'] = context.get('NUM_CARTA_FMT', '')
            # Legacy placeholders from PHP template
            context['USUARIO'] = context.get('USUARIO', '') or ''
            context['USUARIO_DNI'] = context.get('USUARIO_DNI', '') or ''
            context['COMPROBANTE'] = context.get('COMPROBANTE', '') or 'sin'

            doc = DocxTemplate(io.BytesIO(template_bytes))
            doc.render(context)

            buffer = io.BytesIO()
            doc.save(buffer)
            self._save_document_to_r2(buffer, filename)
            return self._create_response(buffer, filename, num_carta, mode)
        except Exception as e:
            traceback.print_exc()
            return self.json_error(500, f"Error generating document: {e}")

    def _get_template_from_r2(self) -> Optional[bytes]:
        s3 = get_s3_client()
        object_key = self._object_key_for_template(self.template_filename)
        try:
            response = s3.get_object(Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'), Key=object_key)
            return response['Body'].read()
        except Exception:
            return None

    def _save_document_to_r2(self, buffer: io.BytesIO, filename: str) -> None:
        s3 = get_s3_client()
        object_key = self._object_key_for_document(filename)
        buffer.seek(0)
        s3.put_object(
            Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'),
            Key=object_key,
            Body=buffer.read(),
        )
        buffer.seek(0)

    def _create_response(self, buffer: Optional[io.BytesIO], filename: str, num_carta: str, mode: str = "download") -> HttpResponse:
        if mode == "open":
            s3 = get_s3_client()
            object_key = self._object_key_for_document(filename)
            try:
                url = s3.generate_presigned_url(
                    'get_object',
                    Params={'Bucket': os.environ.get('CLOUDFLARE_R2_BUCKET'), 'Key': object_key},
                    ExpiresIn=3600,
                )
                response = JsonResponse({
                    'status': 'success', 'mode': 'open', 'url': url,
                    'filename': filename, 'num_carta': num_carta,
                    'message': 'Document is ready to be opened.'
                })
                response['Access-Control-Allow-Origin'] = '*'
                return response
            except Exception as e:
                return HttpResponse(f"Error generating pre-signed URL: {e}", status=500)
        if buffer is None:
            return HttpResponse("Error: Document buffer is missing for download mode.", status=500)
        buffer.seek(0)
        response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'inline; filename="{filename}"'
        response['Content-Length'] = str(buffer.getbuffer().nbytes)
        response['Access-Control-Allow-Origin'] = '*'
        return response

    def _get_notary_data(self) -> Dict[str, str]:
        with connection.cursor() as cursor:
            cursor.execute("SELECT CONCAT(nombre, ' ', apellido) AS notario FROM confinotario")
            row = cursor.fetchone()
            if row:
                return {'NOTARIO': str(row[0]).upper() if row[0] else ''}
        return {'NOTARIO': ''}

    def _get_user_data(self, usuario_imprime: Optional[str]) -> Dict[str, str]:
        if not usuario_imprime:
            return {'USUARIO': '?', 'USUARIO_DNI': '?'}
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT loginusuario, dni FROM usuarios
                WHERE CONCAT(apepat,' ',prinom) = %s
                """,
                [usuario_imprime],
            )
            row = cursor.fetchone()
            if row:
                return {'USUARIO': row[0] or '?', 'USUARIO_DNI': row[1] or '?'}
            return {'USUARIO': '?', 'USUARIO_DNI': '?'}

    def _format_num_carta(self, raw: Optional[str]) -> str:
        if not raw or len(raw) < 6:
            return raw or ''
        # Format as NNNNNN-YYYY like PHP CONCAT(RIGHT(6), '-', LEFT(4))
        return f"{raw[-6:]}-{raw[:4]}"

    def _get_carta_data(self, num_carta: str) -> Dict[str, Any]:
        data: Dict[str, Any] = {}
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT 
                    num_carta,
                    conte_carta,
                    emple_entrega,
                    STR_TO_DATE(fec_entrega, '%%d/%%m/%%Y') AS fecha_diligencia,
                    hora_entrega,
                    STR_TO_DATE(fec_ingreso, '%%d/%%m/%%Y') AS fecha_ingreso
                FROM ingreso_cartas
                WHERE num_carta = %s
                """,
                [num_carta],
            )
            row = cursor.fetchone()
            if not row:
                return {}
            raw_num_carta = row[0]
            contenido = row[1] or ''
            usuario_imprime = row[2] or ''
            fecha_diligencia = row[3]
            hora_entrega = row[4] or ''
            fecha_ingreso = row[5]

            # Prepare replacements in contenido (00/00/0000 -> dd/mm/YYYY, 00:00 -> hora_entrega)
            fecha_diligencia_ddmmyyyy = fecha_diligencia.strftime('%d/%m/%Y') if fecha_diligencia else ''
            contenido_replaced = (
                str(contenido)
                .replace('\u00a0', ' ')
                .replace('00/00/0000', fecha_diligencia_ddmmyyyy)
                .replace('00:00', hora_entrega)
                .strip()
            )

            data.update({
                'NUM_CARTA': raw_num_carta or '',
                'NUM_CARTA_FMT': self._format_num_carta(raw_num_carta),
                'CONTENIDO_CARTA': contenido_replaced,
                'conte_carta': contenido_replaced,
                'USUARIO_IMPRIME': usuario_imprime,
                'FECHA_DILIGENCIA': fecha_diligencia_ddmmyyyy,
                'FECHA_DILIGENCIA_LETRAS': self.letras.date_to_letters(fecha_diligencia).upper() if fecha_diligencia else '',
                'HORA_DILIGENCIA': hora_entrega,
                'FECHA_INGRESO_LETRAS': self.letras.date_to_letters(fecha_ingreso).lower() if fecha_ingreso else '',
            })
        return data 

    def _sanitize_cell_value(self, value):
        """Sanitize cell values to prevent Excel corruption"""
        if value is None:
            return ""
        
        # Convert to string and strip whitespace
        val_str = str(value).strip()
        
        # Remove control characters that can cause XML corruption
        import re
        val_str = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', val_str)
        
        # Limit to Excel's cell character limit
        val_str = val_str[:32767]
        
        return val_str


class CartasNotarialesReportService:
    """Service for generating cartas notariales reports"""
    
    def _sanitize_cell_value(self, value):
        """Sanitize cell values to prevent Excel corruption"""
        if value is None:
            return ""
        
        # Convert to string and strip whitespace
        val_str = str(value).strip()
        
        # Remove control characters that can cause XML corruption
        import re
        val_str = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', val_str)
        
        # Limit to Excel's cell character limit
        val_str = val_str[:32767]
        
        return val_str
    
    def _get_report_data(self, desde, hasta):
        """Fetch data for the report"""
        with connection.cursor() as cursor:
            # Convert YYYY-MM-DD back to DD/MM/YYYY for the SQL query
            desde_formatted = datetime.strptime(desde, '%Y-%m-%d').strftime('%d/%m/%Y')
            hasta_formatted = datetime.strptime(hasta, '%Y-%m-%d').strftime('%d/%m/%Y')
            
            cursor.execute("""
                SELECT
                    ic.num_carta AS num_carta,
                    DATE_FORMAT(STR_TO_DATE(ic.fec_ingreso,'%%d/%%m/%%Y'),'%%d/%%m/%%Y') AS fec_ingreso,
                    ic.fec_entrega AS fec_entrega,
                    ic.hora_entrega AS hora_entrega,
                    ic.nom_destinatario AS destinatario,
                    ic.nom_remitente AS remitente,
                    u.nomdis as zona,
                    ic.dir_destinatario,
                    ic.id_remitente as dni_remitente,
                    ic.dni_destinatario,
                    ic.recepcion,
                    ic.firmo
                FROM ingreso_cartas ic
                INNER JOIN ubigeo u ON u.coddis = ic.zona_destinatario
                WHERE STR_TO_DATE(ic.fec_ingreso,'%%d/%%m/%%Y') 
                      BETWEEN STR_TO_DATE(%s,'%%d/%%m/%%Y') AND STR_TO_DATE(%s,'%%d/%%m/%%Y')
                ORDER BY ic.num_carta ASC
            """, [desde_formatted, hasta_formatted])
            
            result = cursor.fetchall()
            if result:
                return result
    
    def _get_notary_info(self):
        """Get notary configuration info from database"""
        config = get_notary_config()
        return config["nombre"]
    
    def _format_date_in_spanish(self, date_str):
        """Convert date to Spanish format like 'LUNES, 15 DE ENERO DEL 2025'"""
        try:
            # Parse the date string (assuming YYYY-MM-DD format)
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            
            # Spanish day names
            dias = ['LUNES', 'MARTES', 'MIÉRCOLES', 'JUEVES', 'VIERNES', 'SÁBADO', 'DOMINGO']
            # Spanish month names
            meses = ['ENERO', 'FEBRERO', 'MARZO', 'ABRIL', 'MAYO', 'JUNIO', 
                    'JULIO', 'AGOSTO', 'SEPTIEMBRE', 'OCTUBRE', 'NOVIEMBRE', 'DICIEMBRE']
            
            dia_semana = dias[date_obj.weekday()]
            dia = date_obj.day
            mes = meses[date_obj.month - 1]
            anio = date_obj.year
            
            return f"{dia_semana}, {dia} DE {mes} DEL {anio}"
        except:
            return date_str
    
    def _extract_year_from_date(self, date_str):
        """Extract year from date string DD/MM/YYYY or YYYY-MM-DD"""
        try:
            # Try to parse as YYYY-MM-DD first
            if '-' in date_str and len(date_str.split('-')[0]) == 4:
                return date_str.split('-')[0]
            # Try to parse as DD/MM/YYYY
            elif '/' in date_str:
                return date_str.split('/')[-1]
            else:
                return str(datetime.now().year)
        except:
            return str(datetime.now().year)
    
    def generate_excel_report(self, desde, hasta):
        """Generate Excel report matching PHP script format"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            import json
            
            # Get data
            report_data = self._get_report_data(desde, hasta)
            notary_config = get_notary_config()  # Get config from database
            notary_name = self._get_notary_info()
            anio = self._extract_year_from_date(hasta)
            
            # Create workbook and worksheet - simple approach
            wb = Workbook()
            ws = wb.active
            ws.title = "CARTAS NOTARIALES"
            
            # Styles matching PHP script
            title_font = Font(name='Arial', size=18, bold=True)
            header_font = Font(name='Arial', size=13, bold=True)
            data_font = Font(name='Arial', size=13)
            center_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            left_alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
            
            # Borders for data table only
            thin_border = Border(
                left=Side(border_style="thin"),
                right=Side(border_style="thin"),
                top=Side(border_style="thin"),
                bottom=Side(border_style="thin")
            )
            
            # No borders for header section
            no_border = Border(
                left=Side(style=None),
                right=Side(style=None),
                top=Side(style=None),
                bottom=Side(style=None)
            )
            
            # Title section - merge across columns for better appearance
            ws.merge_cells('A1:F1')
            ws['A1'] = 'INDICE CRONOLOGICO - CARTAS NOTARIALES'
            ws['A1'].font = title_font
            ws['A1'].alignment = center_alignment
            ws['A1'].border = no_border
            
            ws.merge_cells('A2:F2')
            ws['A2'] = f'AÑO {anio}'
            ws['A2'].font = title_font
            ws['A2'].alignment = center_alignment
            ws['A2'].border = no_border
            
            # Notary info section - simple layout, no merging
            row = 4
            ws[f'A{row}'] = 'NOTARIA'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].border = no_border
            ws[f'B{row}'] = f': {self._sanitize_cell_value(notary_name)}'
            ws[f'B{row}'].font = data_font
            ws[f'B{row}'].border = no_border
            
            row += 1
            ws[f'A{row}'] = 'DIRECCION'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].border = no_border
            ws[f'B{row}'] = f': {notary_config["direccion"]}'
            ws[f'B{row}'].font = data_font
            ws[f'B{row}'].border = no_border
            ws[f'D{row}'] = 'TELEFONO'
            ws[f'D{row}'].font = header_font
            ws[f'D{row}'].border = no_border
            ws[f'E{row}'] = f': {notary_config["telefono"]}'
            ws[f'E{row}'].font = data_font
            ws[f'E{row}'].border = no_border
            
            row += 1
            ws[f'A{row}'] = 'DEPARTAMENTO'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].border = no_border
            ws[f'B{row}'] = ': PUNO'
            ws[f'B{row}'].font = data_font
            ws[f'B{row}'].border = no_border
            ws[f'D{row}'] = 'RUC'
            ws[f'D{row}'].font = header_font
            ws[f'D{row}'].border = no_border
            ws[f'E{row}'] = f': {notary_config["ruc"]}'
            ws[f'E{row}'].font = data_font
            ws[f'E{row}'].border = no_border
            
            row += 1
            ws[f'A{row}'] = 'PROVINCIA'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].border = no_border
            ws[f'B{row}'] = ': SAN ROMAN'
            ws[f'B{row}'].font = data_font
            ws[f'B{row}'].border = no_border
            ws[f'D{row}'] = 'DESDE'
            ws[f'D{row}'].font = header_font
            ws[f'D{row}'].border = no_border
            ws[f'E{row}'] = f': {self._format_date_in_spanish(desde)}'
            ws[f'E{row}'].font = data_font
            ws[f'E{row}'].border = no_border
            
            row += 1
            ws[f'A{row}'] = 'DISTRITO'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].border = no_border
            ws[f'B{row}'] = ': JULIACA'
            ws[f'B{row}'].font = data_font
            ws[f'B{row}'].border = no_border
            ws[f'D{row}'] = 'HASTA'
            ws[f'D{row}'].font = header_font
            ws[f'D{row}'].border = no_border
            ws[f'E{row}'] = f': {self._format_date_in_spanish(hasta)}'
            ws[f'E{row}'].font = data_font
            ws[f'E{row}'].border = no_border
            
            # Table headers with borders
            row += 2
            headers = ['NRO', 'FEC. INGRESO', 'DIRECCION ENTREGA', 'FEC. ENTREGA', 'RESULTADO', 'FIRMO']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = header_font
                cell.alignment = center_alignment
                cell.border = thin_border
            
            # Set column widths
            ws.column_dimensions['A'].width = 15  # NRO
            ws.column_dimensions['B'].width = 18  # FEC. INGRESO  
            ws.column_dimensions['C'].width = 50  # DIRECCION ENTREGA
            ws.column_dimensions['D'].width = 18  # FEC. ENTREGA
            ws.column_dimensions['E'].width = 45  # RESULTADO (as you requested)
            ws.column_dimensions['F'].width = 20  # FIRMO
            
            # Data rows
            row += 1
            if report_data:
                for data_row in report_data:
                    try:
                        # Extract and sanitize all data
                        num_carta = self._sanitize_cell_value(data_row[0] if len(data_row) > 0 else '')
                        fec_ingreso = self._sanitize_cell_value(data_row[1] if len(data_row) > 1 else '')
                        fec_entrega = self._sanitize_cell_value(data_row[2] if len(data_row) > 2 else '')
                        dir_destinatario = self._sanitize_cell_value(data_row[7] if len(data_row) > 7 else '')
                        recepcion = self._sanitize_cell_value(data_row[10] if len(data_row) > 10 else '')
                        firmo = self._sanitize_cell_value(data_row[11] if len(data_row) > 11 else '')
                        
                        # Extract correlative number safely
                        correlativo = "1"
                        if num_carta and len(str(num_carta)) >= 6:
                            try:
                                correlativo = str(int(str(num_carta)[-6:]))  # Convert to int then back to str
                            except:
                                correlativo = str(num_carta)[-6:] if len(str(num_carta)) >= 6 else str(num_carta)
                        
                        # Main data row - simple cell assignment
                        ws.cell(row=row, column=1, value=correlativo).font = data_font
                        ws.cell(row=row, column=2, value=fec_ingreso).font = data_font
                        ws.cell(row=row, column=3, value=dir_destinatario.upper()).font = data_font
                        ws.cell(row=row, column=4, value=fec_entrega).font = data_font
                        ws.cell(row=row, column=5, value=recepcion.upper()).font = data_font
                        ws.cell(row=row, column=6, value=firmo.upper()).font = data_font
                        
                        # Apply borders and alignment
                        for col in range(1, 7):
                            cell = ws.cell(row=row, column=col)
                            cell.border = thin_border
                            if col in [1, 2, 4, 6]:  # Center align NRO, dates, FIRMO
                                cell.alignment = center_alignment
                            else:  # Left align DIRECCION and RESULTADO
                                cell.alignment = left_alignment
                        
                        # Set row height
                        ws.row_dimensions[row].height = 40
                        
                        # Additional info row - simple approach
                        row += 1
                        remitente = self._sanitize_cell_value(data_row[5] if len(data_row) > 5 else '')
                        destinatario = self._sanitize_cell_value(data_row[4] if len(data_row) > 4 else '')
                        dni_remitente = self._sanitize_cell_value(data_row[8] if len(data_row) > 8 else '')
                        dni_destinatario = self._sanitize_cell_value(data_row[9] if len(data_row) > 9 else '')
                        
                        # Simple cell assignment without complex merging
                        ws.cell(row=row, column=1, value='').font = data_font
                        ws.cell(row=row, column=2, value='REMITENTE:\nDESTINATARIO:').font = data_font
                        ws.cell(row=row, column=3, value=f'{remitente}\n{destinatario}').font = data_font
                        ws.cell(row=row, column=4, value='').font = data_font
                        ws.cell(row=row, column=5, value=f'DNI: {dni_remitente}\nDNI: {dni_destinatario}').font = data_font
                        ws.cell(row=row, column=6, value='').font = data_font
                        
                        # Apply borders to info row
                        for col in range(1, 7):
                            cell = ws.cell(row=row, column=col)
                            cell.border = thin_border
                            if col == 2:  # REMITENTE/DESTINATARIO
                                cell.alignment = center_alignment
                            else:
                                cell.alignment = left_alignment
                        
                        ws.row_dimensions[row].height = 40
                        row += 1
                        
                    except Exception as e:
                        # Error handling - add error row instead of crashing
                        ws.cell(row=row, column=1, value="ERROR").font = data_font
                        ws.cell(row=row, column=2, value=f"Error: {str(e)[:50]}").font = data_font
                        for col in range(1, 7):
                            ws.cell(row=row, column=col).border = thin_border
                        row += 1
            else:
                # No data message
                ws.cell(row=row, column=1, value='No se encontraron registros').font = data_font
                for col in range(1, 7):
                    ws.cell(row=row, column=col).border = thin_border
            
            # Save to BytesIO
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            
            # Create HTTP response
            filename = f"INDICE_CRONOLOGICO_CARTAS_NOTARIALES_{anio}.xlsx"
            response = HttpResponse(
                output.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
            
        except Exception as e:
            return HttpResponse(
                json.dumps({'error': f'Error generating Excel report: {str(e)}'}),
                content_type='application/json',
                status=500
            )
    
    def generate_word_report(self, desde, hasta):
        """Generate Word report matching PHP script format"""
        try:
            # Get data
            report_data = self._get_report_data(desde, hasta)
            notary_config = get_notary_config()  # Get config from database
            notary_name = self._get_notary_info()
            
            # Create a new document
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            # Title - matching PHP font-size:18.5px
            title = doc.add_heading('INDICE CRONOLOGICO - CARTAS NOTARIALES', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(18.5)
                run.font.bold = True
            
            # Year
            anio = self._extract_year_from_date(hasta)
            year_heading = doc.add_heading(f'AÑO {anio}', 0)
            year_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in year_heading.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(18.5)
                run.font.bold = True
            
            # Add spacing
            doc.add_paragraph()
            doc.add_paragraph()
            
            # Notary info table - NO BORDERS
            info_table = doc.add_table(rows=5, cols=6)
            # No table style = no borders
            
            # Row 1: NOTARIA
            row = info_table.rows[0]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = 'NOTARIA'
            row.cells[2].text = f': {notary_name}'
            
            # Row 2: DIRECCION
            row = info_table.rows[1]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = 'DIRECCION'
            row.cells[2].text = f': {notary_config["direccion"]}'
            row.cells[3].text = 'TELEFONO'
            row.cells[4].merge(row.cells[5])
            row.cells[4].text = f': {notary_config["telefono"]}'
            
            # Row 3: DEPARTAMENTO
            row = info_table.rows[2]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = 'DEPARTAMENTO'
            row.cells[2].text = ': PUNO'
            row.cells[3].text = 'RUC'
            row.cells[4].merge(row.cells[5])
            row.cells[4].text = f': {notary_config["ruc"]}'
            
            # Row 4: PROVINCIA
            row = info_table.rows[3]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = 'PROVINCIA'
            row.cells[2].text = ': SAN ROMAN'
            row.cells[3].text = 'DESDE'
            row.cells[4].merge(row.cells[5])
            row.cells[4].text = f': {self._format_date_in_spanish(desde)}'
            
            # Row 5: DISTRITO
            row = info_table.rows[4]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = 'DISTRITO'
            row.cells[2].text = ': JULIACA'
            row.cells[3].text = 'HASTA'
            row.cells[4].merge(row.cells[5])
            row.cells[4].text = f': {self._format_date_in_spanish(hasta)}'
            
            # Style the info table - Simple styling without borders
            for row in info_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if ':' in run.text:
                                run.font.bold = True
                            run.font.size = Pt(12)
            
            # Add spacing
            doc.add_paragraph()
            
            # Main data table - ALWAYS CREATE, even if empty
            # Create table with headers
            headers = ['NRO', 'FEC. INGRESO', 'DIRECCION ENTREGA', 'FEC. ENTREGA', 'RESULTADO', 'FIRMO']
            data_table = doc.add_table(rows=1, cols=6)
            data_table.style = 'Table Grid'
            
            # Add headers
            header_row = data_table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(12)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add data rows if data exists
            if report_data and len(report_data) > 0:
                for data_row in report_data:
                    
                    num_carta = data_row[0] if len(data_row) > 0 else ''
                    fec_ingreso = data_row[1] if len(data_row) > 1 else ''
                    fec_entrega = data_row[2] if len(data_row) > 2 else ''
                    dir_destinatario = data_row[7] if len(data_row) > 7 else ''
                    recepcion = data_row[10] if len(data_row) > 10 else ''  # Fixed: was 11
                    firmo = data_row[11] if len(data_row) > 11 else ''     # Fixed: was 12
                    
                    # Extract correlative number
                    correlativo = str(num_carta)[-6:] if num_carta else ''
                    
                    # Row 1: Main data
                    row = data_table.add_row()
                    row.cells[0].text = correlativo
                    row.cells[1].text = fec_ingreso
                    row.cells[2].text = dir_destinatario.upper()
                    row.cells[3].text = fec_entrega
                    row.cells[4].text = recepcion.upper()
                    row.cells[5].text = firmo.upper()
                    
                    # Center align main data
                    for i in range(6):
                        for paragraph in row.cells[i].paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Row 2: Additional info
                    row = data_table.add_row()
                    remitente = data_row[5] if len(data_row) > 5 else ''      # Fixed: was 5
                    destinatario = data_row[4] if len(data_row) > 4 else ''    # Fixed: was 4
                    dni_remitente = data_row[8] if len(data_row) > 8 else ''  # Fixed: was 8
                    dni_destinatario = data_row[9] if len(data_row) > 9 else '' # Fixed: was 9
                    
                    row.cells[0].text = ''
                    row.cells[1].text = 'REMITENTE:\nDESTINATARIO:'
                    row.cells[1].merge(row.cells[2])
                    row.cells[1].text = f'{remitente}\n{destinatario}'
                    row.cells[3].text = ''
                    row.cells[4].text = f'DNI: {dni_remitente}\nDNI: {dni_destinatario}'
                    row.cells[5].text = ''
                    
                    # Style the info row
                    for i in range(6):
                        if i == 1:  # REMITENTE/DESTINATARIO column
                            for paragraph in row.cells[i].paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif i == 4:  # DNI column
                            for paragraph in row.cells[i].paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            for paragraph in row.cells[i].paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # Add a "no data" row if no data exists
                no_data_row = data_table.add_row()
                no_data_cell = no_data_row.cells[0]
                no_data_cell.merge(no_data_row.cells[5])
                no_data_cell.text = 'No se encontraron registros para el período especificado'
                for paragraph in no_data_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.italic = True
            
            # Style all cells
            for row in data_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(12)
            
            # Save to BytesIO
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            # Create HTTP response
            filename = f"INDICE_CRONOLOGICO_CARTAS_NOTARIALES_{anio}.docx"  # Use correct .docx extension
            response = HttpResponse(
                output.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            return response
            
        except Exception as e:
            return Response(
                {'error': f'Error generating Word report: {str(e)}'}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )