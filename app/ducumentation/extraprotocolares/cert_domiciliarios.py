import os
import io
import traceback
from typing import Dict, Any, Optional

from django.db import connection
from django.http import HttpResponse, JsonResponse
from docxtpl import DocxTemplate

from ..shared.base_r2_documents import get_s3_client, BaseR2DocumentService
from ..utils import NumberToLetterConverter
from ..protocolares.utils import get_notary_config


class CertDomiciliariosDocumentService(BaseR2DocumentService):
    """
    Service to generate and retrieve Certificación/Constatación Domiciliaria documents.

    - Template expected in R2: 'CERTIFICADO DOMICILIARIO BASE.docx'
    - Output filename: '__CDOM__{RIGHT6(num_certificado)}-{LEFT4(num_certificado)}.docx'
    - Stored under: {os.environ.get('CLOUDFLARE_R2_MAIN_URL')}/documentos/
    """

    def __init__(self) -> None:
        self.letras = NumberToLetterConverter()
        self.template_filename = "CERTIFICADO DOMICILIARIO BASE.docx"

    def retrieve_cdom_document(self, num_certificado: str, mode: str = "download") -> HttpResponse:
        try:
            if not num_certificado:
                return HttpResponse("Error: num_certificado is required to retrieve document", status=400)

            formatted = self._format_num_certificado(num_certificado)
            filename = f"__CDOM__{formatted}.docx"

            if mode == "open":
                return self._create_response(None, filename, formatted, mode)

            s3 = get_s3_client()
            object_key = self._object_key_for_document(filename)
            response = s3.get_object(Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'), Key=object_key)
            buffer = io.BytesIO(response['Body'].read())
            return self._create_response(buffer, filename, formatted, mode)
        except Exception as e:
            if hasattr(e, 'response') and isinstance(getattr(e, 'response'), dict):
                if e.response.get('Error', {}).get('Code') == 'NoSuchKey':
                    return HttpResponse(f"Error: Document '{filename}' not found in R2.", status=404)
            traceback.print_exc()
            return HttpResponse(f"Error retrieving document: {e}", status=500)

    def generate_cdom_document(self, num_certificado: str, mode: str = "download", id_domiciliario: Optional[str] = None) -> HttpResponse:
        try:
            if not num_certificado:
                return JsonResponse({'status': 'error', 'message': 'num_certificado is required'}, status=400)

            formatted = self._format_num_certificado(num_certificado)
            filename = f"__CDOM__{formatted}.docx"

            # if self._document_exists_in_r2(filename):
            #     return self.json_error(409, "Document already exists. Use action=retrieve to fetch it.", {
            #         'num_certificado': num_certificado,
            #         'filename': filename,
            #     })

            template_bytes = self._get_template_from_r2()
            if template_bytes is None:
                template_key = self._object_key_for_template(self.template_filename)
                return self.json_error(404, f"Template '{self.template_filename}' not found in '{os.environ.get('CLOUDFLARE_R2_MAIN_URL')}/plantillas/'.")

            cert_data = self._get_cert_data(num_certificado, id_domiciliario)
            if not cert_data:
                return self.json_error(404, f"cert_domiciliario record with num_certificado {num_certificado} not found")

            # Build context
            context: Dict[str, Any] = {}
            context.update(cert_data)
            context.update(self._get_notary_data())

            # Aliases used by templates (lowercase, case-sensitive)
            context['numcrono2'] = formatted
            context['P_NOM'] = context.get('NOMBRE_SOLIC', '')
            # Gendered identity strings
            sex = (context.get('SEXO', '') or 'M').upper()
            tip_doc = context.get('TIP_DOC', '')
            num_doc = context.get('NUM_DOC', '')
            context['P_DOC'] = f"{'IDENTIFICADA' if sex == 'F' else 'IDENTIFICADO'} CON {tip_doc} N°"
            context['DOC'] = f"{tip_doc} N°"
            context['P_IDE'] = num_doc
            
            # Domicile text - match PHP logic exactly
            domicilio = context.get('DIRECCION', '')
            distrito_texto = context.get('DISTRITO_TEXTO', '')
            context['P_DOMICILIO'] = f"CON DOMICILIO EN {domicilio} {distrito_texto}".strip()
            
            # Civil status and nationality with gender handling (matching PHP logic)
            est_civil = context.get('E_CIVIL', '') or ''
            nacionalidad = context.get('NACIONALIDAD', '') or ''
            
            # Gender-based variables (matching PHP)
            if sex == 'F':
                context['EL_LA'] = 'LA'
                context['DEL_A'] = 'DE LA' 
                context['P_O_A'] = 'A'
                context['A_EL'] = 'A LA'
            else:  # sex == 'M'
                context['EL_LA'] = 'EL'
                context['DEL_A'] = 'DEL'
                context['P_O_A'] = 'O'
                context['A_EL'] = 'AL'
            
            # Handle civil status with proper gender endings (PHP logic)
            # Note: E_CIVIL appears to be an ID, we need to map it to actual text
            if est_civil and str(est_civil).strip() and str(est_civil) != '0':
                # For now, provide common civil status mappings
                civil_status_map = {
                    '1': 'SOLTERO',
                    '2': 'CASADO', 
                    '3': 'DIVORCIADO',
                    '4': 'VIUDO',
                    '5': 'CONVIVIENTE'
                }
                est_civil_text = civil_status_map.get(str(est_civil), 'SOLTERO')
                if len(est_civil_text) > 1:
                    context['P_ESTADO_CIVIL'] = est_civil_text[:-1] + ('A' if sex == 'F' else 'O')
                else:
                    context['P_ESTADO_CIVIL'] = est_civil_text
            else:
                context['P_ESTADO_CIVIL'] = 'SOLTERA' if sex == 'F' else 'SOLTERO'
            
            # Handle nationality with proper gender endings (PHP logic)
            if nacionalidad and str(nacionalidad).strip():
                nacionalidad_str = str(nacionalidad).strip().upper()
                if len(nacionalidad_str) > 1:
                    context['P_NACIONALIDAD'] = nacionalidad_str[:-1] + ('A' if sex == 'F' else 'O')
                else:
                    context['P_NACIONALIDAD'] = nacionalidad_str
            else:
                context['P_NACIONALIDAD'] = 'PERUANA' if sex == 'F' else 'PERUANO'
            
            context['P_OCUPACION'] = context.get('PROFESION', '') or ''

            # Testigo line (optional)
            if context.get('NOM_TESTIGO'):
                context['DATOS_TESTIGO'] = (
                    f"INTERVIENE EN CALIDAD DE TESTIGO A RUEGO :{context.get('NOM_TESTIGO', '')}, CON "
                    f"{context.get('TIPDOC_TESTIGO', '')} NUMERO {context.get('NUMDOC_TESTIGO', '')} "
                    f"{context.get('UBIGEO_TESTIGO', '')}"
                )
            else:
                context['DATOS_TESTIGO'] = ''

            # Testigo aliases matching <I_...> tags in legacy template
            context['I_NOM'] = context.get('NOM_TESTIGO', '')
            context['I_NACIONALIDAD'] = ''
            context['I_DOC'] = 'IDENTIFICADO CON'
            context['DOC_I'] = f"{context.get('TIPDOC_TESTIGO', '')} N°".strip()
            context['I_IDE'] = context.get('NUMDOC_TESTIGO', '')
            context['I_OCUPACION'] = ''
            context['I_ESTADO_CIVIL'] = ''
            context['I_DOMICILIO'] = (f"CON DOMICILIO EN {context.get('UBIGEO_TESTIGO', '')}").strip()
            context['I_ROGADO'] = 'QUIEN INTERVIENE EN CALIDAD DE TESTIGO A RUEGO'

            # Date alias for template expecting fec_letras_completa
            context['fec_letras_completa'] = context.get('FECHA_INGRESO_LETRAS', '')
            
            # Additional aliases to match PHP variable names
            context['MOTIVO'] = context.get('MOTIVO', '') or 'trámites varios'
            context['OBSERVACION'] = context.get('OBSERVACION', '') or ''

            # Additional fields for certificate details (matching PHP variable names)
            context['DECLARA_SER'] = context.get('declara_ser', '') or ''
            context['PROPIETARIO'] = context.get('propietario', '') or ''
            context['RECIBIDO_POR'] = context.get('recibido', '') or ''
            context['NRO_RECIBO_SERVICIOS'] = context.get('numero_recibo', '') or ''
            context['MES_RECIBO_SERVICIOS'] = context.get('mes_facturado', '') or ''
            context['RECIBO_SERVICIOS'] = context.get('recibo_empresa', '') or ''
            
            # Handle service type description (matching PHP logic)
            recibo_empresa = context.get('recibo_empresa', '') or ''
            if 'SEDA JULIACA' in recibo_empresa.upper():
                context['RECIBO_SERVICIOS_D'] = 'saneamiento y agua potable'
            elif 'ELECTRO PUNO' in recibo_empresa.upper():
                context['RECIBO_SERVICIOS_D'] = 'servicios de luz eléctrica'
            else:
                context['RECIBO_SERVICIOS_D'] = 'servicios básicos'
            
            # Format occupation date if available
            fecha_ocupa = context.get('fecha_ocupa')
            if fecha_ocupa:
                try:
                    context['FECHA_OCUPA_LETRAS'] = self.letras.date_to_letters(fecha_ocupa).upper()
                except Exception:
                    context['FECHA_OCUPA_LETRAS'] = ''
            else:
                context['FECHA_OCUPA_LETRAS'] = ''
            
            # Signature section logic (matching PHP)
            nom_testigo = context.get('NOM_TESTIGO', '') or ''
            if nom_testigo and nom_testigo.strip():
                # With witness
                context['evalua_firma'] = f"-----------------------------------\nHUELLA DEL SOLICITANTE\n{context.get('NOMBRE_SOLIC', '')}\n{tip_doc} N°: {num_doc}\n\n\n\n"
                context['evalua_firma_testigo'] = f"-----------------------------------\n{nom_testigo}\n{context.get('TIPDOC_TESTIGO', '')} N°: {context.get('NUMDOC_TESTIGO', '')}\n\n\n\n"
            else:
                # Without witness
                context['evalua_firma'] = f"-----------------------------------\n{context.get('NOMBRE_SOLIC', '')}\n{tip_doc} N°: {num_doc}\n\n\n\n"
                context['evalua_firma_testigo'] = ""

            # Render and save
            doc = DocxTemplate(io.BytesIO(template_bytes))
            doc.render(context)

            buffer = io.BytesIO()
            doc.save(buffer)
            self._save_document_to_r2(buffer, filename)
            return self._create_response(buffer, filename, formatted, mode)
        except Exception as e:
            traceback.print_exc()
            return self.json_error(500, f"Error generating document: {e}")

    def _get_template_from_r2(self) -> Optional[bytes]:
        s3 = get_s3_client()
        object_key = self._object_key_for_template(self.template_filename)
        try:
            response = s3.get_object(Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'), Key=object_key)
            return response['Body'].read()
        except Exception as e:
            return None

    def _save_document_to_r2(self, buffer: io.BytesIO, filename: str) -> None:
        s3 = get_s3_client()
        object_key = self._object_key_for_document(filename)
        buffer.seek(0)
        s3.put_object(
            Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'),
            Key=object_key,
            Body=buffer.read(),
        )
        buffer.seek(0)

    def _create_response(self, buffer: Optional[io.BytesIO], filename: str, key_id: str, mode: str = "download") -> HttpResponse:
        if mode == "open":
            s3 = get_s3_client()
            object_key = self._object_key_for_document(filename)
            try:
                url = s3.generate_presigned_url(
                    'get_object',
                    Params={'Bucket': os.environ.get('CLOUDFLARE_R2_BUCKET'), 'Key': object_key},
                    ExpiresIn=3600,
                )
                response = JsonResponse({
                    'status': 'success', 'mode': 'open', 'url': url,
                    'filename': filename, 'num_certificado': key_id,
                    'message': 'Document is ready to be opened.'
                })
                response['Access-Control-Allow-Origin'] = '*'
                return response
            except Exception as e:
                return HttpResponse(f"Error generating pre-signed URL: {e}", status=500)
        if buffer is None:
            return HttpResponse("Error: Document buffer is missing for download mode.", status=500)
        buffer.seek(0)
        response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'inline; filename="{filename}"'
        response['Content-Length'] = str(buffer.getbuffer().nbytes)
        response['Access-Control-Allow-Origin'] = '*'
        return response

    def _format_num_certificado(self, raw: Optional[str]) -> str:
        if not raw or len(raw) < 6:
            return raw or ''
        return f"{raw[-6:]}-{raw[:4]}"

    def _get_notary_data(self) -> Dict[str, str]:
        with connection.cursor() as cursor:
            cursor.execute("SELECT CONCAT(nombre, ' ', apellido) AS notario, direccion, distrito FROM confinotario")
            row = cursor.fetchone()
            if row:
                return {
                    'NOTARIO': str(row[0]).upper() if row[0] else '',
                    'DIRECCION_NOTARIO': str(row[1]).upper() if row[1] else '',
                    'UBIGEO_NOTARIO': str(row[2]).upper() if row[2] else '',
                }
        return {'NOTARIO': '', 'DIRECCION_NOTARIO': '', 'UBIGEO_NOTARIO': ''}

    def _get_cert_data(self, num_certificado: str, id_domiciliario: Optional[str] = None) -> Dict[str, Any]:
        data: Dict[str, Any] = {}
        with connection.cursor() as cursor:
            where_clause = (
                "WHERE cd.id_domiciliario = %s"
                if id_domiciliario
                else """WHERE CONVERT(cd.num_certificado USING utf8mb4) COLLATE utf8mb4_unicode_ci =
                      CONVERT(CAST(%s AS CHAR) USING utf8mb4) COLLATE utf8mb4_unicode_ci"""
            )
            params = [id_domiciliario] if id_domiciliario else [num_certificado]
            cursor.execute(
                f"""
                SELECT 
                    UPPER(cd.num_certificado) AS NUM_CERTI,
                    cd.fec_ingreso AS FEC_INGRESO,
                    UPPER(cd.num_formu) AS NUM_FORMU,
                    CONCAT(c.prinom,' ',c.segnom,' ',c.apepat,' ',c.apemat) AS NOMBRE_SOLIC,
                    UPPER(td.td_abrev) AS TIP_DOC,
                    UPPER(cd.numdoc_solic) AS NUM_DOC,
                    UPPER(cd.domic_solic) AS DIRECCION,
                    UPPER(cd.motivo_solic) AS MOTIVO,
                    UPPER(u.nomdis) AS NOM_DIST,
                    UPPER(cd.texto_cuerpo) AS OBSERVACION,
                    CASE WHEN u.coddis='070101' THEN 'DISTRITO DE CALLAO , PROVINCIA CONSTITUCIONAL DEL CALLAO'
                         ELSE CONCAT('DISTRITO DE ',u.nomdis, ', PROVINCIA DE ',u.nomprov,', DEPARTAMENTO DE ',u.nomdpto) END AS UBIGEO,
                    cd.IDESTCIVIL AS E_CIVIL,
                    cd.profesionc AS IDPROFESION,
                    cd.detprofesionc AS PROFESION,
                    CONCAT(' DEL DISTRITO DE ',u.nomdis,' PROVINCIA DE ',u.nomprov,' Y DEPARTAMENTO DE ',u.nomdpto) AS DISTRITO_TEXTO,
                    cd.id_domiciliario,
                    cd.fecha_ocupa,
                    cd.declara_ser,
                    cd.propietario,
                    cd.recibido,
                    cd.numero_recibo,
                    cd.mes_facturado,
                    cd.recibo_empresa,
                    c.sexo AS SEXO,
                    n.descripcion AS NACIONALIDAD,
                    -- Testigo fields (optional) via correlated selects
                    cd.nom_testigo AS NOM_TESTIGO,
                    (SELECT UPPER(td2.td_abrev)
                     FROM tipodocumento td2
                     WHERE CONVERT(td2.codtipdoc USING utf8mb4) COLLATE utf8mb4_unicode_ci =
                           CONVERT(cd.tdoc_testigo USING utf8mb4) COLLATE utf8mb4_unicode_ci) AS TIPDOC_TESTIGO,
                    cd.ndocu_testigo AS NUMDOC_TESTIGO,
                    CASE WHEN u.coddis='070101' THEN 'DISTRITO DE CALLAO , PROVINCIA CONSTITUCIONAL DEL CALLAO'
                         ELSE CONCAT('DISTRITO DE ',u.nomdis, ', PROVINCIA DE ',u.nomprov,', DEPARTAMENTO DE ',u.nomdpto) END AS UBIGEO_TESTIGO
                FROM cert_domiciliario cd
                LEFT JOIN tipodocumento td ON CONVERT(cd.tipdoc_solic USING utf8mb4) COLLATE utf8mb4_unicode_ci =
                                              CONVERT(td.codtipdoc USING utf8mb4) COLLATE utf8mb4_unicode_ci
                LEFT JOIN cliente c ON CONVERT(cd.numdoc_solic USING utf8mb4) COLLATE utf8mb4_unicode_ci =
                                       CONVERT(c.numdoc USING utf8mb4) COLLATE utf8mb4_unicode_ci
                LEFT JOIN ubigeo u ON CONVERT(u.coddis USING utf8mb4) COLLATE utf8mb4_unicode_ci =
                                      CONVERT(cd.distrito_solic USING utf8mb4) COLLATE utf8mb4_unicode_ci
                LEFT JOIN nacionalidades n ON CAST(c.nacionalidad AS UNSIGNED) = n.idnacionalidad
                {where_clause}
                LIMIT 1
                """,
                params,
            )
            row = cursor.fetchone()
            if not row:
                return {}
            cols = [col[0] for col in cursor.description]
            data = {k: (str(v).upper() if isinstance(v, str) and v is not None else v) for k, v in dict(zip(cols, row)).items()}

            # Additional date formatting
            fec_ingreso = data.get('FEC_INGRESO')
            if fec_ingreso:
                try:
                    data['FECHA_INGRESO_LETRAS'] = self.letras.date_to_letters(fec_ingreso).upper()
                except Exception:
                    data['FECHA_INGRESO_LETRAS'] = ''
            else:
                data['FECHA_INGRESO_LETRAS'] = ''
        return data 


class CertDomiciliariosReportService:
    """Service for generating cert_domiciliario reports matching PHP script format"""
    
    def _sanitize_cell_value(self, value):
        """Sanitize cell values to prevent Excel corruption"""
        if value is None:
            return ""
        
        # Convert to string and strip whitespace
        val_str = str(value).strip()
        
        # Remove control characters that can cause XML corruption
        import re
        val_str = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', val_str)
        
        # Limit to Excel's cell character limit
        val_str = val_str[:32767]
        
        return val_str
    
    def _get_report_data(self, desde, hasta):
        """Fetch data for the report matching PHP query"""
        try:
            from django.db import connection
            from notaria.models import CertDomiciliario
            
            
            # Convert dates to proper format for Django ORM
            if isinstance(desde, str):
                if '-' in desde and len(desde.split('-')[0]) == 4:
                    # It's YYYY-MM-DD format, convert to datetime
                    from datetime import datetime
                    desde_dt = datetime.strptime(desde, '%Y-%m-%d')
                else:
                    # It's DD/MM/YYYY format, convert to datetime
                    from datetime import datetime
                    desde_dt = datetime.strptime(desde, '%d/%m/%Y')
            else:
                desde_dt = desde
            
            if isinstance(hasta, str):
                if '-' in hasta and len(hasta.split('-')[0]) == 4:
                    # It's YYYY-MM-DD format, convert to datetime
                    from datetime import datetime
                    hasta_dt = datetime.strptime(hasta, '%Y-%m-%d')
                else:
                    # It's DD/MM/YYYY format, convert to datetime
                    from datetime import datetime
                    hasta_dt = datetime.strptime(hasta, '%d/%m/%Y')
            else:
                hasta_dt = hasta
            
            
            # Use Django ORM like the working list method
            queryset = CertDomiciliario.objects.filter(
                fec_ingreso__range=(desde_dt, hasta_dt)
            ).order_by('num_certificado')
            
            # Convert to the format expected by the report
            result = []
            for record in queryset:
                row = (
                    record.num_certificado,
                    record.fec_ingreso,
                    record.nombre_solic,
                    record.numdoc_solic,
                    record.domic_solic,
                    record.motivo_solic,
                    record.recibo_empresa,
                    record.numero_recibo
                )
                result.append(row)
            
            return result
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            return []
    
    def _get_notary_info(self):
        """Get notary configuration info from database"""
        config = get_notary_config()
        return config["nombre"]
    
    def _format_date_in_spanish(self, date_input):
        """Convert date to Spanish format like 'LUNES, 15 DE ENERO DEL 2025'"""
        try:
            # Handle both datetime objects and date strings
            if hasattr(date_input, 'strftime'):
                # It's a datetime object
                date_obj = date_input
            else:
                # It's a string, try to parse it
                from datetime import datetime
                # Try different formats
                if '-' in str(date_input) and len(str(date_input).split('-')[0]) == 4:
                    date_obj = datetime.strptime(str(date_input), '%Y-%m-%d')
                elif '/' in str(date_input):
                    date_obj = datetime.strptime(str(date_input), '%d/%m/%Y')
                else:
                    return str(date_input)
            
            # Spanish day names
            dias = ['LUNES', 'MARTES', 'MIÉRCOLES', 'JUEVES', 'VIERNES', 'SÁBADO', 'DOMINGO']
            # Spanish month names
            meses = ['ENERO', 'FEBRERO', 'MARZO', 'ABRIL', 'MAYO', 'JUNIO', 
                    'JULIO', 'AGOSTO', 'SEPTIEMBRE', 'OCTUBRE', 'NOVIEMBRE', 'DICIEMBRE']
            
            dia_semana = dias[date_obj.weekday()]
            dia = date_obj.day
            mes = meses[date_obj.month - 1]
            anio = date_obj.year
            
            return f"{dia_semana}, {dia} DE {mes} DEL {anio}"
        except:
            return str(date_input)
    
    def _extract_year_from_date(self, date_input):
        """Extract year from date string DD/MM/YYYY or YYYY-MM-DD or datetime object"""
        try:
            # Handle datetime objects
            if hasattr(date_input, 'year'):
                return str(date_input.year)
            
            # Handle strings
            date_str = str(date_input)
            # Try to parse as YYYY-MM-DD first
            if '-' in date_str and len(date_str.split('-')[0]) == 4:
                return date_str.split('-')[0]
            # Try to parse as DD/MM/YYYY
            elif '/' in date_str:
                return date_str.split('/')[-1]
            else:
                from datetime import datetime
                return str(datetime.now().year)
        except:
            from datetime import datetime
            return str(datetime.now().year)
    
    def _format_date_for_display(self, date_obj):
        """Format date like PHP script logic"""
        try:
            if hasattr(date_obj, 'strftime'):
                return date_obj.strftime('%d/%m/%Y')
            return str(date_obj)
        except:
            return str(date_obj)
    
    def _format_recibo_type(self, recibo):
        """Format recibo type like PHP script logic"""
        if recibo == 'SEDA JULIACA S.A.':
            return 'RECIBO DE AGUA'
        elif recibo == 'ELECTRO PUNO S.A.A':
            return 'RECIBO DE LUZ'
        else:
            return recibo or ''
    
    def generate_excel_report(self, desde, hasta):
        """Generate Excel report matching PHP script format"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            import json
            
            # Get data
            report_data = self._get_report_data(desde, hasta)
            notary_config = get_notary_config()  # Get config from database
            notary_name = self._get_notary_info()
            anio = self._extract_year_from_date(hasta)
            
            # Create workbook and worksheet
            wb = Workbook()
            ws = wb.active
            ws.title = "CERTIFICADO DOMICILIARIO"
            
            # Styles matching PHP script
            title_font = Font(name='Arial', size=18.5, bold=True)
            header_font = Font(name='Arial', size=13.5, bold=True)
            data_font = Font(name='Arial', size=13.5)
            center_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            left_alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
            
            # Borders for data table
            thin_border = Border(
                left=Side(border_style="thin"),
                right=Side(border_style="thin"),
                top=Side(border_style="thin"),
                bottom=Side(border_style="thin")
            )
            
            # No borders for header section
            no_border = Border(
                left=Side(style=None),
                right=Side(style=None),
                top=Side(style=None),
                bottom=Side(style=None)
            )
            
            # Title section
            ws.merge_cells('A1:I1')
            ws['A1'] = 'INDICE CRONOLOGICO - CERTIFICADO DOMICILIARIO'
            ws['A1'].font = title_font
            ws['A1'].alignment = center_alignment
            ws['A1'].border = no_border
            
            ws.merge_cells('A2:I2')
            ws['A2'] = f'AÑO {anio}'
            ws['A2'].font = title_font
            ws['A2'].alignment = center_alignment
            ws['A2'].border = no_border
            
            # Notary info section
            row = 4
            ws[f'A{row}'] = 'NOTARIA'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].border = no_border
            ws[f'C{row}'] = f': {self._sanitize_cell_value(notary_name)}'
            ws[f'C{row}'].font = data_font
            ws[f'C{row}'].border = no_border
            
            row += 1
            ws[f'A{row}'] = 'DIRECCION'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].border = no_border
            ws[f'C{row}'] = f': {notary_config["direccion"]}'
            ws[f'C{row}'].font = data_font
            ws[f'C{row}'].border = no_border
            ws[f'F{row}'] = 'TELEFONO'
            ws[f'F{row}'].font = header_font
            ws[f'F{row}'].border = no_border
            ws[f'H{row}'] = f': {notary_config["telefono"]}'
            ws[f'H{row}'].font = data_font
            ws[f'H{row}'].border = no_border
            
            row += 1
            ws[f'A{row}'] = 'DEPARTAMENTO'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].border = no_border
            ws[f'C{row}'] = f': {notary_config["departamento"]}'
            ws[f'C{row}'].font = data_font
            ws[f'C{row}'].border = no_border
            ws[f'F{row}'] = 'RUC'
            ws[f'F{row}'].font = header_font
            ws[f'F{row}'].border = no_border
            ws[f'H{row}'] = f': {notary_config["ruc"]}'
            ws[f'H{row}'].font = data_font
            ws[f'H{row}'].border = no_border
            
            row += 1
            ws[f'A{row}'] = 'PROVINCIA'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].border = no_border
            ws[f'C{row}'] = f': {notary_config["provincia"]}'
            ws[f'C{row}'].font = data_font
            ws[f'C{row}'].border = no_border
            ws[f'F{row}'] = 'DESDE'
            ws[f'F{row}'].font = header_font
            ws[f'F{row}'].border = no_border
            ws[f'H{row}'] = f': {self._format_date_in_spanish(desde).upper()}'
            ws[f'H{row}'].font = data_font
            ws[f'H{row}'].border = no_border
            
            row += 1
            ws[f'A{row}'] = 'DISTRITO'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].border = no_border
            ws[f'C{row}'] = f': {notary_config["distrito"]}'
            ws[f'C{row}'].font = data_font
            ws[f'C{row}'].border = no_border
            ws[f'F{row}'] = 'HASTA'
            ws[f'F{row}'].font = header_font
            ws[f'F{row}'].border = no_border
            ws[f'H{row}'] = f': {self._format_date_in_spanish(hasta).upper()}'
            ws[f'H{row}'].font = data_font
            ws[f'H{row}'].border = no_border
            
            # Data table headers
            row += 2
            headers = ['N°', 'FECHA', 'SOLICITANTE', 'N° DNI', 'DOMICILIO', 'MOTIVO', 'DOCUMENTO VERIFICADO', 'COMPROBANTE']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = header_font
                cell.alignment = center_alignment
                cell.border = thin_border
            
            # Data rows
            for data_row in report_data:
                row += 1
                kardex, fecha, solicitante, documento_solicitante, domicilio_solicitante, motivo_solicitante, recibo, numero_recibo = data_row
                
                # Format kardex (substr from position 4 like PHP)
                kardex_formatted = str(kardex)[3:] if len(str(kardex)) > 3 else str(kardex)
                
                # Format recibo type
                recibo_formatted = self._format_recibo_type(recibo)
                
                # Row data
                row_data = [
                    kardex_formatted,
                    self._format_date_for_display(fecha),
                    self._sanitize_cell_value(solicitante),
                    documento_solicitante,
                    self._sanitize_cell_value(domicilio_solicitante),
                    self._sanitize_cell_value(motivo_solicitante),
                    recibo_formatted,
                    numero_recibo
                ]
                
                for col, value in enumerate(row_data, 1):
                    cell = ws.cell(row=row, column=col, value=value)
                    cell.font = data_font
                    cell.border = thin_border
                    
                    # Right align numbers, left align text
                    if col in [1, 2, 4, 8]:  # N°, FECHA, N° DNI, COMPROBANTE
                        cell.alignment = Alignment(horizontal='right', vertical='top', wrap_text=True)
                    else:
                        cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
            
            # Auto-adjust column widths
            for column in ws.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
            
            # Save to buffer
            buffer = io.BytesIO()
            wb.save(buffer)
            buffer.seek(0)
            
            # Create response
            from django.http import HttpResponse
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = f'attachment; filename=INDICE_CRONOLOGICO_CERTIFICADO_DOMICILIARIO_{anio}.xlsx'
            response['Access-Control-Allow-Origin'] = '*'
            
            return response
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            from django.http import HttpResponse
            return HttpResponse(f"Error generating Excel report: {e}", status=500)
    
    def generate_word_report(self, desde, hasta):
        """Generate Word report matching PHP script format"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            
            # Get data
            report_data = self._get_report_data(desde, hasta)
            notary_config = get_notary_config()  # Get config from database
            notary_name = self._get_notary_info()
            anio = self._extract_year_from_date(hasta)
            
            # Create document
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            # Title
            title = doc.add_heading('INDICE CRONOLOGICO - CERTIFICADO DOMICILIARIO', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading(f'AÑO {anio}', 0)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add spacing
            doc.add_paragraph()
            
            # Notary info table
            info_table = doc.add_table(rows=5, cols=9)
            info_table.style = 'Table Grid'
            
            # Row 1: NOTARIA
            row1 = info_table.rows[0]
            row1.cells[0].text = 'NOTARIA'
            row1.cells[0].paragraphs[0].runs[0].bold = True
            row1.cells[2].text = f': {notary_name}'
            
            # Row 2: DIRECCION
            row2 = info_table.rows[1]
            row2.cells[0].text = 'DIRECCION'
            row2.cells[0].paragraphs[0].runs[0].bold = True
            row2.cells[2].text = f': {notary_config["direccion"]}'
            row2.cells[4].text = 'TELEFONO'
            row2.cells[4].paragraphs[0].runs[0].bold = True
            row2.cells[7].text = f': {notary_config["telefono"]}'
            
            # Row 3: DEPARTAMENTO
            row3 = info_table.rows[2]
            row3.cells[0].text = 'DEPARTAMENTO'
            row3.cells[0].paragraphs[0].runs[0].bold = True
            row3.cells[2].text = f': {notary_config["departamento"]}'
            row3.cells[4].text = 'RUC'
            row3.cells[4].paragraphs[0].runs[0].bold = True
            row3.cells[7].text = f': {notary_config["ruc"]}'
            
            # Row 4: PROVINCIA
            row4 = info_table.rows[3]
            row4.cells[0].text = 'PROVINCIA'
            row4.cells[0].paragraphs[0].runs[0].bold = True
            row4.cells[2].text = f': {notary_config["provincia"]}'
            row4.cells[4].text = 'DESDE'
            row4.cells[4].paragraphs[0].runs[0].bold = True
            row4.cells[7].text = f': {self._format_date_in_spanish(desde).upper()}'
            
            # Row 5: DISTRITO
            row5 = info_table.rows[4]
            row5.cells[0].text = 'DISTRITO'
            row5.cells[0].paragraphs[0].runs[0].bold = True
            row5.cells[2].text = f': {notary_config["distrito"]}'
            row5.cells[4].text = 'HASTA'
            row5.cells[4].paragraphs[0].runs[0].bold = True
            row5.cells[7].text = f': {self._format_date_in_spanish(hasta).upper()}'
            
            # Add spacing
            doc.add_paragraph()
            
            # Data table
            if report_data:
                data_table = doc.add_table(rows=1, cols=8)
                data_table.style = 'Table Grid'
                
                # Headers
                header_row = data_table.rows[0]
                headers = ['N°', 'FECHA', 'SOLICITANTE', 'N° DNI', 'DOMICILIO', 'MOTIVO', 'DOCUMENTO VERIFICADO', 'COMPROBANTE']
                for i, header in enumerate(headers):
                    header_row.cells[i].text = header
                    header_row.cells[i].paragraphs[0].runs[0].bold = True
                
                # Data rows
                for data_row in report_data:
                    kardex, fecha, solicitante, documento_solicitante, domicilio_solicitante, motivo_solicitante, recibo, numero_recibo = data_row
                    
                    # Format kardex (substr from position 4 like PHP)
                    kardex_formatted = str(kardex)[3:] if len(str(kardex)) > 3 else str(kardex)
                    
                    # Format recibo type
                    recibo_formatted = self._format_recibo_type(recibo)
                    
                    # Add row
                    row = data_table.add_row()
                    row.cells[0].text = str(kardex_formatted)
                    row.cells[1].text = self._format_date_for_display(fecha)
                    row.cells[2].text = str(solicitante or '')
                    row.cells[3].text = str(documento_solicitante or '')
                    row.cells[4].text = str(domicilio_solicitante or '')
                    row.cells[5].text = str(motivo_solicitante or '')
                    row.cells[6].text = str(recibo_formatted or '')
                    row.cells[7].text = str(numero_recibo or '')
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Create response
            from django.http import HttpResponse
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename=INDICE_CRONOLOGICO_CERTIFICADO_DOMICILIARIO_{anio}.docx'
            response['Access-Control-Allow-Origin'] = '*'
            
            return response
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            from django.http import HttpResponse
            return HttpResponse(f"Error generating Word report: {e}", status=500) 