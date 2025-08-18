import boto3
from botocore.client import Config
from botocore.exceptions import ClientError
from django.conf import settings
import os
import io
from decimal import Decimal
from typing import Dict, Any, List
from django.http import HttpResponse, JsonResponse
from notaria.models import TplTemplate, Cliente, Tipodocumento, Nacionalidades, Tipoestacivil, Profesiones, Ubigeo, PermiViaje, ViajeContratantes
from .utils import NumberToLetterConverter
import time
from django.db import connection
import re
from docxtpl import DocxTemplate, RichText
import traceback
from datetime import datetime
import json



from ..shared.base_r2_documents import get_s3_client, BaseR2DocumentService


class BasePermisoViajeDocumentService(BaseR2DocumentService):
    """
    Base service with common logic for generating Permiso Viaje documents.
    """
    def __init__(self):
        self.letras = NumberToLetterConverter()
        self.template_filename = None  # Must be set by child classes
    
    def retrieve_document(self, id_permiviaje: int, mode: str = "download") -> HttpResponse:
        try:
            permiviaje = PermiViaje.objects.get(id_viaje=id_permiviaje)
            num_kardex = permiviaje.num_kardex
            if not num_kardex:
                return HttpResponse(f"Error: num_kardex is empty for PermiViaje id {id_permiviaje}", status=400)

            anio_kardex = (num_kardex or '')[:4]
            filename = f"__PERMIVIAJE__{id_permiviaje}-{anio_kardex}.docx"

            if mode == "open":
                return self._create_response(None, filename, id_permiviaje, mode)

            s3 = get_s3_client()
            object_key = f"rodriguez-zea/documentos/{filename}"
            response = s3.get_object(Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'), Key=object_key)
            buffer = io.BytesIO(response['Body'].read())
            
            return self._create_response(buffer, filename, id_permiviaje, mode)

        except PermiViaje.DoesNotExist:
            return HttpResponse(f"Error: PermiViaje with id {id_permiviaje} not found", status=404)
        except ClientError as e:
            if e.response['Error']['Code'] == 'NoSuchKey':
                return HttpResponse(f"Error: Document '{filename}' not found in R2.", status=404)
            else:
                traceback.print_exc()
                return HttpResponse(f"Error retrieving document: {e}", status=500)
        except Exception as e:
            traceback.print_exc()
            return HttpResponse(f"Error retrieving document: {e}", status=500)

    def _get_template_from_r2(self) -> bytes:
        if not self.template_filename:
            raise ValueError("template_filename must be set in the child service class.")
        s3 = get_s3_client()
        object_key = f"rodriguez-zea/plantillas/{self.template_filename}"
        try:
            response = s3.get_object(Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'), Key=object_key)
            return response['Body'].read()
        except ClientError as e:
            if e.response['Error']['Code'] == 'NoSuchKey':
                return None
            else:
                raise
        except Exception as e:
            raise

    def _get_licencia_data(self, fecha_ingreso: str) -> Dict[str, str]:
        if not fecha_ingreso:
            return {'licencia': ''}
        with connection.cursor() as cursor:
            cursor.execute("SELECT notario, resolucion, (SELECT CONCAT(nombre, ' ', apellido) FROM confinotario LIMIT 1) as notario_principal, (SELECT direccion FROM confinotario LIMIT 1) as direccion_notario FROM confinotario WHERE %s BETWEEN fechainicio AND fechafin", [fecha_ingreso])
            row = cursor.fetchone()
            if row:
                return {'licencia': f'POR LICENCIA DE LA NOTARIA {row[2]} FIRMA EL NOTARIO {row[0]} SEGUN RESOLUCION N° {row[1]}'}
            else:
                cursor.execute("SELECT CONCAT(nombre, ' ', apellido) as notario, direccion FROM confinotario LIMIT 1")
                notary_info = cursor.fetchone()
                if notary_info:
                    return {'licencia': f'YO {notary_info[0]} ABOGADO - NOTARIO DE PUNO CON OFICIO NOTARIAL EN {notary_info[1]}'}
        return {'licencia': ''}

    def _get_notary_data(self) -> Dict[str, str]:
        with connection.cursor() as cursor:
            cursor.execute("SELECT nombre AS nombres, apellido AS apellidos, CONCAT(nombre,' ',apellido) AS notario, ruc AS ruc_notario, distrito AS distrito_notario FROM confinotario")
            row = cursor.fetchone()
            if row:
                columns = [col[0] for col in cursor.description]
                notary_data = dict(zip(columns, row))
                final_data = {}
                for key, value in notary_data.items():
                    if key == 'ruc_notario' and value:
                        final_data[f'LETRA_{key.upper()}'] = self.letras.number_to_letters(value)
                    final_data[key.upper().strip()] = str(value).upper() if value is not None else '?'
                return final_data
            return {'NOMBRES': '?','APELLIDOS': '?','NOTARIO': '?','RUC_NOTARIO': '?','DISTRITO_NOTARIO': '?','LETRA_RUC_NOTARIO': ''}

    def _get_viaje_data(self, id_permiviaje: int) -> Dict[str, str]:
        try:
            viaje = PermiViaje.objects.get(id_viaje=id_permiviaje)
            data = {
                'ID_VIAJE': str(viaje.id_viaje), 'KARDEX': f"{viaje.num_kardex[4:]}-{viaje.num_kardex[:4]}" if viaje.num_kardex and len(viaje.num_kardex) > 4 else viaje.num_kardex or '?',
                'ASUNTO': viaje.asunto or '?', 'FECHA_INGRESO': viaje.fec_ingreso.strftime('%d/%m/%Y') if viaje.fec_ingreso else '?',
                'FECHA_INGRESO_RAW': viaje.fec_ingreso.strftime('%Y-%m-%d') if viaje.fec_ingreso else None, 'NOMBRE_RECEPCIONISTA': viaje.nom_recep or '?',
                'HORA_RECEPCION': viaje.hora_recep or '?', 'REFERENCIA': viaje.referencia or '?', 'COMUNICARSE': viaje.nom_comu or '?', 'COMUNICARSE_EMAIL': viaje.email_comu or '?',
                'DOCUMENTO': viaje.documento or '?', 'NUMERO_CRONOLOGICO': viaje.num_crono or '?', 'FECHA_CRONOLOGICO': viaje.fecha_crono.strftime('%d/%m/%Y') if viaje.fecha_crono else '?',
                'NUMERO_FORMULARIO': viaje.num_formu or '?', 'DESTINO': viaje.lugar_formu or '?', 'OBSERVACION': viaje.observacion or '?', 'SWT_EST': viaje.swt_est or '?',
                'PARTIDA_E': viaje.partida_e or '?', 'SEDE_REGIS': viaje.sede_regis or '?', 'REFER': viaje.referencia or '?', 'VIA_TRANS': getattr(viaje, 'via', '?') or '?',
                'FEC_DESDE': self.letras.date_to_letters(viaje.fecha_desde) if hasattr(viaje, 'fecha_desde') and viaje.fecha_desde else '?',
                'FEC_HASTA': self.letras.date_to_letters(viaje.fecha_hasta) if hasattr(viaje, 'fecha_hasta') and viaje.fecha_hasta else '?'
            }
            if viaje.fec_ingreso: data['LETRA_FECHA_INGRESO'] = self.letras.date_to_letters(viaje.fec_ingreso)
            return data
        except PermiViaje.DoesNotExist: return {}

    def _get_user_data(self, usuario_imprime: str = None) -> Dict[str, str]:
        if not usuario_imprime: return {'USUARIO': '?','USUARIO_DNI': '?'}
        with connection.cursor() as cursor:
            cursor.execute("SELECT loginusuario, dni FROM usuarios WHERE CONCAT(apepat,' ',prinom) = %s", [usuario_imprime])
            row = cursor.fetchone()
            return {'USUARIO': row[0] or '?','USUARIO_DNI': row[1] or '?'} if row else {'USUARIO': '?','USUARIO_DNI': '?'}

    def _process_document(self, template_bytes: bytes, data: Dict[str, Any]) -> DocxTemplate:
        """
        Renders the document using docxtpl with Jinja2 syntax.
        """
        doc = DocxTemplate(io.BytesIO(template_bytes))
        
        # Temporarily disabling RichText to debug file corruption issue.
        # This will render the document without red color.
        context = {}
        for key, value in data.items():
            if isinstance(value, list):
                new_list = []
                for item in value:
                    if isinstance(item, dict):
                        new_item = {k: RichText(str(v) if v is not None else '', color='#FF0000') for k, v in item.items()}
                        new_list.append(new_item)
                    else:
                        new_list.append(RichText(str(item) if item is not None else '', color='#FF0000'))
                context[key] = new_list
            else:
                context[key] = RichText(str(value) if value is not None else '', color='#FF0000')
        doc.render(context)
        return doc

    def _create_response(self, buffer: io.BytesIO, filename: str, id_permiviaje: int, mode: str = "download"):
        if mode == "open":
            s3 = get_s3_client()
            object_key = f"rodriguez-zea/documentos/{filename}"
            try:
                url = s3.generate_presigned_url(
                    'get_object',
                    Params={'Bucket': os.environ.get('CLOUDFLARE_R2_BUCKET'), 'Key': object_key},
                    ExpiresIn=3600  # URL expires in 1 hour
                )
                response = JsonResponse({
                    'status': 'success',
                    'mode': 'open',
                    'url': url,
                    'filename': filename,
                    'id_permiviaje': id_permiviaje,
                    'message': 'Document is ready to be opened.'
                })
                response['Access-Control-Allow-Origin'] = '*'
                return response
            except Exception as e:
                return HttpResponse(f"Error generating pre-signed URL: {e}", status=500)
        else:
            if buffer is None:
                return HttpResponse("Error: Document buffer is missing for download mode.", status=500)
            buffer.seek(0)
            response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'inline; filename="{filename}"'
            response['Content-Length'] = str(buffer.getbuffer().nbytes)
            response['Access-Control-Allow-Origin'] = '*'
            return response

    def _save_document_to_r2(self, buffer: io.BytesIO, filename: str):
        s3 = get_s3_client()
        object_key = f"rodriguez-zea/documentos/{filename}"
        buffer.seek(0)
        s3.put_object(
            Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'),
            Key=object_key,
            Body=buffer.read()
        )
        buffer.seek(0)

    def _document_exists_in_r2(self, filename: str) -> bool:
        s3 = get_s3_client()
        object_key = f"rodriguez-zea/documentos/{filename}"
        try:
            s3.head_object(Bucket=os.environ.get('CLOUDFLARE_R2_BUCKET'), Key=object_key)
            return True
        except ClientError as e:
            code = e.response.get('Error', {}).get('Code') if hasattr(e, 'response') else None
            if code in ('NoSuchKey', '404'):
                return False
            # Some R2 providers return numeric strings for 404 on head_object
            try:
                status = e.response.get('ResponseMetadata', {}).get('HTTPStatusCode')
                if status == 404:
                    return False
            except Exception:
                pass
            raise
        except Exception:
            raise

class PermisoViajeInteriorDocumentService(BasePermisoViajeDocumentService):
    def __init__(self):
        super().__init__()
        self.template_filename = "AUTORIZACION VIAJE MENOR INTERIOR.docx"

    def generate_permiso_viaje_interior_document(self, id_permiviaje: int, mode: str = "download") -> HttpResponse:
        try:
            permiviaje = PermiViaje.objects.get(id_viaje=id_permiviaje)
            num_kardex = permiviaje.num_kardex
            if not num_kardex:
                return self.json_error(400, f"num_kardex is empty for PermiViaje id {id_permiviaje}")

            anio_kardex = (num_kardex or '')[:4]
            filename = f"__PERMIVIAJE__{id_permiviaje}-{anio_kardex}.docx"
            if self._document_exists_in_r2(filename):
                return self.json_error(409, "Document already exists. Use action=retrieve to fetch it.", {
                    'id_permiviaje': id_permiviaje,
                    'filename': filename,
                })

            template_bytes = self._get_template_from_r2()
            if template_bytes is None:
                return self.json_error(404, f"Template file '{self.template_filename}' not found in 'rodriguez-zea/plantillas/'.")
            
            document_data = self.get_document_data(id_permiviaje)
            doc = self._process_document(template_bytes, document_data)
            buffer = io.BytesIO()
            doc.save(buffer)
            self._save_document_to_r2(buffer, filename)
            
            return self._create_response(buffer, filename, id_permiviaje, mode)

        except PermiViaje.DoesNotExist:
            return self.json_error(404, f"PermiViaje with id {id_permiviaje} not found", {'id_permiviaje': id_permiviaje})
        except Exception as e:
            traceback.print_exc()
            return self.json_error(500, f"Error generating document: {e}")

    def get_document_data(self, id_permiviaje: int) -> Dict[str, Any]:
        notary_data = self._get_notary_data()
        viaje_data = self._get_viaje_data(id_permiviaje)
        user_data = self._get_user_data(viaje_data.get('NOMBRE_RECEPCIONISTA'))
        participants_data, blocks_data = self._get_participants_data(id_permiviaje)
        
        context = {}
        context.update(notary_data)
        context.update(viaje_data)
        context.update(user_data)
        context.update(participants_data)
        context.update(blocks_data)

        context['PADRE_MADRE'] = self._determine_padre_madre(blocks_data)
        context['VACIO'] = ''
        context['CONFIG'] = f"{id_permiviaje}_permiviaje/"
        
        licencia_info = self._get_licencia_data(viaje_data.get('FECHA_INGRESO_RAW'))
        context.update(licencia_info)
        
        c_list = context.get('c', [])
        if c_list:
            first_contractor = c_list[0]
            context['procede'] = first_contractor.get('procede', '')
            context['SOLICITANTE'] = first_contractor.get('SOLICITANTE', '')

        return context

    def _get_participants_data(self, id_permiviaje: int) -> tuple[Dict[str, Any], Dict[str, List[Dict[str, Any]]]]:
        with connection.cursor() as cursor:
            participants_data, blocks_data = {}, {}
            cursor.execute("SELECT COUNT(*) FROM viaje_contratantes WHERE c_condicontrat IN ('001','003','004','005','010') AND id_viaje = %s", [id_permiviaje])
            num_contratantes = cursor.fetchone()[0]
            
            contratantes_query = """SELECT vc.c_condicontrat, CONCAT_WS(' ', c.prinom, c.segnom, c.apepat, c.apemat) AS contratante, td.destipdoc as tipo_documento, td.td_abrev as abreviatura, c.numdoc as numero_documento, n.descripcion as nacionalidad, c.direccion, vc.c_fircontrat, IF(u.coddis='' OR ISNULL(u.coddis), '', CONCAT('DISTRITO DE ',u.nomdis, ', PROVINCIA DE ', u.nomprov,', DEPARTAMENTO DE ',u.nomdpto)) AS ubigeo, tec.desestcivil as estado_civil, IFNULL(p.desprofesion,'') as profesion, IFNULL(vc.codi_podera,'') as codigo_poderado, c.detaprofesion as profesion_cliente, (CASE WHEN vc.condi_edad = 1 THEN CONCAT(vc.edad,' AÑOS') WHEN vc.condi_edad = 2 THEN CONCAT(vc.edad,' MESES') ELSE '' END) as edad, c.sexo FROM viaje_contratantes vc JOIN cliente c ON c.numdoc = vc.c_codcontrat JOIN tipodocumento td ON td.idtipdoc = c.idtipdoc JOIN nacionalidades n ON n.idnacionalidad = c.nacionalidad JOIN tipoestacivil tec ON tec.idestcivil = c.idestcivil LEFT JOIN profesiones p ON p.idprofesion = c.idprofesion LEFT JOIN ubigeo u ON u.coddis = c.idubigeo WHERE vc.c_condicontrat IN ('001','003','004','005','010') AND vc.id_viaje = %s"""
            cursor.execute(contratantes_query, [id_permiviaje])
            
            columns = [col[0] for col in cursor.description]
            contratantes_list = [dict(zip(columns, row)) for row in cursor.fetchall()]

            for p in contratantes_list:
                sex = p.get('sexo', 'M')
                p.update({'identificado': 'IDENTIFICADO' if sex == 'M' else 'IDENTIFICADA', 'domiciliado': 'CON DOMICILIO ', 'senor': 'SEÑOR' if sex == 'M' else 'SEÑORA', 'el': 'EL' if sex == 'M' else 'LA', 'don': 'DON' if sex == 'M' else 'DOÑA'})
                if p.get('nacionalidad'): p['nacionalidad'] = p['nacionalidad'][:-1] + ('O' if sex == 'M' and p['nacionalidad'].endswith('A') else ('A' if sex == 'F' and p['nacionalidad'].endswith('O') else p['nacionalidad'][-1]))
                if num_contratantes > 1: p.update({'SOLICITANTE': 'a los solicitantes', 'procede': 'Los compareciente proceden'})
                else: p.update({'SOLICITANTE': 'al solicitante' if sex == 'M' else 'a la solicitante', 'procede': 'El compareciente procede' if sex == 'M' else 'La compareciente procede'})
            blocks_data['c'] = contratantes_list
            
            max_cols = 3
            signature_rows = []
            for i in range(0, len(contratantes_list), max_cols):
                row = contratantes_list[i:i + max_cols]
                while len(row) < max_cols:
                    row.append(None)
                signature_rows.append(row)
            blocks_data['signature_rows'] = signature_rows

            minors_query = "SELECT CONCAT_WS(' ', c.prinom, c.segnom, c.apepat, c.apemat) AS contratante, (CASE WHEN vc.condi_edad = 1 THEN CONCAT(vc.edad,' AÑOS') WHEN vc.condi_edad = 2 THEN CONCAT(vc.edad,' MESES') ELSE '' END) as edad, c.sexo, td.td_abrev as abreviatura, c.numdoc as numero_documento FROM viaje_contratantes vc JOIN cliente c ON c.numdoc = vc.c_codcontrat JOIN tipodocumento td ON td.idtipdoc=c.idtipdoc WHERE vc.c_condicontrat = '002' AND vc.id_viaje = %s"
            cursor.execute(minors_query, [id_permiviaje])
            
            columns = [col[0] for col in cursor.description]
            minors_list = [dict(zip(columns, row)) for row in cursor.fetchall()]
            all_female = all(p.get('sexo') == 'F' for p in minors_list)
            for i, p in enumerate(minors_list):
                sex = p.get('sexo', 'M')
                p['identificado'] = 'IDENTIFICADO' if sex == 'M' else 'IDENTIFICADA'
                p['y_coma'] = '.' if i == len(minors_list) - 1 else (' Y' if i == len(minors_list) - 2 else ',')
            blocks_data['m'] = minors_list
            blocks_data['f'] = contratantes_list
            
            if len(contratantes_list) > 1: participants_data.update({'A_EL_LOS': 'LOS', 'A_S': 'S', 'A_N': 'N'})
            else: participants_data.update({'A_EL_LOS': 'EL', 'A_S': '', 'A_N': ''})
            
            if len(minors_list) == 1:
                sex = minors_list[0].get('sexo', 'M')
                participants_data.update({'EL_LA_LOS': 'LA' if sex == 'F' else 'EL', 'HIJO': 'HIJA' if sex == 'F' else 'HIJO', 'MENOR': 'SU MENOR', 'AUTORIZA': 'AUTORIZA'})
            else:
                participants_data.update({'EL_LA_LOS': 'LAS' if all_female else 'LOS', 'HIJO': 'HIJAS' if all_female else 'HIJOS', 'MENOR': 'SUS MENORES', 'AUTORIZA': 'AUTORIZAN'})
            
            cursor.execute("SELECT id_condicion, des_condicion FROM c_condiciones WHERE swt_condicion = 'V' AND id_condicion != '002'")
            for id_cond, desc in cursor.fetchall():
                cursor.execute(contratantes_query.replace("WHERE vc.c_condicontrat IN ('001','003','004','005','010')", f"WHERE vc.c_condicontrat = '{id_cond}'"), [id_permiviaje])
                rows = cursor.fetchall()
                if rows:
                    cols = [col[0] for col in cursor.description]
                    participant_dict = dict(zip(cols, rows[0]))
                    for k, v in participant_dict.items():
                        participants_data[f"{desc.lower()}_{k.upper()}"] = v
            
            return participants_data, blocks_data

    def _determine_padre_madre(self, blocks_data: Dict[str, Any]) -> str:
        contratantes = blocks_data.get('c', [])
        has_male = any(p.get('sexo') == 'M' for p in contratantes)
        has_female = any(p.get('sexo') == 'F' for p in contratantes)
        if has_male and has_female: return 'PADRES'
        if has_male: return 'PADRE'
        if has_female: return 'MADRE'
        return ''

class PermisoViajeExteriorDocumentService(BasePermisoViajeDocumentService):
    def __init__(self):
        super().__init__()
        self.template_filename = "AUTORIZACION VIAJE MENOR EXTERIOR.docx"

    def generate_permiso_viaje_exterior_document(self, id_permiviaje: int, mode: str = "download") -> HttpResponse:
        try:
            permiviaje = PermiViaje.objects.get(id_viaje=id_permiviaje)
            num_kardex = permiviaje.num_kardex
            if not num_kardex:
                return self.json_error(400, f"num_kardex is empty for PermiViaje id {id_permiviaje}")

            anio_kardex = (num_kardex or '')[:4]
            filename = f"__PERMIVIAJE__{id_permiviaje}-{anio_kardex}.docx"
            if self._document_exists_in_r2(filename):
                return self.json_error(409, "Document already exists. Use action=retrieve to fetch it.", {
                    'id_permiviaje': id_permiviaje,
                    'filename': filename,
                })

            template_bytes = self._get_template_from_r2()
            if template_bytes is None:
                return self.json_error(404, f"Template file '{self.template_filename}' not found in 'rodriguez-zea/plantillas/'.")
            
            document_data = self.get_document_data(id_permiviaje)
            doc = self._process_document(template_bytes, document_data)
            buffer = io.BytesIO()
            doc.save(buffer)
            self._save_document_to_r2(buffer, filename)
            
            return self._create_response(buffer, filename, id_permiviaje, mode)

        except PermiViaje.DoesNotExist:
            return self.json_error(404, f"PermiViaje with id {id_permiviaje} not found", {'id_permiviaje': id_permiviaje})
        except Exception as e:
            traceback.print_exc()
            return self.json_error(500, f"Error generating document: {e}")

    def get_document_data(self, id_permiviaje: int) -> Dict[str, Any]:
        notary_data = self._get_notary_data()
        viaje_data = self._get_viaje_data(id_permiviaje)
        user_data = self._get_user_data(viaje_data.get('NOMBRE_RECEPCIONISTA'))
        participants_data, blocks_data = self._get_participants_data(id_permiviaje)
        
        context = {}
        context.update(notary_data)
        context.update(viaje_data)
        context.update(user_data)
        context.update(participants_data)
        context.update(blocks_data)

        licencia_info = self._get_licencia_data(viaje_data.get('FECHA_INGRESO_RAW'))
        context.update(licencia_info)
        
        # Pull out standalone values from the first participant for easier access in the template
        c_list = context.get('c', [])
        if c_list:
            context['procede'] = c_list[0].get('procede', '')
            context['SOLICITANTE'] = c_list[0].get('SOLICITANTE', '')

        return context

    def _get_participants_data(self, id_permiviaje: int) -> tuple[Dict[str, Any], Dict[str, List[Dict[str, Any]]]]:
        with connection.cursor() as cursor:
            participants_data = {}
            blocks_data = {}

            # These are all the roles that can sign the Exterior permit
            contratante_conditions = "('001', '003', '004', '005', '010')"
            
            cursor.execute(f"SELECT COUNT(*) FROM viaje_contratantes WHERE c_condicontrat IN {contratante_conditions} AND id_viaje = %s", [id_permiviaje])
            num_contratantes = cursor.fetchone()[0]

            base_query = f"""
                SELECT 
                    vc.c_condicontrat AS id_condicion,
                    CONCAT_WS(' ', c.prinom, c.segnom, c.apepat, c.apemat) AS contratante,
                    td.destipdoc AS tipo_documento,
                    td.td_abrev AS abreviatura,
                    c.numdoc AS numero_documento,
                    IF(c.sexo='M', CONCAT(SUBSTRING(n.descripcion, 1, LENGTH(n.descripcion)-1),'O'), CONCAT(SUBSTRING(n.descripcion, 1, LENGTH(n.descripcion)-1),'A')) AS nacionalidad,
                    CONCAT(' CON DOMICILIO EN ', c.direccion) AS direccion,
                    IF(u.coddis='' OR ISNULL(u.coddis), '', CONCAT('DEL DISTRITO DE ', u.nomdis, ', PROVINCIA DE ', u.nomprov, ', DEPARTAMENTO DE ', u.nomdpto)) AS ubigeo,
                    tec.desestcivil AS estado_civil,
                    IFNULL(p.desprofesion, '') AS profesion,
                    c.sexo
                FROM viaje_contratantes vc
                JOIN cliente c ON c.numdoc = vc.c_codcontrat
                JOIN tipodocumento td ON td.idtipdoc = c.idtipdoc
                JOIN nacionalidades n ON n.idnacionalidad = c.nacionalidad
                JOIN tipoestacivil tec ON tec.idestcivil = c.idestcivil
                LEFT JOIN profesiones p ON p.idprofesion = c.idprofesion
                LEFT JOIN ubigeo u ON u.coddis = c.idubigeo
                WHERE vc.id_viaje = %s AND vc.c_condicontrat IN {contratante_conditions}
            """
            
            cursor.execute(base_query, [id_permiviaje])
            columns = [col[0] for col in cursor.description]
            contratantes_list = [dict(zip(columns, row)) for row in cursor.fetchall()]

            for p in contratantes_list:
                sex = p.get('sexo', 'M')
                p.update({
                    'identificado': 'IDENTIFICADO' if sex == 'M' else 'IDENTIFICADA',
                    'senor': 'SEÑOR' if sex == 'M' else 'SEÑORA',
                    'el': 'EL' if sex == 'M' else 'LA',
                })
                if num_contratantes > 1:
                    p.update({'SOLICITANTE': 'a los solicitantes', 'procede': 'Los comparecientes proceden'})
                else:
                    p.update({'SOLICITANTE': 'al solicitante' if sex == 'M' else 'a la solicitante', 'procede': 'El compareciente procede' if sex == 'M' else 'La compareciente procede'})
            
            blocks_data['c'] = contratantes_list

            # This will be used for the flexible signature block
            max_cols = 2 # Exterior template has 2 signatures per row
            signature_rows = []
            for i in range(0, len(contratantes_list), max_cols):
                row = contratantes_list[i:i + max_cols]
                while len(row) < max_cols:
                    row.append(None)
                signature_rows.append(row)
            blocks_data['signature_rows'] = signature_rows

            minors_query = "SELECT CONCAT_WS(' ', c.prinom, c.segnom, c.apepat, c.apemat) AS contratante, (CASE WHEN vc.condi_edad = 1 THEN CONCAT(vc.edad,' AÑOS') WHEN vc.condi_edad = 2 THEN CONCAT(vc.edad,' MESES') ELSE '' END) as edad, c.sexo, td.td_abrev as abreviatura, c.numdoc as numero_documento FROM viaje_contratantes vc JOIN cliente c ON c.numdoc = vc.c_codcontrat JOIN tipodocumento td ON td.idtipdoc=c.idtipdoc WHERE vc.c_condicontrat = '002' AND vc.id_viaje = %s"
            cursor.execute(minors_query, [id_permiviaje])
            columns = [col[0] for col in cursor.description]
            minors_list = [dict(zip(columns, row)) for row in cursor.fetchall()]
            all_female = all(p.get('sexo') == 'F' for p in minors_list)
            for i, p in enumerate(minors_list):
                sex = p.get('sexo', 'M')
                p['identificado'] = 'IDENTIFICADO' if sex == 'M' else 'IDENTIFICADA'
                p['y_coma'] = '.' if i == len(minors_list) - 1 else (' Y' if i == len(minors_list) - 2 else ',')
            blocks_data['m'] = minors_list

            if len(minors_list) == 1:
                sex = minors_list[0].get('sexo', 'M')
                participants_data.update({'HIJO': 'HIJA' if sex == 'F' else 'HIJO', 'MENOR': 'MENOR', 'AUTORIZA': 'AUTORIZA'})
            else:
                participants_data.update({'HIJO': 'HIJAS' if all_female else 'HIJOS', 'MENOR': 'MENORES', 'AUTORIZA': 'AUTORIZAN'})

            return participants_data, blocks_data 


class PermisosViajeReportService:
    """
    Service for generating Word and Excel reports for Permisos de Viaje.
    """
    
    def __init__(self):
        self.letras = NumberToLetterConverter()
    
    def _get_report_data(self, desde: str, hasta: str) -> List[tuple]:
        """
        Fetch data from permi_viaje and related tables for the date range.
        Parameters desde and hasta should be in YYYY-MM-DD format.
        """
        with connection.cursor() as cursor:
            # Since fecha_crono is a DATE field, use direct date comparison
            query = """
                SELECT
                    pv.id_viaje as cod_viaje,
                    pv.fec_ingreso as fec_ingreso,
                    pv.fecha_crono as fec_crono,
                    pv.num_kardex as kard,
                    (CASE WHEN(pv.asunto='001') THEN 'PERMISO VIAJE AL INTERIOR' ELSE 'PERMISO VIAJE AL EXTERIOR' END) as asunto,
                    pv.lugar_formu as lugar,
                    pv.swt_est as estado,
                    pv.num_kardex AS crono,
                    pv.num_formu AS formulario,
                    pv.via,
                    UPPER(pv.observacion) as observacion
                FROM permi_viaje pv
                WHERE pv.fecha_crono IS NOT NULL 
                AND pv.fecha_crono >= %s 
                AND pv.fecha_crono <= %s
                ORDER BY kard
            """
            
            # Execute with YYYY-MM-DD format for DATE field comparison
            cursor.execute(query, [desde, hasta])
            result = cursor.fetchall()
            return result if result else []
    
    def _get_notary_info(self) -> str:
        """Fetch notary name from confinotario table."""
        with connection.cursor() as cursor:
            cursor.execute("SELECT CONCAT(nombre, ' ', apellido) as notario FROM confinotario LIMIT 1")
            row = cursor.fetchone()
            return row[0] if row else "NOTARIO"
    
    def _get_participants_for_viaje(self, id_viaje: int) -> List[Dict[str, str]]:
        """Fetch participants for a specific viaje."""
        with connection.cursor() as cursor:
            query = """
                SELECT 
                    vc.id_viaje, 
                    vc.c_descontrat,
                    cc.des_condicion,
                    vc.c_codcontrat as doc,
                    td.td_abrev as tipo_documento 
                FROM viaje_contratantes as vc
                LEFT JOIN cliente as c ON c.numdoc=vc.c_codcontrat
                LEFT JOIN tipodocumento as td ON td.idtipdoc=c.idtipdoc
                LEFT JOIN c_condiciones as cc ON vc.c_condicontrat = cc.id_condicion
                WHERE vc.id_viaje=%s 
                GROUP BY vc.c_codcontrat, cc.des_condicion
                ORDER BY vc.id_contratante
                LIMIT 10
            """
            
            try:
                cursor.execute(query, [id_viaje])
                columns = [col[0] for col in cursor.description]
                return [dict(zip(columns, row)) for row in cursor.fetchall()]
            except Exception as e:
                print(f"DEBUG: Error getting participants for viaje {id_viaje}: {e}")
                return []  # Return empty list on error to continue processing

    def _get_participants_batch(self, viaje_ids: List[int]) -> Dict[int, List[Dict[str, str]]]:
        """Fetch participants for multiple viajes in a single query."""
        if not viaje_ids:
            return {}
            
        with connection.cursor() as cursor:
            # Create placeholders for IN clause
            placeholders = ','.join(['%s'] * len(viaje_ids))
            
            query = f"""
                SELECT 
                    vc.id_viaje, 
                    vc.c_descontrat,
                    cc.des_condicion,
                    vc.c_codcontrat as doc,
                    td.td_abrev as tipo_documento 
                FROM viaje_contratantes as vc
                LEFT JOIN cliente as c ON c.numdoc=vc.c_codcontrat
                LEFT JOIN tipodocumento as td ON td.idtipdoc=c.idtipdoc
                LEFT JOIN c_condiciones as cc ON vc.c_condicontrat = cc.id_condicion
                WHERE vc.id_viaje IN ({placeholders})
                GROUP BY vc.id_viaje, vc.c_codcontrat, cc.des_condicion
                ORDER BY vc.id_viaje, vc.id_contratante
                LIMIT 100
            """
            
            try:
                cursor.execute(query, viaje_ids)
                columns = [col[0] for col in cursor.description]
                results = cursor.fetchall()
                
                # Group participants by viaje_id
                participants_by_viaje = {}
                for row in results:
                    viaje_id = row[0]
                    if viaje_id not in participants_by_viaje:
                        participants_by_viaje[viaje_id] = []
                    
                    participant = dict(zip(columns, row))
                    participants_by_viaje[viaje_id].append(participant)
                
                return participants_by_viaje
                
            except Exception as e:
                print(f"DEBUG: Error in batch participant query: {e}")
                return {}  # Return empty dict on error
    
    def _format_date_in_spanish(self, date_str: str) -> str:
        """Convert YYYY-MM-DD date string to Spanish format."""
        if not date_str:
            return ""
        
        try:
            # Handle both DD/MM/YYYY and YYYY-MM-DD formats
            if '/' in date_str:
                # DD/MM/YYYY format
                day, month, year = date_str.split('/')
                date_obj = datetime(int(year), int(month), int(day))
            else:
                # YYYY-MM-DD format
                date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            
            dias = ['LUNES', 'MARTES', 'MIÉRCOLES', 'JUEVES', 'VIERNES', 'SÁBADO', 'DOMINGO']
            meses = ['ENERO', 'FEBRERO', 'MARZO', 'ABRIL', 'MAYO', 'JUNIO', 
                    'JULIO', 'AGOSTO', 'SEPTIEMBRE', 'OCTUBRE', 'NOVIEMBRE', 'DICIEMBRE']
            
            dia_semana = dias[date_obj.weekday()]
            dia = date_obj.day
            mes = meses[date_obj.month - 1]
            año = date_obj.year
            
            return f"{dia_semana}, {dia} DE {mes} DEL {año}"
        except:
            return date_str
    
    def _extract_year_from_date(self, date_str: str) -> str:
        """Extract year from date string."""
        if not date_str:
            return ""
        
        try:
            if '/' in date_str:
                # DD/MM/YYYY format
                return date_str.split('/')[-1]
            else:
                # YYYY-MM-DD format
                return date_str[:4]
        except:
            return ""
    
    def generate_excel_report(self, desde: str, hasta: str) -> HttpResponse:
        """Generate Excel report for Permisos de Viaje."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
            
            report_data = self._get_report_data(desde, hasta)
            notary_name = self._get_notary_info()
            anio = self._extract_year_from_date(hasta)
            
            # Create workbook and worksheet
            wb = Workbook()
            ws = wb.active
            ws.title = "Permisos de Viaje"
            
            # Styles to match PHP formatting
            title_font = Font(name="Arial", size=18, bold=True)  # 18.5px ≈ 18pt, bold
            header_font = Font(name="Arial", size=13, bold=True)  # 13.5px ≈ 13pt, bold
            data_font = Font(name="Arial", size=13)  # 13.5px ≈ 13pt, normal
            center_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            left_alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
            
            # Borders for data table (matching PHP BORDER="1" bordercolor="#333333")
            thin_border = Border(
                left=Side(border_style="thin", color="333333"),
                right=Side(border_style="thin", color="333333"),
                top=Side(border_style="thin", color="333333"),
                bottom=Side(border_style="thin", color="333333")
            )
            
            # No borders for header section (matching PHP border='0')
            no_border = Border(
                left=Side(style=None),
                right=Side(style=None),
                top=Side(style=None),
                bottom=Side(style=None)
            )
            
            # Title - matches PHP font-size:18.5px, bold, center
            ws.merge_cells('A1:G1')
            ws['A1'] = 'INDICE CRONOLOGICO - PERMISOS DE VIAJE'
            ws['A1'].font = title_font
            ws['A1'].alignment = center_alignment
            ws['A1'].border = no_border
            
            ws.merge_cells('A2:G2')
            ws['A2'] = f'AÑO {anio}'
            ws['A2'].font = title_font
            ws['A2'].alignment = center_alignment
            ws['A2'].border = no_border
            
            # Notary info section - matches PHP table with border='0'
            row = 4
            # Row 1: NOTARIA (colspan="2" align="left")
            ws.merge_cells(f'A{row}:B{row}')
            ws[f'A{row}'] = 'NOTARIA'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].alignment = left_alignment
            ws[f'A{row}'].border = no_border
            ws[f'C{row}'] = f': {notary_name}'
            ws[f'C{row}'].font = data_font
            ws[f'C{row}'].alignment = left_alignment
            ws[f'C{row}'].border = no_border
            
            row += 1
            # Row 2: DIRECCION
            ws.merge_cells(f'A{row}:B{row}')
            ws[f'A{row}'] = 'DIRECCION'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].alignment = left_alignment
            ws[f'A{row}'].border = no_border
            ws[f'C{row}'] = ': JR.BOLIVAR NRO. 340'
            ws[f'C{row}'].font = data_font
            ws[f'C{row}'].alignment = left_alignment
            ws[f'C{row}'].border = no_border
            ws[f'E{row}'] = 'TELEFONO'
            ws[f'E{row}'].font = header_font
            ws[f'E{row}'].alignment = left_alignment
            ws[f'E{row}'].border = no_border
            ws.merge_cells(f'F{row}:G{row}')
            ws[f'F{row}'] = ': (051) 326609'
            ws[f'F{row}'].font = data_font
            ws[f'F{row}'].alignment = left_alignment
            ws[f'F{row}'].border = no_border
            
            row += 1
            # Row 3: DEPARTAMENTO
            ws.merge_cells(f'A{row}:B{row}')
            ws[f'A{row}'] = 'DEPARTAMENTO'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].alignment = left_alignment
            ws[f'A{row}'].border = no_border
            ws[f'C{row}'] = ': PUNO'
            ws[f'C{row}'].font = data_font
            ws[f'C{row}'].alignment = left_alignment
            ws[f'C{row}'].border = no_border
            ws[f'E{row}'] = 'RUC'
            ws[f'E{row}'].font = header_font
            ws[f'E{row}'].alignment = left_alignment
            ws[f'E{row}'].border = no_border
            ws.merge_cells(f'F{row}:G{row}')
            ws[f'F{row}'] = ': 10024231572'
            ws[f'F{row}'].font = data_font
            ws[f'F{row}'].alignment = left_alignment
            ws[f'F{row}'].border = no_border
            
            row += 1
            # Row 4: PROVINCIA
            ws.merge_cells(f'A{row}:B{row}')
            ws[f'A{row}'] = 'PROVINCIA'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].alignment = left_alignment
            ws[f'A{row}'].border = no_border
            ws[f'C{row}'] = ': SAN ROMAN'
            ws[f'C{row}'].font = data_font
            ws[f'C{row}'].alignment = left_alignment
            ws[f'C{row}'].border = no_border
            ws[f'E{row}'] = 'DESDE'
            ws[f'E{row}'].font = header_font
            ws[f'E{row}'].alignment = left_alignment
            ws[f'E{row}'].border = no_border
            ws.merge_cells(f'F{row}:G{row}')
            ws[f'F{row}'] = f': {self._format_date_in_spanish(desde)}'
            ws[f'F{row}'].font = data_font
            ws[f'F{row}'].alignment = left_alignment
            ws[f'F{row}'].border = no_border
            
            row += 1
            # Row 5: DISTRITO
            ws.merge_cells(f'A{row}:B{row}')
            ws[f'A{row}'] = 'DISTRITO'
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].alignment = left_alignment
            ws[f'A{row}'].border = no_border
            ws[f'C{row}'] = ': JULIACA'
            ws[f'C{row}'].font = data_font
            ws[f'C{row}'].alignment = left_alignment
            ws[f'C{row}'].border = no_border
            ws[f'E{row}'] = 'HASTA'
            ws[f'E{row}'].font = header_font
            ws[f'E{row}'].alignment = left_alignment
            ws[f'E{row}'].border = no_border
            ws.merge_cells(f'F{row}:G{row}')
            ws[f'F{row}'] = f': {self._format_date_in_spanish(hasta)}'
            ws[f'F{row}'].font = data_font
            ws[f'F{row}'].alignment = left_alignment
            ws[f'F{row}'].border = no_border
            
            # Add spacing
            row += 2
            
            # Data table headers - matches PHP BORDER="1" with borders
            headers = ['NRO.', 'FECHA', 'PARTICIPANTES', 'VIA', 'DESTINO', 'OBSERVACIONES']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col, value=header)
                cell.font = header_font
                cell.alignment = center_alignment
                cell.border = thin_border  # Add borders for data table headers
            
            # Set specific column widths to fit everything in A-F range (6 columns)
            ws.column_dimensions['A'].width = 10  # NRO. (expanded from 8)
            ws.column_dimensions['B'].width = 15  # FECHA (expanded from 12)
            ws.column_dimensions['C'].width = 55  # PARTICIPANTES (expanded from 40)
            ws.column_dimensions['D'].width = 25  # VIA (expanded from 18)
            ws.column_dimensions['E'].width = 35  # DESTINO (expanded from 25)
            ws.column_dimensions['F'].width = 50  # OBSERVACIONES (expanded from 35)
            ws.column_dimensions['G'].width = 5   # Empty column for spacing (no border)
            
            # Add data rows
            row += 1
            
            # Limit to first 20 records for faster testing
            limited_data = report_data[:20] if len(report_data) > 20 else report_data
            
            # Batch query all participants for all viajes at once
            viaje_ids = [data_row[0] for data_row in limited_data if len(data_row) > 0]
            all_participants = self._get_participants_batch(viaje_ids)
            
            for data_row in limited_data:
                # Main data row
                correlativo = str(data_row[7])[-6:] if data_row[7] else ''  # Extract last 6 digits
                # Convert to simple number (1 instead of 000001) - matches PHP (int)substr()
                correlativo_simple = str(int(correlativo)) if correlativo and correlativo.isdigit() else correlativo
                
                fecha_crono = data_row[2] if len(data_row) > 2 else ''
                via = data_row[9] if len(data_row) > 9 else ''
                destino = data_row[5] if len(data_row) > 5 else ''
                observacion = data_row[10] if len(data_row) > 10 else ''
                
                # Convert datetime.date objects to strings
                if hasattr(fecha_crono, 'strftime'):
                    fecha_crono = fecha_crono.strftime('%d/%m/%Y')  # matches fechabd_an() format
                
                # Get participants for this viaje from batch results
                id_viaje = data_row[0] if len(data_row) > 0 else 0
                participants = all_participants.get(id_viaje, [])
                
                # Create participant text exactly like PHP nested table
                participant_text = ""
                for participant in participants:
                    condicion = participant.get('des_condicion', '')
                    nombre = participant.get('c_descontrat', '')
                    tipo_doc = participant.get('tipo_documento', '')
                    num_doc = participant.get('doc', '')
                    # Format exactly like PHP: "CONDICION :NOMBRE" and "TIPO_DOC:NUM_DOC"
                    participant_text += f"{condicion} :{nombre}\n{tipo_doc}:{num_doc}\n"
                
                # Format observations to fit in column F
                if observacion:
                    # Split long observations and add line breaks every 70 characters (increased from 50)
                    formatted_observacion = ""
                    words = observacion.split()
                    current_line = ""
                    for word in words:
                        if len(current_line + word) > 70:  # Increased from 50 to use more of the wide column
                            formatted_observacion += current_line.strip() + "\n"
                            current_line = word + " "
                        else:
                            current_line += word + " "
                    formatted_observacion += current_line.strip()
                    observacion = formatted_observacion
                
                # Main data row - matches PHP table structure with borders
                ws.cell(row=row, column=1, value=correlativo_simple).alignment = center_alignment
                ws.cell(row=row, column=2, value=fecha_crono).alignment = center_alignment
                ws.cell(row=row, column=3, value=participant_text.strip()).alignment = left_alignment
                ws.cell(row=row, column=4, value=via).alignment = center_alignment
                ws.cell(row=row, column=5, value=destino.replace('?', '-').upper() if destino else '').alignment = center_alignment
                ws.cell(row=row, column=6, value=observacion.upper() if observacion else '').alignment = left_alignment
                
                # Set font and borders for data cells (columns A-F get borders)
                for col in range(1, 7):  # A to F (6 columns)
                    cell = ws.cell(row=row, column=col)
                    cell.font = data_font
                    cell.border = thin_border  # Add borders for data cells
                
                # Calculate required row height based on content
                max_lines = 1
                
                # Count lines in participant text
                if participant_text:
                    participant_lines = len(participant_text.strip().split('\n'))
                    max_lines = max(max_lines, participant_lines)
                
                # Count lines in observations (more accurate calculation)
                if observacion:
                    observation_lines = len(observacion.strip().split('\n'))
                    max_lines = max(max_lines, observation_lines)
                    
                    # Also check if any single line in observations is very long
                    for line in observacion.strip().split('\n'):
                        if len(line) > 70:  # If a line is longer than our wrap limit
                            # Calculate how many visual lines this would actually take
                            visual_lines = (len(line) // 70) + 1
                            max_lines = max(max_lines, visual_lines)
                
                # Count lines in via (in case it's long)
                if via and len(str(via)) > 25:
                    via_lines = (len(str(via)) // 25) + 1
                    max_lines = max(max_lines, via_lines)
                
                # Count lines in destino (in case it's long)
                if destino and len(str(destino)) > 35:
                    destino_lines = (len(str(destino)) // 35) + 1
                    max_lines = max(max_lines, destino_lines)
                
                # Set row height based on content with more generous spacing
                # Minimum 60px, then 30px per line for better readability
                calculated_height = max(60, max_lines * 30)
                
                # Add extra padding for very content-heavy rows
                if max_lines > 4:
                    calculated_height += 20  # Extra padding for complex rows
                
                ws.row_dimensions[row].height = calculated_height
                
                row += 1
            
            # Save to buffer
            buffer = io.BytesIO()
            wb.save(buffer)
            buffer.seek(0)
            
            # Create response
            filename = f"INDICE_CRONOLOGICO_PERMISOS_DE_VIAJE_{anio}.xlsx"
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
            
        except Exception as e:
            return HttpResponse(
                json.dumps({'error': f'Error generating Excel report: {str(e)}'}),
                content_type='application/json',
                status=500
            )
    
    def generate_word_report(self, desde: str, hasta: str) -> HttpResponse:
        """Generate Word report for Permisos de Viaje."""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            print(f"DEBUG: Starting Word report generation for {desde} to {hasta}")
            
            report_data = self._get_report_data(desde, hasta)
            print(f"DEBUG: Retrieved {len(report_data)} records")
            
            notary_name = self._get_notary_info()
            print(f"DEBUG: Got notary info: {notary_name}")
            
            anio = self._extract_year_from_date(hasta)
            print(f"DEBUG: Extracted year: {anio}")
            
            # Create a new document
            doc = Document()
            print("DEBUG: Created Word document")
            
            # Title
            title = doc.add_heading('INDICE CRONOLOGICO - PERMISOS DE VIAJE', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading(f'AÑO {anio}', 0)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add spacing
            doc.add_paragraph()
            
            # Notary info table - NO BORDERS
            info_table = doc.add_table(rows=5, cols=6)
            # No table style = no borders
            
            # Row 1: NOTARIA
            row = info_table.rows[0]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = 'NOTARIA'
            row.cells[2].text = f': {notary_name}'
            
            # Row 2: DIRECCION
            row = info_table.rows[1]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = 'DIRECCION'
            row.cells[2].text = ': JR.BOLIVAR NRO. 340'
            row.cells[3].text = 'TELEFONO'
            row.cells[4].merge(row.cells[5])
            row.cells[4].text = ': (051) 326609'
            
            # Row 3: DEPARTAMENTO
            row = info_table.rows[2]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = 'DEPARTAMENTO'
            row.cells[2].text = ': PUNO'
            row.cells[3].text = 'RUC'
            row.cells[4].merge(row.cells[5])
            row.cells[4].text = ': 10024231572'
            
            # Row 4: PROVINCIA
            row = info_table.rows[3]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = 'PROVINCIA'
            row.cells[2].text = ': SAN ROMAN'
            row.cells[3].text = 'DESDE'
            row.cells[4].merge(row.cells[5])
            row.cells[4].text = f': {self._format_date_in_spanish(desde)}'
            
            # Row 5: DISTRITO
            row = info_table.rows[4]
            row.cells[0].merge(row.cells[1])
            row.cells[0].text = 'DISTRITO'
            row.cells[2].text = ': JULIACA'
            row.cells[3].text = 'HASTA'
            row.cells[4].merge(row.cells[5])
            row.cells[4].text = f': {self._format_date_in_spanish(hasta)}'
            
            # Style the info table - Simple styling without borders
            for row in info_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if ':' in run.text:
                                run.font.bold = True
                            run.font.size = Pt(12)
            
            # Add spacing
            doc.add_paragraph()
            print("DEBUG: Added header table")
            
            # Main data table - ALWAYS CREATE, even if empty
            # Create table with headers
            headers = ['NRO.', 'FECHA', 'PARTICIPANTES', 'VIA', 'DESTINO', 'OBSERVACIONES']
            data_table = doc.add_table(rows=1, cols=6)
            data_table.style = 'Table Grid'
            
            # Add headers
            header_row = data_table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(12)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            print("DEBUG: Added table headers")
            
            # Add data rows if data exists
            if report_data and len(report_data) > 0:
                print(f"DEBUG: Processing {len(report_data)} data rows...")
                
                # Limit to first 20 records for faster testing
                limited_data = report_data[:20] if len(report_data) > 20 else report_data
                print(f"DEBUG: Limited to {len(limited_data)} records for testing")
                
                # Batch query all participants for all viajes at once
                viaje_ids = [data_row[0] for data_row in limited_data if len(data_row) > 0]
                print(f"DEBUG: Batch querying participants for {len(viaje_ids)} viajes")
                
                all_participants = self._get_participants_batch(viaje_ids)
                print(f"DEBUG: Retrieved {len(all_participants)} total participants")
                
                for i, data_row in enumerate(limited_data):
                    if i % 5 == 0:  # Log progress every 5 records
                        print(f"DEBUG: Processing record {i+1}/{len(limited_data)}")
                    
                    try:
                        # Main data row
                        correlativo = str(data_row[7])[-6:] if data_row[7] else ''  # Extract last 6 digits
                        fecha_crono = data_row[2] if len(data_row) > 2 else ''
                        via = data_row[9] if len(data_row) > 9 else ''
                        destino = data_row[5] if len(data_row) > 5 else ''
                        observacion = data_row[10] if len(data_row) > 10 else ''
                        
                        print(f"DEBUG: Record {i+1} data - correlativo: {correlativo}, fecha: {fecha_crono}, via: {via}")
                        
                        # Convert datetime.date objects to strings
                        if hasattr(fecha_crono, 'strftime'):
                            fecha_crono = fecha_crono.strftime('%d/%m/%Y')
                            print(f"DEBUG: Converted fecha_crono to: {fecha_crono}")
                        
                        # Get participants for this viaje from batch results
                        id_viaje = data_row[0] if len(data_row) > 0 else 0
                        participants = all_participants.get(id_viaje, [])
                        print(f"DEBUG: Got {len(participants)} participants for viaje {id_viaje}")
                        
                        # Create participant text
                        participant_text = ""
                        for participant in participants:
                            condicion = participant.get('des_condicion', '')
                            nombre = participant.get('c_descontrat', '')
                            tipo_doc = participant.get('tipo_documento', '')
                            num_doc = participant.get('doc', '')
                            participant_text += f"{condicion}: {nombre}\n{tipo_doc}: {num_doc}\n"
                        
                        print(f"DEBUG: Participant text: {participant_text.strip()}")
                        
                        # Add data row
                        row = data_table.add_row()
                        print(f"DEBUG: Added row {i+1} to table")
                        
                        row.cells[0].text = str(correlativo)
                        row.cells[1].text = str(fecha_crono)
                        row.cells[2].text = participant_text.strip()
                        row.cells[3].text = str(via) if via else ''
                        row.cells[4].text = str(destino) if destino else ''
                        row.cells[5].text = str(observacion) if observacion else ''
                        
                        print(f"DEBUG: Populated row {i+1} cells")
                        
                        # Center align NRO, FECHA, VIA, DESTINO
                        for i in [0, 1, 3, 4]:
                            for paragraph in row.cells[i].paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Left align PARTICIPANTES and OBSERVACIONES
                        for i in [2, 5]:
                            for paragraph in row.cells[i].paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        print(f"DEBUG: Finished styling row {i+1}")
                    
                    except Exception as e:
                        print(f"DEBUG: Error processing record {i+1}: {e}")
                        # Add error row instead of crashing
                        row = data_table.add_row()
                        row.cells[0].text = f"ERROR-{i+1}"
                        row.cells[1].text = "ERROR"
                        row.cells[2].text = f"Error processing record: {str(e)}"
                        row.cells[3].text = ""
                        row.cells[4].text = ""
                        row.cells[5].text = ""
                
                print(f"DEBUG: Finished processing data rows. Table now has {len(data_table.rows)} rows")
            else:
                # Add "No se encontraron registros" message
                row = data_table.add_row()
                row.cells[0].merge(row.cells[5])
                row.cells[0].text = "No se encontraron registros para el período especificado"
                for paragraph in row.cells[0].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            print("DEBUG: Saving document to buffer...")
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            print("DEBUG: Document saved successfully")
            
            # Create response
            filename = f"INDICE_CRONOLOGICO_PERMISOS_DE_VIAJE_{anio}.docx"
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            print("DEBUG: Response created successfully")
            return response
            
        except Exception as e:
            print(f"DEBUG: Error in Word report generation: {str(e)}")
            import traceback
            traceback.print_exc()
            return HttpResponse(
                json.dumps({'error': f'Error generating Word report: {str(e)}'}),
                content_type='application/json',
                status=500
            ) 